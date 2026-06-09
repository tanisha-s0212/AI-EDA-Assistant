from __future__ import annotations

import base64
import csv
import hashlib
import io
import json
import logging
import os
import re
import secrets
import time
import traceback
import uuid
import warnings
from concurrent.futures import ThreadPoolExecutor, as_completed
from html import escape
from math import erf, sqrt
from datetime import date, datetime, time as dt_time
from decimal import Decimal
from pathlib import Path
from typing import Any, Literal

import joblib
import matplotlib
import numpy as np
import pandas as pd
import polars as pl
import psycopg
try:
    import pyarrow.parquet as pq
except Exception:  # pragma: no cover - optional runtime dependency with friendly parquet errors
    pq = None
try:
    from dateutil import parser as date_parser
except Exception:  # pragma: no cover - pandas normally installs python-dateutil
    date_parser = None
from fastapi import APIRouter, FastAPI, File, Form, HTTPException, Query, Request, Response, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.gzip import GZipMiddleware
from fastapi.responses import JSONResponse, Response
from pydantic import BaseModel, Field
from psycopg.rows import dict_row
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from sklearn.base import clone
from sklearn.compose import ColumnTransformer
from sklearn.ensemble import GradientBoostingClassifier, GradientBoostingRegressor, RandomForestClassifier, RandomForestRegressor
from sklearn.impute import SimpleImputer
from sklearn.inspection import permutation_importance
from sklearn.linear_model import ElasticNet, Lasso, LinearRegression, LogisticRegression, Ridge
from sklearn.metrics import accuracy_score, f1_score, mean_absolute_error, mean_squared_error, precision_score, r2_score, recall_score
from sklearn.model_selection import KFold, LeaveOneOut, StratifiedKFold, cross_val_score, cross_validate, train_test_split
from sklearn.neighbors import KNeighborsClassifier, KNeighborsRegressor
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import LabelEncoder, OneHotEncoder, StandardScaler
from sklearn.svm import SVC, SVR
from sklearn.tree import DecisionTreeClassifier, DecisionTreeRegressor

try:
    from xgboost import XGBRegressor
except Exception:  # pragma: no cover - optional production dependency
    XGBRegressor = None

try:
    from xgboost import XGBClassifier
except Exception:  # pragma: no cover - optional production dependency
    XGBClassifier = None

try:
    from lightgbm import LGBMRegressor
except Exception:  # pragma: no cover - optional production dependency
    LGBMRegressor = None

try:
    from prophet import Prophet
except Exception:  # pragma: no cover - optional production dependency
    Prophet = None

try:
    from statsmodels.tsa.statespace.sarimax import SARIMAX
except Exception:  # pragma: no cover - optional production dependency
    SARIMAX = None

try:
    from pmdarima import auto_arima
except Exception:  # pragma: no cover - optional production dependency
    auto_arima = None

try:
    import optuna
except Exception:  # pragma: no cover - optional production dependency
    optuna = None

warnings.filterwarnings('ignore')
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from dtype_inference import LARGE_COL_CUTOFF, RANDOM_STATE, SAMPLE_SIZE, dtype_review_flags, dtype_summary_report, infer_universal_dtypes
from ml_forecast_pipeline import run_full_pipeline

BASE_DIR = Path(__file__).resolve().parent
REPO_ROOT = BASE_DIR.parent


def load_env_file(path: Path, *, override: bool = False) -> None:
    if not path.exists():
        return

    for raw_line in path.read_text(encoding='utf-8').splitlines():
        line = raw_line.strip()
        if not line or line.startswith('#') or '=' not in line:
            continue
        key, value = line.split('=', 1)
        key = key.strip()
        value = value.strip()
        if not key or (not override and key in os.environ):
            continue
        if len(value) >= 2 and value[0] == value[-1] and value[0] in {'"', "'"}:
            value = value[1:-1]
        os.environ[key] = value


load_env_file(BASE_DIR / '.env')
load_env_file(REPO_ROOT / 'agentic-layer' / '.env')

MODEL_DIR = BASE_DIR / 'models'
MODEL_DIR.mkdir(exist_ok=True)
DATASET_DIR = BASE_DIR / 'datasets'
DATASET_DIR.mkdir(exist_ok=True)
RUNTIME_TEMP_DIR = BASE_DIR / 'tmp'
RUNTIME_TEMP_DIR.mkdir(exist_ok=True)
LOG_DIR = BASE_DIR / 'logs'
LOG_DIR.mkdir(exist_ok=True)
os.environ.setdefault('TMP', str(RUNTIME_TEMP_DIR))
os.environ.setdefault('TEMP', str(RUNTIME_TEMP_DIR))
os.environ.setdefault('TMPDIR', str(RUNTIME_TEMP_DIR))
ACTIVITY_DATABASE_URL = os.environ.get('ACTIVITY_DATABASE_URL') or os.environ.get('DATABASE_URL') or 'postgresql://postgres:postgres@localhost:5432/ai_eda_assistant'
ACTIVITY_DB_REQUIRED = os.environ.get('ACTIVITY_DB_REQUIRED', 'false').strip().lower() == 'true'
ACTIVITY_DB_CONNECT_TIMEOUT = int(os.environ.get('ACTIVITY_DB_CONNECT_TIMEOUT', '1'))
TRAINING_N_JOBS = 1
ACTIVITY_DB_AVAILABLE = False
EMAIL_REGEX = re.compile(r'^[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}$', re.IGNORECASE)
PROFILE_IMAGE_MAX_BYTES = 1_500_000
PROFILE_IMAGE_CONTENT_TYPES = {'image/png', 'image/jpeg', 'image/webp', 'image/gif'}
PASSWORD_HASH_ITERATIONS = 600_000
SESSION_COOKIE_NAME = 'ai_eda_session'
SESSION_DURATION_SECONDS = 60 * 60 * 24 * 7
SESSION_MAX_AGE = SESSION_DURATION_SECONDS
SESSION_COOKIE_SECURE = os.environ.get('SESSION_COOKIE_SECURE', 'false').strip().lower() == 'true'
SESSION_COOKIE_SAMESITE = os.environ.get('SESSION_COOKIE_SAMESITE', 'lax').strip().lower() or 'lax'
SESSION_COOKIE_DOMAIN = os.environ.get('SESSION_COOKIE_DOMAIN') or None
ENABLE_PLOTLY_STATIC_EXPORT = os.environ.get('ENABLE_PLOTLY_STATIC_EXPORT', 'false').strip().lower() == 'true'

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s %(levelname)s %(message)s',
    handlers=[logging.FileHandler(LOG_DIR / 'backend.log', encoding='utf-8'), logging.StreamHandler()],
)
logger = logging.getLogger(__name__)

def parse_allowed_origins(value: str | None) -> list[str]:
    default_origins = [
        'http://localhost:3000',
        'http://127.0.0.1:3000',
        'http://localhost:3001',
        'http://127.0.0.1:3001',
    ]

    if not value:
        return default_origins

    origins = [origin.strip() for origin in value.split(',') if origin.strip()]
    if '*' in origins:
        return ['*']
    return origins or default_origins


allowed_origins = parse_allowed_origins(os.environ.get('CORS_ALLOWED_ORIGINS'))
allow_credentials = '*' not in allowed_origins

app = FastAPI(title='AI-Assisted EDA & ML Backend', version='3.0.0')
app.add_middleware(GZipMiddleware, minimum_size=1024)
app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=allow_credentials,
    allow_methods=['*'],
    allow_headers=['*'],
)
router = APIRouter(prefix='/api')
MODEL_CACHE: dict[str, dict[str, Any]] = {}
DATASET_CACHE: dict[str, dict[str, Any]] = {}
SESSION_STATE: dict[str, dict[str, Any]] = {}


@app.on_event('startup')
def startup_event() -> None:
    global ACTIVITY_DB_AVAILABLE
    try:
        init_activity_db()
        ACTIVITY_DB_AVAILABLE = True
        logger.info('Activity database is available.')
    except Exception:
        ACTIVITY_DB_AVAILABLE = False
        if ACTIVITY_DB_REQUIRED:
            raise
        logger.exception(
            'Activity database is unavailable. Continuing without persisted activity logging because ACTIVITY_DB_REQUIRED is false.'
        )

ProblemType = Literal['regression', 'classification']
TrainingMode = Literal['fast', 'balanced']

LARGE_DATASET_ROW_THRESHOLD = 20_000
VERY_LARGE_DATASET_ROW_THRESHOLD = 50_000
CV_SAMPLE_LIMIT = 3_000
VERY_LARGE_CV_SAMPLE_LIMIT = 1_500
TRAIN_SAMPLE_LIMIT = 30_000
VERY_LARGE_TRAIN_SAMPLE_LIMIT = 15_000
IMPORTANCE_SAMPLE_LIMIT = 800
VERY_LARGE_IMPORTANCE_SAMPLE_LIMIT = 300
DATASET_PREVIEW_ROW_LIMIT = 5_000
EDA_MAX_MISSINGNESS_COLUMNS = 30
EDA_MISSINGNESS_BUCKETS = 60
UPLOAD_READ_CHUNK_SIZE = 4 * 1024 * 1024
EDA_MAX_NUMERIC_CHARTS = 8
EDA_MAX_CATEGORICAL_CHARTS = 8
EDA_MAX_CATEGORY_BARS = 10
EDA_MAX_INTERACTION_COLUMNS = 40
EDA_MAX_INTERACTION_PAIRS = 3
MAX_UPLOAD_SIZE_BYTES = 512 * 1024 * 1024


def utc_now_iso() -> str:
    return datetime.utcnow().isoformat()


def get_activity_connection() -> psycopg.Connection:
    connection = psycopg.connect(
        ACTIVITY_DATABASE_URL,
        row_factory=dict_row,
        connect_timeout=ACTIVITY_DB_CONNECT_TIMEOUT,
    )
    connection.autocommit = True
    return connection


def init_activity_db() -> None:
    with get_activity_connection() as connection:
        connection.execute(
            '''
            CREATE TABLE IF NOT EXISTS app_users (
                id BIGSERIAL PRIMARY KEY,
                user_id TEXT NOT NULL UNIQUE,
                username TEXT NOT NULL,
                email TEXT NOT NULL UNIQUE,
                password_hash TEXT,
                profile_image_data_url TEXT,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL,
                last_login_at TEXT NOT NULL
            )
            '''
        )
        try:
            connection.execute('ALTER TABLE app_users ADD COLUMN IF NOT EXISTS password_hash TEXT')
        except Exception:
            logger.exception('Failed to ensure password_hash column on app_users.')
        try:
            connection.execute('ALTER TABLE app_users ADD COLUMN IF NOT EXISTS profile_image_data_url TEXT')
        except Exception:
            logger.exception('Failed to ensure profile_image_data_url column on app_users.')
        connection.execute('CREATE INDEX IF NOT EXISTS idx_app_users_email ON app_users (email)')
        connection.execute('CREATE INDEX IF NOT EXISTS idx_app_users_last_login_at ON app_users (last_login_at DESC)')
        connection.execute(
            '''
            CREATE TABLE IF NOT EXISTS app_user_sessions (
                id BIGSERIAL PRIMARY KEY,
                session_id TEXT NOT NULL UNIQUE,
                user_id TEXT NOT NULL,
                session_token_hash TEXT NOT NULL UNIQUE,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL,
                expires_at TEXT NOT NULL,
                revoked_at TEXT,
                client_ip TEXT,
                user_agent TEXT
            )
            '''
        )
        connection.execute('CREATE INDEX IF NOT EXISTS idx_app_user_sessions_user_id ON app_user_sessions (user_id)')
        connection.execute('CREATE INDEX IF NOT EXISTS idx_app_user_sessions_expires_at ON app_user_sessions (expires_at DESC)')
        connection.execute(
            '''
            CREATE TABLE IF NOT EXISTS user_activities (
                id BIGSERIAL PRIMARY KEY,
                activity_id TEXT NOT NULL UNIQUE,
                created_at TEXT NOT NULL,
                client_session_id TEXT,
                server_session_id TEXT,
                dataset_id TEXT,
                model_id TEXT,
                activity_type TEXT NOT NULL,
                action TEXT NOT NULL,
                status TEXT NOT NULL,
                api_path TEXT,
                http_method TEXT,
                status_code INTEGER,
                duration_ms REAL,
                file_name TEXT,
                detail TEXT,
                metadata_json TEXT,
                client_ip TEXT,
                user_agent TEXT
            )
            '''
        )
        connection.execute('CREATE INDEX IF NOT EXISTS idx_user_activities_created_at ON user_activities (created_at DESC)')
        connection.execute('CREATE INDEX IF NOT EXISTS idx_user_activities_client_session_id ON user_activities (client_session_id)')
        connection.execute('CREATE INDEX IF NOT EXISTS idx_user_activities_dataset_id ON user_activities (dataset_id)')
        connection.execute('CREATE INDEX IF NOT EXISTS idx_user_activities_server_session_id ON user_activities (server_session_id)')
        connection.execute('CREATE INDEX IF NOT EXISTS idx_user_activities_action ON user_activities (action)')
        connection.execute(
            '''
            CREATE TABLE IF NOT EXISTS agentic_runs (
                run_id TEXT PRIMARY KEY,
                session_id TEXT NOT NULL,
                status TEXT NOT NULL,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL
            )
            '''
        )
        connection.execute('CREATE INDEX IF NOT EXISTS idx_agentic_runs_session_id ON agentic_runs (session_id)')
        connection.execute('CREATE INDEX IF NOT EXISTS idx_agentic_runs_updated_at ON agentic_runs (updated_at DESC)')
        connection.execute(
            '''
            CREATE TABLE IF NOT EXISTS agentic_steps (
                step_id TEXT PRIMARY KEY,
                run_id TEXT NOT NULL REFERENCES agentic_runs(run_id) ON DELETE CASCADE,
                step_name TEXT NOT NULL,
                status TEXT NOT NULL,
                result_json JSONB,
                executed_at TEXT NOT NULL
            )
            '''
        )
        connection.execute('CREATE INDEX IF NOT EXISTS idx_agentic_steps_run_id ON agentic_steps (run_id)')
        connection.execute('CREATE INDEX IF NOT EXISTS idx_agentic_steps_executed_at ON agentic_steps (executed_at DESC)')
        connection.execute(
            '''
            CREATE TABLE IF NOT EXISTS agentic_decisions (
                decision_id TEXT PRIMARY KEY,
                run_id TEXT NOT NULL REFERENCES agentic_runs(run_id) ON DELETE CASCADE,
                step_id TEXT,
                decision TEXT NOT NULL,
                reason TEXT,
                decided_at TEXT NOT NULL
            )
            '''
        )
        connection.execute('ALTER TABLE agentic_decisions ADD COLUMN IF NOT EXISTS decision_id TEXT')
        connection.execute('ALTER TABLE agentic_decisions ADD COLUMN IF NOT EXISTS run_id TEXT')
        connection.execute('ALTER TABLE agentic_decisions ADD COLUMN IF NOT EXISTS step_id TEXT')
        connection.execute('ALTER TABLE agentic_decisions ADD COLUMN IF NOT EXISTS reason TEXT')
        connection.execute('ALTER TABLE agentic_decisions ADD COLUMN IF NOT EXISTS decided_at TEXT')
        connection.execute('CREATE UNIQUE INDEX IF NOT EXISTS idx_agentic_decisions_decision_id ON agentic_decisions (decision_id)')
        connection.execute('CREATE INDEX IF NOT EXISTS idx_agentic_decisions_run_id ON agentic_decisions (run_id)')
        connection.execute('CREATE INDEX IF NOT EXISTS idx_agentic_decisions_decided_at ON agentic_decisions (decided_at DESC)')
        connection.execute(
            '''
            CREATE TABLE IF NOT EXISTS agentic_audit (
                audit_id TEXT PRIMARY KEY,
                run_id TEXT NOT NULL REFERENCES agentic_runs(run_id) ON DELETE CASCADE,
                event_type TEXT NOT NULL,
                payload_json JSONB,
                created_at TEXT NOT NULL
            )
            '''
        )
        connection.execute('CREATE INDEX IF NOT EXISTS idx_agentic_audit_run_id ON agentic_audit (run_id)')
        connection.execute('CREATE INDEX IF NOT EXISTS idx_agentic_audit_created_at ON agentic_audit (created_at DESC)')
        connection.execute(
            '''
            CREATE TABLE IF NOT EXISTS loss_forecast_results (
                id TEXT PRIMARY KEY,
                session_id TEXT NOT NULL,
                period DATE NOT NULL,
                revenue_loss DOUBLE PRECISION NOT NULL DEFAULT 0,
                operational_loss DOUBLE PRECISION NOT NULL DEFAULT 0,
                inventory_loss DOUBLE PRECISION NOT NULL DEFAULT 0,
                discount_loss DOUBLE PRECISION NOT NULL DEFAULT 0,
                total_loss DOUBLE PRECISION NOT NULL DEFAULT 0,
                lower_bound DOUBLE PRECISION,
                upper_bound DOUBLE PRECISION,
                loss_risk_score DOUBLE PRECISION NOT NULL DEFAULT 0,
                risk_label TEXT NOT NULL,
                segment TEXT,
                created_at TEXT NOT NULL
            )
            '''
        )
        connection.execute('CREATE INDEX IF NOT EXISTS idx_loss_forecast_session_period ON loss_forecast_results (session_id, period)')
        connection.execute(
            '''
            CREATE TABLE IF NOT EXISTS profit_forecast_results (
                id TEXT PRIMARY KEY,
                session_id TEXT NOT NULL,
                period DATE NOT NULL,
                forecasted_revenue DOUBLE PRECISION NOT NULL DEFAULT 0,
                forecasted_cogs DOUBLE PRECISION NOT NULL DEFAULT 0,
                gross_profit DOUBLE PRECISION NOT NULL DEFAULT 0,
                operating_expenses DOUBLE PRECISION NOT NULL DEFAULT 0,
                total_losses DOUBLE PRECISION NOT NULL DEFAULT 0,
                net_profit DOUBLE PRECISION NOT NULL DEFAULT 0,
                gross_margin_pct DOUBLE PRECISION NOT NULL DEFAULT 0,
                net_margin_pct DOUBLE PRECISION NOT NULL DEFAULT 0,
                scenario TEXT NOT NULL,
                created_at TEXT NOT NULL
            )
            '''
        )
        connection.execute('CREATE INDEX IF NOT EXISTS idx_profit_forecast_session_scenario_period ON profit_forecast_results (session_id, scenario, period)')

        # ---- TS Forecast migrations ----
        connection.execute('ALTER TABLE agentic_runs ADD COLUMN IF NOT EXISTS ts_best_model VARCHAR(100)')
        connection.execute('ALTER TABLE agentic_runs ADD COLUMN IF NOT EXISTS ts_forecast_mae FLOAT')
        connection.execute('ALTER TABLE agentic_runs ADD COLUMN IF NOT EXISTS ts_forecast_smape FLOAT')
        connection.execute('ALTER TABLE agentic_runs ADD COLUMN IF NOT EXISTS ts_stationarity_status VARCHAR(50)')
        connection.execute(
            '''
            CREATE TABLE IF NOT EXISTS ts_forecast_results (
                id                  SERIAL PRIMARY KEY,
                dataset_id          VARCHAR(255) NOT NULL,
                best_model          VARCHAR(100),
                mae                 FLOAT,
                rmse                FLOAT,
                mape                FLOAT,
                smape               FLOAT,
                model_comparison    JSONB,
                future_forecast     JSONB,
                stationarity_report JSONB,
                insight             JSONB,
                created_at          TIMESTAMP DEFAULT NOW()
            )
            '''
        )
        connection.execute('CREATE INDEX IF NOT EXISTS idx_ts_forecast_results_dataset_id ON ts_forecast_results (dataset_id)')
        connection.execute(
            '''
            CREATE TABLE IF NOT EXISTS workspace_context (
                id            SERIAL PRIMARY KEY,
                dataset_id    VARCHAR(255) NOT NULL,
                context_key   VARCHAR(100) NOT NULL,
                context_value JSONB,
                updated_at    TIMESTAMP DEFAULT NOW(),
                UNIQUE(dataset_id, context_key)
            )
            '''
        )
        connection.execute('CREATE INDEX IF NOT EXISTS idx_workspace_context_dataset_key ON workspace_context (dataset_id, context_key)')


def sanitize_metadata(metadata: dict[str, Any] | None) -> str | None:
    if not metadata:
        return None
    return json.dumps(safe_serialize(metadata), default=str)


def get_client_session_id(request: Request | None) -> str | None:
    if request is None:
        return None
    return request.headers.get('x-client-session-id') or None


# ── TS Forecast helper utilities ──────────────────────────────

def get_workspace_context(dataset_id: str, context_key: str) -> Any:
    """Read shared context from workspace_context table."""
    if not ACTIVITY_DB_AVAILABLE:
        return None
    try:
        with get_activity_connection() as conn:
            row = conn.execute(
                'SELECT context_value FROM workspace_context WHERE dataset_id = %s AND context_key = %s',
                [dataset_id, context_key]
            ).fetchone()
            return row['context_value'] if row else None
    except Exception:
        logger.exception('get_workspace_context failed dataset_id=%s key=%s', dataset_id, context_key)
        return None


def save_workspace_context(dataset_id: str, context_key: str, context_value: dict[str, Any]) -> None:
    """Write shared context to workspace_context table."""
    if not ACTIVITY_DB_AVAILABLE:
        return
    try:
        with get_activity_connection() as conn:
            conn.execute(
                '''INSERT INTO workspace_context (dataset_id, context_key, context_value)
                   VALUES (%s, %s, %s)
                   ON CONFLICT (dataset_id, context_key)
                   DO UPDATE SET context_value = EXCLUDED.context_value, updated_at = NOW()''',
                [dataset_id, context_key, json.dumps(context_value)]
            )
    except Exception:
        logger.exception('save_workspace_context failed dataset_id=%s key=%s', dataset_id, context_key)


def get_ts_output_dir(dataset_id: str) -> str:
    """Resolve ts_forecast_output/ folder path relative to active dataset."""
    dataset_entry = DATASET_CACHE.get(dataset_id)
    if dataset_entry is None:
        base_dir = str(BASE_DIR / 'ts_output')
    else:
        parquet_path = dataset_entry.get('parquet_path') or dataset_entry.get('frame_path')
        base_dir = str(Path(str(parquet_path)).parent) if parquet_path else str(BASE_DIR / 'ts_output')
    output_dir = os.path.join(base_dir, 'ts_forecast_output')
    os.makedirs(output_dir, exist_ok=True)
    return output_dir


def load_ts_dataset(dataset_id: str) -> tuple[pd.DataFrame, str, str]:
    """Load active dataset, auto-detect date and target columns."""
    dataset_entry = DATASET_CACHE.get(dataset_id)
    if dataset_entry is None:
        raise HTTPException(status_code=400, detail='Cached dataset not found. Please upload the file again.')
    available_columns = list(dataset_entry['columns'])
    date_candidates = [c for c in available_columns if any(k in c.lower() for k in ['date', 'week', 'month', 'period', 'start', 'time'])]
    target_candidates = [c for c in available_columns if any(k in c.lower() for k in ['sale_free', 'sale_value', 'total_value', 'revenue'])]
    numeric_cols = [c for c in available_columns if any(
        col['name'] == c and col.get('role') == 'numeric' for col in dataset_entry.get('column_info', []) if isinstance(col, dict)
    )]
    if not date_candidates:
        date_candidates = [c for c in available_columns if any(
            col['name'] == c and col.get('role') in ('datetime', 'date') for col in dataset_entry.get('column_info', []) if isinstance(col, dict)
        )]
    if not date_candidates:
        raise ValueError(f'No date column found. Available columns: {available_columns}')
    if not target_candidates:
        target_candidates = numeric_cols if numeric_cols else [c for c in available_columns if c not in date_candidates]
    if not target_candidates:
        raise ValueError(f'No target column found. Available columns: {available_columns}')
    date_col = date_candidates[0]
    target_col = target_candidates[0]
    series_frame, freq, period_label = prepare_sales_series_from_cached_dataset(dataset_entry, date_col, target_col)
    series_frame.rename(columns={'period': date_col, 'sales': target_col}, inplace=True)
    return series_frame, date_col, target_col


# ── Frequency detection ───────────────────────────────────────

def detect_ts_frequency(df: pd.DataFrame, date_col: str) -> tuple[str, int]:
    """Detect weekly or monthly from median date gaps."""
    gaps = df[date_col].sort_values().diff().dt.days.dropna()
    median_gap = gaps.median() if not gaps.empty else 30
    if pd.isna(median_gap):
        median_gap = 30
    if median_gap <= 10:
        return 'weekly', 52
    elif median_gap <= 40:
        return 'monthly', 12
    else:
        logger.warning('Unusual frequency: median_gap=%s days. Falling back to monthly.', median_gap)
        return 'monthly', 12


# ── Stationarity check ────────────────────────────────────────

def check_stationarity(series: pd.Series, frequency: str) -> dict[str, Any]:
    """Run ADF + KPSS tests and return stationarity report dict."""
    from statsmodels.tsa.stattools import adfuller, kpss
    clean = series.dropna()
    clean = clean[clean > 0]
    if len(clean) < 4:
        return {
            'status': 'insufficient_data',
            'adf_pvalue': None, 'kpss_pvalue': None,
            'note': 'Insufficient data for stationarity tests (need >= 4 non-zero values).',
            'recommended_model': 'Prophet',
            'differencing_required': True
        }
    try:
        adf_result = adfuller(clean, autolag='AIC')
        adf_p = float(adf_result[1])
    except Exception as exc:
        logger.warning('ADF test failed: %s', exc)
        adf_p = 0.5
    try:
        kpss_result = kpss(clean, regression='c', nlags='auto')
        kpss_p = float(kpss_result[1])
    except Exception as exc:
        logger.warning('KPSS test failed: %s', exc)
        kpss_p = 0.5
    if adf_p < 0.05 and kpss_p > 0.05:
        status = 'stationary'
        note = 'Series is stationary. SARIMA, Prophet and Holt-Winters all applicable.'
        recommendation = 'SARIMA'
    elif adf_p >= 0.05 and kpss_p <= 0.05:
        status = 'non_stationary'
        note = 'Series is non-stationary. Prophet recommended; SARIMA will auto-difference.'
        recommendation = 'Prophet'
    else:
        status = 'trend_stationary'
        note = 'Trend-stationary series detected. Holt-Winters with damped trend recommended.'
        recommendation = 'HoltWinters'
    return {
        'status': status,
        'adf_pvalue': round(float(adf_p), 4),
        'kpss_pvalue': round(float(kpss_p), 4),
        'note': note,
        'recommended_model': recommendation,
        'differencing_required': adf_p >= 0.05
    }


# ── Metrics ────────────────────────────────────────────────────

def compute_ts_metrics(actual: np.ndarray, predicted: np.ndarray) -> dict[str, float | None]:
    """Compute MAE, RMSE, MAPE, SMAPE."""
    actual_a = np.array(actual, dtype=float)
    predicted_a = np.array(predicted, dtype=float)
    mae = float(np.mean(np.abs(actual_a - predicted_a)))
    rmse = float(np.sqrt(np.mean((actual_a - predicted_a) ** 2)))
    mask = actual_a != 0
    if mask.sum() > 0:
        mape = float(np.mean(np.abs((actual_a[mask] - predicted_a[mask]) / actual_a[mask])) * 100)
    else:
        mape = None
    denom = np.abs(actual_a) + np.abs(predicted_a)
    smape_val = float(np.mean(2 * np.abs(actual_a - predicted_a) / np.where(denom == 0, 1e-8, denom)) * 100)
    return {'mae': round(mae, 2), 'rmse': round(rmse, 2), 'mape': round(mape, 2) if mape is not None else None, 'smape': round(smape_val, 2)}


# ── 3-model training ──────────────────────────────────────────

def train_all_ts_models(df: pd.DataFrame, target_col: str, date_col: str, frequency: str, freq_period: int, training_split: float = 0.8, horizon: int = 3) -> tuple[dict[str, Any], np.ndarray, np.ndarray, pd.DataFrame, pd.DataFrame, pd.DataFrame, str]:
    """Train SARIMA, Prophet, HoltWinters. Returns results dict and split data."""
    first_nonzero_idx = df[df[target_col] > 0].index[0]
    clean_df = df.loc[first_nonzero_idx:].copy().reset_index(drop=True)
    first_nonzero_date = str(clean_df[date_col].iloc[0])[:10]
    if len(clean_df) < 24:
        raise ValueError(f'Insufficient data after trimming: {len(clean_df)} rows. Min 24 required.')
    split_idx = int(len(clean_df) * training_split)
    train = clean_df.iloc[:split_idx]
    test = clean_df.iloc[split_idx:]
    y_train = train[target_col].values.astype(float)
    y_test = test[target_col].values.astype(float)
    max_period = len(y_train) // 2
    if freq_period > max_period:
        logger.warning('freq_period %s reduced to %s - insufficient data.', freq_period, max_period)
        freq_period = max_period
    results: dict[str, Any] = {}

    # SARIMA
    try:
        import pmdarima as pm
        sarima_model = pm.auto_arima(
            y_train, seasonal=True, m=freq_period,
            stepwise=True, suppress_warnings=True,
            error_action='ignore', max_p=3, max_q=3,
            max_P=2, max_Q=2, information_criterion='aic',
            random_state=42
        )
        sarima_pred, conf_int = sarima_model.predict(n_periods=len(y_test), return_conf_int=True)
        results['SARIMA'] = {
            'status': 'completed', 'model_object': sarima_model,
            'predictions': sarima_pred.tolist(),
            'conf_int_lower': conf_int[:, 0].tolist(),
            'conf_int_upper': conf_int[:, 1].tolist(),
            'order': str(sarima_model.order),
            'seasonal_order': str(sarima_model.seasonal_order),
            'aic': round(float(sarima_model.aic()), 2),
            **compute_ts_metrics(y_test, sarima_pred)
        }
    except Exception as exc:
        logger.warning('SARIMA training failed: %s', exc)
        results['SARIMA'] = {'status': 'failed', 'error': str(exc), 'mae': None, 'rmse': None, 'mape': None, 'smape': None}

    # Prophet
    try:
        from prophet import Prophet
        prophet_train = pd.DataFrame({'ds': train[date_col].values, 'y': y_train})
        prophet_model = Prophet(seasonality_mode='multiplicative', yearly_seasonality=True, weekly_seasonality=(frequency == 'weekly'), changepoint_prior_scale=0.05, seasonality_prior_scale=10, interval_width=0.95)
        prophet_model.fit(prophet_train)
        freq_str = 'W' if frequency == 'weekly' else 'MS'
        future = prophet_model.make_future_dataframe(periods=len(y_test), freq=freq_str)
        fc = prophet_model.predict(future)
        prophet_pred = fc['yhat'].iloc[-len(y_test):].values
        results['Prophet'] = {
            'status': 'completed', 'model_object': prophet_model,
            'predictions': prophet_pred.tolist(),
            'conf_int_lower': fc['yhat_lower'].iloc[-len(y_test):].tolist(),
            'conf_int_upper': fc['yhat_upper'].iloc[-len(y_test):].tolist(),
            **compute_ts_metrics(y_test, prophet_pred)
        }
    except Exception as exc:
        logger.warning('Prophet training failed: %s', exc)
        results['Prophet'] = {'status': 'failed', 'error': str(exc), 'mae': None, 'rmse': None, 'mape': None, 'smape': None}

    # Holt-Winters
    try:
        from statsmodels.tsa.holtwinters import ExponentialSmoothing
        has_zeros = (y_train == 0).any()
        has_neg = (y_train < 0).any()
        seasonal_type = 'add' if (has_zeros or has_neg) else 'mul'
        if has_zeros or has_neg:
            logger.warning('HoltWinters: additive mode - zeros/negatives detected.')
        hw_model = ExponentialSmoothing(y_train, trend='add', seasonal=seasonal_type, seasonal_periods=freq_period, damped_trend=True).fit(optimized=True)
        hw_pred = hw_model.forecast(len(y_test))
        sim = hw_model.simulate(len(y_test), repetitions=100, error='add', random_errors='bootstrap')
        lower = sim.quantile(0.025, axis=1).values
        upper = sim.quantile(0.975, axis=1).values
        results['HoltWinters'] = {
            'status': 'completed', 'model_object': hw_model,
            'predictions': hw_pred.tolist(),
            'conf_int_lower': lower.tolist(),
            'conf_int_upper': upper.tolist(),
            'alpha': round(float(hw_model.params['smoothing_level']), 4),
            'beta': round(float(hw_model.params['smoothing_trend']), 4),
            'gamma': round(float(hw_model.params['smoothing_seasonal']), 4),
            **compute_ts_metrics(y_test, hw_pred)
        }
    except Exception as exc:
        logger.warning('HoltWinters training failed: %s', exc)
        results['HoltWinters'] = {'status': 'failed', 'error': str(exc), 'mae': None, 'rmse': None, 'mape': None, 'smape': None}

    return results, y_train, y_test, train, test, clean_df, first_nonzero_date


# ── Auto-select best model ─────────────────────────────────────

def auto_select_ts_model(results: dict[str, Any]) -> tuple[str, dict[str, Any], str]:
    """Select best model by lowest SMAPE. Tiebreak: SARIMA > HoltWinters > Prophet."""
    completed = {k: v for k, v in results.items() if v['status'] == 'completed' and v.get('smape') is not None}
    if not completed:
        raise ValueError('All 3 TS models failed. Check data quality and logs.')
    ranked = sorted(completed.items(), key=lambda x: x[1]['smape'])
    best_name = ranked[0][0]
    best_smape = ranked[0][1]['smape']
    if len(ranked) > 1:
        second_smape = ranked[1][1]['smape']
        if abs(best_smape - second_smape) < 2.0:
            priority = ['SARIMA', 'HoltWinters', 'Prophet']
            for p in priority:
                if p in [r[0] for r in ranked[:2]]:
                    best_name = p
                    break
    others = ', '.join([f'{k}={v["smape"]}%' for k, v in completed.items() if k != best_name])
    reason = f'{best_name} selected by lowest SMAPE ({results[best_name]["smape"]}%). Compared: {others}'
    return best_name, results[best_name], reason


# ── Future forecast generation ─────────────────────────────────

def generate_ts_future_forecast(best_model_name: str, clean_df: pd.DataFrame, target_col: str, date_col: str, frequency: str, freq_period: int, horizon: int = 3) -> list[dict[str, Any]]:
    """Retrain best model on full dataset and generate horizon periods."""
    import pmdarima as pm
    from prophet import Prophet
    from statsmodels.tsa.holtwinters import ExponentialSmoothing
    y_full = clean_df[target_col].values.astype(float)
    last_date = clean_df[date_col].max()
    freq_str = 'W' if frequency == 'weekly' else 'MS'
    future_dates = pd.date_range(start=last_date, periods=horizon + 1, freq=freq_str)[1:]
    forecast_vals: list[float] = []
    lower_bounds: list[float] = []
    upper_bounds: list[float] = []
    if best_model_name == 'SARIMA':
        model = pm.auto_arima(y_full, seasonal=True, m=freq_period, stepwise=True, suppress_warnings=True, error_action='ignore', random_state=42)
        preds, ci = model.predict(n_periods=horizon, return_conf_int=True)
        forecast_vals = preds.tolist()
        lower_bounds = ci[:, 0].tolist()
        upper_bounds = ci[:, 1].tolist()
    elif best_model_name == 'Prophet':
        prophet_df = pd.DataFrame({'ds': clean_df[date_col].values, 'y': y_full})
        model = Prophet(seasonality_mode='multiplicative', yearly_seasonality=True, weekly_seasonality=(frequency == 'weekly'), interval_width=0.95)
        model.fit(prophet_df)
        future = model.make_future_dataframe(periods=horizon, freq=freq_str)
        fc = model.predict(future)
        forecast_vals = fc['yhat'].iloc[-horizon:].tolist()
        lower_bounds = fc['yhat_lower'].iloc[-horizon:].tolist()
        upper_bounds = fc['yhat_upper'].iloc[-horizon:].tolist()
    elif best_model_name == 'HoltWinters':
        has_zeros = (y_full == 0).any()
        has_neg = (y_full < 0).any()
        s_type = 'add' if (has_zeros or has_neg) else 'mul'
        model = ExponentialSmoothing(y_full, trend='add', seasonal=s_type, seasonal_periods=freq_period, damped_trend=True).fit(optimized=True)
        forecast_vals = model.forecast(horizon).tolist()
        sim = model.simulate(horizon, repetitions=100, error='add', random_errors='bootstrap')
        lower_bounds = sim.quantile(0.025, axis=1).tolist()
        upper_bounds = sim.quantile(0.975, axis=1).tolist()
    fmt = '%Y-%m-%d' if frequency == 'weekly' else '%Y-%m'
    label_t = 'Week of {}' if frequency == 'weekly' else 'Month of {}'
    return [{'period': label_t.format(d.strftime(fmt)), 'forecast': round(f, 2), 'lower': round(l, 2), 'upper': round(u, 2)} for d, f, l, u in zip(future_dates, forecast_vals, lower_bounds, upper_bounds)]


# ── Programmatic insight ──────────────────────────────────────

def generate_ts_insight(best_model_name: str, metrics: dict[str, Any], stationarity: dict[str, Any], model_comparison: list[dict[str, Any]]) -> dict[str, Any]:
    """Generate insight text programmatically from metric values. No LLM call."""
    smape = metrics.get('smape', 0)
    mae = metrics.get('mae', 0)
    if smape < 15:
        confidence = 'high'
        quality_note = 'strong forecast accuracy'
    elif smape < 30:
        confidence = 'medium'
        quality_note = 'moderate forecast accuracy'
    else:
        confidence = 'low'
        quality_note = 'high volatility detected'
    risk_flag = None
    if stationarity.get('status') == 'non_stationary':
        risk_flag = 'Non-stationary series — forecast confidence decreases over longer horizons.'
    elif smape > 50:
        risk_flag = f'SMAPE {smape}% indicates high variance — use forecast range not point values.'
    insight_text = f'{best_model_name} selected for sales forecasting with {quality_note} (SMAPE: {smape}%, MAE: {round(mae / 1e6, 2)}M). Series stationarity: {stationarity.get("status", "unknown")}.'
    return {
        'insight_text': insight_text,
        'risk_flag': risk_flag,
        'confidence': confidence,
        'best_model': best_model_name,
        'selection_metric': 'SMAPE',
        'stationarity_note': stationarity.get('note', ''),
        'top_model_smape': smape,
        'audit_trail': 'SARIMA, Prophet, HoltWinters compared via walk-forward validation using MAE, RMSE, MAPE, SMAPE. Best model selected by lowest SMAPE. Tiebreak priority: SARIMA > HoltWinters > Prophet.'
    }


# ── Write TS output files ─────────────────────────────────────

def write_ts_output_files(output_dir: str, frequency: str, freq_period: int, target_col: str, date_col: str, df: pd.DataFrame, clean_df: pd.DataFrame, train: pd.DataFrame, test: pd.DataFrame, stationarity: dict[str, Any], results: dict[str, Any], best_name: str, best_metrics: dict[str, Any], reason: str, future_forecast: list[dict[str, Any]], insight: dict[str, Any]) -> None:
    """Write all 7 JSON output files to ts_forecast_output/."""
    save_dir = output_dir
    with open(os.path.join(save_dir, 'ts_metadata.json'), 'w') as f:
        json.dump({'frequency': frequency, 'frequency_period': freq_period, 'target_col': target_col, 'date_col': date_col, 'total_rows': len(df), 'clean_rows': len(clean_df), 'training_rows': len(train), 'test_rows': len(test), 'training_split_pct': 80, 'horizon_periods': 3}, f, indent=2)
    with open(os.path.join(save_dir, 'stationarity_report.json'), 'w') as f:
        json.dump(stationarity, f, indent=2)
    comparison = []
    for model_name, v in results.items():
        row = {'model': model_name, 'status': v['status'], 'mae': v.get('mae'), 'rmse': v.get('rmse'), 'mape': v.get('mape'), 'smape': v.get('smape'), 'aic': v.get('aic'), 'order': v.get('order'), 'note': v.get('error', 'completed')}
        if model_name == 'HoltWinters':
            row['alpha'] = v.get('alpha')
            row['beta'] = v.get('beta')
            row['gamma'] = v.get('gamma')
        comparison.append(row)
    with open(os.path.join(save_dir, 'model_comparison.json'), 'w') as f:
        json.dump(comparison, f, indent=2)
    with open(os.path.join(save_dir, 'selected_model.json'), 'w') as f:
        json.dump({'model_name': best_name, 'smape': best_metrics['smape'], 'mae': best_metrics['mae'], 'rmse': best_metrics['rmse'], 'mape': best_metrics.get('mape'), 'selection_reason': reason, 'selection_metric': 'SMAPE', 'stationarity_status': stationarity.get('status'), 'recommended_by_stationarity': stationarity.get('recommended_model')}, f, indent=2)
    forecast_line = []
    for _, row in clean_df.iterrows():
        forecast_line.append({'period': str(row[date_col])[:10], 'actual': float(row[target_col]), 'backtest': None, 'forecast': None, 'lower': None, 'upper': None, 'type': 'actual'})
    best_preds = results[best_name].get('predictions', [])
    best_lower = results[best_name].get('conf_int_lower', [])
    best_upper = results[best_name].get('conf_int_upper', [])
    for i, row in enumerate(test.itertuples()):
        dt_val = str(getattr(row, date_col))[:10] if hasattr(row, date_col) else str(i)
        forecast_line.append({'period': dt_val, 'actual': None, 'backtest': round(best_preds[i], 2) if i < len(best_preds) else None, 'forecast': None, 'lower': round(best_lower[i], 2) if i < len(best_lower) else None, 'upper': round(best_upper[i], 2) if i < len(best_upper) else None, 'type': 'backtest'})
    for item in future_forecast:
        forecast_line.append({'period': item['period'], 'actual': None, 'backtest': None, 'forecast': item['forecast'], 'lower': item['lower'], 'upper': item['upper'], 'type': 'forecast'})
    with open(os.path.join(save_dir, 'forecast_line.json'), 'w') as f:
        json.dump(forecast_line, f, indent=2)
    values = [r['forecast'] for r in future_forecast]
    peak_val = max(values) if values else 0
    peak_period = future_forecast[values.index(peak_val)]['period'] if values else ''
    with open(os.path.join(save_dir, 'future_forecast_table.json'), 'w') as f:
        json.dump({'horizon_avg': round(sum(values) / len(values), 2) if values else 0, 'peak_value': peak_val, 'peak_period': peak_period, 'first_period': future_forecast[0]['period'] if future_forecast else '', 'last_period': future_forecast[-1]['period'] if future_forecast else '', 'latest_actual': float(clean_df[target_col].iloc[-1]) if len(clean_df) else 0, 'horizon_label': 'week horizon' if frequency == 'weekly' else 'month horizon', 'rows': future_forecast}, f, indent=2)
    with open(os.path.join(save_dir, 'ts_insight.json'), 'w') as f:
        json.dump(insight, f, indent=2)


def write_ts_to_postgres(dataset_id: str, best_name: str, best_metrics: dict[str, Any], results: dict[str, Any], stationarity: dict[str, Any], future_forecast: list[dict[str, Any]], insight: dict[str, Any]) -> None:
    """Write TS results to PostgreSQL."""
    if not ACTIVITY_DB_AVAILABLE:
        return
    try:
        with get_activity_connection() as conn:
            conn.execute(
                '''UPDATE agentic_runs SET
                   ts_best_model = %s, ts_forecast_mae = %s,
                   ts_forecast_smape = %s, ts_stationarity_status = %s,
                   updated_at = %s
                   WHERE session_id = %s''',
                [best_name, best_metrics.get('mae'), best_metrics.get('smape'),
                 stationarity.get('status'), datetime.utcnow().isoformat(), dataset_id]
            )
            model_comparison_json = [{'model': k, 'status': v['status'], 'mae': v.get('mae'), 'rmse': v.get('rmse'), 'mape': v.get('mape'), 'smape': v.get('smape'), 'note': v.get('error', 'completed')} for k, v in results.items()]
            conn.execute(
                '''INSERT INTO ts_forecast_results
                   (dataset_id, best_model, mae, rmse, mape, smape,
                    model_comparison, future_forecast, stationarity_report, insight)
                   VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)''',
                [dataset_id, best_name, best_metrics.get('mae'), best_metrics.get('rmse'),
                 best_metrics.get('mape'), best_metrics.get('smape'),
                 json.dumps(model_comparison_json), json.dumps(future_forecast),
                 json.dumps(stationarity), json.dumps(insight)]
            )
    except Exception:
        logger.exception('write_ts_to_postgres failed dataset_id=%s', dataset_id)


def record_activity(
    *,
    action: str,
    status: str,
    activity_type: str = 'workflow',
    request: Request | None = None,
    dataset_id: str | None = None,
    model_id: str | None = None,
    server_session_id: str | None = None,
    file_name: str | None = None,
    detail: str | None = None,
    metadata: dict[str, Any] | None = None,
    api_path: str | None = None,
    http_method: str | None = None,
    status_code: int | None = None,
    duration_ms: float | None = None,
) -> None:
    global ACTIVITY_DB_AVAILABLE

    if not ACTIVITY_DB_AVAILABLE:
        return

    try:
        with get_activity_connection() as connection:
            connection.execute(
                '''
                INSERT INTO user_activities (
                    activity_id,
                    created_at,
                    client_session_id,
                    server_session_id,
                    dataset_id,
                    model_id,
                    activity_type,
                    action,
                    status,
                    api_path,
                    http_method,
                    status_code,
                    duration_ms,
                    file_name,
                    detail,
                    metadata_json,
                    client_ip,
                    user_agent
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                ''',
                (
                    uuid.uuid4().hex,
                    utc_now_iso(),
                    get_client_session_id(request),
                    server_session_id,
                    dataset_id,
                    model_id,
                    activity_type,
                    action,
                    status,
                    api_path or (str(request.url.path) if request is not None else None),
                    http_method or (request.method if request is not None else None),
                    status_code,
                    duration_ms,
                    file_name,
                    detail,
                    sanitize_metadata(metadata),
                    request.client.host if request is not None and request.client is not None else None,
                    request.headers.get('user-agent') if request is not None else None,
                ),
            )
    except Exception:
        ACTIVITY_DB_AVAILABLE = False
        logger.exception('Failed to persist activity action=%s status=%s', action, status)


def get_session_id(dataset_id: str | None, session_id: str | None = None) -> str:
    if session_id:
        return session_id
    if dataset_id:
        return dataset_id
    return f'adhoc-{uuid.uuid4().hex[:8]}'


def normalize_email(value: str) -> str:
    return value.strip().lower()


def normalize_username(value: str) -> str:
    return ' '.join(value.strip().split())


def validate_login_payload(username: str, email: str) -> tuple[str, str]:
    normalized_username = normalize_username(username)
    normalized_email = normalize_email(email)

    if len(normalized_username) < 3:
        raise HTTPException(status_code=400, detail='Username must be at least 3 characters long.')
    if len(normalized_username) > 80:
        raise HTTPException(status_code=400, detail='Username must be 80 characters or fewer.')
    if not re.fullmatch(r"[A-Za-z0-9][A-Za-z0-9 ._'-]{1,79}", normalized_username):
        raise HTTPException(
            status_code=400,
            detail='Username may contain letters, numbers, spaces, periods, apostrophes, underscores, and hyphens.',
        )
    if not EMAIL_REGEX.fullmatch(normalized_email):
        raise HTTPException(status_code=400, detail='Enter a valid email address.')

    return normalized_username, normalized_email


def validate_password(password: str) -> str:
    if len(password) < 8:
        raise HTTPException(status_code=400, detail='Password must be at least 8 characters long.')
    if len(password) > 128:
        raise HTTPException(status_code=400, detail='Password must be 128 characters or fewer.')
    return password


def hash_session_token(token: str) -> str:
    return hashlib.sha256(token.encode('utf-8')).hexdigest()


def hash_password(password: str) -> str:
    salt = secrets.token_bytes(16)
    derived_key = hashlib.pbkdf2_hmac(
        'sha256',
        password.encode('utf-8'),
        salt,
        PASSWORD_HASH_ITERATIONS,
    )
    return f'pbkdf2_sha256${PASSWORD_HASH_ITERATIONS}${base64.b64encode(salt).decode("ascii")}${base64.b64encode(derived_key).decode("ascii")}'


def verify_password(password: str, encoded_hash: str | None) -> bool:
    if not encoded_hash:
        return False

    try:
        algorithm, iterations_text, salt_b64, hash_b64 = encoded_hash.split('$', 3)
        if algorithm != 'pbkdf2_sha256':
            return False
        iterations = int(iterations_text)
        salt = base64.b64decode(salt_b64.encode('ascii'))
        expected_hash = base64.b64decode(hash_b64.encode('ascii'))
    except Exception:
        return False

    candidate_hash = hashlib.pbkdf2_hmac(
        'sha256',
        password.encode('utf-8'),
        salt,
        iterations,
    )
    return secrets.compare_digest(candidate_hash, expected_hash)


def build_user_payload(row: dict[str, Any]) -> dict[str, str | None]:
    return {
        'userId': row['user_id'],
        'username': row['username'],
        'email': row['email'],
        'profileImageDataUrl': row.get('profile_image_data_url'),
        'createdAt': row['created_at'],
        'updatedAt': row['updated_at'],
        'lastLoginAt': row['last_login_at'],
    }


def get_user_by_email(email: str) -> dict[str, Any] | None:
    with get_activity_connection() as connection:
        return connection.execute(
            '''
            SELECT user_id, username, email, password_hash, profile_image_data_url, created_at, updated_at, last_login_at
            FROM app_users
            WHERE email = %s
            ''',
            (normalize_email(email),),
        ).fetchone()


def create_authenticated_session(*, user_id: str, request: Request) -> tuple[str, str]:
    session_token = secrets.token_urlsafe(48)
    session_id = uuid.uuid4().hex
    timestamp = utc_now_iso()
    expires_at = datetime.utcfromtimestamp(time.time() + SESSION_DURATION_SECONDS).isoformat()

    with get_activity_connection() as connection:
        connection.execute(
            '''
            INSERT INTO app_user_sessions (
                session_id,
                user_id,
                session_token_hash,
                created_at,
                updated_at,
                expires_at,
                revoked_at,
                client_ip,
                user_agent
            )
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            ''',
            (
                session_id,
                user_id,
                hash_session_token(session_token),
                timestamp,
                timestamp,
                expires_at,
                None,
                request.client.host if request.client is not None else None,
                request.headers.get('user-agent'),
            ),
        )

    return session_id, session_token


def set_session_cookie(response: Response, session_token: str) -> None:
    same_site = SESSION_COOKIE_SAMESITE if SESSION_COOKIE_SAMESITE in {'lax', 'strict', 'none'} else 'lax'
    response.set_cookie(
        key=SESSION_COOKIE_NAME,
        value=session_token,
        httponly=True,
        secure=SESSION_COOKIE_SECURE,
        samesite=same_site,
        max_age=SESSION_MAX_AGE,
        path='/',
        domain=SESSION_COOKIE_DOMAIN,
    )


def clear_session_cookie(response: Response) -> None:
    same_site = SESSION_COOKIE_SAMESITE if SESSION_COOKIE_SAMESITE in {'lax', 'strict', 'none'} else 'lax'
    response.delete_cookie(
        key=SESSION_COOKIE_NAME,
        path='/',
        httponly=True,
        samesite=same_site,
        secure=SESSION_COOKIE_SECURE,
        domain=SESSION_COOKIE_DOMAIN,
    )


def revoke_session(session_token: str | None) -> None:
    if not session_token:
        return

    with get_activity_connection() as connection:
        connection.execute(
            '''
            UPDATE app_user_sessions
            SET revoked_at = %s, updated_at = %s
            WHERE session_token_hash = %s AND revoked_at IS NULL
            ''',
            (utc_now_iso(), utc_now_iso(), hash_session_token(session_token)),
        )


def get_authenticated_user(request: Request) -> dict[str, Any]:
    session_token = request.cookies.get(SESSION_COOKIE_NAME)
    if not session_token:
        raise HTTPException(status_code=401, detail='Authentication required.')

    current_timestamp = utc_now_iso()
    with get_activity_connection() as connection:
        row = connection.execute(
            '''
            SELECT
                u.user_id,
                u.username,
                u.email,
                u.profile_image_data_url,
                u.created_at,
                u.updated_at,
                u.last_login_at,
                s.session_id,
                s.expires_at
            FROM app_user_sessions s
            INNER JOIN app_users u ON u.user_id = s.user_id
            WHERE s.session_token_hash = %s
              AND s.revoked_at IS NULL
              AND s.expires_at > %s
            ''',
            (hash_session_token(session_token), current_timestamp),
        ).fetchone()

    if row is None:
        raise HTTPException(status_code=401, detail='Session expired or invalid.')

    return row


def create_app_user(*, username: str, email: str, password: str) -> dict[str, str | None]:
    normalized_username, normalized_email = validate_login_payload(username, email)
    validated_password = validate_password(password)

    if get_user_by_email(normalized_email) is not None:
        raise HTTPException(status_code=409, detail='An account with this email already exists.')

    with get_activity_connection() as connection:
        timestamp = utc_now_iso()
        user_id = uuid.uuid4().hex
        row = connection.execute(
            '''
            INSERT INTO app_users (
                user_id,
                username,
                email,
                password_hash,
                created_at,
                updated_at,
                last_login_at
            )
            VALUES (%s, %s, %s, %s, %s, %s, %s)
            RETURNING user_id, username, email, profile_image_data_url, created_at, updated_at, last_login_at
            ''',
            (
                user_id,
                normalized_username,
                normalized_email,
                hash_password(validated_password),
                timestamp,
                timestamp,
                timestamp,
            ),
        ).fetchone()

    if row is None:
        raise HTTPException(status_code=500, detail='Failed to store user details.')

    return build_user_payload(row)


def authenticate_user(*, email: str, password: str) -> dict[str, str | None]:
    normalized_email = normalize_email(email)
    validated_password = validate_password(password)
    row = get_user_by_email(normalized_email)

    if row is None or not verify_password(validated_password, row.get('password_hash')):
        raise HTTPException(status_code=401, detail='Invalid email or password.')

    timestamp = utc_now_iso()
    with get_activity_connection() as connection:
        updated_row = connection.execute(
            '''
            UPDATE app_users
            SET updated_at = %s, last_login_at = %s
            WHERE email = %s
            RETURNING user_id, username, email, profile_image_data_url, created_at, updated_at, last_login_at
            ''',
            (timestamp, timestamp, normalized_email),
        ).fetchone()

    if updated_row is None:
        raise HTTPException(status_code=500, detail='Failed to load authenticated user.')

    return build_user_payload(updated_row)


async def build_profile_image_data_url(profile_image: UploadFile | None) -> str | None:
    if profile_image is None:
        return None

    content_type = (profile_image.content_type or '').lower()
    if content_type not in PROFILE_IMAGE_CONTENT_TYPES:
        raise HTTPException(status_code=400, detail='Profile image must be a PNG, JPEG, WEBP, or GIF file.')

    image_bytes = await profile_image.read()
    if not image_bytes:
        return None
    if len(image_bytes) > PROFILE_IMAGE_MAX_BYTES:
        raise HTTPException(status_code=400, detail='Profile image must be 1.5 MB or smaller.')

    encoded = base64.b64encode(image_bytes).decode('ascii')
    return f'data:{content_type};base64,{encoded}'


async def update_authenticated_user_profile(
    *,
    request: Request,
    username: str,
    email: str,
    profile_image: UploadFile | None,
) -> dict[str, str | None]:
    current_user = get_authenticated_user(request)
    normalized_username = normalize_username(username)
    normalized_email = normalize_email(email)
    if len(normalized_username) < 3:
        raise HTTPException(status_code=400, detail='Name must be at least 3 characters long.')
    if len(normalized_username) > 80:
        raise HTTPException(status_code=400, detail='Name must be 80 characters or fewer.')
    if not EMAIL_REGEX.fullmatch(normalized_email):
        raise HTTPException(status_code=400, detail='Enter a valid email address.')

    existing_user = get_user_by_email(normalized_email)
    if existing_user is not None and existing_user.get('user_id') != current_user['user_id']:
        raise HTTPException(status_code=409, detail='An account with this email already exists.')

    profile_image_data_url = await build_profile_image_data_url(profile_image)
    timestamp = utc_now_iso()

    with get_activity_connection() as connection:
        if profile_image_data_url is None:
            row = connection.execute(
                '''
                UPDATE app_users
                SET username = %s, email = %s, updated_at = %s
                WHERE user_id = %s
                RETURNING user_id, username, email, profile_image_data_url, created_at, updated_at, last_login_at
                ''',
                (normalized_username, normalized_email, timestamp, current_user['user_id']),
            ).fetchone()
        else:
            row = connection.execute(
                '''
                UPDATE app_users
                SET username = %s, email = %s, profile_image_data_url = %s, updated_at = %s
                WHERE user_id = %s
                RETURNING user_id, username, email, profile_image_data_url, created_at, updated_at, last_login_at
                ''',
                (normalized_username, normalized_email, profile_image_data_url, timestamp, current_user['user_id']),
            ).fetchone()

    if row is None:
        raise HTTPException(status_code=500, detail='Failed to update user profile.')

    return build_user_payload(row)


def ensure_session_state(session_id: str) -> dict[str, Any]:
    if session_id not in SESSION_STATE:
        SESSION_STATE[session_id] = {
            'forecast_steps': {'ts': False, 'ml': False, 'loss': False, 'profit': False},
            'time_series_result': None,
            'ml_forecast_result': None,
            'updated_at': datetime.utcnow().isoformat(),
        }
    return SESSION_STATE[session_id]


def normalize_column_name(name: str) -> str:
    normalized = ''.join(ch.lower() if ch.isalnum() else '_' for ch in name)
    return re.sub(r'_+', '_', normalized).strip('_')


def make_unique_column_names(columns: list[Any]) -> list[str]:
    seen: dict[str, int] = {}
    unique_columns: list[str] = []
    for index, column in enumerate(columns):
        base_name = normalize_column_name(str(column)) or f'column_{index + 1}'
        next_count = seen.get(base_name, 0) + 1
        seen[base_name] = next_count
        unique_columns.append(base_name if next_count == 1 else f'{base_name}_{next_count}')
    return unique_columns


def dataset_file_path(dataset_id: str, suffix: str = '.parquet') -> Path:
    return DATASET_DIR / f'{dataset_id}{suffix}'


def write_dataset_file(dataset_id: str, content: bytes, suffix: str = '.parquet') -> Path:
    target = dataset_file_path(dataset_id, suffix)
    target.write_bytes(content)
    return target


async def write_uploaded_file(upload_file: UploadFile, dataset_id: str, suffix: str) -> tuple[Path, int]:
    target = dataset_file_path(dataset_id, suffix)
    total_bytes = 0

    with target.open('wb') as handle:
        while True:
            chunk = await upload_file.read(UPLOAD_READ_CHUNK_SIZE)
            if not chunk:
                break
            total_bytes += len(chunk)
            if total_bytes > MAX_UPLOAD_SIZE_BYTES:
                handle.close()
                try:
                    target.unlink(missing_ok=True)
                except Exception:
                    logger.exception('Failed to remove oversized uploaded dataset file %s', target)
                raise HTTPException(status_code=400, detail='File exceeds 512MB limit.')
            handle.write(chunk)

    await upload_file.seek(0)
    return target, total_bytes


def write_cached_frame(dataset_id: str, frame: pd.DataFrame) -> Path:
    target = dataset_file_path(dataset_id, '.joblib')
    joblib.dump(frame, target)
    return target


def load_dataset(filepath: str) -> pd.DataFrame:
    path = Path(filepath)
    ext = path.suffix.lower()
    loaders = {
        '.csv': lambda f: pd.read_csv(f),
        '.tsv': lambda f: pd.read_csv(f, sep='\t'),
        '.txt': lambda f: pd.read_csv(f, sep=None, engine='python'),
        '.xlsx': lambda f: pd.read_excel(f, engine='openpyxl'),
        '.xls': lambda f: pd.read_excel(f, engine='xlrd'),
        '.xlsm': lambda f: pd.read_excel(f, engine='openpyxl'),
        '.parquet': lambda f: pd.read_parquet(f),
        '.json': lambda f: pd.read_json(f),
        '.jsonl': lambda f: pd.read_json(f, lines=True),
        '.joblib': lambda f: joblib.load(f),
        '.pkl': lambda f: joblib.load(f),
        '.pickle': lambda f: joblib.load(f),
    }
    loader = loaders.get(ext)
    if loader is None:
        raise ValueError(f"Unsupported file format: {ext}. Supported: {list(loaders.keys())}")
    try:
        frame = loader(path)
        if not isinstance(frame, pd.DataFrame):
            frame = pd.DataFrame(frame)
        if frame.empty:
            raise ValueError('Loaded dataset is empty.')
        return normalize_dataframe(frame)
    except Exception as error:
        raise ValueError(f'Failed to load {filepath}: {error}') from error


class IngestionFormatError(ValueError):
    def __init__(self, message: str, *, issue: str = 'format_error') -> None:
        super().__init__(message)
        self.issue = issue
        self.public_message = message


TEXT_ENCODINGS = ['utf-8-sig', 'utf-8', 'utf-16', 'utf-16-le', 'utf-16-be', 'cp1252', 'latin1']
DELIMITERS = [',', '\t', ';', '|']


def friendly_format_error(error: Exception, file_kind: str = 'dataset') -> str:
    if isinstance(error, IngestionFormatError):
        return error.public_message
    if isinstance(error, HTTPException):
        return str(error.detail)
    raw = str(error).strip()
    if not raw:
        raw = type(error).__name__
    return f'Could not read this {file_kind}. Check that the file is not corrupted, password-protected, or using an unsupported format. Details: {raw}'


def detect_text_encoding(path: Path) -> str:
    sample = path.read_bytes()[:131_072]
    if not sample:
        raise IngestionFormatError('This delimited file is empty. Upload a CSV or TSV with a header row and at least one data row.', issue='empty_file')
    for encoding in TEXT_ENCODINGS:
        try:
            decoded = sample.decode(encoding)
        except UnicodeDecodeError:
            continue
        if decoded.count('\ufffd') <= max(1, len(decoded) // 500):
            return encoding
    raise IngestionFormatError('Could not detect the text encoding for this CSV/TSV file. Save it as UTF-8, UTF-16, or Windows-1252 and try again.', issue='encoding_detection_failed')


def detect_delimiter_from_sample(sample: str, fallback: str = ',') -> str:
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=DELIMITERS)
        if dialect.delimiter in DELIMITERS:
            return dialect.delimiter
    except Exception:
        pass

    lines = [line for line in sample.splitlines()[:30] if line.strip()]
    if not lines:
        return fallback
    counts = {delimiter: sum(max(0, line.count(delimiter)) for line in lines) for delimiter in DELIMITERS}
    delimiter, count = max(counts.items(), key=lambda item: item[1])
    return delimiter if count > 0 else fallback


def sniff_delimited_options(path: Path) -> dict[str, Any]:
    encoding = detect_text_encoding(path)
    try:
        sample = path.read_text(encoding=encoding, errors='replace')[:65_536]
    except Exception as error:
        raise IngestionFormatError(f'Could not decode this CSV/TSV file with detected encoding {encoding}.') from error
    fallback = '\t' if path.suffix.lower() == '.tsv' else ','
    separator = detect_delimiter_from_sample(sample, fallback=fallback)
    header_row = detect_delimited_header_row(sample, separator)
    return {'encoding': encoding, 'separator': separator, 'header_row': header_row}


def detect_delimited_header_row(sample: str, separator: str) -> int:
    rows = list(csv.reader(io.StringIO(sample), delimiter=separator))
    rows = rows[:40]
    best_index = 0
    best_score = float('-inf')
    for index, row in enumerate(rows):
        cells = [str(cell).strip() for cell in row]
        non_empty = [cell for cell in cells if cell]
        if len(non_empty) < 2:
            continue
        next_rows = [
            [str(cell).strip() for cell in rows[next_index]]
            for next_index in range(index + 1, min(len(rows), index + 6))
            if len([cell for cell in rows[next_index] if str(cell).strip()]) >= 2
        ]
        if not next_rows:
            continue
        width_matches = sum(1 for next_row in next_rows if abs(len(next_row) - len(cells)) <= 1)
        unique_ratio = len(set(non_empty)) / max(1, len(non_empty))
        alpha_ratio = sum(bool(re.search(r'[A-Za-z_]', cell)) for cell in non_empty) / len(non_empty)
        numeric_ratio = sum(is_likely_numeric_text(cell) for cell in non_empty) / len(non_empty)
        score = (width_matches * 3) + (unique_ratio * 2) + (alpha_ratio * 2) - (numeric_ratio * 2) - (index * 0.25)
        if score > best_score:
            best_index = index
            best_score = score
    return best_index


def is_likely_numeric_text(value: Any) -> bool:
    text = str(value).strip()
    if not text:
        return False
    return bool(re.fullmatch(r'\(?\s*[-+]?\s*[$₹€£]?\s*\d{1,3}(?:,\d{2,3})*(?:\.\d+)?\s*%?\s*\)?|\(?\s*[-+]?\s*[$₹€£]?\s*\d+(?:\.\d+)?\s*%?\s*\)?', text))


def parse_numeric_text(value: Any) -> float | None:
    if value is None or pd.isna(value):
        return np.nan
    text = str(value).strip()
    if not text:
        return np.nan
    negative = text.startswith('(') and text.endswith(')')
    cleaned = re.sub(r'[$₹€£,%\s()]', '', text.replace(',', ''))
    if cleaned in {'', '+', '-'}:
        return np.nan
    try:
        number = float(cleaned)
    except ValueError:
        return None
    return -number if negative else number


def looks_like_date_value(value: Any) -> bool:
    text = str(value).strip()
    if not text or len(text) < 5:
        return False
    if re.fullmatch(r'[-+]?\d+(?:\.\d+)?', text):
        return False
    return bool(re.search(r'\d', text) and (re.search(r'[-/.\s]', text) or re.search(r'[A-Za-z]{3,}', text)))


def parse_dateutil_value(value: Any, *, dayfirst: bool) -> pd.Timestamp | pd.NaT:
    if value is None or pd.isna(value):
        return pd.NaT
    if isinstance(value, (datetime, date, pd.Timestamp, np.datetime64)):
        return pd.to_datetime(value, errors='coerce')
    text = str(value).strip()
    if not looks_like_date_value(text):
        return pd.NaT
    if date_parser is not None:
        try:
            return pd.Timestamp(date_parser.parse(text, dayfirst=dayfirst, fuzzy=True))
        except Exception:
            return pd.NaT
    return pd.to_datetime(text, errors='coerce', dayfirst=dayfirst)


def choose_dayfirst(sample: pd.Series) -> bool:
    first_token_gt_12 = 0
    second_token_gt_12 = 0
    for value in sample.astype(str).head(100):
        match = re.match(r'\s*(\d{1,2})[-/.](\d{1,2})[-/.](\d{2,4})', value)
        if not match:
            continue
        first, second = int(match.group(1)), int(match.group(2))
        if first > 12:
            first_token_gt_12 += 1
        if second > 12:
            second_token_gt_12 += 1
    return first_token_gt_12 > second_token_gt_12


def postprocess_ingested_frame(frame: pd.DataFrame) -> pd.DataFrame:
    cleaned = pd.DataFrame(frame).copy()
    cleaned = cleaned.dropna(axis=0, how='all').dropna(axis=1, how='all')
    seen_columns: dict[str, int] = {}
    next_columns: list[str] = []
    for index, column in enumerate(cleaned.columns):
        base_name = f'column_{index + 1}' if str(column).strip().lower().startswith('unnamed:') or not str(column).strip() else str(column).strip()
        count = seen_columns.get(base_name, 0) + 1
        seen_columns[base_name] = count
        next_columns.append(base_name if count == 1 else f'{base_name}_{count}')
    cleaned.columns = next_columns

    for column in cleaned.columns:
        series = cleaned[column]
        if pd.api.types.is_object_dtype(series) or pd.api.types.is_string_dtype(series):
            text_series = series.astype(object).where(series.notna(), np.nan)
            text_series = text_series.replace(r'^\s*$', np.nan, regex=True)
            sample = text_series.dropna().astype(str).head(200)
            if not sample.empty:
                numeric_ratio = float(sample.map(is_likely_numeric_text).mean())
                if numeric_ratio >= 0.8:
                    converted = text_series.map(parse_numeric_text)
                    if pd.Series(converted).notna().sum() >= max(1, int(text_series.notna().sum() * 0.7)):
                        cleaned[column] = pd.to_numeric(converted, errors='coerce')
                        continue

                column_name = str(column).lower()
                date_candidate = any(token in column_name for token in ['date', 'time', 'month', 'year', 'period']) or float(sample.map(looks_like_date_value).mean()) >= 0.7
                if date_candidate:
                    dayfirst = choose_dayfirst(sample)
                    parsed = text_series.map(lambda value: parse_dateutil_value(value, dayfirst=dayfirst))
                    parsed = pd.to_datetime(parsed, errors='coerce')
                    if parsed.notna().sum() >= max(1, int(text_series.notna().sum() * 0.65)):
                        cleaned[column] = parsed
                        continue
            cleaned[column] = text_series
    return normalize_dataframe(cleaned)


def read_delimited_frame(path: Path, *, n_rows: int | None = None, columns: list[str] | None = None, options: dict[str, Any] | None = None) -> pd.DataFrame:
    resolved = options or sniff_delimited_options(path)
    try:
        frame = pd.read_csv(
            path,
            sep=resolved['separator'],
            encoding=resolved['encoding'],
            skiprows=int(resolved.get('header_row') or 0),
            nrows=n_rows,
            low_memory=False,
        )
    except UnicodeDecodeError as error:
        raise IngestionFormatError('Could not decode this CSV/TSV file. Try saving it as UTF-8 and upload again.', issue='encoding_decode_failed') from error
    except pd.errors.EmptyDataError as error:
        raise IngestionFormatError('This CSV/TSV file does not contain a readable table.', issue='empty_file') from error
    except Exception as error:
        raise IngestionFormatError(f'Could not parse the CSV/TSV file. Detected delimiter "{resolved.get("separator")}" and encoding "{resolved.get("encoding")}". {error}', issue='delimited_parse_failed') from error
    frame = postprocess_ingested_frame(frame)
    if columns is not None:
        missing_columns = [column for column in columns if column not in frame.columns]
        if missing_columns:
            raise HTTPException(status_code=400, detail=f'Missing columns: {missing_columns}')
        frame = frame.loc[:, columns]
    return frame


def count_delimited_rows_from_path(path: Path, options: dict[str, Any]) -> int:
    try:
        row_count = 0
        for chunk in pd.read_csv(
            path,
            sep=options['separator'],
            encoding=options['encoding'],
            skiprows=int(options.get('header_row') or 0),
            chunksize=100_000,
            low_memory=False,
        ):
            row_count += len(chunk.dropna(how='all'))
        return int(row_count)
    except Exception as error:
        raise IngestionFormatError(f'Could not count rows in this CSV/TSV file after detecting its format. {error}', issue='delimited_count_failed') from error


def read_cached_frame(dataset_entry: dict[str, Any], columns: list[str] | None = None, n_rows: int | None = None) -> pd.DataFrame:
    frame_path = dataset_entry.get('frame_path')
    if not frame_path:
        raise HTTPException(status_code=400, detail='Cached dataset frame path is missing. Please upload the file again.')

    try:
        frame = load_dataset(str(frame_path))
        if columns is not None:
            frame = frame.loc[:, columns]
        if n_rows is not None and n_rows > 0:
            frame = frame.head(n_rows)
        return frame.copy()
    except Exception as error:
        raise HTTPException(status_code=400, detail=f'Failed to load cached dataset: {error}') from error


def read_cached_parquet(dataset_entry: dict[str, Any], **kwargs: Any) -> pl.DataFrame:
    parquet_path = dataset_entry.get('parquet_path')
    if not parquet_path:
        raise HTTPException(status_code=400, detail='Cached parquet dataset path is missing. Please upload the file again.')
    if pq is None:
        raise HTTPException(status_code=500, detail='Parquet support is unavailable because pyarrow is not installed in the backend environment.')

    try:
        return pl.read_parquet(parquet_path, **kwargs)
    except Exception as error:
        raise HTTPException(status_code=400, detail=f'Failed to load cached parquet dataset: {error}') from error


def get_delimited_separator(dataset_entry: dict[str, Any]) -> str:
    separator = dataset_entry.get('separator')
    if separator in {',', '\t', ';', '|'}:
        return str(separator)

    csv_path = str(dataset_entry.get('csv_path') or '').lower()
    if csv_path.endswith('.tsv'):
        return '\t'
    return ','


def sniff_delimited_separator(path: Path, fallback: str = ',') -> str:
    if path.suffix.lower() == '.tsv':
        return '\t'
    try:
        sample = path.read_text(encoding='utf-8-sig', errors='ignore')[:8192]
        dialect = csv.Sniffer().sniff(sample, delimiters=[',', '\t', ';', '|'])
        delimiter = dialect.delimiter
        return delimiter if delimiter in {',', '\t', ';', '|'} else fallback
    except Exception:
        return fallback


def read_cached_csv_preview(dataset_entry: dict[str, Any], n_rows: int | None = None) -> pd.DataFrame:
    csv_path = dataset_entry.get('csv_path')
    if not csv_path:
        raise HTTPException(status_code=400, detail='Cached CSV dataset path is missing. Please upload the file again.')

    try:
        options = {
            'separator': get_delimited_separator(dataset_entry),
            'encoding': dataset_entry.get('encoding') or 'utf-8-sig',
            'header_row': int(dataset_entry.get('header_row') or 0),
        }
        return read_delimited_frame(Path(str(csv_path)), n_rows=n_rows, options=options)
    except Exception as error:
        raise HTTPException(status_code=400, detail=friendly_format_error(error, 'CSV/TSV preview')) from error


def read_cached_csv(dataset_entry: dict[str, Any], columns: list[str] | None = None, n_rows: int | None = None) -> pd.DataFrame:
    csv_path = dataset_entry.get('csv_path')
    if not csv_path:
        raise HTTPException(status_code=400, detail='Cached CSV dataset path is missing. Please upload the file again.')

    try:
        options = {
            'separator': get_delimited_separator(dataset_entry),
            'encoding': dataset_entry.get('encoding') or 'utf-8-sig',
            'header_row': int(dataset_entry.get('header_row') or 0),
        }
        return read_delimited_frame(Path(str(csv_path)), n_rows=n_rows, columns=columns, options=options)
    except Exception as error:
        raise HTTPException(status_code=400, detail=friendly_format_error(error, 'CSV/TSV dataset')) from error


def read_cached_excel(dataset_entry: dict[str, Any], columns: list[str] | None = None, n_rows: int | None = None) -> pd.DataFrame:
    excel_path = dataset_entry.get('excel_path')
    if not excel_path:
        raise HTTPException(status_code=400, detail='Cached Excel dataset path is missing. Please upload the file again.')

    try:
        selected_sheets = [str(sheet) for sheet in (dataset_entry.get('selected_sheets') or []) if str(sheet).strip()]
        merge_mode = str(dataset_entry.get('merge_mode') or 'single').lower()
        if merge_mode not in {'single', 'stack'}:
            merge_mode = 'single'

        if not selected_sheets:
            active_sheet = str(dataset_entry.get('active_sheet') or '').strip()
            selected_sheets = [active_sheet] if active_sheet else []

        if not selected_sheets:
            sheet_summaries = dataset_entry.get('workbook_sheets') or build_excel_sheet_summaries(Path(str(excel_path)))
            selected_sheets = [str(sheet_summaries[0]['name'])]

        if merge_mode == 'single' or len(selected_sheets) == 1:
            return read_excel_sheet_frame(Path(str(excel_path)), selected_sheets[0], n_rows=n_rows, columns=columns)

        frames: list[pd.DataFrame] = []
        base_columns: list[str] | None = None
        for sheet_name in selected_sheets:
            sheet_frame = read_excel_sheet_frame(Path(str(excel_path)), sheet_name, n_rows=n_rows)
            current_columns = [str(col) for col in sheet_frame.columns]
            if base_columns is None:
                base_columns = current_columns
            elif current_columns != base_columns:
                raise HTTPException(
                    status_code=400,
                    detail=f'Cannot stack sheets with different schemas. Sheet "{sheet_name}" does not match the first selected sheet columns.',
                )

            if columns is not None:
                missing_columns = [column for column in columns if column not in sheet_frame.columns]
                if missing_columns:
                    raise HTTPException(
                        status_code=400,
                        detail=f'Sheet "{sheet_name}" is missing selected columns: {missing_columns}',
                    )
                sheet_frame = sheet_frame.loc[:, columns]

            frames.append(sheet_frame)

        if not frames:
            return pd.DataFrame(columns=columns or [])
        return pd.concat(frames, ignore_index=True)
    except Exception as error:
        raise HTTPException(status_code=400, detail=friendly_format_error(error, 'Excel workbook')) from error


def load_cached_preview(dataset_entry: dict[str, Any], limit: int = DATASET_PREVIEW_ROW_LIMIT) -> tuple[pd.DataFrame | pl.DataFrame, bool]:
    if dataset_entry.get('parquet_path'):
        return read_cached_parquet(dataset_entry, n_rows=limit, low_memory=True), True
    if dataset_entry.get('csv_path'):
        return read_cached_csv_preview(dataset_entry, n_rows=limit), False
    if dataset_entry.get('excel_path'):
        return read_cached_excel(dataset_entry, n_rows=limit), False
    if dataset_entry.get('frame_path'):
        return read_cached_frame(dataset_entry, n_rows=limit), False
    raise HTTPException(status_code=400, detail='Cached dataset storage is missing. Please upload the file again.')


def load_cached_analysis_frame(dataset_entry: dict[str, Any]) -> tuple[pd.DataFrame, int]:
    total_rows = int(dataset_entry.get('row_count') or 0)

    if dataset_entry.get('parquet_path'):
        frame = read_cached_parquet(dataset_entry, low_memory=True)
        return normalize_dataframe(frame.to_pandas(use_pyarrow_extension_array=False)), total_rows

    if dataset_entry.get('csv_path'):
        return normalize_dataframe(read_cached_csv(dataset_entry)), total_rows

    if dataset_entry.get('excel_path'):
        return normalize_dataframe(read_cached_excel(dataset_entry)), total_rows

    if dataset_entry.get('frame_path'):
        frame = read_cached_frame(dataset_entry)
        return normalize_dataframe(frame), int(len(frame))

    raise HTTPException(status_code=400, detail='Cached dataset storage is missing. Please upload the file again.')


def count_csv_rows(buffer: io.BytesIO, sep: str = ',') -> int:
    buffer.seek(0)
    row_count = 0
    try:
        for chunk in pd.read_csv(buffer, sep=sep, low_memory=True, chunksize=100_000):
            row_count += len(chunk)
    except Exception as error:
        raise HTTPException(status_code=400, detail=f'Failed to determine CSV row count: {error}') from error
    finally:
        buffer.seek(0)
    return row_count


def count_csv_rows_from_path(path: Path, sep: str = ',') -> int:
    row_count = 0
    try:
        for chunk in pd.read_csv(path, sep=sep, low_memory=True, chunksize=100_000):
            row_count += len(chunk)
    except Exception as error:
        raise HTTPException(status_code=400, detail=f'Failed to determine CSV row count: {error}') from error
    return row_count


def count_excel_rows(buffer: io.BytesIO, filename: str) -> int:
    buffer.seek(0)
    suffix = Path(filename).suffix.lower()

    if suffix == '.xlsx':
        try:
            import openpyxl
        except ImportError as error:
            raise HTTPException(status_code=500, detail='openpyxl is required to count rows in .xlsx files.') from error

        workbook = openpyxl.load_workbook(buffer, read_only=True, data_only=True)
        row_count = workbook.active.max_row - 1
        buffer.seek(0)
        return max(0, row_count)

    if suffix == '.xls':
        try:
            import xlrd
        except ImportError as error:
            raise HTTPException(status_code=500, detail='xlrd is required to count rows in .xls files.') from error

        workbook = xlrd.open_workbook(file_contents=buffer.read())
        sheet = workbook.sheet_by_index(0)
        row_count = sheet.nrows - 1
        buffer.seek(0)
        return max(0, row_count)

    buffer.seek(0)
    return 0


def count_excel_rows_from_path(path: Path) -> int:
    suffix = path.suffix.lower()

    if suffix == '.xlsx':
        try:
            import openpyxl
        except ImportError as error:
            raise HTTPException(status_code=500, detail='openpyxl is required to count rows in .xlsx files.') from error

        workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
        return max(0, workbook.active.max_row - 1)

    if suffix == '.xls':
        try:
            import xlrd
        except ImportError as error:
            raise HTTPException(status_code=500, detail='xlrd is required to count rows in .xls files.') from error

        workbook = xlrd.open_workbook(path)
        sheet = workbook.sheet_by_index(0)
        return max(0, sheet.nrows - 1)

    return 0


def get_excel_sheet_names(path: Path) -> list[str]:
    suffix = path.suffix.lower()
    if suffix == '.xlsx':
        try:
            import openpyxl
        except ImportError as error:
            raise HTTPException(status_code=500, detail='openpyxl is required to read .xlsx sheet names.') from error
        workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
        return [str(name) for name in workbook.sheetnames]

    if suffix == '.xls':
        try:
            import xlrd
        except ImportError as error:
            raise HTTPException(status_code=500, detail='xlrd is required to read .xls sheet names.') from error
        workbook = xlrd.open_workbook(path)
        return [str(sheet.name) for sheet in workbook.sheets()]

    return []


def detect_excel_header_row(path: Path, sheet_name: str) -> int:
    engine = 'openpyxl' if path.suffix.lower() == '.xlsx' else 'xlrd'
    try:
        raw = pd.read_excel(path, engine=engine, sheet_name=sheet_name, header=None, nrows=40)
    except Exception as error:
        raise IngestionFormatError(f'Could not inspect worksheet "{sheet_name}". {error}', issue='excel_sheet_inspection_failed') from error
    best_index = 0
    best_score = float('-inf')
    for index in range(len(raw)):
        cells = [str(value).strip() for value in raw.iloc[index].tolist() if pd.notna(value) and str(value).strip()]
        if len(cells) < 2:
            continue
        next_rows = []
        for next_index in range(index + 1, min(len(raw), index + 6)):
            next_cells = [str(value).strip() for value in raw.iloc[next_index].tolist() if pd.notna(value) and str(value).strip()]
            if len(next_cells) >= 2:
                next_rows.append(next_cells)
        if not next_rows:
            continue
        width_matches = sum(1 for next_row in next_rows if abs(len(next_row) - len(cells)) <= 1)
        unique_ratio = len(set(cells)) / max(1, len(cells))
        alpha_ratio = sum(bool(re.search(r'[A-Za-z_]', cell)) for cell in cells) / len(cells)
        numeric_ratio = sum(is_likely_numeric_text(cell) for cell in cells) / len(cells)
        score = (width_matches * 3) + (unique_ratio * 2) + (alpha_ratio * 2) - (numeric_ratio * 2) - (index * 0.25)
        if score > best_score:
            best_index = index
            best_score = score
    return best_index


def read_excel_sheet_frame(path: Path, sheet_name: str, *, n_rows: int | None = None, columns: list[str] | None = None) -> pd.DataFrame:
    engine = 'openpyxl' if path.suffix.lower() == '.xlsx' else 'xlrd'
    header_row = detect_excel_header_row(path, sheet_name)
    try:
        frame = pd.read_excel(path, engine=engine, sheet_name=sheet_name, skiprows=header_row, nrows=n_rows)
    except Exception as error:
        raise IngestionFormatError(f'Could not parse worksheet "{sheet_name}". {error}', issue='excel_sheet_parse_failed') from error
    frame = postprocess_ingested_frame(frame)
    if columns is not None:
        missing_columns = [column for column in columns if column not in frame.columns]
        if missing_columns:
            raise HTTPException(status_code=400, detail=f'Sheet "{sheet_name}" is missing selected columns: {missing_columns}')
        frame = frame.loc[:, columns]
    return frame


def build_excel_sheet_summaries(path: Path) -> list[dict[str, Any]]:
    summaries: list[dict[str, Any]] = []
    for sheet_name in get_excel_sheet_names(path):
        try:
            preview = read_excel_sheet_frame(path, sheet_name, n_rows=min(50, DATASET_PREVIEW_ROW_LIMIT))
            if preview.empty or len(preview.columns) == 0:
                continue
            total_rows = int(count_excel_rows_for_sheet(path, sheet_name))
            if total_rows <= 0:
                total_rows = int(len(preview))
            summaries.append({
                'name': sheet_name,
                'rowCount': total_rows,
                'columnCount': int(len(preview.columns)),
                'columns': [str(column) for column in preview.columns],
            })
        except IngestionFormatError:
            continue
        except Exception:
            logger.warning('Skipping unreadable or empty workbook sheet %s in %s', sheet_name, path.name, exc_info=True)
            continue
    if not summaries:
        raise IngestionFormatError('No non-empty worksheets were found in this workbook. Add data to at least one sheet and upload again.', issue='excel_no_data_sheets')
    return summaries


def count_excel_rows_for_sheet(path: Path, sheet_name: str) -> int:
    frame = read_excel_sheet_frame(path, sheet_name)
    return int(len(frame))


def resolve_selected_excel_sheets(selected_sheets: list[str], available_sheets: list[str]) -> list[str]:
    normalized_available = {sheet.strip().casefold(): sheet for sheet in available_sheets}
    if not selected_sheets:
        if not available_sheets:
            raise HTTPException(status_code=400, detail='No worksheets are available in this workbook.')
        return [available_sheets[0]]

    resolved: list[str] = []
    for raw_name in selected_sheets:
        candidate = str(raw_name).strip()
        if not candidate:
            continue
        matched = normalized_available.get(candidate.casefold())
        if matched is None:
            raise HTTPException(status_code=400, detail=f'Sheet "{candidate}" was not found in this workbook.')
        if matched not in resolved:
            resolved.append(matched)

    if not resolved:
        raise HTTPException(status_code=400, detail='Select at least one worksheet to continue.')
    return resolved


def build_excel_selection_payload(
    *,
    excel_path: Path,
    selected_sheets: list[str],
    merge_mode: Literal['single', 'stack'],
) -> dict[str, Any]:
    resolved_merge_mode: Literal['single', 'stack'] = merge_mode if merge_mode in {'single', 'stack'} else 'single'

    if resolved_merge_mode == 'single':
        sheet_name = selected_sheets[0]
        preview_frame = read_excel_sheet_frame(excel_path, sheet_name, n_rows=DATASET_PREVIEW_ROW_LIMIT)
        total_rows = count_excel_rows_for_sheet(excel_path, sheet_name)
    else:
        preview_frames: list[pd.DataFrame] = []
        base_columns: list[str] | None = None
        total_rows = 0
        for sheet_name in selected_sheets:
            preview_sheet = read_excel_sheet_frame(excel_path, sheet_name, n_rows=DATASET_PREVIEW_ROW_LIMIT)
            current_columns = [str(col) for col in preview_sheet.columns]
            if base_columns is None:
                base_columns = current_columns
            elif current_columns != base_columns:
                raise HTTPException(
                    status_code=400,
                    detail=f'Cannot stack sheets with different schemas. Sheet "{sheet_name}" does not match the first selected sheet columns.',
                )
            preview_frames.append(preview_sheet)
            total_rows += count_excel_rows_for_sheet(excel_path, sheet_name)

        preview_frame = pd.concat(preview_frames, ignore_index=True) if preview_frames else pd.DataFrame()

    loaded_row_count = len(preview_frame)
    preview_loaded = int(total_rows) > loaded_row_count
    column_info = build_column_info_from_frame(preview_frame)
    preview_rows = preview_frame.where(pd.notna(preview_frame), None).to_dict(orient='records')
    duplicate_rows = int(max(0, len(preview_frame) - len(preview_frame.drop_duplicates())))

    return {
        'frame': preview_frame,
        'rows': safe_serialize(preview_rows),
        'column_info': column_info,
        'total_rows': int(total_rows),
        'loaded_row_count': int(loaded_row_count),
        'preview_loaded': bool(preview_loaded),
        'duplicate_rows': duplicate_rows,
    }


def resolve_requested_columns(requested_columns: list[str], available_columns: list[str]) -> dict[str, str]:
    exact_matches = {column: column for column in available_columns}
    normalized_matches: dict[str, str] = {}
    for column in available_columns:
        normalized = normalize_column_name(column)
        normalized_matches.setdefault(normalized, column)

    resolved: dict[str, str] = {}
    missing: list[str] = []
    for requested in requested_columns:
        if requested in exact_matches:
            resolved[requested] = requested
            continue

        normalized_requested = normalize_column_name(requested)
        matched = normalized_matches.get(normalized_requested)
        if matched is None:
            missing.append(requested)
            continue
        resolved[requested] = matched

    if missing:
        raise HTTPException(status_code=400, detail=f'Missing columns: {missing}')

    return resolved


def sample_training_rows(
    X: pd.DataFrame,
    y: pd.Series,
    max_rows: int,
    random_state: int,
    stratify: pd.Series | None = None,
) -> tuple[pd.DataFrame, pd.Series]:
    if len(X) <= max_rows:
        return X, y

    split_kwargs: dict[str, Any] = {
        'train_size': max_rows / len(X),
        'random_state': random_state,
    }
    if stratify is not None:
        split_kwargs['stratify'] = stratify

    sampled_X, _, sampled_y, _ = train_test_split(X, y, **split_kwargs)
    return sampled_X, sampled_y


def build_training_profile(row_count: int, requested_cv_folds: int, training_mode: TrainingMode) -> dict[str, int | bool | str]:
    profile = {
        'training_mode': training_mode,
        'cv_folds': min(requested_cv_folds, 5),
        'cv_sample_limit': 0,
        'train_sample_limit': 0,
        'importance_sample_limit': 0,
        'importance_repeats': 3,
        'skip_cv_for_large_dataset': False,
    }

    if row_count >= VERY_LARGE_DATASET_ROW_THRESHOLD:
        profile['cv_folds'] = min(requested_cv_folds, 3)
        profile['train_sample_limit'] = VERY_LARGE_TRAIN_SAMPLE_LIMIT
        profile['importance_sample_limit'] = VERY_LARGE_IMPORTANCE_SAMPLE_LIMIT
        profile['importance_repeats'] = 1
        if training_mode == 'fast':
            profile['cv_sample_limit'] = 0
            profile['skip_cv_for_large_dataset'] = True
        else:
            profile['cv_sample_limit'] = max(VERY_LARGE_CV_SAMPLE_LIMIT, 3000)
    elif row_count >= LARGE_DATASET_ROW_THRESHOLD:
        profile['cv_folds'] = min(requested_cv_folds, 4)
        profile['train_sample_limit'] = TRAIN_SAMPLE_LIMIT
        profile['importance_sample_limit'] = IMPORTANCE_SAMPLE_LIMIT
        profile['importance_repeats'] = 2
        if training_mode == 'fast':
            profile['cv_sample_limit'] = 0
            profile['skip_cv_for_large_dataset'] = True
        else:
            profile['cv_sample_limit'] = max(CV_SAMPLE_LIMIT, 5000)

    return profile


class TrainRequest(BaseModel):
    data: list[dict[str, Any]] = Field(default_factory=list)
    dataset_id: str | None = None
    target_column: str
    feature_columns: list[str]
    problem_type: ProblemType
    model_type: str
    test_size: float = Field(default=0.2, ge=0.1, le=0.4)
    random_state: int = 42
    cv_folds: int = Field(default=5, ge=2, le=10)
    training_mode: TrainingMode = 'balanced'


class PredictRequest(BaseModel):
    model_id: str
    features: dict[str, Any]


class DatasetCacheRequest(BaseModel):
    file_name: str
    data: list[dict[str, Any]] = Field(default_factory=list)


class DatasetSheetSelectionRequest(BaseModel):
    dataset_id: str
    selected_sheets: list[str] = Field(default_factory=list)
    merge_mode: Literal['single', 'stack'] = 'single'


class LoginRequest(BaseModel):
    email: str
    password: str


class RegisterRequest(BaseModel):
    username: str
    email: str
    password: str


class AdvancedEdaRequest(BaseModel):
    data: list[dict[str, Any]] = Field(default_factory=list)
    dataset_id: str | None = None


class SalesForecastRequest(BaseModel):
    data: list[dict[str, Any]] = Field(default_factory=list)
    dataset_id: str | None = None
    session_id: str | None = None
    date_column: str
    target_column: str
    forecast_periods: int = Field(default=3, ge=1, le=24)
    test_percentage: int = Field(default=20, ge=10, le=50)
    test_periods: int | None = Field(default=None, ge=1, le=24)
    lag_periods: int = Field(default=3, ge=1, le=12)
    model_type: str | None = None
    feature_groups: list[str] = Field(default_factory=lambda: ['trend', 'calendar', 'seasonality', 'lags', 'rolling'])


class TimeSeriesForecastRequest(BaseModel):
    data: list[dict[str, Any]] = Field(default_factory=list)
    dataset_id: str | None = None
    session_id: str | None = None
    date_column: str
    target_column: str
    forecast_periods: int = Field(default=3, ge=1, le=24)
    test_percentage: int = Field(default=20, ge=10, le=50)
    model_type: str = Field(default='auto')
    require_quality_gate: bool = True


class TsStationarityRequest(BaseModel):
    dataset_id: str


class TsForecastRunRequest(BaseModel):
    dataset_id: str
    horizon: int = Field(default=3, ge=1, le=24)
    training_split: float = Field(default=0.8, ge=0.5, le=0.95)


class MlForecastRequest(BaseModel):
    data: list[dict[str, Any]] = Field(default_factory=list)
    dataset_id: str | None = None
    session_id: str | None = None
    date_column: str
    target_column: str
    forecast_periods: int = Field(default=3, ge=1, le=24)
    test_percentage: int = Field(default=20, ge=10, le=50)
    lag_periods: int = Field(default=3, ge=1, le=12)
    model_type: str = Field(default='auto')
    feature_groups: list[str] = Field(default_factory=lambda: ['trend', 'calendar', 'lags', 'rolling'])
    require_quality_gate: bool = True


class ForecastRunRequest(BaseModel):
    session_id: str
    forecast_periods: int = Field(default=30, ge=1, le=180)
    confirmed_assumptions: bool = False
    column_mapping: dict[str, str] = Field(default_factory=dict)
    scenario_parameters: dict[str, dict[str, float]] = Field(default_factory=dict)


class ReportConfigPayload(BaseModel):
    includeLoss: bool = True
    includeProfit: bool = True
    scenario: Literal['optimistic', 'baseline', 'pessimistic'] = 'baseline'


class CleaningLog(BaseModel):
    action: str
    detail: str
    timestamp: str


class CleaningJustificationRequest(BaseModel):
    logs: list[CleaningLog]
    totalRows: int
    totalColumns: int
    fileName: str | None = None
    loadedRowCount: int | None = None
    previewLoaded: bool = False

class ParquetCleaningRequest(BaseModel):
    dataset_id: str
    remove_duplicates: bool = True
    handle_missing: bool = True
    convert_dates: bool = True
    standardize_names: bool = True
    infer_dtypes: bool = True


class DtypeInferenceRequest(BaseModel):
    data: list[dict[str, Any]] = Field(default_factory=list)
    dataset_id: str | None = None
    persist: bool = False


class ColumnInfo(BaseModel):
    name: str
    dtype: str
    nonNull: int
    nullCount: int
    uniqueCount: int
    role: str


class PredictionHistoryItem(BaseModel):
    id: str
    prediction: str | float | int
    confidence: float | None = None
    features: dict[str, str | float | int] | None = None
    timestamp: str


class EdaStats(BaseModel):
    numericColumns: list[str] = Field(default_factory=list)
    categoricalColumns: list[str] = Field(default_factory=list)
    stats: dict[str, dict[str, float]] = Field(default_factory=dict)
    correlations: list[dict[str, float | str]] = Field(default_factory=list)



class UploadedModelPayload(BaseModel):
    name: str
    type: str
    target: str
    problem: str
    trainedAt: str
    metrics: dict[str, float] = Field(default_factory=dict)
    features: list[str] = Field(default_factory=list)


class ForecastPointPayload(BaseModel):
    period: str
    actual: float | None = None
    predicted: float | None = None
    lower: float | None = None
    upper: float | None = None


class ForecastMetricsPayload(BaseModel):
    mae: float
    rmse: float
    mape: float


class ForecastTrainingSummaryPayload(BaseModel):
    model_name: str
    total_periods: int
    train_periods: int
    test_periods: int
    train_percentage: float
    test_percentage: float
    forecast_periods: int
    lag_periods: int = 0
    train_start: str
    train_end: str
    test_start: str
    test_end: str
    last_observed_period: str


class DatasetProfilePayload(BaseModel):
    detected_frequency: str
    usable_periods: int
    volatility: float
    zero_value_share: float


class StationarityCheckPayload(BaseModel):
    test_name: str
    p_value: float
    verdict: str
    note: str


class TimeSeriesForecastResultPayload(BaseModel):
    date_column: str
    target_column: str
    frequency: str | None = None
    period_label: str | None = None
    dataset_profile: DatasetProfilePayload
    stationarity_check: StationarityCheckPayload
    history: list[ForecastPointPayload] = Field(default_factory=list)
    test_forecast: list[ForecastPointPayload] = Field(default_factory=list)
    future_forecast: list[ForecastPointPayload] = Field(default_factory=list)
    metrics: ForecastMetricsPayload
    training_summary: ForecastTrainingSummaryPayload
    recommended_models: list[dict[str, Any]] = Field(default_factory=list)
    model_details: dict[str, Any] = Field(default_factory=dict)
    analysis: str


class MlForecastResultPayload(BaseModel):
    date_column: str
    target_column: str
    frequency: str | None = None
    period_label: str | None = None
    dataset_profile: DatasetProfilePayload
    generated_features: list[str] = Field(default_factory=list)
    feature_preview_rows: list[dict[str, Any]] = Field(default_factory=list)
    history: list[ForecastPointPayload] = Field(default_factory=list)
    test_forecast: list[ForecastPointPayload] = Field(default_factory=list)
    future_forecast: list[ForecastPointPayload] = Field(default_factory=list)
    metrics: ForecastMetricsPayload
    training_summary: dict[str, Any]
    shap_feature_importance: list[dict[str, Any]] = Field(default_factory=list)
    recommended_models: list[dict[str, Any]] = Field(default_factory=list)
    model_details: dict[str, Any] = Field(default_factory=dict)
    analysis: str


class ReportPayload(BaseModel):
    datasetId: str | None = None
    sessionId: str | None = None
    fileName: str
    totalRows: int
    previewLoaded: bool = False
    loadedRowCount: int | None = None
    columns: list[ColumnInfo]
    duplicates: int
    memoryUsage: str
    cleaningLogs: list[CleaningLog]
    cleaningDone: bool
    cleanedRowCount: int
    targetColumn: str | None = None
    problemType: str
    selectedFeatures: list[str]
    selectedModel: str | None = None
    modelMetrics: dict[str, float] | None = None
    featureImportance: list[dict[str, Any]] | None = None
    aiInsights: str | None = None
    uploadedModel: UploadedModelPayload | None = None
    timeSeriesForecastResult: TimeSeriesForecastResultPayload | None = None
    mlForecastResult: MlForecastResultPayload | None = None
    lossForecast: list[dict[str, Any]] = Field(default_factory=list)
    profitForecast: list[dict[str, Any]] = Field(default_factory=list)
    lossSegments: list[dict[str, Any]] = Field(default_factory=list)
    scenarios: dict[str, list[dict[str, Any]]] | None = None
    breakevenPeriod: str | None = None
    reportConfig: ReportConfigPayload = Field(default_factory=ReportConfigPayload)
    forecastingStepsCompleted: list[int] = Field(default_factory=list)
    predictionResult: str | float | int | None = None
    predictionAnalysis: str | None = None
    predictionProbabilities: dict[str, float] | None = None
    predictionHistory: list[PredictionHistoryItem] = Field(default_factory=list)
    edaStats: EdaStats = Field(default_factory=EdaStats)


class EdaPdfPayload(BaseModel):
    datasetId: str | None = None
    fileName: str
    totalRows: int
    loadedRowCount: int | None = None
    previewLoaded: bool = False
    columns: list[ColumnInfo] = Field(default_factory=list)
    edaStats: EdaStats = Field(default_factory=EdaStats)
    advancedAnalysis: dict[str, Any] | None = None


REGRESSION_MODELS: dict[str, tuple[str, Any, dict[str, Any]]] = {
    'ridge_regression': ('Ridge Regression', Ridge, {'alpha': 1.0}),
    'lasso_regression': ('Lasso Regression', Lasso, {'alpha': 0.001, 'max_iter': 5000}),
    'elasticnet': ('Elastic Net', ElasticNet, {'alpha': 0.001, 'l1_ratio': 0.5, 'max_iter': 5000}),
    'random_forest': ('Random Forest', RandomForestRegressor, {'n_estimators': 100, 'min_samples_leaf': 2, 'n_jobs': TRAINING_N_JOBS}),
    'gradient_boosting': ('Gradient Boosting', GradientBoostingRegressor, {'n_estimators': 100, 'learning_rate': 0.05, 'max_depth': 3}),
    'svr': ('Support Vector Regression', SVR, {'kernel': 'rbf', 'C': 1.0}),
    'decision_tree': ('Decision Tree', DecisionTreeRegressor, {'max_depth': 8}),
    'knn_regressor': ('K-Nearest Neighbors', KNeighborsRegressor, {'n_neighbors': 7, 'weights': 'distance'}),
}

CLASSIFICATION_MODELS: dict[str, tuple[str, Any, dict[str, Any]]] = {
    'logistic_regression': ('Logistic Regression', LogisticRegression, {'max_iter': 2000, 'solver': 'lbfgs'}),
    'random_forest': ('Random Forest', RandomForestClassifier, {'n_estimators': 100, 'min_samples_leaf': 2, 'n_jobs': TRAINING_N_JOBS}),
    'gradient_boosting': ('Gradient Boosting', GradientBoostingClassifier, {'n_estimators': 100, 'learning_rate': 0.05, 'max_depth': 3}),
    'svm': ('Support Vector Machine', SVC, {'kernel': 'rbf', 'C': 1.0, 'probability': True}),
    'decision_tree': ('Decision Tree', DecisionTreeClassifier, {'max_depth': 8}),
    'knn': ('K-Nearest Neighbors', KNeighborsClassifier, {'n_neighbors': 7, 'weights': 'distance'}),
}


def safe_serialize(value: Any) -> Any:
    if isinstance(value, dict):
        return {k: safe_serialize(v) for k, v in value.items()}
    if isinstance(value, list):
        return [safe_serialize(v) for v in value]
    if isinstance(value, tuple):
        return [safe_serialize(v) for v in value]
    if isinstance(value, (datetime, date, dt_time)):
        return value.isoformat()
    if isinstance(value, Decimal):
        return float(value)
    if value is pd.NA or value is pd.NaT:
        return None
    if isinstance(value, (np.datetime64,)):
        if np.isnat(value):
            return None
        return str(value)
    if isinstance(value, (np.integer,)):
        return int(value)
    if isinstance(value, (np.floating,)):
        if np.isnan(value):
            return None
        return float(value)
    if isinstance(value, (np.bool_,)):
        return bool(value)
    if isinstance(value, np.ndarray):
        return value.tolist()
    if isinstance(value, float):
        if np.isnan(value) or np.isinf(value):
            return None
        return value
    return value


def normalize_feature_frame(frame: pd.DataFrame) -> pd.DataFrame:
    normalized = frame.copy()
    for column in normalized.columns:
        series = normalized[column]
        if pd.api.types.is_datetime64_any_dtype(series):
            normalized[column] = series.dt.strftime('%Y-%m-%d %H:%M:%S').where(series.notna(), np.nan).astype(object)
        elif pd.api.types.is_object_dtype(series) or pd.api.types.is_string_dtype(series):
            cleaned = series.astype(object)
            cleaned = cleaned.where(pd.notna(cleaned), np.nan)
            cleaned = cleaned.replace(r'^\s*$', np.nan, regex=True)
            normalized[column] = cleaned
    return normalized


def build_preprocessor(frame: pd.DataFrame) -> ColumnTransformer:
    numeric_features = frame.select_dtypes(include=[np.number, 'bool']).columns.tolist()
    categorical_features = [column for column in frame.columns if column not in numeric_features]

    transformers: list[tuple[str, Pipeline, list[str]]] = []
    if numeric_features:
        transformers.append(
            (
                'numeric',
                Pipeline([
                    ('imputer', SimpleImputer(strategy='median')),
                    ('scaler', StandardScaler()),
                ]),
                numeric_features,
            )
        )
    if categorical_features:
        transformers.append(
            (
                'categorical',
                Pipeline([
                    ('imputer', SimpleImputer(strategy='most_frequent')),
                    ('onehot', OneHotEncoder(handle_unknown='infrequent_if_exist', min_frequency=10, max_categories=100)),
                ]),
                categorical_features,
            )
        )

    if not transformers:
        raise HTTPException(status_code=400, detail='No usable features found for training.')

    return ColumnTransformer(transformers=transformers)


def build_estimator(problem_type: ProblemType, model_type: str, random_state: int):
    registry = REGRESSION_MODELS if problem_type == 'regression' else CLASSIFICATION_MODELS
    if model_type not in registry:
        raise HTTPException(status_code=400, detail=f"Model '{model_type}' is not available for {problem_type}.")

    model_name, estimator_cls, params = registry[model_type]
    params = dict(params)
    if 'random_state' in estimator_cls().get_params().keys():
        params['random_state'] = random_state
    estimator = estimator_cls(**params)
    return model_name, estimator


def normalize_dataframe(frame: pd.DataFrame) -> pd.DataFrame:
    normalized = frame.copy()
    for column in normalized.columns:
        if normalized[column].dtype == 'object':
            normalized[column] = normalized[column].replace(r'^\s*$', np.nan, regex=True)
    return normalized


def persist_inferred_dataset_frame(dataset_id: str, source_entry: dict[str, Any], frame: pd.DataFrame) -> Path:
    cached_path = write_cached_frame(dataset_id, frame)
    duplicate_rows = int(max(0, len(frame) - len(frame.drop_duplicates())))
    DATASET_CACHE[dataset_id] = {
        'frame_path': str(cached_path),
        'filename': source_entry.get('filename') or 'dataset',
        'row_count': int(len(frame)),
        'column_count': int(len(frame.columns)),
        'columns': list(frame.columns),
        'duplicate_count': duplicate_rows,
    }
    return cached_path


def build_dtype_inference_payload(frame: pd.DataFrame) -> tuple[pd.DataFrame, dict[str, Any]]:
    inferred_frame, cast_log = infer_universal_dtypes(normalize_dataframe(pd.DataFrame(frame)))
    summary = dtype_summary_report(cast_log)
    review_flags = dtype_review_flags(cast_log)
    memory_saved = int(pd.DataFrame(cast_log)['memory_delta_bytes'].sum()) if cast_log else 0
    payload = {
        'memorySavedBytes': memory_saved,
        'memorySavedKb': float(memory_saved / 1024.0),
        'report': safe_serialize(summary.to_dict(orient='records')),
        'audit': safe_serialize(cast_log),
        'reviewFlags': safe_serialize(review_flags.to_dict(orient='records')),
    }
    return inferred_frame, payload


def load_dataset_frame(dataset_id: str | None, data: list[dict[str, Any]], required_columns: list[str]) -> pd.DataFrame:
    if dataset_id:
        dataset_entry = DATASET_CACHE.get(dataset_id)
        if dataset_entry is None:
            raise HTTPException(status_code=400, detail='Cached dataset not found. Please upload the file again.')

        available_columns = list(dataset_entry['columns'])
        resolved_columns = resolve_requested_columns(required_columns, available_columns)
        resolved_selected_columns = [resolved_columns[column] for column in required_columns]

        if dataset_entry.get('frame_path'):
            frame = read_cached_frame(dataset_entry, columns=resolved_selected_columns)
            frame.columns = required_columns
            return normalize_dataframe(frame)

        if dataset_entry.get('parquet_path'):
            parquet_frame = read_cached_parquet(dataset_entry, columns=resolved_selected_columns, low_memory=True)
            parquet_frame.columns = required_columns
            return normalize_dataframe(parquet_frame.to_pandas(use_pyarrow_extension_array=False))

        if dataset_entry.get('csv_path'):
            frame = read_cached_csv(dataset_entry, columns=resolved_selected_columns)
            frame.columns = required_columns
            return normalize_dataframe(frame)

        if dataset_entry.get('excel_path'):
            frame = read_cached_excel(dataset_entry, columns=resolved_selected_columns)
            frame.columns = required_columns
            return normalize_dataframe(frame)

    if not data:
        raise HTTPException(status_code=400, detail='Dataset rows are required.')

    frame = normalize_dataframe(pd.DataFrame(data))
    missing_columns = [column for column in required_columns if column not in frame.columns]
    if missing_columns:
        raise HTTPException(status_code=400, detail=f'Missing columns: {missing_columns}')
    return frame[required_columns].copy()


def load_full_dataset_frame(dataset_id: str | None, data: list[dict[str, Any]]) -> pd.DataFrame:
    if dataset_id:
        dataset_entry = DATASET_CACHE.get(dataset_id)
        if dataset_entry is None:
            raise HTTPException(status_code=400, detail='Cached dataset not found. Please upload the file again.')

        if dataset_entry.get('frame_path'):
            return normalize_dataframe(read_cached_frame(dataset_entry))

        if dataset_entry.get('parquet_path'):
            parquet_frame = read_cached_parquet(dataset_entry, low_memory=True)
            return normalize_dataframe(parquet_frame.to_pandas(use_pyarrow_extension_array=False))

        if dataset_entry.get('csv_path'):
            return normalize_dataframe(read_cached_csv(dataset_entry))

        if dataset_entry.get('excel_path'):
            return normalize_dataframe(read_cached_excel(dataset_entry))

    if not data:
        raise HTTPException(status_code=400, detail='Dataset rows are required.')

    frame = normalize_dataframe(pd.DataFrame(data))
    if frame.empty or frame.shape[1] == 0:
        raise HTTPException(status_code=400, detail='Dataset must contain at least one row and one column.')
    return frame


def safe_numeric_series(series: pd.Series) -> pd.Series:
    numeric = pd.to_numeric(series, errors='coerce')
    numeric = numeric.replace([np.inf, -np.inf], np.nan).dropna()
    return numeric.astype(float)


def is_identifier_like_name(column_name: str) -> bool:
    normalized = re.sub(r'[^a-z0-9]+', '_', str(column_name).strip().lower())
    return bool(re.search(r'(^|_)(id|uuid|guid|token|hash|key|session)(_|$)', normalized))


def is_identifier_like_categorical(series: pd.Series, column_name: str) -> bool:
    if is_identifier_like_name(column_name):
        return True
    cleaned = series.dropna().astype(str).replace(r'^\s*$', np.nan, regex=True).dropna()
    if cleaned.empty:
        return False
    total_rows = max(1, len(cleaned))
    unique_count = int(cleaned.nunique())
    unique_ratio = unique_count / total_rows
    average_length = float(cleaned.str.len().mean()) if not cleaned.empty else 0.0

    # Heuristic 1: High cardinality + high uniqueness ratio (adapts to any dataset size)
    if unique_count >= 50 and unique_ratio > 0.90:
        return True

    # Heuristic 2: Alphanumeric ID pattern (e.g., "ORD12345", "INV-2024-001", "AZ-001")
    alpha_num_ratio = cleaned.str.match(r'^[A-Za-z]+[-_ ]?[0-9]+$', na=False).mean()
    if alpha_num_ratio > 0.50:
        return True

    # Heuristic 3: Mostly-digit strings (phone numbers, zip codes, account numbers stored as strings)
    mostly_digit_ratio = cleaned.str.replace(r'[\s\-\(\)\.]', '', regex=True, n=0).str.match(r'^\d{5,}$', na=False).mean()
    if mostly_digit_ratio > 0.50:
        return True

    # Heuristic 4: UUID / GUID patterns
    uuid_ratio = cleaned.str.match(
        r'^[0-9a-fA-F]{8}[-]?[0-9a-fA-F]{4}[-]?[0-9a-fA-F]{4}[-]?[0-9a-fA-F]{4}[-]?[0-9a-fA-F]{12}$',
        na=False,
    ).mean()
    if uuid_ratio > 0.50:
        return True

    # Heuristic 5: Long average length combined with moderate uniqueness (code-like fields)
    if unique_count >= 30 and unique_ratio >= 0.80 and average_length >= 12:
        return True

    return False


def is_identifier_like_numeric(series: pd.Series, column_name: str) -> bool:
    if is_identifier_like_name(column_name):
        return True
    values = safe_numeric_series(series)
    if values.empty:
        return False
    unique_count = int(values.nunique())
    unique_ratio = unique_count / max(1, len(values))
    return unique_count >= 50 and unique_ratio >= 0.98


def figure_to_base64(figure: go.Figure, *, width: int = 1400, height: int = 700) -> str | None:
    if not ENABLE_PLOTLY_STATIC_EXPORT:
        return None
    try:
        figure.update_layout(
            template='plotly_white',
            paper_bgcolor='white',
            plot_bgcolor='white',
            margin=dict(l=40, r=30, t=60, b=40),
        )
        image_bytes = figure.to_image(format='png', width=width, height=height, scale=2)
        return f"data:image/png;base64,{base64.b64encode(image_bytes).decode('ascii')}"
    except Exception:
        logger.exception('Advanced EDA chart rendering failed')
        return None


def matplotlib_figure_to_base64(fig: plt.Figure) -> str | None:
    try:
        buffer = io.BytesIO()
        fig.savefig(buffer, format='png', dpi=180, bbox_inches='tight', facecolor='white')
        buffer.seek(0)
        return f"data:image/png;base64,{base64.b64encode(buffer.read()).decode('ascii')}"
    except Exception:
        logger.exception('Matplotlib chart rendering failed')
        return None
    finally:
        plt.close(fig)


def build_missingness_chart_matplotlib(matrix: list[list[float]], columns: list[str], y_labels: list[str]) -> str | None:
    try:
        fig_width = max(10, min(20, 4 + len(columns) * 0.7))
        fig_height = max(4.5, min(12, 2.8 + len(y_labels) * 0.12))
        fig, ax = plt.subplots(figsize=(fig_width, fig_height))
        image = ax.imshow(np.array(matrix), aspect='auto', cmap='BuGn', vmin=0, vmax=100)
        ax.set_title('Missingness intensity across row groups', fontsize=12, fontweight='bold')
        ax.set_xlabel('Columns with missing values')
        ax.set_ylabel('Row index percentile groups')
        ax.set_xticks(np.arange(len(columns)))
        ax.set_xticklabels(columns, rotation=35, ha='right', fontsize=8)
        ax.set_yticks(np.arange(len(y_labels)))
        ax.set_yticklabels(y_labels, fontsize=8)
        colorbar = fig.colorbar(image, ax=ax, fraction=0.025, pad=0.03)
        colorbar.set_label('Missing %')
        fig.tight_layout()
        return matplotlib_figure_to_base64(fig)
    except Exception:
        logger.exception('Missingness matplotlib fallback failed')
        return None


def build_distribution_chart_matplotlib(frame: pd.DataFrame, selected_columns: list[str]) -> str | None:
    try:
        fig, axes = plt.subplots(len(selected_columns), 2, figsize=(14, max(4, len(selected_columns) * 2.8)))
        axes_array = np.atleast_2d(axes)
        for row_index, column in enumerate(selected_columns):
            values = safe_numeric_series(frame[column])
            hist_ax = axes_array[row_index, 0]
            box_ax = axes_array[row_index, 1]
            if values.empty:
                hist_ax.text(0.5, 0.5, 'No numeric data', ha='center', va='center', fontsize=9, color='#64748b')
                box_ax.text(0.5, 0.5, 'No numeric data', ha='center', va='center', fontsize=9, color='#64748b')
                hist_ax.set_axis_off()
                box_ax.set_axis_off()
                continue
            raw_values = values.to_numpy(dtype=float)
            hist_ax.hist(raw_values, bins=min(40, max(12, int(np.sqrt(raw_values.size)))), density=True, color='#38bdf8', alpha=0.75, edgecolor='white')
            kde = estimate_kde(raw_values)
            if kde is not None:
                hist_ax.plot(kde[0], kde[1], color='#7c3aed', linewidth=2)
            hist_ax.set_title(f'{column} distribution', fontsize=10, fontweight='bold')
            hist_ax.grid(alpha=0.18)
            box_ax.boxplot(raw_values, vert=False, patch_artist=True, boxprops=dict(facecolor='#10b981', alpha=0.55), medianprops=dict(color='#065f46', linewidth=2))
            box_ax.set_title(f'{column} outlier view', fontsize=10, fontweight='bold')
            box_ax.grid(alpha=0.18)
        fig.tight_layout()
        return matplotlib_figure_to_base64(fig)
    except Exception:
        logger.exception('Distribution matplotlib fallback failed')
        return None


def build_distribution_chart_for_column(frame: pd.DataFrame, column: str) -> str | None:
    values = safe_numeric_series(frame[column])
    if values.empty:
        return None
    raw_values = values.to_numpy(dtype=float)
    figure = make_subplots(
        rows=1,
        cols=2,
        subplot_titles=(f'{column} distribution', f'{column} outlier view'),
        horizontal_spacing=0.1,
    )
    figure.add_trace(
        go.Histogram(
            x=raw_values,
            nbinsx=min(40, max(12, int(np.sqrt(raw_values.size)))),
            histnorm='probability density',
            marker=dict(color='rgba(14,165,233,0.78)'),
            name=f'{column} histogram',
            showlegend=False,
        ),
        row=1,
        col=1,
    )
    kde = estimate_kde(raw_values)
    if kde is not None:
        figure.add_trace(
            go.Scatter(
                x=kde[0],
                y=kde[1],
                mode='lines',
                line=dict(color='#7c3aed', width=2),
                name=f'{column} KDE',
                showlegend=False,
            ),
            row=1,
            col=1,
        )
    figure.add_trace(
        go.Box(
            x=raw_values,
            orientation='h',
            marker=dict(color='#10b981'),
            line=dict(color='#047857'),
            boxmean='sd',
            name=f'{column} boxplot',
            showlegend=False,
        ),
        row=1,
        col=2,
    )
    figure.update_layout(
        title=f'{column}: distribution and outlier screening',
        title_x=0.5,
        width=1380,
        height=520,
        margin=dict(l=90, r=50, t=90, b=80),
        template='plotly_white',
    )
    figure.update_xaxes(title_text='Value', automargin=True, tickangle=0)
    figure.update_yaxes(automargin=True)
    chart_base64 = figure_to_base64(figure, width=1380, height=520)
    if chart_base64 is not None:
        return chart_base64
    return build_distribution_chart_matplotlib(frame, [column])


def build_categorical_chart_matplotlib(counts: pd.Series, title: str) -> str | None:
    try:
        labels = [str(value) for value in counts.index[::-1]]
        values = counts.values[::-1]
        fig_height = max(4, 1.2 + len(labels) * 0.5)
        fig, ax = plt.subplots(figsize=(10, fig_height))
        ax.barh(labels, values, color='#0ea5e9', alpha=0.85)
        ax.set_title(title, fontsize=12, fontweight='bold')
        ax.set_xlabel('Count')
        ax.grid(axis='x', alpha=0.18)
        fig.tight_layout()
        return matplotlib_figure_to_base64(fig)
    except Exception:
        logger.exception('Categorical matplotlib fallback failed')
        return None


def build_interaction_chart_matplotlib(x_values: np.ndarray, y_values: np.ndarray, title: str, x_label: str, y_label: str) -> str | None:
    try:
        fig, ax = plt.subplots(figsize=(10, 4.8))
        ax.scatter(x_values, y_values, s=22, alpha=0.7, color='#0ea5e9')
        if np.unique(x_values).size > 1:
            slope, intercept = np.polyfit(x_values, y_values, 1)
            line_x = np.linspace(float(x_values.min()), float(x_values.max()), 100)
            line_y = slope * line_x + intercept
            ax.plot(line_x, line_y, color='#7c3aed', linewidth=2)
        ax.set_title(title, fontsize=12, fontweight='bold')
        ax.set_xlabel(x_label)
        ax.set_ylabel(y_label)
        ax.grid(alpha=0.18)
        fig.tight_layout()
        return matplotlib_figure_to_base64(fig)
    except Exception:
        logger.exception('Interaction matplotlib fallback failed')
        return None


def estimate_kde(values: np.ndarray, point_count: int = 160) -> tuple[np.ndarray, np.ndarray] | None:
    if values.size < 2:
        return None
    std = float(np.std(values))
    min_value = float(np.min(values))
    max_value = float(np.max(values))
    if not np.isfinite(std) or std == 0 or min_value == max_value:
        return None
    bandwidth = 1.06 * std * (values.size ** (-1 / 5))
    if not np.isfinite(bandwidth) or bandwidth <= 0:
        bandwidth = max((max_value - min_value) / 25, 1e-6)
    x_points = np.linspace(min_value, max_value, point_count)
    scaled = (x_points[:, None] - values[None, :]) / bandwidth
    density = np.exp(-0.5 * scaled ** 2).sum(axis=1) / (values.size * bandwidth * np.sqrt(2 * np.pi))
    return x_points, density


def build_missingness_payload(frame: pd.DataFrame) -> dict[str, Any]:
    total_missing = int(frame.isna().sum().sum())
    if total_missing == 0:
        return {
            'status': 'success',
            'message': 'Data Quality: No missing values detected in the dataset.',
            'total_missing': 0,
            'chart_base64': None,
            'columns_analyzed': [],
            'row_groups': 0,
        }

    missing_columns = [str(column) for column in frame.columns if frame[column].isna().any()]
    display_columns = missing_columns[:EDA_MAX_MISSINGNESS_COLUMNS]
    display_frame = frame[display_columns].isna().astype(float).reset_index(drop=True)
    bucket_count = min(EDA_MISSINGNESS_BUCKETS, max(8, len(display_frame)))
    bucket_edges = np.linspace(0, len(display_frame), num=bucket_count + 1, dtype=int)

    matrix: list[list[float]] = []
    y_labels: list[str] = []
    for index in range(bucket_count):
        start = int(bucket_edges[index])
        end = int(bucket_edges[index + 1])
        if end <= start:
            continue
        segment = display_frame.iloc[start:end]
        matrix.append((segment.mean(axis=0) * 100).round(2).tolist())
        start_pct = int(round((start / max(len(display_frame), 1)) * 100))
        end_pct = int(round((end / max(len(display_frame), 1)) * 100))
        y_labels.append(f'{start_pct}-{end_pct}%')

    figure = go.Figure(
        data=go.Heatmap(
            z=matrix,
            x=display_columns,
            y=y_labels,
            colorscale='Teal',
            colorbar=dict(title='Missing %'),
            hovertemplate='Row Group %{y}<br>Column %{x}<br>Missing %{z:.2f}%<extra></extra>',
        )
    )
    figure.update_layout(
        title='Missingness intensity across row groups',
        xaxis_title='Columns with missing values',
        yaxis_title='Row index percentile groups',
        margin=dict(l=100, r=40, t=80, b=120),
        xaxis=dict(tickangle=-45, automargin=True),
        yaxis=dict(automargin=True),
    )
    chart_base64 = figure_to_base64(
        figure,
        width=max(1100, 180 + len(display_columns) * 85),
        height=max(500, 220 + len(y_labels) * 8),
    )
    if chart_base64 is None:
        chart_base64 = build_missingness_chart_matplotlib(matrix, display_columns, y_labels)
    if chart_base64 is None:
        return {
            'status': 'error',
            'message': 'Missingness visualization could not be generated for this dataset.',
            'total_missing': total_missing,
            'chart_base64': None,
            'columns_analyzed': display_columns,
            'row_groups': len(y_labels),
        }
    note = None
    if len(missing_columns) > len(display_columns):
        note = f'Displaying the first {len(display_columns)} columns with missing values to keep the view stable.'
    return {
        'status': 'chart',
        'message': note,
        'total_missing': total_missing,
        'chart_base64': chart_base64,
        'columns_analyzed': display_columns,
        'row_groups': len(y_labels),
    }


def build_distribution_payload(frame: pd.DataFrame) -> dict[str, Any]:
    numeric_columns = [
        str(column) for column in frame.columns
        if pd.api.types.is_numeric_dtype(frame[column]) and not pd.api.types.is_bool_dtype(frame[column])
    ]
    if not numeric_columns:
        return {
            'status': 'empty',
            'message': 'No numeric columns available for this analysis.',
            'chart_base64': None,
            'columns_analyzed': [],
            'charts': [],
        }

    selected_columns = numeric_columns[:EDA_MAX_NUMERIC_CHARTS]
    chart_payloads = [
        {
            'column': column,
            'chart_base64': build_distribution_chart_for_column(frame, column),
        }
        for column in selected_columns
    ]
    chart_payloads = [payload for payload in chart_payloads if payload['chart_base64']]
    if not chart_payloads:
        return {
            'status': 'error',
            'message': 'Distribution charts could not be generated for this dataset.',
            'chart_base64': None,
            'columns_analyzed': selected_columns,
            'charts': [],
        }
    return {
        'status': 'chart',
        'message': None if len(numeric_columns) <= len(selected_columns) else f'Displaying the first {len(selected_columns)} numeric columns to prevent oversized browser rendering.',
        'chart_base64': chart_payloads[0]['chart_base64'],
        'columns_analyzed': selected_columns,
        'charts': chart_payloads,
    }


def build_categorical_payload(frame: pd.DataFrame) -> dict[str, Any]:
    all_categorical_columns = [
        str(column) for column in frame.columns
        if (
            pd.api.types.is_object_dtype(frame[column])
            or pd.api.types.is_string_dtype(frame[column])
            or pd.api.types.is_categorical_dtype(frame[column])
            or pd.api.types.is_bool_dtype(frame[column])
        )
    ]
    excluded_columns: list[str] = []
    categorical_columns: list[str] = []
    for column in all_categorical_columns:
        if is_identifier_like_categorical(frame[column], column):
            excluded_columns.append(column)
        else:
            categorical_columns.append(column)

    total_categorical = len(all_categorical_columns)
    excluded_count = len(excluded_columns)
    selected_count = len(categorical_columns)

    logger.info(
        'Categorical analysis: %d total categorical columns, %d excluded as identifier-like, %d selected for charting',
        total_categorical, excluded_count, selected_count,
    )

    if not categorical_columns:
        if excluded_columns:
            excluded_list = ', '.join(f"'{c}'" for c in excluded_columns[:8])
            remainder = excluded_count - 8
            if remainder > 0:
                excluded_list += f' and {remainder} more'
            message = (
                f'All {excluded_count} categorical-looking {"column" if excluded_count == 1 else "columns"} '
                f'appear to be identifier-like fields ({excluded_list}) and were excluded from charting. '
                'Identifier-like fields are detected using uniqueness ratio, cardinality, and value patterns.'
            )
        else:
            message = 'No categorical columns available for this analysis.'
        return {
            'status': 'empty',
            'message': message,
            'charts': [],
            'warnings': [],
            'excluded_columns': excluded_columns,
            'total_categorical': total_categorical,
            'excluded_count': excluded_count,
            'selected_count': selected_count,
        }

    warnings_payload: list[dict[str, Any]] = []
    for column in categorical_columns:
        unique_count = int(frame[column].dropna().astype(str).replace(r'^\s*$', np.nan, regex=True).dropna().nunique())
        if unique_count > 20:
            warnings_payload.append({
                'column': column,
                'unique_count': unique_count,
                'message': f"High Cardinality: Column '{column}' has {unique_count} unique values. Consider encoding strategies before ML.",
            })

    chart_payloads: list[dict[str, Any]] = []
    for column in categorical_columns[:EDA_MAX_CATEGORICAL_CHARTS]:
        series = frame[column].copy()
        series = series.astype(object).where(pd.notna(series), 'Missing')
        labels = pd.Series(series).astype(str).replace(r'^\s*$', '(blank)', regex=True)
        counts = labels.value_counts().head(EDA_MAX_CATEGORY_BARS)
        if counts.empty:
            continue
        longest_label = max((len(str(value)) for value in counts.index), default=0)
        left_margin = min(420, max(220, 110 + longest_label * 7))
        figure = go.Figure(
            data=go.Bar(
                x=counts.values[::-1],
                y=[str(value) for value in counts.index[::-1]],
                orientation='h',
                marker=dict(color='#0ea5e9'),
                hovertemplate='%{y}<br>Count %{x}<extra></extra>',
            )
        )
        figure.update_layout(
            title=f'Top categories for {column}',
            xaxis_title='Count',
            yaxis_title='Category',
            margin=dict(l=left_margin, r=50, t=80, b=60),
            yaxis=dict(automargin=True, tickfont=dict(size=11)),
            xaxis=dict(automargin=True, tickfont=dict(size=11)),
        )
        chart_base64 = figure_to_base64(figure, width=1440, height=max(480, 140 + len(counts) * 42))
        if chart_base64 is None:
            chart_base64 = build_categorical_chart_matplotlib(counts, f'Top categories for {column}')
        chart_payloads.append({
            'column': column,
            'unique_count': int(labels.nunique()),
            'chart_base64': chart_base64,
        })

    status = 'chart' if chart_payloads else 'error'
    message_parts: list[str] = []
    if excluded_count > 0:
        excluded_list = ', '.join(f"'{c}'" for c in excluded_columns[:5])
        remainder = excluded_count - 5
        if remainder > 0:
            excluded_list += f' and {remainder} more'
        message_parts.append(f'Excluded {excluded_count} identifier-like categorical {"column" if excluded_count == 1 else "columns"} ({excluded_list}) from charting.')
    if len(categorical_columns) > EDA_MAX_CATEGORICAL_CHARTS:
        message_parts.append(f'Displaying the first {EDA_MAX_CATEGORICAL_CHARTS} categorical columns to keep the analysis responsive.')
    message = ' '.join(message_parts) or None
    if status == 'error':
        message = 'Categorical charts could not be generated for this dataset.'
    return {
        'status': status,
        'message': message,
        'charts': chart_payloads,
        'warnings': warnings_payload[:12],
        'excluded_columns': excluded_columns,
        'total_categorical': total_categorical,
        'excluded_count': excluded_count,
        'selected_count': selected_count,
    }


def build_interaction_payload(frame: pd.DataFrame) -> dict[str, Any]:
    numeric_columns = [
        str(column) for column in frame.columns
        if pd.api.types.is_numeric_dtype(frame[column]) and not pd.api.types.is_bool_dtype(frame[column])
    ]
    if len(numeric_columns) < 2:
        return {
            'status': 'empty',
            'message': 'Need at least 2 numeric columns.',
            'plots': [],
        }

    eligible_numeric_columns = [column for column in numeric_columns if not is_identifier_like_numeric(frame[column], column)]
    numeric_frame = frame[eligible_numeric_columns].apply(pd.to_numeric, errors='coerce')
    numeric_frame = numeric_frame.replace([np.inf, -np.inf], np.nan)
    non_constant_columns = [column for column in numeric_frame.columns if numeric_frame[column].dropna().nunique() > 1]
    numeric_frame = numeric_frame[non_constant_columns[:EDA_MAX_INTERACTION_COLUMNS]]
    if numeric_frame.shape[1] < 2:
        excluded_identifier_like = len(numeric_columns) - len(eligible_numeric_columns)
        message = 'Need at least 2 numeric columns.'
        if excluded_identifier_like > 0:
            message = f'Need at least 2 non-identifier numeric columns. Excluded {excluded_identifier_like} identifier-like numeric column{"s" if excluded_identifier_like != 1 else ""}.'
        return {
            'status': 'empty',
            'message': message,
            'plots': [],
        }

    correlation_matrix = numeric_frame.corr().fillna(0)
    pairs: list[dict[str, Any]] = []
    for index, left in enumerate(correlation_matrix.columns):
        for right in correlation_matrix.columns[index + 1:]:
            correlation = float(correlation_matrix.loc[left, right])
            if np.isnan(correlation):
                correlation = 0.0
            pairs.append({'x': str(left), 'y': str(right), 'correlation': correlation})
    pairs.sort(key=lambda item: abs(item['correlation']), reverse=True)

    plots: list[dict[str, Any]] = []
    for pair in pairs[:EDA_MAX_INTERACTION_PAIRS]:
        pair_frame = numeric_frame[[pair['x'], pair['y']]].dropna()
        if len(pair_frame) < 2:
            continue
        x_values = pair_frame[pair['x']].to_numpy(dtype=float)
        y_values = pair_frame[pair['y']].to_numpy(dtype=float)
        figure = go.Figure()
        figure.add_trace(
            go.Scatter(
                x=x_values,
                y=y_values,
                mode='markers',
                marker=dict(color='#0ea5e9', size=8, opacity=0.7),
                name='Observed values',
                showlegend=False,
            )
        )
        if np.unique(x_values).size > 1:
            slope, intercept = np.polyfit(x_values, y_values, 1)
            line_x = np.linspace(float(x_values.min()), float(x_values.max()), 100)
            line_y = slope * line_x + intercept
            figure.add_trace(
                go.Scatter(
                    x=line_x,
                    y=line_y,
                    mode='lines',
                    line=dict(color='#7c3aed', width=2),
                    name='OLS trend',
                    showlegend=False,
                )
            )
        figure.update_layout(
            title=f"{pair['x']} vs {pair['y']}",
            xaxis_title=pair['x'],
            yaxis_title=pair['y'],
            margin=dict(l=130, r=50, t=80, b=100),
            xaxis=dict(tickangle=-20, automargin=True, tickfont=dict(size=11)),
            yaxis=dict(automargin=True, tickfont=dict(size=11)),
        )
        chart_base64 = figure_to_base64(figure, width=1400, height=560)
        if chart_base64 is None:
            chart_base64 = build_interaction_chart_matplotlib(x_values, y_values, f"{pair['x']} vs {pair['y']}", pair['x'], pair['y'])
        plots.append({
            'pair': f"{pair['x']} vs {pair['y']}",
            'correlation': round(pair['correlation'], 4),
            'chart_base64': chart_base64,
        })

    if not plots:
        return {
            'status': 'empty',
            'message': 'Need at least 2 numeric columns.',
            'plots': [],
        }
    message_parts: list[str] = []
    excluded_identifier_like = len(numeric_columns) - len(eligible_numeric_columns)
    if excluded_identifier_like > 0:
        message_parts.append(f'Excluded {excluded_identifier_like} identifier-like numeric column{"s" if excluded_identifier_like != 1 else ""} from interaction analysis.')
    if len(non_constant_columns) > EDA_MAX_INTERACTION_COLUMNS:
        message_parts.append(f'Interaction search was capped to the first {EDA_MAX_INTERACTION_COLUMNS} non-constant numeric columns for stability.')
    return {
        'status': 'chart',
        'message': ' '.join(message_parts) or None,
        'plots': plots,
    }


def build_automated_insights(frame: pd.DataFrame) -> dict[str, Any]:
    insights: list[str] = []

    try:
        numeric_columns = [
            str(column) for column in frame.columns
            if pd.api.types.is_numeric_dtype(frame[column]) and not pd.api.types.is_bool_dtype(frame[column])
        ]
    except Exception:
        numeric_columns = []

    for column in numeric_columns[:EDA_MAX_INTERACTION_COLUMNS]:
        try:
            values = safe_numeric_series(frame[column])
            if len(values) < 8:
                continue
            skewness = float(values.skew())
            if np.isfinite(skewness) and abs(skewness) > 1:
                insights.append(f"'{column}' is highly skewed (Skew: {skewness:.2f}). Consider Log/Box-Cox transformation.")
        except Exception:
            continue

    try:
        if len(numeric_columns) >= 2:
            numeric_frame = frame[numeric_columns[:EDA_MAX_INTERACTION_COLUMNS]].apply(pd.to_numeric, errors='coerce')
            numeric_frame = numeric_frame.replace([np.inf, -np.inf], np.nan)
            numeric_frame = numeric_frame[[column for column in numeric_frame.columns if numeric_frame[column].dropna().nunique() > 1]]
            if numeric_frame.shape[1] >= 2:
                corr_matrix = numeric_frame.corr().fillna(0)
                for index, left in enumerate(corr_matrix.columns):
                    for right in corr_matrix.columns[index + 1:]:
                        corr_value = float(corr_matrix.loc[left, right])
                        if abs(corr_value) > 0.95:
                            insights.append(f"'{left}' and '{right}' are highly correlated (>0.95). Consider dropping one to prevent multicollinearity.")
    except Exception:
        pass

    for column in numeric_columns[:EDA_MAX_INTERACTION_COLUMNS]:
        try:
            values = safe_numeric_series(frame[column])
            if len(values) < 8:
                continue
            q1 = float(values.quantile(0.25))
            q3 = float(values.quantile(0.75))
            iqr = q3 - q1
            if not np.isfinite(iqr) or iqr <= 0:
                continue
            lower = q1 - 1.5 * iqr
            upper = q3 + 1.5 * iqr
            outlier_count = int(((values < lower) | (values > upper)).sum())
            if outlier_count >= max(3, int(len(values) * 0.01)):
                insights.append(f"Extreme outliers detected in '{column}'.")
        except Exception:
            continue

    deduped_insights = list(dict.fromkeys(insights))
    return {
        'status': 'success',
        'message': 'No major statistical anomalies detected.' if not deduped_insights else None,
        'insights': deduped_insights[:20],
    }


def build_advanced_eda_payload(request: AdvancedEdaRequest) -> dict[str, Any]:
    if request.dataset_id:
        dataset_entry = DATASET_CACHE.get(request.dataset_id)
        if dataset_entry is None:
            raise HTTPException(status_code=400, detail='Cached dataset not found. Please upload the file again.')
        analysis_frame, total_rows = load_cached_analysis_frame(dataset_entry)
        if analysis_frame.empty or analysis_frame.shape[1] == 0:
            raise HTTPException(status_code=400, detail='Dataset must contain at least one row and one column.')
        row_count = total_rows if total_rows > 0 else int(len(analysis_frame))
        column_count = int(dataset_entry.get('column_count') or len(analysis_frame.columns))
    else:
        frame = load_full_dataset_frame(request.dataset_id, request.data)
        if frame.empty or frame.shape[1] == 0:
            raise HTTPException(status_code=400, detail='Dataset must contain at least one row and one column.')
        analysis_frame = frame
        row_count = int(len(frame))
        column_count = int(len(frame.columns))

    return {
        'row_count': row_count,
        'sampled_row_count': int(len(analysis_frame)),
        'column_count': column_count,
        'missingness': build_missingness_payload(analysis_frame),
        'distributions': build_distribution_payload(analysis_frame),
        'categorical': build_categorical_payload(analysis_frame),
        'interactions': build_interaction_payload(analysis_frame),
        'insights': build_automated_insights(analysis_frame),
    }


def build_polars_datetime_expr(column_name: str, dtype: pl.DataType) -> pl.Expr:
    column = pl.col(column_name)
    if dtype in pl.TEMPORAL_DTYPES:
        return column.cast(pl.Datetime, strict=False)

    text_column = column.cast(pl.String, strict=False).str.strip_chars()
    return pl.coalesce([
        text_column.str.strptime(pl.Datetime, '%Y-%m-%d %H:%M:%S%.f', strict=False),
        text_column.str.strptime(pl.Datetime, '%Y-%m-%d %H:%M:%S', strict=False),
        text_column.str.strptime(pl.Datetime, '%Y-%m-%dT%H:%M:%S%.f', strict=False),
        text_column.str.strptime(pl.Datetime, '%Y-%m-%dT%H:%M:%S', strict=False),
        text_column.str.strptime(pl.Datetime, '%Y-%m-%d', strict=False),
        text_column.str.strptime(pl.Datetime, '%d-%m-%Y', strict=False),
        text_column.str.strptime(pl.Datetime, '%m-%d-%Y', strict=False),
        text_column.str.strptime(pl.Datetime, '%d/%m/%Y', strict=False),
        text_column.str.strptime(pl.Datetime, '%m/%d/%Y', strict=False),
        text_column.str.strptime(pl.Datetime, '%Y/%m/%d', strict=False),
        text_column.str.strptime(pl.Datetime, '%b %d, %Y', strict=False),
        text_column.str.strptime(pl.Datetime, '%B %d, %Y', strict=False),
        text_column.str.strptime(pl.Datetime, '%b-%Y', strict=False),
        text_column.str.strptime(pl.Datetime, '%B-%Y', strict=False),
        text_column.str.strptime(pl.Datetime, '%b %Y', strict=False),
        text_column.str.strptime(pl.Datetime, '%B %Y', strict=False),
        text_column.str.strptime(pl.Datetime, '%Y-%m', strict=False),
        text_column.str.strptime(pl.Datetime, '%m-%Y', strict=False),
    ])


def prepare_sales_series_from_parquet(dataset_entry: dict[str, Any], date_column: str, target_column: str) -> tuple[pd.DataFrame, str, str]:
    parquet_path = dataset_entry.get('parquet_path')
    if not parquet_path:
        raise HTTPException(status_code=400, detail='Cached parquet dataset path is missing. Please upload the file again.')

    try:
        lazy_frame = pl.scan_parquet(parquet_path).select([
            pl.col(date_column).alias(date_column),
            pl.col(target_column).alias(target_column),
        ])
        schema = lazy_frame.collect_schema()
        parsed_date_expr = build_polars_datetime_expr(date_column, schema[date_column])
        sample_dates = lazy_frame.select(parsed_date_expr.alias('__parsed_date')).limit(5000).collect(streaming=True).to_series()
        parsed_sample = pd.to_datetime(pd.Series(sample_dates.to_list()), errors='coerce')
        parsed_sample = parsed_sample.dropna()
        if parsed_sample.empty:
            raise HTTPException(status_code=400, detail='No valid rows remained after parsing the date and sales columns.')

        freq, period_label = infer_sales_time_frequency(parsed_sample)
        period_freq = {'MS': '1mo', 'QS': '1q', 'YS': '1y', 'D': '1d', 'W-MON': '1w'}.get(freq, '1mo')

        aggregated = (
            lazy_frame
            .with_columns([
                parsed_date_expr.alias('__parsed_date'),
                pl.col(target_column).cast(pl.Float64, strict=False).alias('__parsed_sales'),
            ])
            .drop_nulls(['__parsed_date'])
            .group_by_dynamic('__parsed_date', every=period_freq, label='left')
            .agg([
                pl.col('__parsed_sales').sum().alias('sales'),
                pl.col('__parsed_sales').is_not_null().sum().alias('__valid_sales_count'),
            ])
            .with_columns(
                pl.when(pl.col('__valid_sales_count') > 0)
                .then(pl.col('sales'))
                .otherwise(None)
                .alias('sales')
            )
            .drop('__valid_sales_count')
            .sort('__parsed_date')
            .collect(streaming=True)
        )
    except HTTPException:
        raise
    except Exception as error:
        raise HTTPException(status_code=400, detail=f'Failed to prepare parquet sales series: {error}') from error

    if aggregated.height == 0:
        raise HTTPException(status_code=400, detail='No valid rows remained after parsing the date and sales columns.')

    series_frame = aggregated.rename({'__parsed_date': 'period'}).to_pandas(use_pyarrow_extension_array=False)
    full_range = pd.date_range(series_frame['period'].min(), series_frame['period'].max(), freq=freq)
    series_frame = series_frame.set_index('period').reindex(full_range).rename_axis('period').reset_index()
    series_frame['sales'] = series_frame['sales'].astype(float)
    series_frame = repair_zero_period_values_with_seasonal_interpolation(series_frame, 'sales')

    minimum_periods = 2 * seasonal_period_for_label(period_label)
    if len(series_frame) < minimum_periods:
        needed = minimum_periods - len(series_frame)
        raise HTTPException(status_code=422, detail=f'Forecasting needs at least {minimum_periods} {period_label} rows after aggregation. Add {needed} more {period_label} row{"s" if needed != 1 else ""} and rerun the forecast.')

    return series_frame, freq, period_label


def prepare_sales_series_from_cached_dataset(dataset_entry: dict[str, Any], date_column: str, target_column: str) -> tuple[pd.DataFrame, str, str]:
    required_columns = [date_column, target_column]
    available_columns = list(dataset_entry['columns'])
    resolved_columns = resolve_requested_columns(required_columns, available_columns)
    resolved_date_column = resolved_columns[date_column]
    resolved_target_column = resolved_columns[target_column]

    if dataset_entry.get('parquet_path'):
        return prepare_sales_series_from_parquet(dataset_entry, resolved_date_column, resolved_target_column)

    if dataset_entry.get('frame_path'):
        frame = read_cached_frame(dataset_entry, columns=[resolved_date_column, resolved_target_column])
    elif dataset_entry.get('csv_path'):
        frame = read_cached_csv(dataset_entry, columns=[resolved_date_column, resolved_target_column])
    elif dataset_entry.get('excel_path'):
        frame = read_cached_excel(dataset_entry, columns=[resolved_date_column, resolved_target_column])
    else:
        raise HTTPException(status_code=400, detail='Cached dataset storage is missing. Please upload the file again.')

    frame.columns = required_columns
    frame = normalize_dataframe(frame)
    return prepare_sales_series(frame, date_column, target_column)


def infer_sales_time_frequency(dates: pd.Series) -> tuple[str, str]:
    ordered = pd.Series(pd.to_datetime(dates, errors='coerce')).dropna().sort_values().drop_duplicates()
    if len(ordered) < 2:
        return 'MS', 'month'

    deltas = ordered.diff().dropna().dt.total_seconds() / 86400.0
    median_days = float(deltas.median()) if not deltas.empty else 30.0

    if median_days <= 2:
        return 'D', 'day'
    if median_days <= 10:
        return 'W-MON', 'week'
    if median_days <= 45:
        return 'MS', 'month'
    if median_days <= 120:
        return 'QS', 'quarter'
    return 'YS', 'year'


def seasonal_period_for_label(period_label: str) -> int:
    return {'day': 365, 'week': 52, 'month': 12, 'quarter': 4, 'year': 1}.get(period_label, 12)


def statsmodels_frequency_for_label(period_label: str) -> str:
    return {'day': 'D', 'week': 'W', 'month': 'MS', 'quarter': 'QS', 'year': 'YS'}.get(period_label, 'MS')


def format_forecast_period(value: pd.Timestamp, period_label: str) -> str:
    timestamp = pd.Timestamp(value)
    if period_label == 'day':
        return timestamp.strftime('%Y-%m-%d')
    if period_label == 'week':
        return f"Week of {timestamp.strftime('%Y-%m-%d')}"
    if period_label == 'month':
        return timestamp.strftime('%Y-%m')
    if period_label == 'quarter':
        quarter = ((timestamp.month - 1) // 3) + 1
        return f"{timestamp.year}-Q{quarter}"
    return timestamp.strftime('%Y')


def prepare_sales_series(frame: pd.DataFrame, date_column: str, target_column: str) -> tuple[pd.DataFrame, str, str]:
    working = frame[[date_column, target_column]].copy()
    working[date_column] = pd.to_datetime(working[date_column], errors='coerce')
    working[target_column] = pd.to_numeric(working[target_column], errors='coerce')
    working = working.dropna(subset=[date_column])

    if working.empty:
        raise HTTPException(status_code=400, detail='No valid rows remained after parsing the date and sales columns.')

    freq, period_label = infer_sales_time_frequency(working[date_column])
    if freq == 'W-MON':
        period_index = working[date_column].dt.to_period('W').dt.start_time
    elif freq == 'D':
        period_index = working[date_column].dt.floor('D')
    else:
        period_freq = {'MS': 'M', 'QS': 'Q', 'YS': 'Y'}.get(freq, freq)
        period_index = working[date_column].dt.to_period(period_freq).dt.to_timestamp()
    working = working.assign(period=period_index)
    series_frame = working.groupby('period', as_index=False)[target_column].sum(min_count=1).sort_values('period')
    series_frame = series_frame.rename(columns={target_column: 'sales'})

    full_range = pd.date_range(series_frame['period'].min(), series_frame['period'].max(), freq=freq)
    series_frame = series_frame.set_index('period').reindex(full_range).rename_axis('period').reset_index()
    series_frame['sales'] = series_frame['sales'].astype(float)
    series_frame = repair_zero_period_values_with_seasonal_interpolation(series_frame, 'sales')

    minimum_periods = 2 * seasonal_period_for_label(period_label)
    if len(series_frame) < minimum_periods:
        needed = minimum_periods - len(series_frame)
        raise HTTPException(status_code=422, detail=f'Forecasting needs at least {minimum_periods} {period_label} rows after aggregation. Add {needed} more {period_label} row{"s" if needed != 1 else ""} and rerun the forecast.')

    return series_frame, freq, period_label


def build_forecast_feature_row(history: list[float], current_period: pd.Timestamp, lag_periods: int) -> dict[str, float]:
    if len(history) < lag_periods:
        raise ValueError('Not enough history to build forecast features.')

    month_number = float(current_period.month)
    quarter_number = float(current_period.quarter)
    day_of_month = float(current_period.day)
    day_of_week = float(current_period.dayofweek)

    row: dict[str, float] = {
        'trend_index': float(len(history) + 1),
        'month_number': month_number,
        'quarter_number': quarter_number,
        'day_of_month': day_of_month,
        'day_of_week': day_of_week,
        'month_sin': float(np.sin(2 * np.pi * month_number / 12)),
        'month_cos': float(np.cos(2 * np.pi * month_number / 12)),
        'quarter_sin': float(np.sin(2 * np.pi * quarter_number / 4)),
        'quarter_cos': float(np.cos(2 * np.pi * quarter_number / 4)),
        'lag_mean': float(np.mean(history[-lag_periods:])),
        'lag_last_3_mean': float(np.mean(history[-min(3, len(history)):])),
    }

    for lag_index in range(1, lag_periods + 1):
        row[f'lag_{lag_index}'] = float(history[-lag_index])

    return row


def build_forecast_training_frame(series_frame: pd.DataFrame, lag_periods: int) -> tuple[pd.DataFrame, pd.Series]:
    rows: list[dict[str, float]] = []
    targets: list[float] = []
    periods = series_frame['period'].tolist()
    values = series_frame['sales'].astype(float).tolist()

    for index in range(lag_periods, len(values)):
        history = values[:index]
        rows.append(build_forecast_feature_row(history, periods[index], lag_periods))
        targets.append(float(values[index]))

    if len(rows) < 3:
        raise HTTPException(status_code=400, detail='Not enough history is available to train the time-series model.')

    return pd.DataFrame(rows), pd.Series(targets)


def recursive_forecast(model: LinearRegression, history: list[float], start_period: pd.Timestamp, periods: int, lag_periods: int, freq: str, period_label: str) -> list[dict[str, Any]]:
    forecasts: list[dict[str, Any]] = []
    running_history = list(history)
    current_period = pd.Timestamp(start_period)

    for _ in range(periods):
        features = build_forecast_feature_row(running_history, current_period, lag_periods)
        prediction = float(model.predict(pd.DataFrame([features]))[0])
        prediction = max(prediction, 0.0)
        forecasts.append({
            'period': format_forecast_period(current_period, period_label),
            'predicted': round(prediction, 2),
        })
        running_history.append(prediction)
        current_period = current_period + pd.tseries.frequencies.to_offset(freq)

    return forecasts



def assess_overfitting(problem_type: ProblemType, metrics: dict[str, Any]) -> dict[str, Any]:
    train_metric_name = 'train_r2' if problem_type == 'regression' else 'train_accuracy'
    test_metric_name = 'test_r2' if problem_type == 'regression' else 'test_accuracy'

    train_score = float(metrics.get(train_metric_name, 0.0) or 0.0)
    test_score = float(metrics.get(test_metric_name, 0.0) or 0.0)
    cv_mean = float(metrics.get('cv_mean', 0.0) or 0.0)
    cv_available = bool(metrics.get('cv_scores'))

    generalization_gap = round(train_score - test_score, 6)
    cv_gap = round(train_score - cv_mean, 6) if cv_available else None

    status = 'healthy'
    explanation = 'Train and validation performance are reasonably aligned.'

    if problem_type == 'regression':
        if train_score >= 0.85 and (generalization_gap >= 0.15 or (cv_gap is not None and cv_gap >= 0.12)):
            status = 'detected'
            explanation = 'Training R2 is much higher than test/CV R2, suggesting the model is memorizing the training set.'
        elif train_score >= 0.75 and (generalization_gap >= 0.08 or (cv_gap is not None and cv_gap >= 0.08)):
            status = 'watch'
            explanation = 'There is a noticeable train-to-validation gap. Review feature leakage, model complexity, or test size.'
    else:
        if train_score >= 0.95 and (generalization_gap >= 0.08 or (cv_gap is not None and cv_gap >= 0.08)):
            status = 'detected'
            explanation = 'Training accuracy is materially above test/CV accuracy, which is a strong overfitting signal.'
        elif train_score >= 0.85 and (generalization_gap >= 0.04 or (cv_gap is not None and cv_gap >= 0.04)):
            status = 'watch'
            explanation = 'The model performs better on training than on held-out data. Monitor for overfitting.'

    if not cv_available and status == 'healthy':
        explanation = 'No strong overfitting signal was found from the train/test comparison.'

    return {
        'status': status,
        'detected': status == 'detected',
        'generalization_gap': generalization_gap,
        'cv_gap': cv_gap,
        'train_score': round(train_score, 6),
        'test_score': round(test_score, 6),
        'explanation': explanation,
    }


def calculate_forecast_metrics(actual: list[float], predicted: list[float]) -> dict[str, float]:
    if not actual or not predicted:
        return {'mae': 0.0, 'rmse': 0.0, 'mape': 0.0}

    mae = float(mean_absolute_error(actual, predicted))
    rmse = float(np.sqrt(mean_squared_error(actual, predicted)))
    percentage_errors = [abs((a - p) / a) for a, p in zip(actual, predicted) if a != 0]
    mape = float(np.mean(percentage_errors) * 100) if percentage_errors else 0.0
    return {
        'mae': round(mae, 2),
        'rmse': round(rmse, 2),
        'mape': round(mape, 2),
    }


def build_forecast_data_quality(series_frame: pd.DataFrame, period_label: str) -> dict[str, Any]:
    values = pd.to_numeric(series_frame['sales'], errors='coerce')
    total = max(1, len(series_frame))
    missing_share = float(values.isna().sum() / total)
    zero_share = float((values.fillna(0) <= 0).sum() / total)
    usable_periods = int(values.notna().sum())
    volatility = 0.0
    if usable_periods > 1 and float(abs(values.fillna(0).mean())) > 0:
        volatility = float(values.fillna(0).std() / abs(values.fillna(0).mean()))

    minimum_required = 2 * seasonal_period_for_label(period_label)
    period_score = min(1.0, usable_periods / minimum_required)
    completeness_score = max(0.0, 1.0 - missing_share)
    signal_score = max(0.0, 1.0 - min(0.7, zero_share))
    stability_score = max(0.2, 1.0 - min(0.8, volatility / 2.5))
    score = round(100 * ((0.4 * period_score) + (0.25 * completeness_score) + (0.2 * signal_score) + (0.15 * stability_score)), 1)
    issues: list[str] = []
    if usable_periods < minimum_required:
        issues.append(f'{usable_periods} usable {period_label} periods found; production forecasting recommends at least {minimum_required}.')
    if missing_share > 0:
        issues.append(f'{missing_share:.1%} of aggregated target periods are missing.')
    if zero_share > 0.25:
        issues.append(f'{zero_share:.1%} of aggregated target periods are zero or negative.')

    return {
        'score': score,
        'status': 'pass' if score >= 70 and usable_periods >= minimum_required else 'warning' if usable_periods >= 6 else 'fail',
        'minimum_required_periods': minimum_required,
        'usable_periods': usable_periods,
        'missing_share': round(missing_share, 4),
        'zero_or_negative_share': round(zero_share, 4),
        'volatility': round(volatility, 4),
        'issues': issues,
    }


def ensure_forecast_data_sufficiency(series_frame: pd.DataFrame, period_label: str, require_quality_gate: bool = True) -> dict[str, Any]:
    quality = build_forecast_data_quality(series_frame, period_label)
    minimum_required = 2 * seasonal_period_for_label(period_label)
    if len(series_frame) < minimum_required:
        needed = minimum_required - len(series_frame)
        raise HTTPException(status_code=422, detail=f'Forecasting needs at least {minimum_required} {period_label} rows after aggregation. Add {needed} more {period_label} row{"s" if needed != 1 else ""} and rerun the forecast.')
    if require_quality_gate and quality['status'] == 'fail':
        raise HTTPException(status_code=422, detail=f'Data quality gate failed: {"; ".join(quality["issues"]) or "insufficient usable signal"}.')
    return quality


def naive_forecast_step(history: list[float]) -> float:
    return max(0.0, float(history[-1]) if history else 0.0)


def append_interval(point: dict[str, Any], residual_std: float) -> dict[str, Any]:
    lower, upper = build_confidence_bounds(float(point.get('predicted') or 0), residual_std)
    return {**point, 'lower': lower, 'upper': upper}


def append_forecast_version(session_id: str, step: str, payload: dict[str, Any]) -> None:
    state = ensure_session_state(session_id)
    versions = state.setdefault('forecast_history_versions', [])
    versions.append({
        'id': uuid.uuid4().hex,
        'step': step,
        'created_at': utc_now_iso(),
        'model': (payload.get('model_details') or {}).get('model_name') or payload.get('status') or step,
        'metrics': payload.get('metrics') or {},
        'data_quality': payload.get('data_quality') or {},
        'assumptions_audit': payload.get('assumptions_audit') or [],
    })
    del versions[:-25]


def normal_tail_probability(value: float) -> float:
    return 0.5 * (1 - erf(abs(value) / sqrt(2)))


def compute_stationarity_check(values: list[float]) -> dict[str, Any]:
    if len(values) < 6:
        return {
            'test_name': 'Dickey-Fuller',
            'p_value': 1.0,
            'verdict': 'Insufficient history',
            'note': 'At least 6 periods are needed for a meaningful stationarity diagnostic.',
        }

    y_t = np.array(values[1:], dtype=float)
    y_tm1 = np.array(values[:-1], dtype=float)
    x = np.column_stack([np.ones(len(y_tm1)), y_tm1])
    beta, *_ = np.linalg.lstsq(x, y_t, rcond=None)
    residuals = y_t - x @ beta
    dof = max(len(y_t) - x.shape[1], 1)
    sigma2 = float(np.sum(residuals ** 2) / dof)
    try:
        cov = sigma2 * np.linalg.inv(x.T @ x)
    except np.linalg.LinAlgError:
        cov = sigma2 * np.linalg.pinv(x.T @ x)
    se = float(np.sqrt(max(cov[1, 1], 1e-9)))
    test_stat = float((beta[1] - 1.0) / se)
    p_value = round(min(1.0, max(0.0, 2 * normal_tail_probability(test_stat))), 4)
    stationary = p_value < 0.05
    return {
        'test_name': 'Dickey-Fuller',
        'p_value': p_value,
        'verdict': 'Likely stationary' if stationary else 'Trend or seasonality still present',
        'note': 'The series looks stationary enough for difference-based models.' if stationary else 'The series still shows a trend or seasonal structure, so seasonal statistical models are favored.',
    }


def with_fitted_order_note(stationarity: dict[str, Any], model_name: str, order: tuple[int, int, int] | None, seasonal_order: tuple[int, int, int, int] | None) -> dict[str, Any]:
    order_text = order or 'unknown'
    seasonal_text = seasonal_order or 'unknown'
    return {
        **stationarity,
        'note': f'{stationarity.get("note", "Stationarity diagnostic completed.")} Auto-selected {model_name} with fitted order {order_text} and seasonal order {seasonal_text}.',
        'fitted_order': order,
        'fitted_seasonal_order': seasonal_order,
    }


def build_dataset_profile(series_frame: pd.DataFrame, period_label: str) -> dict[str, Any]:
    values = series_frame['sales'].astype(float).tolist()
    mean = float(np.mean(values)) if values else 0.0
    volatility = 0.0 if mean == 0 else float(np.std(values) / abs(mean))
    zero_share = float(sum(value == 0 for value in values) / len(values)) if values else 0.0
    return {
        'detected_frequency': period_label,
        'usable_periods': int(len(series_frame)),
        'volatility': round(volatility, 4),
        'zero_value_share': round(zero_share, 4),
    }


def infer_season_length(period_label: str, total_periods: int) -> int:
    default = seasonal_period_for_label(period_label)
    return max(1, min(default, max(1, total_periods // 2)))


def resolve_time_series_model_name(model_type: str) -> str:
    mapping = {'sarima': 'SARIMA', 'prophet': 'Prophet', 'arima': 'ARIMA'}
    return mapping.get(model_type, 'SARIMA')


def statistical_forecast_step(history: list[float], season_length: int, model_type: str) -> float:
    recent = history[-min(3, len(history)):]
    recent_mean = float(np.mean(recent))
    trend = float(np.mean(np.diff(history[-min(5, len(history)):])) if len(history) > 1 else 0.0)
    seasonal_value = float(history[-season_length]) if season_length > 1 and len(history) >= season_length else recent_mean

    if model_type == 'arima':
        prediction = recent_mean + (0.35 * trend)
    elif model_type == 'prophet':
        prediction = (0.55 * recent_mean) + (0.45 * seasonal_value) + (0.65 * trend)
    else:
        prediction = (0.4 * recent_mean) + (0.6 * seasonal_value) + (0.4 * trend)

    return max(0.0, float(prediction))


def build_confidence_bounds(prediction: float, residual_std: float) -> tuple[float, float]:
    interval = 1.96 * residual_std
    lower = max(0.0, prediction - interval)
    upper = max(lower, prediction + interval)
    return round(lower, 2), round(upper, 2)


def build_ml_forecast_feature_row(
    history: list[float],
    current_period: pd.Timestamp,
    lag_periods: int,
    feature_groups: list[str],
) -> dict[str, float]:
    if len(history) < lag_periods:
        raise ValueError('Not enough history to build forecast features.')

    row: dict[str, float] = {}
    if 'trend' in feature_groups:
        row['trend_index'] = float(len(history) + 1)
    if 'calendar' in feature_groups:
        row['month_number'] = float(current_period.month)
        row['quarter_number'] = float(current_period.quarter)
        row['weekday_number'] = float(current_period.dayofweek)
        row['is_month_end'] = float(int(current_period.is_month_end))
    if 'lags' in feature_groups:
        for lag_index in range(1, lag_periods + 1):
            row[f'lag_{lag_index}'] = float(history[-lag_index])
    if 'rolling' in feature_groups:
        row['rolling_mean_3'] = float(np.mean(history[-min(3, len(history)):]))
        row['rolling_mean_6'] = float(np.mean(history[-min(6, len(history)):]))
        row['rolling_std_3'] = float(np.std(history[-min(3, len(history)):]))
    return row


def build_ml_forecast_training_frame(series_frame: pd.DataFrame, lag_periods: int, feature_groups: list[str]) -> tuple[pd.DataFrame, pd.Series]:
    rows: list[dict[str, float]] = []
    targets: list[float] = []
    periods = series_frame['period'].tolist()
    values = series_frame['sales'].astype(float).tolist()
    for index in range(lag_periods, len(values)):
        rows.append(build_ml_forecast_feature_row(values[:index], periods[index], lag_periods, feature_groups))
        targets.append(float(values[index]))
    if len(rows) < 3:
        raise HTTPException(status_code=400, detail='Not enough history is available to train the ML forecasting model.')
    return pd.DataFrame(rows), pd.Series(targets)


def build_forecast_regressor(model_type: str):
    if model_type == 'xgboost' and XGBRegressor is not None:
        return XGBRegressor(
            n_estimators=160,
            learning_rate=0.05,
            max_depth=3,
            objective='reg:squarederror',
            random_state=42,
            n_jobs=TRAINING_N_JOBS,
        )
    if model_type == 'lightgbm' and LGBMRegressor is not None:
        return LGBMRegressor(
            n_estimators=160,
            learning_rate=0.05,
            max_depth=3,
            random_state=42,
            n_jobs=TRAINING_N_JOBS,
            verbose=-1,
        )
    if model_type == 'random_forest':
        return RandomForestRegressor(n_estimators=160, random_state=42, min_samples_leaf=2, n_jobs=TRAINING_N_JOBS)
    if model_type == 'ridge_regression':
        return Ridge(alpha=1.0, random_state=42)
    return GradientBoostingRegressor(random_state=42, n_estimators=140, learning_rate=0.05, max_depth=3)


def recursive_ml_forecast(
    model: Any,
    history: list[float],
    start_period: pd.Timestamp,
    periods: int,
    lag_periods: int,
    feature_groups: list[str],
    freq: str,
    period_label: str,
) -> list[dict[str, Any]]:
    forecasts: list[dict[str, Any]] = []
    running_history = list(history)
    current_period = pd.Timestamp(start_period)
    for _ in range(periods):
        features = build_ml_forecast_feature_row(running_history, current_period, lag_periods, feature_groups)
        prediction = max(0.0, float(model.predict(pd.DataFrame([features]))[0]))
        forecasts.append({'period': format_forecast_period(current_period, period_label), 'predicted': round(prediction, 2)})
        running_history.append(prediction)
        current_period = current_period + pd.tseries.frequencies.to_offset(freq)
    return forecasts


def calculate_shap_like_importance(model: Any, feature_names: list[str]) -> list[dict[str, Any]]:
    if hasattr(model, 'feature_importances_'):
        values = np.asarray(model.feature_importances_, dtype=float)
    elif hasattr(model, 'coef_'):
        values = np.abs(np.asarray(model.coef_, dtype=float).reshape(-1))
    else:
        values = np.ones(len(feature_names), dtype=float)

    importance = [
        {'name': feature_names[index], 'importance': round(float(value), 4)}
        for index, value in enumerate(values[:len(feature_names)])
    ]
    importance.sort(key=lambda item: item['importance'], reverse=True)
    return importance


def production_model_name(model_type: str) -> str:
    return {
        'arima': 'ARIMA',
        'prophet': 'Prophet',
        'sarima': 'SARIMA',
        'gradient_boosting': 'Gradient Boosting',
        'xgboost': 'XGBoost',
        'lightgbm': 'LightGBM',
    }.get(model_type, model_type.replace('_', ' ').title())


def model_availability_note(model_type: str) -> str:
    if model_type == 'prophet':
        return 'Prophet package available.' if Prophet is not None else 'Prophet package unavailable; candidate skipped.'
    if model_type == 'arima':
        if SARIMAX is None:
            return 'statsmodels SARIMAX unavailable; candidate skipped.'
        return 'pmdarima.auto_arima available for ARIMA order selection.' if auto_arima is not None else 'pmdarima.auto_arima unavailable; simplified ARIMA fallback will be attempted.'
    if model_type == 'sarima':
        if SARIMAX is None:
            return 'statsmodels SARIMAX unavailable; candidate skipped.'
        return 'pmdarima.auto_arima available for SARIMA order selection.' if auto_arima is not None else 'pmdarima.auto_arima unavailable; simplified SARIMA fallback will be attempted.'
    if model_type == 'xgboost' and XGBRegressor is None:
        return 'XGBoost package unavailable; candidate skipped.'
    if model_type == 'lightgbm' and LGBMRegressor is None:
        return 'LightGBM unavailable after install; candidate failed. Rebuild the backend image or check native LightGBM dependencies.'
    return 'Candidate available.'


def fit_arima_family_forecast(history: list[float], forecast_periods: int, period_label: str, *, seasonal: bool) -> tuple[list[float], dict[str, Any]]:
    if SARIMAX is None:
        raise RuntimeError('statsmodels SARIMAX is unavailable.')

    seasonal_period = seasonal_period_for_label(period_label)
    mapped_freq = statsmodels_frequency_for_label(period_label)
    index = pd.date_range(start='2000-01-01', periods=len(history), freq=mapped_freq)
    endog = pd.Series(np.asarray(history, dtype=float), index=index)
    attempts: list[dict[str, Any]] = []
    model_label = 'SARIMA' if seasonal else 'ARIMA'

    if auto_arima is not None:
        try:
            auto_model = auto_arima(
                endog,
                seasonal=seasonal and seasonal_period > 1,
                m=seasonal_period if seasonal else 1,
                start_p=0,
                start_q=0,
                max_p=3,
                max_q=3,
                max_P=2 if seasonal else 0,
                max_Q=2 if seasonal else 0,
                d=None,
                D=None if seasonal else 0,
                trace=False,
                error_action='raise',
                suppress_warnings=True,
                stepwise=True,
            )
            order = tuple(int(value) for value in auto_model.order)
            seasonal_order = tuple(int(value) for value in auto_model.seasonal_order) if seasonal else (0, 0, 0, 0)
            fitted = SARIMAX(endog, order=order, seasonal_order=seasonal_order, freq=mapped_freq, enforce_stationarity=False, enforce_invertibility=False).fit(disp=False)
            predictions = [max(0.0, float(value)) for value in fitted.forecast(forecast_periods)]
            return predictions, {'order': order, 'seasonal_order': seasonal_order, 'selection': 'auto_arima', 'attempts': attempts}
        except (ValueError, np.linalg.LinAlgError) as error:
            attempts.append({'selection': 'auto_arima', 'error': str(error)})
        except Exception as error:
            attempts.append({'selection': 'auto_arima', 'error': str(error)})

    try:
        fallback_order = (1, 1, 1)
        fallback_seasonal_order = (1, 1, 1, seasonal_period) if seasonal and seasonal_period > 1 else (0, 0, 0, 0)
        fitted = SARIMAX(endog, order=fallback_order, seasonal_order=fallback_seasonal_order, freq=mapped_freq, enforce_stationarity=False, enforce_invertibility=False).fit(disp=False)
        predictions = [max(0.0, float(value)) for value in fitted.forecast(forecast_periods)]
        return predictions, {'order': fallback_order, 'seasonal_order': fallback_seasonal_order, 'selection': 'simplified_retry', 'attempts': attempts}
    except (ValueError, np.linalg.LinAlgError) as error:
        attempts.append({'selection': 'simplified_retry', 'error': str(error)})
        raise RuntimeError(f'{model_label} fitting failed after auto_arima and simplified retry: {attempts}') from error
    except Exception as error:
        attempts.append({'selection': 'simplified_retry', 'error': str(error)})
        raise RuntimeError(f'{model_label} fitting failed after auto_arima and simplified retry: {attempts}') from error


def fit_sarima_forecast(history: list[float], forecast_periods: int, period_label: str) -> tuple[list[float], dict[str, Any]]:
    return fit_arima_family_forecast(history, forecast_periods, period_label, seasonal=True)


def fit_arima_forecast(history: list[float], forecast_periods: int, period_label: str) -> tuple[list[float], dict[str, Any]]:
    return fit_arima_family_forecast(history, forecast_periods, period_label, seasonal=False)


def walk_forward_splits(total_periods: int, requested_test_periods: int, lag_periods: int) -> list[tuple[int, int]]:
    min_train = max(lag_periods + 3, int(total_periods * 0.5))
    max_start = max(min_train, total_periods - max(1, requested_test_periods))
    starts = sorted(set([min_train, int(total_periods * 0.65), max_start]))
    return [(start, min(requested_test_periods, total_periods - start)) for start in starts if total_periods - start >= 1]


def evaluate_statistical_candidate(
    series_frame: pd.DataFrame,
    model_type: str,
    requested_test_periods: int,
    freq: str,
    period_label: str,
) -> dict[str, Any]:
    if model_type == 'prophet' and Prophet is None:
        return {'model_type': model_type, 'model_name': 'Prophet', 'status': 'skipped', 'skip_reason': model_availability_note(model_type)}
    if model_type in {'arima', 'sarima'} and SARIMAX is None:
        return {'model_type': model_type, 'model_name': production_model_name(model_type), 'status': 'skipped', 'skip_reason': model_availability_note(model_type)}

    values = series_frame['sales'].astype(float).tolist()
    periods = series_frame['period'].tolist()
    season_length = infer_season_length(period_label, len(values))
    actuals: list[float] = []
    predictions: list[float] = []
    test_rows: list[dict[str, Any]] = []
    for train_end, fold_periods in walk_forward_splits(len(values), requested_test_periods, 1):
        history = values[:train_end]
        current_period = pd.Timestamp(periods[train_end])
        fitted_predictions: list[float] = []
        if model_type == 'prophet':
            try:
                prophet_frame = pd.DataFrame({'ds': pd.to_datetime(periods[:train_end]), 'y': history})
                prophet_model = Prophet(interval_width=0.95, daily_seasonality=period_label == 'day', weekly_seasonality=period_label in {'day', 'week'}, yearly_seasonality=period_label in {'month', 'quarter'})
                prophet_model.fit(prophet_frame)
                future_dates = [pd.Timestamp(current_period) + (pd.tseries.frequencies.to_offset(freq) * index) for index in range(fold_periods)]
                forecast = prophet_model.predict(pd.DataFrame({'ds': future_dates}))
                fitted_predictions = [max(0.0, float(value)) for value in forecast['yhat'].tolist()]
            except Exception as error:
                raise RuntimeError(f'Prophet training failed during walk-forward validation: {error}') from error
        elif model_type == 'sarima':
            try:
                fitted_predictions, tuning_details = fit_sarima_forecast(history, fold_periods, period_label)
            except Exception as error:
                raise RuntimeError(f'SARIMA training failed during walk-forward validation: {error}') from error
        elif model_type == 'arima':
            try:
                fitted_predictions, tuning_details = fit_arima_forecast(history, fold_periods, period_label)
            except Exception as error:
                raise RuntimeError(f'ARIMA training failed during walk-forward validation: {error}') from error
        if len(fitted_predictions) < fold_periods:
            raise RuntimeError(f'{production_model_name(model_type)} produced {len(fitted_predictions)} predictions for a {fold_periods}-period validation fold.')
        for offset in range(fold_periods):
            actual = float(values[train_end + offset])
            predicted = fitted_predictions[offset]
            actuals.append(actual)
            predictions.append(predicted)
            test_rows.append({
                'period': format_forecast_period(current_period, period_label),
                'actual': round(actual, 2),
                'predicted': round(predicted, 2),
            })
            history.append(actual)
            current_period = current_period + pd.tseries.frequencies.to_offset(freq)
    metrics = calculate_forecast_metrics(actuals, predictions)
    residuals = [actual - predicted for actual, predicted in zip(actuals, predictions)]
    return {
        'model_type': model_type,
        'model_name': production_model_name(model_type),
        'status': 'completed',
        'metrics': metrics,
        'residual_std': float(np.std(residuals)) if residuals else 0.0,
        'test_forecast': test_rows[-requested_test_periods:],
        'feature_importance': [],
        'generated_features': [],
        'feature_preview_rows': [],
        'tuning': tuning_details if model_type in {'arima', 'sarima'} else {'enabled': optuna is not None, 'note': 'Optuna available.' if optuna is not None else 'Optuna package unavailable; stable model defaults used.'},
        'availability_note': f'{production_model_name(model_type)} trained successfully on walk-forward validation folds.',
    }


def evaluate_ml_candidate(
    series_frame: pd.DataFrame,
    model_type: str,
    requested_test_periods: int,
    lag_periods: int,
    feature_groups: list[str],
    freq: str,
    period_label: str,
) -> dict[str, Any]:
    if model_type == 'xgboost' and XGBRegressor is None:
        return {'model_type': model_type, 'model_name': 'XGBoost', 'status': 'skipped', 'skip_reason': model_availability_note(model_type)}
    if model_type == 'lightgbm' and LGBMRegressor is None:
        return {'model_type': model_type, 'model_name': 'LightGBM', 'status': 'failed', 'skip_reason': model_availability_note(model_type)}

    values = series_frame['sales'].astype(float).tolist()
    periods = series_frame['period'].tolist()
    actuals: list[float] = []
    predictions: list[float] = []
    test_rows: list[dict[str, Any]] = []
    last_model: Any = None
    last_X: pd.DataFrame | None = None
    for train_end, fold_periods in walk_forward_splits(len(values), requested_test_periods, lag_periods):
        train_frame = series_frame.iloc[:train_end].copy()
        train_X, train_y = build_ml_forecast_training_frame(train_frame, lag_periods, feature_groups)
        non_zero_feature_count = int((train_X.var(numeric_only=True) > 1e-9).sum())
        if non_zero_feature_count < 3:
            raise RuntimeError(f'Only {non_zero_feature_count} generated features have non-zero variance; at least 3 are required for ML forecasting.')
        model = build_forecast_regressor(model_type)
        model.fit(train_X, train_y)
        fold_predictions = recursive_ml_forecast(
            model,
            train_frame['sales'].astype(float).tolist(),
            pd.Timestamp(periods[train_end]),
            fold_periods,
            lag_periods,
            feature_groups,
            freq,
            period_label,
        )
        for offset, point in enumerate(fold_predictions):
            actual = float(values[train_end + offset])
            predicted = float(point['predicted'])
            actuals.append(actual)
            predictions.append(predicted)
            test_rows.append({**point, 'actual': round(actual, 2)})
        last_model = model
        last_X = train_X

    metrics = calculate_forecast_metrics(actuals, predictions)
    residuals = [actual - predicted for actual, predicted in zip(actuals, predictions)]
    feature_names = last_X.columns.tolist() if last_X is not None else []
    return {
        'model_type': model_type,
        'model_name': production_model_name(model_type),
        'status': 'completed',
        'metrics': metrics,
        'residual_std': float(np.std(residuals)) if residuals else 0.0,
        'test_forecast': test_rows[-requested_test_periods:],
        'feature_importance': calculate_shap_like_importance(last_model, feature_names) if last_model is not None else [],
        'generated_features': feature_names,
        'feature_preview_rows': safe_serialize(last_X.head(5).round(3).to_dict(orient='records')) if last_X is not None else [],
        'tuning': {
            'enabled': optuna is not None,
            'note': 'Optuna is available for future expanded search; current bounded run uses stable production defaults.' if optuna is not None else 'Optuna package unavailable; stable defaults used.',
        },
        'availability_note': f'{production_model_name(model_type)} trained successfully on walk-forward validation folds.',
    }


def build_future_for_selected_model(
    selected: dict[str, Any],
    series_frame: pd.DataFrame,
    forecast_periods: int,
    lag_periods: int,
    feature_groups: list[str],
    freq: str,
    period_label: str,
) -> list[dict[str, Any]]:
    model_type = selected['model_type']
    residual_std = float(selected.get('residual_std') or 0.0)
    current_period = pd.Timestamp(series_frame.iloc[-1]['period']) + pd.tseries.frequencies.to_offset(freq)
    if model_type in {'prophet', 'arima', 'sarima'}:
        season_length = infer_season_length(period_label, len(series_frame))
        history = series_frame['sales'].astype(float).tolist()
        if model_type == 'prophet' and Prophet is not None:
            try:
                prophet_frame = pd.DataFrame({'ds': pd.to_datetime(series_frame['period']), 'y': history})
                prophet_model = Prophet(interval_width=0.95, daily_seasonality=period_label == 'day', weekly_seasonality=period_label in {'day', 'week'}, yearly_seasonality=period_label in {'month', 'quarter'})
                prophet_model.fit(prophet_frame)
                future_dates = [pd.Timestamp(current_period) + (pd.tseries.frequencies.to_offset(freq) * index) for index in range(forecast_periods)]
                forecast = prophet_model.predict(pd.DataFrame({'ds': future_dates}))
                return [
                    {
                        'period': format_forecast_period(pd.Timestamp(row['ds']), period_label),
                        'predicted': round(max(0.0, float(row['yhat'])), 2),
                        'lower': round(max(0.0, float(row.get('yhat_lower', row['yhat']))), 2),
                        'upper': round(max(0.0, float(row.get('yhat_upper', row['yhat']))), 2),
                    }
                    for _, row in forecast.iterrows()
                ]
            except Exception as error:
                raise RuntimeError(f'Prophet training failed while building the final future forecast: {error}') from error
        if model_type == 'sarima' and SARIMAX is not None:
            try:
                predictions, tuning_details = fit_sarima_forecast(history, forecast_periods, period_label)
                selected['tuning'] = tuning_details
                rows = []
                for prediction in predictions:
                    rows.append(append_interval({'period': format_forecast_period(current_period, period_label), 'predicted': round(prediction, 2)}, residual_std))
                    current_period = current_period + pd.tseries.frequencies.to_offset(freq)
                return rows
            except Exception as error:
                raise RuntimeError(f'SARIMA training failed while building the final future forecast: {error}') from error
        if model_type == 'arima' and SARIMAX is not None:
            try:
                predictions, tuning_details = fit_arima_forecast(history, forecast_periods, period_label)
                selected['tuning'] = tuning_details
                rows = []
                for prediction in predictions:
                    rows.append(append_interval({'period': format_forecast_period(current_period, period_label), 'predicted': round(prediction, 2)}, residual_std))
                    current_period = current_period + pd.tseries.frequencies.to_offset(freq)
                return rows
            except Exception as error:
                raise RuntimeError(f'ARIMA training failed while building the final future forecast: {error}') from error
        raise RuntimeError(f'{production_model_name(model_type)} is unavailable and cannot build a production forecast.')

    full_X, full_y = build_ml_forecast_training_frame(series_frame, lag_periods, feature_groups)
    non_zero_feature_count = int((full_X.var(numeric_only=True) > 1e-9).sum())
    if non_zero_feature_count < 3:
        raise RuntimeError(f'Only {non_zero_feature_count} generated features have non-zero variance; at least 3 are required for ML forecasting.')
    model = build_forecast_regressor(model_type)
    model.fit(full_X, full_y)
    return [append_interval(point, residual_std) for point in recursive_ml_forecast(
        model,
        series_frame['sales'].astype(float).tolist(),
        current_period,
        forecast_periods,
        lag_periods,
        feature_groups,
        freq,
        period_label,
    )]


def auto_select_forecast_model(
    series_frame: pd.DataFrame,
    forecast_periods: int,
    requested_test_periods: int,
    lag_periods: int,
    feature_groups: list[str],
    freq: str,
    period_label: str,
    candidates: list[str] | None = None,
) -> dict[str, Any]:
    candidates = candidates or ['sarima', 'arima', 'prophet']
    comparison: list[dict[str, Any]] = []
    with ThreadPoolExecutor(max_workers=min(4, len(candidates))) as executor:
        futures = {}
        for candidate in [item for item in candidates if item in {'prophet', 'arima', 'sarima'}]:
            futures[executor.submit(evaluate_statistical_candidate, series_frame, candidate, requested_test_periods, freq, period_label)] = candidate
        for candidate in [item for item in candidates if item not in {'prophet', 'arima', 'sarima'}]:
            futures[executor.submit(evaluate_ml_candidate, series_frame, candidate, requested_test_periods, lag_periods, feature_groups, freq, period_label)] = candidate
        for future in as_completed(futures):
            try:
                result = future.result()
                metrics = result.get('metrics') or {}
                if result.get('status') == 'completed' and float(metrics.get('mae') or 0) == 0.0 and float(metrics.get('rmse') or 0) == 0.0:
                    result = {
                        **result,
                        'status': 'failed',
                        'metrics': None,
                        'skip_reason': 'Training produced zero MAE and RMSE; metrics are unreliable, so this candidate was rejected.',
                    }
                comparison.append(result)
            except Exception as error:
                candidate = futures[future]
                comparison.append({'model_type': candidate, 'model_name': production_model_name(candidate), 'status': 'failed', 'skip_reason': str(error)})

    completed = [
        item for item in comparison
        if item.get('status') == 'completed'
        and not (float((item.get('metrics') or {}).get('mae') or 0) == 0.0 and float((item.get('metrics') or {}).get('rmse') or 0) == 0.0)
    ]
    if not completed:
        raise HTTPException(status_code=422, detail='No forecast candidate could be trained with reliable validation metrics. Select a different target column or provide more history.')

    selected = min(completed, key=lambda item: (float(item['metrics'].get('mae', np.inf)), float(item['metrics'].get('rmse', np.inf)), float(item['metrics'].get('mape', np.inf))))
    future = build_future_for_selected_model(selected, series_frame, forecast_periods, lag_periods, feature_groups, freq, period_label)
    test_forecast = [append_interval(point, float(selected.get('residual_std') or 0.0)) for point in selected.get('test_forecast', [])]

    naive_actuals: list[float] = []
    naive_predictions: list[float] = []
    values = series_frame['sales'].astype(float).tolist()
    for train_end, fold_periods in walk_forward_splits(len(values), requested_test_periods, 1):
        history = values[:train_end]
        for offset in range(fold_periods):
            naive_actuals.append(float(values[train_end + offset]))
            naive_predictions.append(naive_forecast_step(history))
            history.append(float(values[train_end + offset]))
    naive_metrics = calculate_forecast_metrics(naive_actuals, naive_predictions)
    selected_mae = float(selected['metrics'].get('mae') or 0)
    naive_mae = float(naive_metrics.get('mae') or 0)
    target_mean = float(series_frame['sales'].astype(float).mean() or 0.0)
    validation_warnings: list[str] = []
    if selected_mae < 0.001 and target_mean > 1.0:
        validation_warnings.append('Model metrics appear unreliable — check target column selection.')

    return {
        'selected': selected,
        'future_forecast': future,
        'test_forecast': test_forecast,
        'model_comparison': sorted(comparison, key=lambda item: (item.get('metrics') or {}).get('mae', np.inf) if item.get('status') == 'completed' else np.inf),
        'naive_baseline': {
            'model_name': 'Naive last-observation baseline',
            'metrics': naive_metrics,
            'mae_improvement_pct': round(((naive_mae - selected_mae) / naive_mae) * 100, 2) if naive_mae else 0.0,
        },
        'validation_warnings': validation_warnings,
    }


def build_time_series_model_recommendations(profile: dict[str, Any], stationarity: dict[str, Any]) -> list[dict[str, Any]]:
    frequency = str(profile.get('detected_frequency') or 'period')
    is_stationary = stationarity.get('verdict') == 'Likely stationary'
    return [
        {
            'model_type': 'sarima',
            'model_name': 'SARIMA',
            'recommended': not is_stationary and frequency in {'week', 'month', 'quarter'},
            'recommendation_reason': 'Best first choice when repeating seasonal patterns are visible; pmdarima.auto_arima selects the seasonal and non-seasonal orders.',
        },
        {
            'model_type': 'arima',
            'model_name': 'ARIMA',
            'recommended': is_stationary,
            'recommendation_reason': 'Useful for shorter or mostly non-seasonal series where recent autocorrelation and differencing explain the target.',
        },
        {
            'model_type': 'prophet',
            'model_name': 'Prophet',
            'recommended': not is_stationary and frequency in {'day', 'week', 'month'},
            'recommendation_reason': 'Flexible trend and seasonality model that handles changing growth patterns and wider forecast intervals.',
        },
    ]


def build_ml_model_recommendations(feature_names: list[str]) -> list[dict[str, Any]]:
    feature_count = len(feature_names)
    return [
        {
            'model_type': 'gradient_boosting',
            'model_name': 'Gradient Boosting',
            'recommended': feature_count >= 6,
            'recommendation_reason': f'Excellent for capturing non-linear patterns across the {feature_count} generated features.',
        },
        {
            'model_type': 'xgboost',
            'model_name': 'XGBoost',
            'recommended': feature_count < 6,
            'recommendation_reason': 'Strong tree boosting candidate when engineered lag and calendar features interact.',
        },
        {
            'model_type': 'lightgbm',
            'model_name': 'LightGBM',
            'recommended': False,
            'recommendation_reason': 'Fast gradient boosting candidate; import failures are surfaced in the comparison table.',
        },
        {
            'model_type': 'prophet',
            'model_name': 'Prophet',
            'recommended': False,
            'recommendation_reason': 'Trend and seasonality candidate used for comparison against engineered ML learners.',
        },
    ]


def make_analysis(problem_type: ProblemType, model_name: str, metrics: dict[str, Any], importances: list[dict[str, Any]]) -> str:
    top_features = ', '.join(item['name'] for item in importances[:3]) or 'No dominant features detected'
    if problem_type == 'regression':
        return (
            f"### {model_name} Summary\n"
            f"- R2: {metrics['primary'].get('R2', 0):.4f}\n"
            f"- RMSE: {metrics['primary'].get('RMSE', 0):.4f}\n"
            f"- MAE: {metrics['primary'].get('MAE', 0):.4f}\n"
            f"- CV Mean: {metrics.get('cv_mean', 0):.4f}\n"
            f"- Top features: {top_features}"
        )
    return (
        f"### {model_name} Summary\n"
        f"- Accuracy: {metrics['primary'].get('Accuracy', 0):.4f}\n"
        f"- Precision: {metrics['primary'].get('Precision', 0):.4f}\n"
        f"- Recall: {metrics['primary'].get('Recall', 0):.4f}\n"
        f"- F1 Score: {metrics['primary'].get('F1 Score', 0):.4f}\n"
        f"- Top features: {top_features}"
    )


def save_model_bundle(model_id: str, bundle: dict[str, Any]) -> None:
    joblib.dump(bundle, MODEL_DIR / f'{model_id}.joblib')
    MODEL_CACHE[model_id] = bundle


def load_model_bundle(model_id: str) -> dict[str, Any]:
    if model_id in MODEL_CACHE:
        return MODEL_CACHE[model_id]
    model_path = MODEL_DIR / f'{model_id}.joblib'
    bundle_path = MODEL_DIR / f'{model_id}_model_bundle.joblib'
    if not model_path.exists() and bundle_path.exists():
        model_path = bundle_path
    if not model_path.exists():
        raise HTTPException(status_code=404, detail=f"Model '{model_id}' was not found.")
    bundle = joblib.load(model_path)
    MODEL_CACHE[model_id] = bundle
    return bundle


def structured_training_failure(error: Exception, *, step: str, request: TrainRequest | None = None) -> JSONResponse:
    detail = ''.join(traceback.format_exception(type(error), error, error.__traceback__))
    logger.error('ML training failed at %s: %s\n%s', step, error, detail)
    if request is not None:
        try:
            record_activity(
                action='train_model',
                status='failed',
                dataset_id=request.dataset_id,
                detail=f'{step}: {error}',
                metadata={
                    'target_column': request.target_column,
                    'feature_count': len(request.feature_columns),
                    'model_type': request.model_type,
                    'problem_type': request.problem_type,
                },
            )
        except Exception:
            logger.exception('Failed to record training failure activity.')
    return JSONResponse(
        status_code=200,
        content={
            'status': 'failed',
            'step': step,
            'error': str(getattr(error, 'detail', None) or error) or 'Training failed.',
            'detail': detail,
        },
    )


def get_cv_strategy(n_rows: int):
    if n_rows < 30:
        return LeaveOneOut()
    if n_rows < 100:
        return KFold(n_splits=min(5, n_rows), shuffle=True, random_state=42)
    return KFold(n_splits=min(10, n_rows), shuffle=True, random_state=42)


def select_candidate_models(task_type: ProblemType, n_rows: int, n_features: int, random_state: int) -> list[tuple[str, Any]]:
    del n_features
    if task_type == 'regression':
        if n_rows < 50:
            return [('Ridge', Ridge(alpha=1.0)), ('Lasso', Lasso(alpha=1.0)), ('Linear Regression', LinearRegression())]
        if n_rows < 500:
            return [
                ('Random Forest', RandomForestRegressor(n_estimators=100, random_state=random_state)),
                ('Gradient Boosting', GradientBoostingRegressor(random_state=random_state)),
                ('Ridge', Ridge(alpha=1.0)),
            ]
        large_candidates: list[tuple[str, Any]] = [
            ('Random Forest', RandomForestRegressor(n_estimators=200, random_state=random_state)),
            ('Gradient Boosting', GradientBoostingRegressor(n_estimators=200, random_state=random_state)),
        ]
        if XGBRegressor is not None:
            large_candidates.append(('XGBoost', XGBRegressor(n_estimators=200, random_state=random_state, verbosity=0)))
        else:
            large_candidates.append(('Gradient Boosting Fallback', GradientBoostingRegressor(n_estimators=200, random_state=random_state)))
        return large_candidates

    if n_rows < 50:
        return [('Logistic Regression', LogisticRegression(max_iter=1000)), ('SVC', SVC(probability=True))]
    if n_rows < 500:
        return [
            ('Random Forest', RandomForestClassifier(n_estimators=100, random_state=random_state)),
            ('Logistic Regression', LogisticRegression(max_iter=1000)),
        ]
    large_classifiers: list[tuple[str, Any]] = [
        ('Random Forest', RandomForestClassifier(n_estimators=200, random_state=random_state)),
        ('Gradient Boosting', GradientBoostingClassifier(random_state=random_state)),
    ]
    if XGBClassifier is not None:
        large_classifiers.append(('XGBoost', XGBClassifier(n_estimators=200, random_state=random_state, verbosity=0)))
    else:
        large_classifiers.append(('Gradient Boosting Fallback', GradientBoostingClassifier(random_state=random_state)))
    return large_classifiers


def is_datetime_feature(series: pd.Series) -> bool:
    if pd.api.types.is_datetime64_any_dtype(series):
        return True
    if not (pd.api.types.is_object_dtype(series) or pd.api.types.is_string_dtype(series)):
        return False
    sample = series.dropna().astype(str).head(50)
    if sample.empty:
        return False
    parsed = pd.to_datetime(sample, errors='coerce')
    return bool(parsed.notna().mean() >= 0.8)


def prepare_universal_training_data(
    frame: pd.DataFrame,
    request: TrainRequest,
) -> tuple[np.ndarray, pd.Series, dict[str, Any], list[str], list[str], LabelEncoder | None]:
    warnings_list: list[str] = []
    target_column = request.target_column
    if target_column not in frame.columns:
        raise ValueError(f'Target column "{target_column}" was not found in the dataset.')
    missing_features = [feature for feature in request.feature_columns if feature not in frame.columns]
    if missing_features:
        raise ValueError(f'Missing selected feature columns: {missing_features}')

    model_frame = frame[[*request.feature_columns, target_column]].copy()
    model_frame = model_frame.dropna(subset=[target_column])
    if len(model_frame) < 3:
        raise ValueError('At least 3 valid target rows are required for universal cross-validation training.')

    y_raw = model_frame[target_column]
    label_encoder: LabelEncoder | None = None
    if request.problem_type == 'classification':
        label_encoder = LabelEncoder()
        y = pd.Series(label_encoder.fit_transform(y_raw.astype(str)), index=model_frame.index)
        if y.nunique() < 2:
            raise ValueError('Classification requires at least two target classes.')
    else:
        y = pd.to_numeric(y_raw, errors='coerce').replace([np.inf, -np.inf], np.nan)
        valid = y.notna()
        model_frame = model_frame.loc[valid]
        y = y.loc[valid].astype(float)
        if len(y) < 3:
            raise ValueError('Regression requires at least 3 numeric target rows.')
        zero_pct = float((y == 0).mean() * 100)
        if zero_pct > 5:
            warnings_list.append(
                f'Target column contains {zero_pct:.1f}% zero values which may affect accuracy. These have been kept as-is. If zeros represent missing data, clean them in the Data Cleaning tab first.'
            )

    X = model_frame[request.feature_columns].copy().replace([np.inf, -np.inf], np.nan)
    dropped: dict[str, str] = {}
    for column in list(X.columns):
        series = X[column]
        non_null = series.dropna()
        if non_null.nunique(dropna=True) <= 1:
            dropped[column] = 'Dropped because all values are identical.'
            X = X.drop(columns=[column])
            continue
        if is_datetime_feature(series):
            dropped[column] = 'Dropped because datetime features are excluded unless explicitly engineered first.'
            X = X.drop(columns=[column])
            continue
        if (pd.api.types.is_object_dtype(series) or pd.api.types.is_string_dtype(series)) and len(non_null) > 0:
            unique_pct = float(non_null.astype(str).nunique() / max(len(series), 1))
            if unique_pct > 0.95:
                dropped[column] = 'Dropped because it behaves like a text identifier with more than 95% unique values.'
                X = X.drop(columns=[column])

    if X.empty:
        raise ValueError('No usable features remain after removing identifiers, constant columns, and datetime columns.')

    numeric_features = X.select_dtypes(include=[np.number, 'bool']).columns.tolist()
    categorical_features = [column for column in X.columns if column not in numeric_features]
    imputers: dict[str, Any] = {'numeric': {}, 'categorical': {}}
    categories: dict[str, list[str]] = {}
    X_imputed = pd.DataFrame(index=X.index)

    for column in numeric_features:
        numeric = pd.to_numeric(X[column], errors='coerce').replace([np.inf, -np.inf], np.nan)
        median = float(numeric.median()) if numeric.notna().any() else 0.0
        X_imputed[column] = numeric.fillna(median)
        imputers['numeric'][column] = median

    for column in categorical_features:
        categorical = X[column].astype(object).where(pd.notna(X[column]), np.nan)
        mode = categorical.dropna().astype(str).mode()
        fill_value = str(mode.iloc[0]) if not mode.empty else 'missing'
        filled = categorical.fillna(fill_value).astype(str)
        X_imputed[column] = filled
        imputers['categorical'][column] = fill_value
        categories[column] = sorted(filled.unique().tolist())

    if X_imputed.isna().any().any():
        raise ValueError('Validation failed: missing values remain after imputation.')

    X_encoded = pd.get_dummies(X_imputed, columns=categorical_features, dummy_na=False, dtype=float)
    final_feature_names = X_encoded.columns.tolist()
    if not final_feature_names:
        raise ValueError('No usable encoded features are available for training.')

    if len(final_feature_names) >= len(X_encoded):
        y_for_corr = pd.Series(y, index=X_encoded.index).astype(float)
        correlations = X_encoded.apply(lambda column: abs(float(column.corr(y_for_corr))) if column.nunique() > 1 else 0.0)
        correlations = correlations.replace([np.inf, -np.inf], np.nan).fillna(0.0)
        keep_count = min(len(final_feature_names), max(1, len(X_encoded) - 1), max(3, len(X_encoded) // 3))
        kept = correlations.sort_values(ascending=False).head(keep_count).index.tolist()
        warnings_list.append(
            f'Dataset has {len(X_encoded)} rows. Auto-selected top {len(kept)} features by correlation with target to prevent overfitting.'
        )
        X_encoded = X_encoded.loc[:, kept]
        final_feature_names = kept

    scaler = StandardScaler()
    scaled = scaler.fit_transform(X_encoded.astype(float))
    if np.isnan(scaled).any() or np.isinf(scaled).any():
        raise ValueError('Validation failed: NaN or Inf values remain after scaling.')

    preprocessing = {
        'raw_feature_columns': list(X.columns),
        'numeric_features': numeric_features,
        'categorical_features': categorical_features,
        'imputers': imputers,
        'categories': categories,
        'encoded_feature_names': final_feature_names,
        'scaler': scaler,
        'dropped_features': dropped,
        'warnings': warnings_list,
    }
    return scaled, y.reset_index(drop=True), preprocessing, final_feature_names, warnings_list, label_encoder


def transform_features_for_bundle(features: dict[str, Any], bundle: dict[str, Any]) -> np.ndarray:
    preprocessing = bundle.get('preprocessing') or {}
    raw_features = preprocessing.get('raw_feature_columns') or bundle.get('feature_columns') or []
    missing = [feature for feature in raw_features if features.get(feature) in [None, '']]
    if missing:
        raise ValueError(f'Missing features: {missing}')

    row = pd.DataFrame([{feature: features.get(feature) for feature in raw_features}])
    pieces: dict[str, Any] = {}
    for column in preprocessing.get('numeric_features', []):
        value = pd.to_numeric(pd.Series([row.at[0, column]]), errors='coerce').replace([np.inf, -np.inf], np.nan).iloc[0]
        if pd.isna(value):
            value = preprocessing.get('imputers', {}).get('numeric', {}).get(column, 0.0)
        pieces[column] = float(value)
    for column in preprocessing.get('categorical_features', []):
        value = row.at[0, column]
        if value in [None, ''] or pd.isna(value):
            value = preprocessing.get('imputers', {}).get('categorical', {}).get(column, 'missing')
        pieces[column] = str(value)

    encoded = pd.get_dummies(pd.DataFrame([pieces]), columns=preprocessing.get('categorical_features', []), dtype=float)
    encoded = encoded.reindex(columns=preprocessing.get('encoded_feature_names', []), fill_value=0.0)
    scaler = preprocessing.get('scaler') or bundle.get('scaler')
    if scaler is None:
        return encoded.to_numpy(dtype=float)
    return scaler.transform(encoded.astype(float))


def train_universal_model(request: TrainRequest, http_request: Request) -> dict[str, Any]:
    selected_columns = [*request.feature_columns, request.target_column]
    try:
        data_frame = load_dataset_frame(request.dataset_id, request.data, selected_columns)
    except Exception as error:
        raise RuntimeError(f'Data Loading failed: {error}') from error

    try:
        X, y, preprocessing, feature_names, warnings_list, label_encoder = prepare_universal_training_data(data_frame, request)
    except Exception as error:
        raise RuntimeError(f'Validation failed: {error}') from error

    candidate_models = select_candidate_models(request.problem_type, len(y), len(feature_names), request.random_state)
    cv = get_cv_strategy(len(y))
    scoring = (
        {'rmse': 'neg_root_mean_squared_error', 'r2': 'r2'}
        if request.problem_type == 'regression'
        else {'accuracy': 'accuracy', 'f1_weighted': 'f1_weighted'}
    )
    primary_metric = 'rmse' if request.problem_type == 'regression' else 'accuracy'
    model_failures: list[str] = []
    best: dict[str, Any] | None = None

    for model_name, estimator in candidate_models:
        try:
            scores = cross_validate(clone(estimator), X, y, cv=cv, scoring=scoring, n_jobs=TRAINING_N_JOBS, error_score='raise')
            primary_values = scores[f'test_{primary_metric}']
            ranking_score = float(np.mean(primary_values))
            summary = {
                metric: {
                    'mean': float(np.mean(values if metric != 'rmse' else -values)),
                    'std': float(np.std(values if metric != 'rmse' else -values)),
                }
                for metric, values in ((key, scores[f'test_{key}']) for key in scoring.keys())
            }
            if best is None or ranking_score > best['ranking_score']:
                best = {'name': model_name, 'estimator': clone(estimator), 'scores': summary, 'ranking_score': ranking_score}
        except Exception as error:
            model_failures.append(f'{model_name}: {error}')
            logger.exception('Candidate model failed during training model=%s', model_name)

    if best is None:
        raise RuntimeError('Training failed for all candidate models. ' + ' | '.join(model_failures))

    try:
        start_time = time.perf_counter()
        best['estimator'].fit(X, y)
        training_time = round(time.perf_counter() - start_time, 4)
    except Exception as error:
        raise RuntimeError(f'Training model failed: {error}') from error

    try:
        fitted_predictions = best['estimator'].predict(X)
        if request.problem_type == 'regression':
            rmse = float(np.sqrt(mean_squared_error(y, fitted_predictions)))
            metrics_primary = {
                'R2': round(float(r2_score(y, fitted_predictions)), 6),
                'RMSE': round(rmse, 6),
                'MAE': round(float(mean_absolute_error(y, fitted_predictions)), 6),
            }
        else:
            metrics_primary = {
                'Accuracy': round(float(accuracy_score(y, fitted_predictions)), 6),
                'Precision': round(float(precision_score(y, fitted_predictions, average='weighted', zero_division=0)), 6),
                'Recall': round(float(recall_score(y, fitted_predictions, average='weighted', zero_division=0)), 6),
                'F1 Score': round(float(f1_score(y, fitted_predictions, average='weighted', zero_division=0)), 6),
            }
            rmse = None
    except Exception:
        logger.exception('Failed to compute fitted metrics; returning CV metrics only.')
        metrics_primary = {}
        rmse = None

    feature_importance = []
    importances = getattr(best['estimator'], 'feature_importances_', None)
    coefficients = getattr(best['estimator'], 'coef_', None)
    if importances is not None:
        values = np.asarray(importances).ravel()
    elif coefficients is not None:
        values = np.abs(np.asarray(coefficients)).mean(axis=0) if np.asarray(coefficients).ndim > 1 else np.abs(np.asarray(coefficients).ravel())
    else:
        values = np.zeros(len(feature_names))
    for index, feature_name in enumerate(feature_names):
        feature_importance.append({'name': feature_name, 'importance': round(float(values[index]) if index < len(values) else 0.0, 6)})
    feature_importance.sort(key=lambda item: item['importance'], reverse=True)

    model_id = request.dataset_id or str(uuid.uuid4())[:8]
    trained_at = datetime.utcnow().isoformat()
    bundle = {
        'model': best['estimator'],
        'pipeline': best['estimator'],
        'scaler': preprocessing['scaler'],
        'encoder': preprocessing['categories'],
        'preprocessing': preprocessing,
        'feature_names': feature_names,
        'feature_columns': preprocessing['raw_feature_columns'],
        'target_name': request.target_column,
        'target_column': request.target_column,
        'task_type': request.problem_type,
        'problem_type': request.problem_type,
        'model_type': best['name'].lower().replace(' ', '_'),
        'model_name': best['name'],
        'label_encoder': label_encoder,
        'cv_score_mean': float(best['scores'][primary_metric]['mean']),
        'cv_score_std': float(best['scores'][primary_metric]['std']),
        'cv_scores': best['scores'],
        'rmse': rmse,
        'trained_at': trained_at,
        'n_rows': int(len(y)),
        'n_features': int(len(feature_names)),
        'file_format': Path(str((DATASET_CACHE.get(request.dataset_id or '') or {}).get('filename') or '')).suffix.lower(),
    }
    try:
        save_model_bundle(model_id, bundle)
        joblib.dump(bundle, MODEL_DIR / f'{model_id}_model_bundle.joblib')
    except Exception as error:
        raise RuntimeError(f'Saving failed: {error}') from error

    if request.problem_type == 'classification' and label_encoder is not None:
        actual_values = label_encoder.inverse_transform(y.astype(int))
        predicted_values = label_encoder.inverse_transform(pd.Series(fitted_predictions).astype(int))
    else:
        actual_values = y.tolist()
        predicted_values = pd.Series(fitted_predictions).tolist()
    sample_predictions = [
        {'actual': safe_serialize(actual_values[index]), 'predicted': safe_serialize(predicted_values[index])}
        for index in range(min(10, len(predicted_values)))
    ]

    full_metrics = {
        'primary': metrics_primary,
        'cv_scores': [round(float(v), 6) for v in np.atleast_1d(best['scores'][primary_metric]['mean'])],
        'cv_mean': round(float(best['scores'][primary_metric]['mean']), 6),
        'cv_std': round(float(best['scores'][primary_metric]['std']), 6),
        'cv_folds_used': int(getattr(cv, 'n_splits', len(y))),
        'cv_rows_evaluated': int(len(y)),
    }
    response = {
        'status': 'success',
        'model': best['name'],
        'model_id': model_id,
        'model_name': best['name'],
        'problem_type': request.problem_type,
        'scores': best['scores'],
        'metrics': metrics_primary,
        'full_metrics': full_metrics,
        'features_used': feature_names,
        'warnings': warnings_list,
        'feature_importance': feature_importance,
        'sample_predictions': sample_predictions,
        'analysis': make_analysis(request.problem_type, best['name'], full_metrics, feature_importance),
        'training_time': training_time,
        'trained_at': trained_at,
        'cv_scores': [round(float(best['scores'][primary_metric]['mean']), 6)],
        'overfitting_detected': False,
        'overfitting_status': 'healthy',
        'overfitting_explanation': 'Model quality is reported from cross-validation mean and standard deviation.',
        'generalization_gap': 0,
        'cv_gap': None,
        'optimization': {
            'training_rows_available': int(len(data_frame)),
            'training_rows_used': int(len(y)),
            'training_sampled': False,
            'cv_rows_evaluated': int(len(y)),
            'cv_folds_used': int(getattr(cv, 'n_splits', len(y))),
            'cv_sampled': False,
            'training_mode': request.training_mode,
            'importance_rows_evaluated': int(len(y)),
            'importance_repeats': 0,
        },
    }
    record_activity(
        request=http_request,
        action='train_model',
        status='success',
        dataset_id=request.dataset_id,
        model_id=model_id,
        detail=f'Trained {best["name"]} for a {request.problem_type} task.',
        metadata={'target_column': request.target_column, 'feature_count': len(feature_names), 'primary_metrics': metrics_primary},
    )
    return response


def infer_problem_type_from_estimator(estimator: Any) -> ProblemType:
    estimator_type = getattr(estimator, '_estimator_type', None)
    if estimator_type == 'classifier' or hasattr(estimator, 'predict_proba'):
        return 'classification'
    return 'regression'


def normalize_uploaded_bundle(raw_bundle: Any, filename: str) -> dict[str, Any]:
    if isinstance(raw_bundle, dict):
        pipeline = raw_bundle.get('pipeline') or raw_bundle.get('model') or raw_bundle.get('estimator')
        feature_columns = raw_bundle.get('feature_columns') or raw_bundle.get('features') or []
        target_column = raw_bundle.get('target_column') or raw_bundle.get('target') or 'prediction_target'
        problem_type = raw_bundle.get('problem_type')
        model_type = raw_bundle.get('model_type')
        model_name = raw_bundle.get('model_name')
        label_encoder = raw_bundle.get('label_encoder')
        trained_at = raw_bundle.get('trained_at') or datetime.utcnow().isoformat()
    else:
        pipeline = raw_bundle
        feature_columns = getattr(raw_bundle, 'feature_names_in_', [])
        target_column = 'prediction_target'
        problem_type = None
        model_type = type(raw_bundle).__name__.lower()
        model_name = type(raw_bundle).__name__
        label_encoder = None
        trained_at = datetime.utcnow().isoformat()
    if pipeline is None or not hasattr(pipeline, 'predict'):
        raise HTTPException(status_code=400, detail='Uploaded file must contain a scikit-learn compatible model or pipeline with a predict() method.')
    if not problem_type:
        problem_type = infer_problem_type_from_estimator(getattr(pipeline, 'named_steps', {}).get('model', pipeline))
    if not model_name:
        model_name = type(getattr(pipeline, 'named_steps', {}).get('model', pipeline)).__name__
    if not model_type:
        model_type = model_name.lower().replace(' ', '_')
    return {
        'pipeline': pipeline,
        'feature_columns': list(feature_columns),
        'target_column': target_column,
        'problem_type': problem_type,
        'model_type': model_type,
        'model_name': model_name,
        'label_encoder': label_encoder,
        'trained_at': trained_at,
    }


@router.post('/train')
def train_model(request: TrainRequest, http_request: Request) -> JSONResponse:
    try:
        logger.info(
            'Train request received problem_type=%s model_type=%s dataset_id=%s feature_count=%s training_mode=%s',
            request.problem_type,
            request.model_type,
            request.dataset_id,
            len(request.feature_columns),
            request.training_mode,
        )
        response = train_universal_model(request, http_request)
        logger.info('Train request completed successfully model_id=%s model_name=%s', response.get('model_id'), response.get('model_name'))
        return JSONResponse(status_code=200, content=safe_serialize(response))
    except Exception as error:
        message = str(error)
        failed_step = 'Training'
        for candidate_step in ('Data Loading', 'Validation', 'Saving'):
            if candidate_step.lower() in message.lower():
                failed_step = candidate_step
                break
        return structured_training_failure(error, step=failed_step, request=request)


@router.post('/sales-forecast')
def sales_forecast(request: SalesForecastRequest, http_request: Request) -> JSONResponse:
    required_columns = [request.date_column, request.target_column]
    if request.dataset_id:
        dataset_entry = DATASET_CACHE.get(request.dataset_id)
        if dataset_entry is None:
            raise HTTPException(status_code=400, detail='Cached dataset not found. Please upload the file again.')
        available_columns = list(dataset_entry['columns'])
        resolved_columns = resolve_requested_columns(required_columns, available_columns)
        resolved_date_column = resolved_columns[request.date_column]
        resolved_target_column = resolved_columns[request.target_column]
        if dataset_entry.get('parquet_path'):
            series_frame, freq, period_label = prepare_sales_series_from_parquet(dataset_entry, resolved_date_column, resolved_target_column)
        else:
            frame = load_dataset_frame(request.dataset_id, request.data, required_columns)
            series_frame, freq, period_label = prepare_sales_series(frame, request.date_column, request.target_column)
    else:
        frame = load_dataset_frame(request.dataset_id, request.data, required_columns)
        series_frame, freq, period_label = prepare_sales_series(frame, request.date_column, request.target_column)

    total_periods = len(series_frame)
    derived_test_periods = int(round(total_periods * (request.test_percentage / 100)))
    requested_test_periods = request.test_periods if request.test_periods is not None else derived_test_periods
    effective_test_periods = min(max(1, requested_test_periods), max(1, total_periods - 4))
    train_periods = total_periods - effective_test_periods
    effective_train_percentage = round((train_periods / total_periods) * 100, 1)
    effective_test_percentage = round((effective_test_periods / total_periods) * 100, 1)
    effective_lag_periods = min(request.lag_periods, max(1, train_periods - 1), max(1, total_periods - 2))
    train_series = series_frame.iloc[:train_periods].copy()
    test_series = series_frame.iloc[train_periods:].copy()

    train_X, train_y = build_forecast_training_frame(train_series, effective_lag_periods)
    model = LinearRegression()
    model.fit(train_X, train_y)

    test_start_period = pd.Timestamp(test_series.iloc[0]['period'])
    historical_train_values = train_series['sales'].astype(float).tolist()
    test_predictions = recursive_forecast(
        model,
        historical_train_values,
        test_start_period,
        effective_test_periods,
        effective_lag_periods,
        freq,
        period_label,
    )

    actual_test_values = test_series['sales'].astype(float).tolist()
    predicted_test_values = [float(item['predicted']) for item in test_predictions]
    metrics = calculate_forecast_metrics(actual_test_values, predicted_test_values)

    full_X, full_y = build_forecast_training_frame(series_frame, effective_lag_periods)
    full_model = LinearRegression()
    full_model.fit(full_X, full_y)
    future_start_period = pd.Timestamp(series_frame.iloc[-1]['period']) + pd.tseries.frequencies.to_offset(freq)
    future_predictions = recursive_forecast(
        full_model,
        series_frame['sales'].astype(float).tolist(),
        future_start_period,
        request.forecast_periods,
        effective_lag_periods,
        freq,
        period_label,
    )

    history = [
        {
            'period': format_forecast_period(pd.Timestamp(row['period']), period_label),
            'actual': round(float(row['sales']), 2),
        }
        for _, row in series_frame.iterrows()
    ]

    test_results = [
        {
            'period': format_forecast_period(pd.Timestamp(test_series.iloc[index]['period']), period_label),
            'actual': round(float(actual_test_values[index]), 2),
            'predicted': round(float(predicted_test_values[index]), 2),
        }
        for index in range(len(test_predictions))
    ]

    plural_label = period_label if request.forecast_periods == 1 else f'{period_label}s'
    analysis = (
        f"Sales forecasting used a time-series regression model trained on {train_periods} historical {plural_label if train_periods != 1 else period_label} "
        f"({effective_train_percentage}% of the dataset) and backtested on {effective_test_periods} {plural_label if effective_test_periods != 1 else period_label} "
        f"({effective_test_percentage}%). The system detected a {period_label}-level pattern from your dataset and projected the next "
        f"{request.forecast_periods} {plural_label}. Backtest MAE is {metrics['mae']}, RMSE is {metrics['rmse']}, and MAPE is {metrics['mape']}%."
    )

    response = {
        'date_column': request.date_column,
        'target_column': request.target_column,
        'frequency': freq,
        'period_label': period_label,
        'history': history,
        'test_forecast': test_results,
        'future_forecast': future_predictions,
        'metrics': metrics,
        'training_summary': {
            'model_name': 'Time-series regression',
            'total_periods': total_periods,
            'train_periods': train_periods,
            'test_periods': effective_test_periods,
            'train_percentage': effective_train_percentage,
            'test_percentage': effective_test_percentage,
            'forecast_periods': request.forecast_periods,
            'lag_periods': effective_lag_periods,
            'train_start': history[0]['period'],
            'train_end': history[train_periods - 1]['period'],
            'test_start': history[train_periods]['period'],
            'test_end': history[-1]['period'],
            'last_observed_period': history[-1]['period'],
        },
        'analysis': analysis,
    }
    server_session_id = get_session_id(request.dataset_id, request.session_id)
    record_activity(
        request=http_request,
        action='sales_forecast',
        status='success',
        dataset_id=request.dataset_id,
        server_session_id=server_session_id,
        detail=f'Generated {request.forecast_periods} future {period_label} sale forecasts.',
        metadata={
            'date_column': request.date_column,
            'target_column': request.target_column,
            'forecast_periods': request.forecast_periods,
            'lag_periods': effective_lag_periods,
            'metrics': metrics,
        },
    )
    return JSONResponse(content=safe_serialize(response))


@router.post('/forecast/ts/run')
def forecast_time_series(request: TimeSeriesForecastRequest, http_request: Request) -> JSONResponse:
    required_columns = [request.date_column, request.target_column]
    if request.dataset_id:
        dataset_entry = DATASET_CACHE.get(request.dataset_id)
        if dataset_entry is None:
            raise HTTPException(status_code=400, detail='Cached dataset not found. Please upload the file again.')
        series_frame, freq, period_label = prepare_sales_series_from_cached_dataset(dataset_entry, request.date_column, request.target_column)
    else:
        frame = load_dataset_frame(request.dataset_id, request.data, required_columns)
        series_frame, freq, period_label = prepare_sales_series(frame, request.date_column, request.target_column)

    total_periods = len(series_frame)
    data_quality = ensure_forecast_data_sufficiency(series_frame, period_label, request.require_quality_gate)
    effective_test_periods = min(max(1, int(round(total_periods * (request.test_percentage / 100)))), max(1, total_periods - 4))
    train_periods = total_periods - effective_test_periods
    stationarity = compute_stationarity_check(series_frame['sales'].astype(float).tolist())
    effective_lag_periods = min(3, max(1, train_periods - 1), max(1, total_periods - 2))
    requested_model_type = (request.model_type or 'auto').strip().lower()
    allowed_time_series_models = {'auto', 'sarima', 'arima', 'prophet'}
    if requested_model_type not in allowed_time_series_models:
        raise HTTPException(status_code=400, detail='Time-series model_type must be auto, sarima, arima, or prophet.')
    time_series_candidates = ['sarima', 'arima', 'prophet'] if requested_model_type == 'auto' else [requested_model_type]
    auto_result = auto_select_forecast_model(
        series_frame,
        request.forecast_periods,
        effective_test_periods,
        effective_lag_periods,
        ['trend', 'calendar', 'lags', 'rolling'],
        freq,
        period_label,
        time_series_candidates,
    )
    selected = auto_result['selected']
    model_name = selected['model_name']
    stationarity = with_fitted_order_note(stationarity, model_name, tuple(selected.get('tuning', {}).get('order') or ()) or None, tuple(selected.get('tuning', {}).get('seasonal_order') or ()) or None)
    backtest = auto_result['test_forecast']
    future_forecast = auto_result['future_forecast']

    history = [{'period': format_forecast_period(pd.Timestamp(row['period']), period_label), 'actual': round(float(row['sales']), 2)} for _, row in series_frame.iterrows()]
    metrics = selected['metrics']
    profile = build_dataset_profile(series_frame, period_label)
    session_id = get_session_id(request.dataset_id, request.session_id)
    session_state = ensure_session_state(session_id)
    assumptions = [
        'SARIMA, ARIMA, and Prophet are compared with walk-forward validation before the best model is trained on the full series.',
        'Every forecast point includes an empirical confidence interval based on walk-forward residual dispersion.',
        'A naive last-observation baseline is always calculated for comparison.',
    ]

    response = {
        'date_column': request.date_column,
        'target_column': request.target_column,
        'frequency': freq,
        'period_label': period_label,
        'dataset_profile': profile,
        'data_quality': data_quality,
        'stationarity_check': stationarity,
        'history': history,
        'test_forecast': backtest,
        'future_forecast': future_forecast,
        'metrics': metrics,
        'training_summary': {
            'model_name': model_name,
            'total_periods': total_periods,
            'train_periods': train_periods,
            'test_periods': effective_test_periods,
            'train_percentage': round((train_periods / total_periods) * 100, 1),
            'test_percentage': round((effective_test_periods / total_periods) * 100, 1),
            'forecast_periods': request.forecast_periods,
            'train_start': history[0]['period'],
            'train_end': history[train_periods - 1]['period'],
            'test_start': history[train_periods]['period'],
            'test_end': history[-1]['period'],
            'last_observed_period': history[-1]['period'],
        },
        'model_comparison': auto_result['model_comparison'],
        'naive_baseline': auto_result['naive_baseline'],
        'validation_warnings': auto_result.get('validation_warnings', []),
        'assumptions_audit': assumptions,
        'recommended_models': build_time_series_model_recommendations(profile, stationarity),
        'model_details': {
            'model_type': selected['model_type'],
            'model_name': model_name,
            'rationale': f'{model_name} had the best validation metrics among the requested time-series candidates and was retrained on the full dataset for the final forecast.',
        },
        'analysis': (
            f'{model_name} was auto-selected and forecasted {request.forecast_periods} future {period_label}{"s" if request.forecast_periods != 1 else ""}. '
            f'The series shows {stationarity["verdict"].lower()}, and the backtest produced MAE {metrics["mae"]}, RMSE {metrics["rmse"]}, and MAPE {metrics["mape"]}%.'
        ),
    }

    session_state['forecast_steps']['ts'] = True
    session_state['time_series_result'] = safe_serialize(response)
    append_forecast_version(session_id, 'time_series_forecast', response)
    session_state['updated_at'] = utc_now_iso()
    record_activity(
        request=http_request,
        action='forecast_time_series',
        status='success',
        dataset_id=request.dataset_id,
        server_session_id=session_id,
        detail=f'Ran {model_name} time-series forecast for {request.forecast_periods} future {period_label} periods.',
        metadata={
            'date_column': request.date_column,
            'target_column': request.target_column,
            'model_type': selected['model_type'],
            'forecast_periods': request.forecast_periods,
            'metrics': metrics,
            'data_quality': data_quality,
        },
    )
    return JSONResponse(content=safe_serialize(response))


# ── TS Forecast: Stationarity endpoint ─────────────────────────

@router.post('/api/ts-forecast/stationarity')
def get_ts_stationarity(request: TsStationarityRequest) -> JSONResponse:
    """Load stationarity when TS Models step opens."""
    dataset_id = request.dataset_id
    if not dataset_id:
        raise HTTPException(status_code=400, detail='dataset_id is required.')
    df, date_col, target_col = load_ts_dataset(dataset_id)
    frequency, freq_period = detect_ts_frequency(df, date_col)
    result = check_stationarity(df[target_col], frequency)
    return JSONResponse(content=result)


# ── TS Forecast: Run (train all 3 + auto-select) ───────────────

@router.post('/api/ts-forecast/run')
def run_ts_forecast(request: TsForecastRunRequest) -> JSONResponse:
    """Train all 3 models, auto-select best, generate future forecast."""
    dataset_id = request.dataset_id
    horizon = request.horizon
    split = request.training_split
    if not dataset_id:
        raise HTTPException(status_code=400, detail='dataset_id is required.')
    df, date_col, target_col = load_ts_dataset(dataset_id)
    frequency, freq_period = detect_ts_frequency(df, date_col)
    stationarity = check_stationarity(df[target_col], frequency)
    try:
        results, y_train, y_test, train, test, clean_df, first_nonzero_date = train_all_ts_models(df, target_col, date_col, frequency, freq_period, split, horizon)
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc))
    best_name, best_metrics, reason = auto_select_ts_model(results)
    future_forecast = generate_ts_future_forecast(best_name, clean_df, target_col, date_col, frequency, freq_period, horizon)
    model_comparison_list = [{'model': k, 'smape': v.get('smape'), 'mae': v.get('mae'), 'status': v['status']} for k, v in results.items()]
    insight = generate_ts_insight(best_name, best_metrics, stationarity, model_comparison_list)
    output_dir = get_ts_output_dir(dataset_id)
    try:
        write_ts_output_files(output_dir, frequency, freq_period, target_col, date_col, df, clean_df, train, test, stationarity, results, best_name, best_metrics, reason, future_forecast, insight)
    except Exception as exc:
        logger.warning('write_ts_output_files failed: %s', exc)
    try:
        write_ts_to_postgres(dataset_id, best_name, best_metrics, results, stationarity, future_forecast, insight)
    except Exception as exc:
        logger.warning('write_ts_to_postgres failed: %s', exc)
    save_workspace_context(dataset_id, 'ts_forecast_context', {'frequency': frequency, 'freq_period': freq_period, 'date_col': date_col, 'target_col': target_col, 'clean_start_date': first_nonzero_date, 'training_split': split, 'pipeline_status': 'completed', 'next_tab': 'ml_forecast', 'retrain_available': True})
    response_data = {
        'status': 'completed',
        'best_model': best_name,
        'smape': best_metrics['smape'],
        'mae': best_metrics['mae'],
        'rmse': best_metrics.get('rmse'),
        'mape': best_metrics.get('mape'),
        'reason': reason,
        'stationarity': stationarity,
        'future_forecast': future_forecast,
        'insight': insight,
        'model_comparison': [{'model': k, 'status': v['status'], 'mae': v.get('mae'), 'rmse': v.get('rmse'), 'mape': v.get('mape'), 'smape': v.get('smape'), 'note': v.get('error', 'completed')} for k, v in results.items()]
    }
    return JSONResponse(content=response_data)


@router.post('/forecast/ml/run')
def forecast_ml(request: MlForecastRequest, http_request: Request) -> JSONResponse:
    session_id = get_session_id(request.dataset_id, request.session_id)
    required_columns = [request.date_column, request.target_column]
    pipeline_frame: pd.DataFrame | None = None
    input_path: Path
    if request.dataset_id:
        dataset_entry = DATASET_CACHE.get(request.dataset_id)
        if dataset_entry is None:
            raise HTTPException(status_code=400, detail='Cached dataset not found. Please upload the file again.')
        if dataset_entry.get('parquet_path'):
            input_path = Path(str(dataset_entry['parquet_path']))
        else:
            pipeline_frame = load_dataset_frame(request.dataset_id, [], required_columns)
            input_path = Path(str(dataset_entry.get('frame_path') or dataset_entry.get('csv_path') or dataset_entry.get('excel_path') or (BASE_DIR / 'datasets')))
    else:
        pipeline_frame = load_dataset_frame(request.dataset_id, request.data, required_columns)
        input_path = BASE_DIR / 'datasets'

    try:
        pipeline_result = run_full_pipeline(
            input_path,
            target_col=request.target_column,
            date_col=request.date_column,
            horizon=request.forecast_periods,
            frequency='auto',
            frame=pipeline_frame,
        )
    except ValueError as error:
        raise HTTPException(status_code=422, detail=str(error)) from error
    except Exception as error:
        raise HTTPException(status_code=500, detail=f'ML forecasting pipeline failed: {error}') from error

    metadata = pipeline_result['metadata']
    selected_model = pipeline_result['selected_model']
    metrics = {
        'mae': selected_model['mae'],
        'rmse': selected_model['rmse'],
        'mape': selected_model['mape'],
        'smape': selected_model['smape'],
    }
    frequency_label = 'week' if metadata['frequency'] == 'weekly' else 'month'
    actual_rows = [row for row in pipeline_result['forecast_line'] if row['type'] == 'actual']
    backtest_rows = [row for row in pipeline_result['forecast_line'] if row['type'] == 'backtest']
    forecast_rows = [row for row in pipeline_result['forecast_line'] if row['type'] == 'forecast']
    history = [{'period': row['period'], 'actual': row['actual']} for row in actual_rows]
    test_results = [{'period': row['period'], 'actual': None, 'predicted': row['backtest']} for row in backtest_rows]
    future_predictions = [{'period': row['period'], 'predicted': row['forecast']} for row in forecast_rows]
    generated_features = pipeline_result['feature_table_sample']['columns']
    importance = [{'name': row['feature'], 'importance': row['importance']} for row in pipeline_result['shap_importance']]
    total_periods = len(history)
    test_periods = len(test_results)
    train_periods = max(0, total_periods - test_periods)
    data_quality = {
        'score': selected_model['data_quality_score'],
        'status': selected_model['data_quality_status'],
        'minimum_required_periods': 20 if metadata['frequency'] == 'weekly' else 12,
        'usable_periods': metadata['usable_periods'],
        'missing_share': 0,
        'zero_or_negative_share': 0,
        'volatility': metadata['volatility'],
        'issues': [] if selected_model['data_quality_status'] == 'pass' else ['Data quality score below production pass threshold.'],
    }
    model_comparison = [
        {
            'model_type': row['model'].lower().replace(' ', '_'),
            'model_name': row['model'],
            'status': row['status'],
            'metrics': {'mae': row['mae'], 'rmse': row['rmse'], 'mape': row['mape'], 'smape': row['smape']} if row['status'] == 'completed' else None,
            'skip_reason': row['note'] if row['status'] != 'completed' else None,
            'availability_note': row['note'],
        }
        for row in pipeline_result['model_comparison']
    ]
    assumptions = [
        'Gradient Boosting, Prophet, XGBoost, and LightGBM candidates are compared with walk-forward validation using MAE, RMSE, MAPE, and SMAPE.',
        'SMAPE is the primary model selection metric to reduce near-zero denominator inflation.',
        'LightGBM availability is reported as a named failure if unavailable.',
        'A naive last-observation baseline is always calculated for comparison.',
    ]

    response = {
        'date_column': metadata['date_col'],
        'target_column': metadata['target_col'],
        'frequency': 'W-MON' if metadata['frequency'] == 'weekly' else 'MS',
        'period_label': frequency_label,
        'dataset_profile': {
            'detected_frequency': frequency_label,
            'usable_periods': metadata['usable_periods'],
            'volatility': metadata['volatility'],
            'zero_value_share': 0,
        },
        'data_quality': data_quality,
        'generated_features': generated_features,
        'feature_preview_rows': pipeline_result['feature_table_sample']['rows'],
        'history': history,
        'test_forecast': test_results,
        'future_forecast': future_predictions,
        'metrics': metrics,
        'training_summary': {
            'model_name': selected_model['model_name'],
            'total_periods': total_periods,
            'train_periods': train_periods,
            'test_periods': test_periods,
            'train_percentage': 80,
            'test_percentage': 20,
            'forecast_periods': request.forecast_periods,
            'lag_periods': len([feature for feature in generated_features if feature.startswith('lag_')]),
            'train_start': history[0]['period'] if history else None,
            'train_end': history[max(0, train_periods - 1)]['period'] if history and train_periods else None,
            'test_start': history[train_periods]['period'] if history and train_periods < len(history) else None,
            'test_end': history[-1]['period'] if history else None,
            'last_observed_period': history[-1]['period'] if history else None,
        },
        'shap_feature_importance': importance,
        'model_comparison': model_comparison,
        'naive_baseline': {
            'model_name': 'Naive last-observation baseline',
            'metrics': {'mae': selected_model['naive_baseline_mae']},
            'mae_improvement_pct': selected_model['mae_improvement_pct'],
        },
        'validation_warnings': [],
        'assumptions_audit': assumptions,
        'recommended_models': build_ml_model_recommendations(generated_features),
        'model_details': {
            'model_type': selected_model['model_name'].lower().replace(' ', '_'),
            'model_name': selected_model['model_name'],
            'rationale': selected_model['selection_note'],
            'selection_metric': selected_model['selection_metric'],
        },
        'artifact_output_dir': pipeline_result['output_dir'],
        'pipeline_status': metadata['pipeline_status'],
        'next_tab': metadata['next_tab'],
        'retrain_available': metadata['retrain_available'],
        'analysis': pipeline_result['forecast_insight']['insight_text'],
    }

    session_state = ensure_session_state(session_id)
    session_state['forecast_steps']['ml'] = True
    session_state['ml_forecast_result'] = safe_serialize(response)
    append_forecast_version(session_id, 'ml_forecast', response)
    session_state['updated_at'] = utc_now_iso()
    record_activity(
        request=http_request,
        action='forecast_ml',
        status='success',
        dataset_id=request.dataset_id,
        server_session_id=session_id,
        detail=f'Ran ML forecasting with {selected_model["model_name"]} over {request.forecast_periods} future {frequency_label} periods.',
        metadata={
            'date_column': request.date_column,
            'target_column': request.target_column,
            'forecast_periods': request.forecast_periods,
            'lag_periods': response['training_summary']['lag_periods'],
            'feature_groups': request.feature_groups,
            'model_type': response['model_details']['model_type'],
            'metrics': metrics,
            'data_quality': data_quality,
        },
    )
    return JSONResponse(content=safe_serialize(response))


LOSS_COLUMN_PATTERNS = {
    'revenue_loss': re.compile(r'revenue_loss|lost_revenue|missed_revenue|sales_loss|lost_sales|lost_sale', re.IGNORECASE),
    'returns': re.compile(r'return|refund|return_qty', re.IGNORECASE),
    'inventory_value': re.compile(r'inventory[_\s-]*(value|amt|amount|cost)|stock[_\s-]*(value|amt|amount|cost)', re.IGNORECASE),
    'discount': re.compile(r'discount|promo|markdown', re.IGNORECASE),
    'waste': re.compile(r'waste|spoil|damage|scrap', re.IGNORECASE),
    'stockout': re.compile(r'stockout|stock_out|out_of_stock', re.IGNORECASE),
    'quantity': re.compile(r'quantity|qty|units', re.IGNORECASE),
    'unit_cost': re.compile(r'unit_cost|cost_per_unit|unit_cogs|cost_each', re.IGNORECASE),
    'cost': re.compile(r'\bcogs\b|cost_of_goods|total_cost|direct_cost|cost_amount|purchase_cost|material_cost|cost\b', re.IGNORECASE),
    'universal_cost': re.compile(r'cost|cogs|expense|opex|operational|discount|inventory|stock', re.IGNORECASE),
    'operating_cost': re.compile(r'operating_cost|operating_expense|opex|expense|overhead|admin_cost|fixed_cost', re.IGNORECASE),
    'gross_profit': re.compile(r'gross_profit|gross_margin_value|gross_income', re.IGNORECASE),
    'net_profit': re.compile(r'net_profit|profit_after|net_income|earnings', re.IGNORECASE),
    'margin_pct': re.compile(r'gross_margin_pct|gross_margin_percent|margin_pct|margin_percent|margin_rate', re.IGNORECASE),
    'category': re.compile(r'category|product|segment|sku|item', re.IGNORECASE),
    'region': re.compile(r'region|market|territory|location|state|city', re.IGNORECASE),
    'revenue': re.compile(r'^(?!.*(?:loss|lost|missed)).*(revenue|sales|amount|total|net_sales)', re.IGNORECASE),
    'date': re.compile(r'date|month|period|time', re.IGNORECASE),
    'price': re.compile(r'price|unit_price|rate', re.IGNORECASE),
}


def column_matches_pattern(column_name: str | None, pattern: re.Pattern[str]) -> bool:
    return bool(column_name and pattern.search(str(column_name)))


def first_matching_column(frame: pd.DataFrame, pattern: re.Pattern[str], numeric: bool | None = None) -> str | None:
    for column in frame.columns:
        if not pattern.search(str(column)):
            continue
        if numeric is True:
            converted = pd.to_numeric(frame[column], errors='coerce')
            if converted.notna().sum() == 0:
                continue
        return str(column)
    return None


def matching_numeric_columns(
    frame: pd.DataFrame,
    pattern: re.Pattern[str],
    exclude: set[str] | None = None,
) -> list[str]:
    excluded = exclude or set()
    columns: list[str] = []
    for column in frame.columns:
        column_name = str(column)
        if column_name in excluded or not pattern.search(column_name):
            continue
        converted = pd.to_numeric(frame[column], errors='coerce')
        if converted.notna().sum() == 0:
            continue
        columns.append(column_name)
    return columns


def bounded_ratio(value: float, fallback: float, lower: float = 0.08, upper: float = 0.92) -> float:
    if not np.isfinite(value) or value <= 0:
        return fallback
    return min(upper, max(lower, float(value)))


def numeric_series(frame: pd.DataFrame, column: str | None, default: float = 0.0) -> pd.Series:
    if not column or column not in frame.columns:
        return pd.Series(default, index=frame.index, dtype='float64')
    return pd.to_numeric(frame[column], errors='coerce').fillna(default).astype(float)


def numeric_columns_sum(frame: pd.DataFrame, columns: list[str], default: float = 0.0) -> pd.Series:
    if not columns:
        return pd.Series(default, index=frame.index, dtype='float64')
    total = pd.Series(0.0, index=frame.index, dtype='float64')
    for column in columns:
        total = total.add(numeric_series(frame, column, 0.0), fill_value=0.0)
    return total.fillna(default).astype(float)


def repair_zero_period_values_with_seasonal_interpolation(
    frame: pd.DataFrame,
    value_column: str,
    period_column: str = 'period',
) -> pd.DataFrame:
    repaired_frame = frame.copy()
    if repaired_frame.empty or value_column not in repaired_frame.columns or period_column not in repaired_frame.columns:
        return repaired_frame

    values = pd.to_numeric(repaired_frame[value_column], errors='coerce').astype(float)
    missing_like = values.isna() | (values <= 0)
    if not missing_like.any():
        repaired_frame[value_column] = values.fillna(0.0)
        return repaired_frame

    tail_size = max(1, int(np.ceil(len(values) * 0.2)))
    tail_start_boundary = max(0, len(values) - tail_size)
    trailing_positions = np.flatnonzero(missing_like.to_numpy())
    if trailing_positions.size == 0 or trailing_positions[-1] != len(values) - 1:
        repaired_frame[value_column] = values.fillna(0.0)
        return repaired_frame

    run_start = int(trailing_positions[-1])
    while run_start > 0 and bool(missing_like.iloc[run_start - 1]):
        run_start -= 1
    if run_start < tail_start_boundary:
        repaired_frame[value_column] = values.fillna(0.0)
        return repaired_frame

    prior_history = values.iloc[:run_start]
    if prior_history.empty or prior_history.isna().any() or (prior_history <= 0).any():
        repaired_frame[value_column] = values.fillna(0.0)
        return repaired_frame

    positive_values = prior_history[prior_history > 0]
    if positive_values.empty:
        repaired_frame[value_column] = values.fillna(0.0)
        return repaired_frame

    periods = pd.to_datetime(repaired_frame[period_column], errors='coerce')
    period_label = infer_sales_time_frequency(periods.dropna())[1] if periods.notna().sum() >= 2 else 'month'
    repair_season_length = seasonal_period_for_label(period_label)
    season_length = max(1, min(repair_season_length, max(1, len(repaired_frame) - 1)))

    repaired = values.copy()
    repaired.iloc[run_start:] = np.nan
    if period_label in {'day', 'week', 'month', 'quarter'} and season_length > 1 and len(prior_history) >= season_length:
        seasonal_positions = np.arange(len(prior_history)) % season_length
        seasonal_means = prior_history.groupby(seasonal_positions).mean()
        overall_mean = max(float(prior_history.mean()), 1e-9)
        seasonal_index = (seasonal_means / overall_mean).replace([np.inf, -np.inf], np.nan).fillna(1.0)
        deseasonalized = prior_history / pd.Series(
            [float(seasonal_index.get(position % season_length, 1.0)) for position in range(len(prior_history))],
            index=prior_history.index,
        )
        trend_values = deseasonalized.rolling(min(3, len(deseasonalized)), min_periods=1).mean()
        recent_trend = float(trend_values.tail(min(season_length, len(trend_values))).median())
        recent_level = float(prior_history.tail(min(season_length, len(prior_history))).median())
        if not np.isfinite(recent_trend) or recent_trend <= 0:
            recent_trend = recent_level
        for position in range(run_start, len(repaired)):
            season_factor = float(seasonal_index.get(position % season_length, 1.0))
            repaired.iloc[position] = max(recent_trend * season_factor, 0.0)
    else:
        repaired = repaired.interpolate(method='linear', limit_direction='forward')
        recent = prior_history.tail(min(4, len(prior_history))).astype(float)
        slope = float(np.mean(np.diff(recent))) if len(recent) > 1 else 0.0
        last_value = float(prior_history.iloc[-1])
        for offset, position in enumerate(range(run_start, len(repaired)), start=1):
            repaired.iloc[position] = max(last_value + (slope * offset), 0.0)

    repaired = repaired.fillna(float(positive_values.median())).clip(lower=0)
    repaired_frame[value_column] = repaired.astype(float)
    return repaired_frame


def distribute_repaired_period_totals(
    frame: pd.DataFrame,
    date_column: str,
    value_column: str,
    period_freq: str,
) -> pd.Series:
    values = numeric_series(frame, value_column, 0.0).clip(lower=0)
    work = frame.copy()
    work['_repair_period'] = pd.to_datetime(work[date_column], errors='coerce').dt.to_period(period_freq).dt.to_timestamp()
    valid_periods = work['_repair_period'].notna()
    if not valid_periods.any():
        return values

    grouped = (
        work.loc[valid_periods]
        .assign(_repair_value=values.loc[valid_periods])
        .groupby('_repair_period', as_index=False)['_repair_value']
        .sum()
        .rename(columns={'_repair_period': 'period', '_repair_value': value_column})
        .sort_values('period')
    )
    repaired_grouped = repair_zero_period_values_with_seasonal_interpolation(grouped, value_column)
    repaired_totals = repaired_grouped.set_index('period')[value_column].to_dict()
    original_totals = grouped.set_index('period')[value_column].to_dict()
    period_counts = work.loc[valid_periods].groupby('_repair_period').size().to_dict()

    repaired_values = values.copy()
    for period_value, row_indexes in work.loc[valid_periods].groupby('_repair_period').groups.items():
        original_total = float(original_totals.get(period_value, 0.0) or 0.0)
        repaired_total = float(repaired_totals.get(period_value, original_total) or 0.0)
        if original_total > 0:
            repaired_values.loc[row_indexes] = values.loc[row_indexes] * (repaired_total / original_total)
        elif repaired_total > 0:
            repaired_values.loc[row_indexes] = repaired_total / max(int(period_counts.get(period_value, 1)), 1)
    return repaired_values.clip(lower=0)


def column_name_suggests_amount(column_name: str | None) -> bool:
    return bool(column_name and re.search(r'(^|[_\s-])(amt|amount|value|inr|rs|rupee|rupees)($|[_\s-])', str(column_name), re.IGNORECASE))


def repair_revenue_column_for_forecast_context(
    frame: pd.DataFrame,
    date_column: str | None,
    revenue_column: str | None,
    period_label: str | None,
) -> pd.DataFrame:
    if not date_column or not revenue_column or date_column not in frame.columns or revenue_column not in frame.columns:
        return frame
    period_freq = 'D' if period_label == 'day' else 'W' if period_label == 'week' else 'Q' if period_label == 'quarter' else 'M'
    repaired = frame.copy()
    repaired[revenue_column] = distribute_repaired_period_totals(repaired, date_column, revenue_column, period_freq)
    return repaired


def repair_forecast_revenue_from_history(
    forecast_frame: pd.DataFrame,
    history_frame: pd.DataFrame,
    date_column: str | None,
    revenue_column: str | None,
    period_label: str | None,
) -> pd.DataFrame:
    if forecast_frame.empty or 'period' not in forecast_frame.columns or not date_column or not revenue_column:
        return forecast_frame
    if date_column not in history_frame.columns or revenue_column not in history_frame.columns:
        return forecast_frame

    period_freq = 'D' if period_label == 'day' else 'W' if period_label == 'week' else 'Q' if period_label == 'quarter' else 'M'
    history = history_frame.copy()
    history['_period'] = pd.to_datetime(history[date_column], errors='coerce').dt.to_period(period_freq).dt.to_timestamp()
    history['_revenue'] = numeric_series(history, revenue_column, 0.0).clip(lower=0)
    history = history.dropna(subset=['_period'])
    if history.empty:
        return forecast_frame

    history_totals = history.groupby('_period', as_index=False)['_revenue'].sum().rename(columns={'_period': 'period', '_revenue': 'revenue'})
    history_totals = repair_zero_period_values_with_seasonal_interpolation(history_totals.sort_values('period'), 'revenue')
    positive_history = history_totals[history_totals['revenue'] > 0].copy()
    if positive_history.empty:
        return forecast_frame

    positive_history['_month'] = pd.to_datetime(positive_history['period']).dt.month
    positive_history['_quarter'] = pd.to_datetime(positive_history['period']).dt.quarter
    month_baseline = positive_history.groupby('_month')['revenue'].median().to_dict()
    quarter_baseline = positive_history.groupby('_quarter')['revenue'].median().to_dict()
    recent_baseline = float(positive_history['revenue'].tail(min(6, len(positive_history))).median())

    repaired_forecast = forecast_frame.copy()
    if 'forecasted_revenue' not in repaired_forecast.columns:
        return repaired_forecast
    revenue_values = pd.to_numeric(repaired_forecast['forecasted_revenue'], errors='coerce').fillna(0.0).astype(float)
    for index, value in revenue_values.items():
        if value > 0:
            continue
        period = pd.to_datetime(repaired_forecast.at[index, 'period'], errors='coerce')
        if pd.isna(period):
            repaired_value = recent_baseline
        elif period_label == 'quarter':
            repaired_value = float(quarter_baseline.get(int(period.quarter), recent_baseline))
        elif period_label in {'month', 'day', 'week'}:
            repaired_value = float(month_baseline.get(int(period.month), recent_baseline))
        else:
            repaired_value = recent_baseline
        repaired_forecast.at[index, 'forecasted_revenue'] = max(repaired_value, 0.0)
    return repaired_forecast


def resolve_cogs_series(frame: pd.DataFrame, revenue_column: str | None = None) -> tuple[pd.Series, str, float]:
    revenue = numeric_series(frame, revenue_column, 0.0).clip(lower=0)
    quantity_column = first_matching_column(frame, LOSS_COLUMN_PATTERNS['quantity'], numeric=True)
    unit_cost_column = first_matching_column(frame, LOSS_COLUMN_PATTERNS['unit_cost'], numeric=True)
    excluded_cost_columns = {
        column for column in frame.columns
        if column_matches_pattern(str(column), LOSS_COLUMN_PATTERNS['operating_cost'])
        or column_matches_pattern(str(column), LOSS_COLUMN_PATTERNS['inventory_value'])
        or column_matches_pattern(str(column), LOSS_COLUMN_PATTERNS['discount'])
        or str(column) == str(revenue_column)
    }
    cost_columns = matching_numeric_columns(frame, LOSS_COLUMN_PATTERNS['cost'], exclude={str(column) for column in excluded_cost_columns})
    gross_profit_column = first_matching_column(frame, LOSS_COLUMN_PATTERNS['gross_profit'], numeric=True)
    margin_column = first_matching_column(frame, LOSS_COLUMN_PATTERNS['margin_pct'], numeric=True)

    if cost_columns:
        cogs = numeric_columns_sum(frame, cost_columns, 0.0).clip(lower=0)
        ratio = bounded_ratio(float(cogs.sum() / max(revenue.sum(), 1.0)), 0.58)
        joined_columns = '", "'.join(cost_columns)
        return cogs, f'mapped cost column(s) "{joined_columns}"', ratio

    if quantity_column and unit_cost_column:
        cogs = (numeric_series(frame, quantity_column, 0.0).clip(lower=0) * numeric_series(frame, unit_cost_column, 0.0).clip(lower=0)).clip(lower=0)
        if float(cogs.sum()) > 0:
            ratio = bounded_ratio(float(cogs.sum() / max(revenue.sum(), 1.0)), 0.58)
            return cogs, f'quantity "{quantity_column}" x unit cost "{unit_cost_column}"', ratio

    if revenue_column and gross_profit_column:
        cogs = (revenue - numeric_series(frame, gross_profit_column, 0.0)).clip(lower=0)
        if float(cogs.sum()) > 0:
            ratio = bounded_ratio(float(cogs.sum() / max(revenue.sum(), 1.0)), 0.58)
            return cogs, f'revenue minus gross profit "{gross_profit_column}"', ratio

    if revenue_column and margin_column:
        margin = numeric_series(frame, margin_column, 0.0)
        margin_ratio = margin.where(margin <= 1, margin / 100).clip(lower=0, upper=0.95)
        cogs = (revenue * (1 - margin_ratio)).clip(lower=0)
        if float(cogs.sum()) > 0:
            ratio = bounded_ratio(float(cogs.sum() / max(revenue.sum(), 1.0)), 0.58)
            return cogs, f'gross margin column "{margin_column}"', ratio

    universal_cost_columns = matching_numeric_columns(frame, LOSS_COLUMN_PATTERNS['universal_cost'], exclude={str(revenue_column)} if revenue_column else set())
    if universal_cost_columns:
        cogs = numeric_columns_sum(frame, universal_cost_columns, 0.0).clip(lower=0)
        if float(cogs.sum()) > 0:
            ratio = bounded_ratio(float(cogs.sum() / max(revenue.sum(), 1.0)), 0.60)
            joined_columns = '", "'.join(universal_cost_columns)
            return cogs, f'universal cost scan column(s) "{joined_columns}"', ratio

    return (revenue * 0.60).clip(lower=0), 'fallback assumption: approximate cost is 60% of forecasted revenue', 0.60


def resolve_operating_expense_series(frame: pd.DataFrame, revenue_column: str | None, gross_profit: pd.Series | None = None) -> tuple[pd.Series, str, float]:
    revenue = numeric_series(frame, revenue_column, 0.0).clip(lower=0)
    operating_columns = matching_numeric_columns(frame, LOSS_COLUMN_PATTERNS['operating_cost'])
    if operating_columns:
        opex = numeric_columns_sum(frame, operating_columns, 0.0).clip(lower=0)
        ratio = bounded_ratio(float(opex.sum() / max(revenue.sum(), 1.0)), 0.12, lower=0.03, upper=0.45)
        joined_columns = '", "'.join(operating_columns)
        return opex, f'mapped operating expense column(s) "{joined_columns}"', ratio

    net_profit_column = first_matching_column(frame, LOSS_COLUMN_PATTERNS['net_profit'], numeric=True)
    if net_profit_column is not None and gross_profit is not None:
        opex = (gross_profit - numeric_series(frame, net_profit_column, 0.0)).clip(lower=0)
        if float(opex.sum()) > 0:
            ratio = bounded_ratio(float(opex.sum() / max(revenue.sum(), 1.0)), 0.12, lower=0.03, upper=0.45)
            return opex, f'gross profit minus net profit "{net_profit_column}"', ratio

    return (revenue * 0.12).clip(lower=0), 'standard 12% operating expense assumption', 0.12


def normalize_period_value(value: Any) -> date:
    if isinstance(value, date) and not isinstance(value, datetime):
        return value

    text_value = str(value).strip()
    quarter_match = re.fullmatch(r'(\d{4})-Q([1-4])', text_value, flags=re.IGNORECASE)
    if quarter_match:
        year = int(quarter_match.group(1))
        month = (int(quarter_match.group(2)) - 1) * 3 + 1
        return date(year, month, 1)

    week_match = re.fullmatch(r'Week of (\d{4}-\d{2}-\d{2})', text_value, flags=re.IGNORECASE)
    if week_match:
        parsed_week = pd.to_datetime(week_match.group(1), errors='coerce')
        if not pd.isna(parsed_week):
            return pd.Timestamp(parsed_week).date()

    if re.fullmatch(r'\d{4}', text_value):
        return date(int(text_value), 1, 1)

    if re.fullmatch(r'\d{4}-\d{2}', text_value):
        parsed_month = pd.to_datetime(f'{text_value}-01', errors='coerce')
        if not pd.isna(parsed_month):
            return pd.Timestamp(parsed_month).date()

    parsed = pd.to_datetime(value, errors='coerce')
    if pd.isna(parsed):
        return datetime.utcnow().date()
    return pd.Timestamp(parsed).date()


def forecast_periods_to_frame(points: list[dict[str, Any]], value_name: str) -> pd.DataFrame:
    rows = []
    for point in points or []:
        period = point.get('period')
        if period is None:
            continue
        rows.append({
            'period': normalize_period_value(period),
            value_name: float(point.get('predicted') or point.get('actual') or 0),
            f'{value_name}_lower': float(point.get('lower') or point.get('predicted') or 0),
            f'{value_name}_upper': float(point.get('upper') or point.get('predicted') or 0),
        })
    return pd.DataFrame(rows)


def get_cached_clean_frame(session_id: str) -> pd.DataFrame:
    dataset_entry = DATASET_CACHE.get(session_id)
    if dataset_entry is None:
        raise HTTPException(
            status_code=422,
            detail='Cleaned dataset cache is unavailable for this session. Reopen the dataset or rerun Data Upload and Data Cleaning.',
        )
    frame = load_full_dataset_frame(session_id, [])
    if frame.empty:
        raise HTTPException(status_code=422, detail='The cached dataset is empty. Upload a dataset with revenue, cost, and date columns.')
    return frame.copy()


def fetch_upstream_forecasts(session_id: str, workflow_name: str = 'Loss Forecast') -> tuple[pd.DataFrame, dict[str, Any], dict[str, Any], pd.DataFrame]:
    session_state = ensure_session_state(session_id)
    ts_result = session_state.get('time_series_result')
    ml_result = session_state.get('ml_forecast_result')
    if not ts_result or not ml_result:
        raise HTTPException(status_code=422, detail='Complete Time Series Forecasting and Machine Learning Forecasting before running Loss or Profit Forecast.')

    clean_frame = get_cached_clean_frame(session_id).fillna(0)
    ts_future = forecast_periods_to_frame(ts_result.get('future_forecast', []), 'forecasted_revenue')
    ml_future = forecast_periods_to_frame(ml_result.get('future_forecast', []), 'ml_forecast_value')
    if ts_future.empty or ml_future.empty:
        raise HTTPException(status_code=422, detail='Upstream forecast results do not contain future periods. Rerun the TS and ML forecast tabs.')

    merged = pd.merge(ts_future, ml_future, on='period', how='outer').sort_values('period').fillna(0)
    revenue_column = first_matching_column(clean_frame, LOSS_COLUMN_PATTERNS['revenue'], numeric=True)
    date_column = first_matching_column(clean_frame, LOSS_COLUMN_PATTERNS['date'])
    period_label = ts_result.get('period_label') or 'month'
    clean_frame = repair_revenue_column_for_forecast_context(clean_frame, date_column, revenue_column, period_label)
    merged = repair_forecast_revenue_from_history(merged, clean_frame, date_column, revenue_column, period_label)
    if merged.empty or not (pd.to_numeric(merged.get('forecasted_revenue', pd.Series(dtype='float64')), errors='coerce').fillna(0.0) > 0).any():
        raise HTTPException(status_code=422, detail=f'{workflow_name} requires valid revenue forecast values. Please re-run TS or ML Forecast with a revenue or sales column as the target.')
    cogs_series, cogs_source, cogs_ratio = resolve_cogs_series(clean_frame, revenue_column)
    ml_target_column = str(ml_result.get('target_column') or '')
    if column_matches_pattern(ml_target_column, LOSS_COLUMN_PATTERNS['cost']) or column_matches_pattern(ml_target_column, LOSS_COLUMN_PATTERNS['unit_cost']):
        merged['forecasted_cogs'] = merged['ml_forecast_value'].clip(lower=0)
        cogs_source = f'ML forecast target "{ml_target_column}"'
    else:
        merged['forecasted_cogs'] = merged['forecasted_revenue'].clip(lower=0) * cogs_ratio
    median_cogs = float(merged['forecasted_cogs'].replace(0, np.nan).median() or 0)
    median_revenue = float(merged['forecasted_revenue'].replace(0, np.nan).median() or 0)
    if median_revenue > 0 and median_cogs > median_revenue * 1.4:
        merged['forecasted_cogs'] = merged['forecasted_revenue'].clip(lower=0) * cogs_ratio
    merged.attrs['cogs_source'] = cogs_source
    merged.attrs['cogs_ratio'] = cogs_ratio
    merged.attrs['historical_cogs_total'] = float(cogs_series.sum())
    merged = merged.drop(columns=['ml_forecast_value'], errors='ignore')
    return merged, ts_result, ml_result, clean_frame


def mapped_or_matching_column(
    frame: pd.DataFrame,
    mapping: dict[str, str] | None,
    key: str,
    pattern: re.Pattern[str],
    numeric: bool | None = None,
) -> str | None:
    mapped = (mapping or {}).get(key)
    if mapped and mapped in frame.columns:
        if numeric is True and pd.to_numeric(frame[mapped], errors='coerce').notna().sum() == 0:
            return None
        return mapped
    return first_matching_column(frame, pattern, numeric=numeric)


def build_loss_base_frame(
    session_id: str,
    forecast_periods: int,
    column_mapping: dict[str, str] | None = None,
    confirmed_assumptions: bool = False,
) -> tuple[pd.DataFrame, list[dict[str, Any]], dict[str, float], list[str]]:
    forecast_frame, ts_result, _ml_result, clean_frame = fetch_upstream_forecasts(session_id, 'Loss Forecast')
    forecast_frame = forecast_frame.head(forecast_periods).copy()
    if forecast_frame.empty:
        raise HTTPException(status_code=422, detail='No forecast periods are available for loss forecasting.')

    date_column = mapped_or_matching_column(clean_frame, column_mapping, 'date', LOSS_COLUMN_PATTERNS['date'])
    revenue_column = mapped_or_matching_column(clean_frame, column_mapping, 'revenue', LOSS_COLUMN_PATTERNS['revenue'], numeric=True)
    missing = []
    if not date_column:
        missing.append('date / period')
    if not revenue_column:
        missing.append('revenue / sales amount')
    if missing:
        raise HTTPException(
            status_code=422,
            detail=f"Missing required columns for loss forecasting: {', '.join(missing)}. Remap or rename these fields in Data Upload, then rerun forecasting.",
        )

    audit_trail = [
        f'Date column mapped to "{date_column}".',
        f'Revenue column mapped to "{revenue_column}".',
        'Loss forecast uses upstream TS/ML future periods and historical loss-driver rates.',
    ]
    returns_column = first_matching_column(clean_frame, LOSS_COLUMN_PATTERNS['returns'], numeric=True)
    revenue_loss_column = first_matching_column(clean_frame, LOSS_COLUMN_PATTERNS['revenue_loss'], numeric=True)
    inventory_value_columns = matching_numeric_columns(clean_frame, LOSS_COLUMN_PATTERNS['inventory_value'])
    discount_columns = matching_numeric_columns(clean_frame, LOSS_COLUMN_PATTERNS['discount'])
    waste_column = first_matching_column(clean_frame, LOSS_COLUMN_PATTERNS['waste'], numeric=True)
    stockout_column = first_matching_column(clean_frame, LOSS_COLUMN_PATTERNS['stockout'], numeric=True)
    quantity_column = first_matching_column(clean_frame, LOSS_COLUMN_PATTERNS['quantity'], numeric=True)
    unit_cost_column = first_matching_column(clean_frame, LOSS_COLUMN_PATTERNS['unit_cost'], numeric=True)
    category_column = first_matching_column(clean_frame, LOSS_COLUMN_PATTERNS['category'])
    region_column = first_matching_column(clean_frame, LOSS_COLUMN_PATTERNS['region'])
    price_column = first_matching_column(clean_frame, LOSS_COLUMN_PATTERNS['price'], numeric=True)

    work = clean_frame.copy()
    work['_period'] = pd.to_datetime(work[date_column], errors='coerce')
    work = work.dropna(subset=['_period'])
    if work.empty:
        raise HTTPException(status_code=422, detail='No usable dates were found for loss forecasting. Check the mapped date column.')
    period_label = ts_result.get('period_label') or 'month'
    period_freq = 'D' if period_label == 'day' else 'W' if period_label == 'week' else 'Q' if period_label == 'quarter' else 'M'
    work['_period'] = work['_period'].dt.to_period(period_freq).dt.to_timestamp().dt.date
    work[revenue_column] = distribute_repaired_period_totals(work, date_column, revenue_column, period_freq)

    revenue = numeric_series(work, revenue_column)
    quantity = numeric_series(work, quantity_column, 1.0).clip(lower=0)
    actual_cost, cogs_source, cogs_ratio = resolve_cogs_series(work, revenue_column)
    gross_profit = (revenue - actual_cost).clip(lower=0)
    operating_cost, operating_source, operating_ratio = resolve_operating_expense_series(work, revenue_column, gross_profit)
    returns = numeric_series(work, returns_column, 0.0).clip(lower=0)
    raw_revenue_loss = numeric_series(work, revenue_loss_column, 0.0).clip(lower=0)
    discount_amount_columns = [column for column in discount_columns if column_name_suggests_amount(column)]
    discount_rate_columns = [column for column in discount_columns if column not in discount_amount_columns]
    discount_amounts = numeric_columns_sum(work, discount_amount_columns, 0.0).clip(lower=0)
    discount_rates = numeric_columns_sum(work, discount_rate_columns, 0.0).clip(lower=0)
    inventory_value = numeric_columns_sum(work, inventory_value_columns, 0.0).clip(lower=0)
    waste = numeric_series(work, waste_column, 0.0).clip(lower=0)
    stockout = numeric_series(work, stockout_column, 0.0).clip(lower=0)
    price = numeric_series(work, price_column, 0.0)
    if price.sum() == 0:
        price = revenue / quantity.replace(0, np.nan)
        price = price.replace([np.inf, -np.inf], np.nan).fillna(0)

    discount_pct = discount_rates.where(discount_rates <= 1, discount_rates / 100).clip(lower=0, upper=0.95)
    baseline_cost = operating_cost.rolling(7, min_periods=1).mean() * 1.2
    inferred_unit_cost = actual_cost / quantity.replace(0, np.nan)
    inferred_unit_cost = inferred_unit_cost.replace([np.inf, -np.inf], np.nan).fillna(0)
    unit_cost = numeric_series(work, unit_cost_column, 0.0).clip(lower=0) if unit_cost_column else inferred_unit_cost
    if float(unit_cost.sum()) == 0:
        unit_cost = inferred_unit_cost
    work['_actual_revenue'] = revenue
    period_revenue_baseline = revenue.rolling(3, min_periods=1).mean().shift(1).fillna(revenue.expanding(min_periods=1).mean())
    inferred_revenue_loss = (period_revenue_baseline - revenue).clip(lower=0)
    work['_revenue_loss'] = raw_revenue_loss.where(raw_revenue_loss > 0, inferred_revenue_loss)
    work['_operational_loss'] = (operating_cost - baseline_cost).clip(lower=0)
    inventory_event_loss = (waste * unit_cost) + (stockout * price)
    inventory_value_loss = inventory_value * 0.02
    work['_inventory_loss'] = inventory_event_loss.where(inventory_event_loss > 0, inventory_value_loss)
    work['_discount_loss'] = (discount_amounts + (revenue * discount_pct)).clip(lower=0)
    work['_return_loss'] = returns.where(returns > 1, returns * price).clip(lower=0)
    historical = work.groupby('_period', as_index=False).agg({
        '_actual_revenue': 'sum',
        '_revenue_loss': 'sum',
        '_operational_loss': 'sum',
        '_inventory_loss': 'sum',
        '_discount_loss': 'sum',
        '_return_loss': 'sum',
    })

    average_actual_revenue = max(float(historical['_actual_revenue'].mean() or 0), 1.0)
    driver_totals = {
        'Revenue Loss': float(historical['_revenue_loss'].sum()),
        'Operational Loss': float(historical['_operational_loss'].sum()),
        'Inventory Loss': float(historical['_inventory_loss'].sum()),
        'Discount Loss': float(historical['_discount_loss'].sum()),
        'Returns / Refunds': float(historical['_return_loss'].sum()),
    }
    total_driver = sum(driver_totals.values()) or 1.0
    driver_weights = {key: value / total_driver for key, value in driver_totals.items()}
    driver_weights['COGS Basis'] = cogs_ratio
    driver_weights['Operating Expense Basis'] = operating_ratio
    inventory_loss_rate = bounded_ratio(
        float(historical['_inventory_loss'].sum() / max(historical['_actual_revenue'].sum(), 1.0)),
        0.018,
        lower=0.002,
        upper=0.18,
    )
    discount_loss_rate = bounded_ratio(
        float(historical['_discount_loss'].sum() / max(historical['_actual_revenue'].sum(), 1.0)),
        0.02,
        lower=0.001,
        upper=0.35,
    )
    driver_weights['Inventory Loss Basis'] = inventory_loss_rate
    driver_weights['Discount Loss Basis'] = discount_loss_rate
    assumption_notes = [
        f'COGS basis: {cogs_source}.',
        f'Operating expense basis: {operating_source}.',
        'Inventory loss basis falls back to inferred waste/stockout or bounded historical exposure when explicit fields are absent.',
        'Discount loss basis falls back to bounded historical exposure when explicit discount fields are absent.',
    ]
    audit_trail.extend(assumption_notes)
    requires_confirmation = any('standard' in note.lower() or 'falls back' in note.lower() or 'fallback assumption' in note.lower() for note in assumption_notes)
    if requires_confirmation and not confirmed_assumptions:
        raise HTTPException(
            status_code=428,
            detail='Loss forecast requires confirmation of calculation assumptions before running.',
            headers={'X-Assumptions-Required': 'true'},
        )

    rows: list[dict[str, Any]] = []
    for index, item in forecast_frame.reset_index(drop=True).iterrows():
        forecasted_revenue = max(float(item.get('forecasted_revenue') or 0), 0.0)
        pressure = 1 + (index * 0.015)
        historical_revenue_loss = float(historical['_revenue_loss'].mean() or 0)
        forecast_shortfall_loss = max(0.0, average_actual_revenue - forecasted_revenue) * 0.12
        revenue_loss = max(historical_revenue_loss, forecast_shortfall_loss, forecasted_revenue * 0.015) * pressure
        operational_loss = max(float(historical['_operational_loss'].mean() or 0), forecasted_revenue * 0.025) * pressure
        inventory_loss = max(float(historical['_inventory_loss'].mean() or 0), forecasted_revenue * inventory_loss_rate) * pressure
        discount_loss = max(float(historical['_discount_loss'].mean() or 0), forecasted_revenue * discount_loss_rate) * pressure
        return_loss = max(float(historical['_return_loss'].mean() or 0), 0.0) * pressure
        total_loss = revenue_loss + operational_loss + inventory_loss + discount_loss + return_loss
        risk_score = min(1.0, total_loss / max(forecasted_revenue, 1.0))
        risk_label = 'Low' if risk_score < 0.05 else 'Medium' if risk_score <= 0.15 else 'High'
        lower = max(0.0, total_loss * 0.82)
        upper = total_loss * 1.18
        rows.append({
            'id': uuid.uuid4().hex,
            'session_id': session_id,
            'period': item['period'],
            'revenue_loss': round(revenue_loss + return_loss, 2),
            'operational_loss': round(operational_loss, 2),
            'inventory_loss': round(inventory_loss, 2),
            'discount_loss': round(discount_loss, 2),
            'total_loss': round(total_loss, 2),
            'lower_bound': round(lower, 2),
            'upper_bound': round(upper, 2),
            'loss_risk_score': round(risk_score, 4),
            'risk_label': risk_label,
            'segment': 'All Business',
            'created_at': utc_now_iso(),
        })

    for row in rows:
        inventory_column_names = '", "'.join(inventory_value_columns)
        mapped_inventory = f'inventory value column(s) "{inventory_column_names}"' if inventory_value_columns else 'waste/stockout-derived inventory exposure'
        discount_basis_parts = []
        if discount_amount_columns:
            discount_amount_names = '", "'.join(discount_amount_columns)
            discount_basis_parts.append(f'amount column(s) "{discount_amount_names}"')
        if discount_rate_columns:
            discount_rate_names = '", "'.join(discount_rate_columns)
            discount_basis_parts.append(f'rate column(s) "{discount_rate_names}"')
        mapped_discount = '; '.join(discount_basis_parts) if discount_basis_parts else 'standard discount exposure'
        row['basis_note'] = f'COGS: {cogs_source}; OpEx: {operating_source}; Inventory: {mapped_inventory}; Discount: {mapped_discount}'

    segments: list[dict[str, Any]] = []
    for column, segment_type in [(category_column, 'category'), (region_column, 'region')]:
        if not column or column not in work.columns:
            continue
        grouped = work.assign(
            _segment_loss=work['_revenue_loss'] + work['_operational_loss'] + work['_inventory_loss'] + work['_discount_loss'] + work['_return_loss'],
            _segment_revenue=work['_actual_revenue'],
        ).groupby(column, dropna=False).agg({'_segment_loss': 'sum', '_segment_revenue': 'sum'}).reset_index()
        for _, segment_row in grouped.sort_values('_segment_loss', ascending=False).head(12).iterrows():
            risk_score = min(1.0, float(segment_row['_segment_loss']) / max(float(segment_row['_segment_revenue']), 1.0))
            segments.append({
                'segment': str(segment_row[column] or 'Unassigned'),
                'segment_type': segment_type,
                'total_loss': round(float(segment_row['_segment_loss']), 2),
                'risk_score': round(risk_score, 4),
                'risk_label': 'Low' if risk_score < 0.05 else 'Medium' if risk_score <= 0.15 else 'High',
            })
    if not segments:
        segments = [{'segment': 'All Business', 'segment_type': 'portfolio', 'total_loss': round(sum(row['total_loss'] for row in rows), 2), 'risk_score': round(float(np.mean([row['loss_risk_score'] for row in rows])), 4), 'risk_label': rows[0]['risk_label'] if rows else 'Low'}]

    return pd.DataFrame(rows), segments, driver_weights, audit_trail


def persist_loss_forecast(session_id: str, rows: list[dict[str, Any]], segments: list[dict[str, Any]]) -> None:
    if not ACTIVITY_DB_AVAILABLE:
        return
    with get_activity_connection() as connection:
        connection.execute('DELETE FROM loss_forecast_results WHERE session_id = %s', (session_id,))
        for row in rows:
            connection.execute(
                '''
                INSERT INTO loss_forecast_results (
                    id, session_id, period, revenue_loss, operational_loss, inventory_loss, discount_loss, total_loss,
                    lower_bound, upper_bound, loss_risk_score, risk_label, segment, created_at
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                ''',
                (
                    row['id'], session_id, row['period'], row['revenue_loss'], row['operational_loss'],
                    row['inventory_loss'], row['discount_loss'], row['total_loss'], row.get('lower_bound'),
                    row.get('upper_bound'), row['loss_risk_score'], row['risk_label'], row.get('segment'), row['created_at'],
                ),
            )


def query_loss_results(session_id: str) -> list[dict[str, Any]]:
    session_state = ensure_session_state(session_id)
    if session_state.get('loss_forecast_result'):
        return list(session_state['loss_forecast_result'])
    if not ACTIVITY_DB_AVAILABLE:
        return []
    with get_activity_connection() as connection:
        rows = connection.execute(
            '''
            SELECT id, session_id, period::text, revenue_loss, operational_loss, inventory_loss, discount_loss,
                   total_loss, lower_bound, upper_bound, loss_risk_score, risk_label, segment, created_at
            FROM loss_forecast_results
            WHERE session_id = %s
            ORDER BY period ASC
            ''',
            (session_id,),
        ).fetchall()
    return [dict(row) for row in rows]


def build_loss_summary(rows: list[dict[str, Any]], driver_weights: dict[str, float] | None = None) -> dict[str, Any]:
    if not rows:
        return {'total_loss': 0, 'highest_risk_period': None, 'average_risk_score': 0, 'top_loss_driver': 'N/A'}
    highest = max(rows, key=lambda row: float(row.get('loss_risk_score') or 0))
    driver_totals = {
        'Revenue Loss': sum(float(row.get('revenue_loss') or 0) for row in rows),
        'Operational Loss': sum(float(row.get('operational_loss') or 0) for row in rows),
        'Inventory Loss': sum(float(row.get('inventory_loss') or 0) for row in rows),
        'Discount Loss': sum(float(row.get('discount_loss') or 0) for row in rows),
    }
    top_driver, top_value = max(driver_totals.items(), key=lambda item: item[1])
    total_loss = sum(float(row.get('total_loss') or 0) for row in rows)
    share = (top_value / total_loss * 100) if total_loss else 0
    return {
        'total_loss': round(total_loss, 2),
        'highest_risk_period': highest.get('period'),
        'average_risk_score': round(sum(float(row.get('loss_risk_score') or 0) for row in rows) / len(rows), 4),
        'top_loss_driver': f'{top_driver} - {share:.0f}%',
        'driver_weights': driver_weights or {},
    }


def build_profit_rows(
    session_id: str,
    forecast_periods: int,
    scenario_parameters: dict[str, dict[str, float]] | None = None,
    column_mapping: dict[str, str] | None = None,
    confirmed_assumptions: bool = False,
) -> tuple[dict[str, list[dict[str, Any]]], dict[str, Any], list[str]]:
    forecast_frame, _ts_result, _ml_result, clean_frame = fetch_upstream_forecasts(session_id, 'Profit Forecast')
    loss_rows = query_loss_results(session_id)
    if not loss_rows:
        loss_frame, segments, driver_weights, loss_audit = build_loss_base_frame(session_id, forecast_periods, column_mapping, confirmed_assumptions)
        loss_rows = safe_serialize(loss_frame.to_dict(orient='records'))
        persist_loss_forecast(session_id, loss_rows, segments)
        state = ensure_session_state(session_id)
        state['loss_forecast_result'] = loss_rows
        state['loss_segments'] = segments
        state['loss_summary'] = build_loss_summary(loss_rows, driver_weights)
    else:
        loss_audit = ['Existing confirmed loss forecast rows reused for profit calculations.']

    loss_by_period = {normalize_period_value(row['period']): float(row.get('total_loss') or 0) for row in loss_rows}
    forecast_frame = forecast_frame.head(forecast_periods).copy()
    revenue_column = mapped_or_matching_column(clean_frame, column_mapping, 'revenue', LOSS_COLUMN_PATTERNS['revenue'], numeric=True)
    cogs_series, _cogs_source, cogs_ratio = resolve_cogs_series(clean_frame, revenue_column)
    revenue_series = numeric_series(clean_frame, revenue_column, 0.0).clip(lower=0)
    gross_profit_series = (revenue_series - cogs_series).clip(lower=0)
    _opex_series, _opex_source, operating_expense_ratio = resolve_operating_expense_series(clean_frame, revenue_column, gross_profit_series)
    date_column = first_matching_column(clean_frame, LOSS_COLUMN_PATTERNS['date'])
    historical_cogs_ratios: list[float] = []
    seasonal_cogs_ratios: dict[int, float] = {}
    if date_column:
        margin_work = clean_frame.copy()
        margin_work['_period'] = pd.to_datetime(margin_work[date_column], errors='coerce')
        margin_work['_revenue'] = revenue_series
        margin_work['_cogs'] = cogs_series
        margin_work = margin_work.dropna(subset=['_period'])
        if not margin_work.empty:
            period_costs = margin_work.groupby(margin_work['_period'].dt.to_period('M')).agg({'_revenue': 'sum', '_cogs': 'sum'}).reset_index()
            period_costs['_ratio'] = period_costs.apply(
                lambda row: bounded_ratio(float(row['_cogs']) / float(row['_revenue']), cogs_ratio) if float(row['_revenue']) else cogs_ratio,
                axis=1,
            )
            historical_cogs_ratios = [float(value) for value in period_costs['_ratio'].tail(max(1, forecast_periods)).tolist()]
            period_costs['_month'] = period_costs['_period'].dt.month
            seasonal_cogs_ratios = {
                int(row['_month']): bounded_ratio(float(row['_ratio']), cogs_ratio)
                for _, row in period_costs.groupby('_month', as_index=False)['_ratio'].mean().iterrows()
            }
    uses_ml_cogs_forecast = 'ML forecast target' in str(forecast_frame.attrs.get('cogs_source') or '')

    scenario_config: dict[str, dict[str, float]] = {
        'optimistic': {'revenue': 1.15, 'cogs': 0.97, 'loss': 0.80},
        'baseline': {'revenue': 1.0, 'cogs': 1.0, 'loss': 1.0},
        'pessimistic': {'revenue': 0.85, 'cogs': 1.05, 'loss': 1.20},
    }
    for scenario, overrides in (scenario_parameters or {}).items():
        if scenario in scenario_config:
            for key in ('revenue', 'cogs', 'loss'):
                if key in overrides and np.isfinite(float(overrides[key])):
                    scenario_config[scenario][key] = float(overrides[key])
    audit_trail = [
        *loss_audit,
        f'Profit revenue column mapped to "{revenue_column or "not available"}".',
        f'COGS ratio basis: {cogs_ratio:.4f}.',
        f'Operating expense ratio basis: {operating_expense_ratio:.4f}.',
        f'Scenario multipliers: {json.dumps(scenario_config, sort_keys=True)}.',
    ]
    if not confirmed_assumptions:
        raise HTTPException(
            status_code=428,
            detail='Profit forecast requires confirmation of scenario and calculation assumptions before running.',
            headers={'X-Assumptions-Required': 'true'},
        )
    scenarios: dict[str, list[dict[str, Any]]] = {}
    for scenario, multipliers in scenario_config.items():
        rows: list[dict[str, Any]] = []
        for index, item in forecast_frame.reset_index(drop=True).iterrows():
            period = normalize_period_value(item['period'])
            revenue = max(float(item.get('forecasted_revenue') or 0) * multipliers['revenue'], 0.0)
            period_date = pd.to_datetime(period, errors='coerce')
            period_month = int(period_date.month) if pd.notna(period_date) else 0
            period_cogs_ratio = seasonal_cogs_ratios.get(
                period_month,
                historical_cogs_ratios[index % len(historical_cogs_ratios)] if historical_cogs_ratios else cogs_ratio,
            )
            base_cogs = float(item.get('forecasted_cogs') or 0) if uses_ml_cogs_forecast else revenue * period_cogs_ratio
            cogs = max(base_cogs * multipliers['cogs'], 0.0)
            if cogs <= 0 and revenue > 0:
                cogs = revenue * cogs_ratio
            if cogs > revenue * 0.92:
                cogs = revenue * min(0.72, max(0.42, cogs_ratio))
            losses = max(loss_by_period.get(period, 0.0) * multipliers['loss'], 0.0)
            operating_expenses = revenue * operating_expense_ratio
            gross_profit = revenue - cogs
            net_profit = gross_profit - operating_expenses - losses
            gross_margin = (gross_profit / revenue * 100) if revenue else 0.0
            net_margin = (net_profit / revenue * 100) if revenue else 0.0
            rows.append({
                'id': uuid.uuid4().hex,
                'session_id': session_id,
                'period': period,
                'forecasted_revenue': round(revenue, 2),
                'forecasted_cogs': round(cogs, 2),
                'gross_profit': round(gross_profit, 2),
                'operating_expenses': round(operating_expenses, 2),
                'total_losses': round(losses, 2),
                'net_profit': round(net_profit, 2),
                'gross_margin_pct': round(gross_margin, 2),
                'net_margin_pct': round(net_margin, 2),
                'scenario': scenario,
                'created_at': utc_now_iso(),
            })
        scenarios[scenario] = rows

    baseline = scenarios.get('baseline', [])
    breakeven_index = next((index for index, row in enumerate(baseline) if float(row['net_profit']) >= 0), None)
    breakeven = {
        'breakeven_period': baseline[breakeven_index]['period'] if breakeven_index is not None else None,
        'periods_to_breakeven': breakeven_index + 1 if breakeven_index is not None else None,
    }
    return scenarios, breakeven, audit_trail


def persist_profit_forecast(session_id: str, scenarios: dict[str, list[dict[str, Any]]]) -> None:
    if not ACTIVITY_DB_AVAILABLE:
        return
    with get_activity_connection() as connection:
        connection.execute('DELETE FROM profit_forecast_results WHERE session_id = %s', (session_id,))
        for scenario_rows in scenarios.values():
            for row in scenario_rows:
                connection.execute(
                    '''
                    INSERT INTO profit_forecast_results (
                        id, session_id, period, forecasted_revenue, forecasted_cogs, gross_profit, operating_expenses,
                        total_losses, net_profit, gross_margin_pct, net_margin_pct, scenario, created_at
                    ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                    ''',
                    (
                        row['id'], session_id, row['period'], row['forecasted_revenue'], row['forecasted_cogs'],
                        row['gross_profit'], row['operating_expenses'], row['total_losses'], row['net_profit'],
                        row['gross_margin_pct'], row['net_margin_pct'], row['scenario'], row['created_at'],
                    ),
                )


def query_profit_results(session_id: str) -> dict[str, list[dict[str, Any]]]:
    state = ensure_session_state(session_id)
    if state.get('profit_scenarios'):
        return dict(state['profit_scenarios'])
    if not ACTIVITY_DB_AVAILABLE:
        return {'optimistic': [], 'baseline': [], 'pessimistic': []}
    with get_activity_connection() as connection:
        rows = connection.execute(
            '''
            SELECT id, session_id, period::text, forecasted_revenue, forecasted_cogs, gross_profit, operating_expenses,
                   total_losses, net_profit, gross_margin_pct, net_margin_pct, scenario, created_at
            FROM profit_forecast_results
            WHERE session_id = %s
            ORDER BY scenario ASC, period ASC
            ''',
            (session_id,),
        ).fetchall()
    scenarios = {'optimistic': [], 'baseline': [], 'pessimistic': []}
    for row in rows:
        scenarios.setdefault(row['scenario'], []).append(dict(row))
    return scenarios


@router.post('/loss-forecast/run')
def run_loss_forecast(request: ForecastRunRequest, http_request: Request) -> JSONResponse:
    try:
        session_id = request.session_id
        frame, segments, driver_weights, audit_trail = build_loss_base_frame(
            session_id,
            request.forecast_periods,
            request.column_mapping,
            request.confirmed_assumptions,
        )
        rows = safe_serialize(frame.to_dict(orient='records'))
        persist_loss_forecast(session_id, rows, segments)
        state = ensure_session_state(session_id)
        state['forecast_steps']['loss'] = True
        state['loss_forecast_result'] = rows
        state['loss_segments'] = segments
        state['loss_summary'] = build_loss_summary(rows, driver_weights)
        state['loss_assumptions_audit'] = audit_trail
        append_forecast_version(session_id, 'loss_forecast', {'status': 'success', 'metrics': state['loss_summary'], 'assumptions_audit': audit_trail})
        state['updated_at'] = utc_now_iso()
        record_activity(request=http_request, action='loss_forecast', status='success', dataset_id=session_id, server_session_id=session_id, detail=f'Generated {len(rows)} loss forecast rows.')
        return JSONResponse(content=safe_serialize({'status': 'success', 'loss_forecast': rows, 'segments': segments, 'summary': state['loss_summary'], 'assumptions_audit': audit_trail}))
    except HTTPException:
        raise
    except Exception as error:
        logger.exception('Loss forecast failed session_id=%s', request.session_id)
        raise HTTPException(status_code=422, detail=f'Loss forecast failed: {error}') from error


@router.get('/loss-forecast/results/{session_id}')
def get_loss_forecast_results(session_id: str, limit: int = Query(default=250, ge=1, le=1000), offset: int = Query(default=0, ge=0)) -> JSONResponse:
    try:
        rows = query_loss_results(session_id)
        return JSONResponse(content=safe_serialize({'results': rows[offset:offset + limit], 'count': len(rows)}))
    except Exception as error:
        logger.exception('Loss forecast result lookup failed session_id=%s', session_id)
        raise HTTPException(status_code=422, detail=f'Unable to fetch loss forecast results: {error}') from error


@router.get('/loss-forecast/segments/{session_id}')
def get_loss_forecast_segments(session_id: str) -> JSONResponse:
    try:
        state = ensure_session_state(session_id)
        segments = state.get('loss_segments') or []
        if not segments:
            loss_rows = query_loss_results(session_id)
            if loss_rows:
                _frame, segments, _driver_weights, _audit_trail = build_loss_base_frame(session_id, len(loss_rows), confirmed_assumptions=True)
                state['loss_segments'] = segments
        return JSONResponse(content=safe_serialize({'segments': segments}))
    except Exception as error:
        raise HTTPException(status_code=422, detail=f'Unable to fetch loss segments: {error}') from error


@router.post('/profit-forecast/run')
def run_profit_forecast(request: ForecastRunRequest, http_request: Request) -> JSONResponse:
    try:
        session_id = request.session_id
        scenarios, breakeven, audit_trail = build_profit_rows(
            session_id,
            request.forecast_periods,
            request.scenario_parameters,
            request.column_mapping,
            request.confirmed_assumptions,
        )
        serialized = safe_serialize(scenarios)
        persist_profit_forecast(session_id, serialized)
        state = ensure_session_state(session_id)
        state['forecast_steps']['profit'] = True
        state['profit_scenarios'] = serialized
        state['breakeven'] = safe_serialize(breakeven)
        state['profit_assumptions_audit'] = audit_trail
        append_forecast_version(session_id, 'profit_forecast', {'status': 'success', 'metrics': breakeven, 'assumptions_audit': audit_trail})
        state['updated_at'] = utc_now_iso()
        record_activity(request=http_request, action='profit_forecast', status='success', dataset_id=session_id, server_session_id=session_id, detail='Generated profit forecast scenarios.')
        return JSONResponse(content=safe_serialize({'status': 'success', 'scenarios': serialized, 'breakeven': breakeven, 'assumptions_audit': audit_trail}))
    except HTTPException:
        raise
    except Exception as error:
        logger.exception('Profit forecast failed session_id=%s', request.session_id)
        raise HTTPException(status_code=422, detail=f'Profit forecast failed: {error}') from error


@router.get('/profit-forecast/results/{session_id}')
def get_profit_forecast_results(session_id: str) -> JSONResponse:
    try:
        scenarios = query_profit_results(session_id)
        return JSONResponse(content=safe_serialize({'scenarios': scenarios}))
    except Exception as error:
        raise HTTPException(status_code=422, detail=f'Unable to fetch profit forecast results: {error}') from error


@router.get('/profit-forecast/breakeven/{session_id}')
def get_profit_breakeven(session_id: str) -> JSONResponse:
    try:
        state = ensure_session_state(session_id)
        breakeven = state.get('breakeven')
        if not breakeven:
            scenarios = query_profit_results(session_id)
            baseline = scenarios.get('baseline', [])
            index = next((idx for idx, row in enumerate(baseline) if float(row.get('net_profit') or 0) >= 0), None)
            breakeven = {
                'breakeven_period': baseline[index]['period'] if index is not None else None,
                'periods_to_breakeven': index + 1 if index is not None else None,
            }
        return JSONResponse(content=safe_serialize(breakeven))
    except Exception as error:
        raise HTTPException(status_code=422, detail=f'Unable to fetch break-even analysis: {error}') from error


@router.post('/predict')
def predict(request: PredictRequest, http_request: Request) -> JSONResponse:
    try:
        bundle = load_model_bundle(request.model_id)
    except HTTPException:
        return JSONResponse(
            status_code=200,
            content={
                'status': 'failed',
                'error': 'No trained model found. Please complete ML Assistant training first.',
                'detail': f'Model bundle not found for id {request.model_id}.',
            },
        )
    except Exception as error:
        logger.exception('Prediction bundle loading failed model_id=%s', request.model_id)
        return JSONResponse(status_code=200, content={'status': 'failed', 'error': 'Unable to load the trained model.', 'detail': str(error)})

    try:
        if bundle.get('preprocessing'):
            frame_or_array = transform_features_for_bundle(request.features, bundle)
            estimator = bundle.get('model') or bundle.get('pipeline')
        else:
            missing = [feature for feature in bundle['feature_columns'] if request.features.get(feature) in [None, '']]
            if missing:
                raise ValueError(f'Missing features: {missing}')
            frame_or_array = normalize_feature_frame(pd.DataFrame([{feature: request.features.get(feature) for feature in bundle['feature_columns']}]))
            estimator = bundle['pipeline']

        raw_prediction = estimator.predict(frame_or_array)[0]
        prediction: Any = raw_prediction
        if bundle['problem_type'] == 'regression':
            prediction = int(round(float(raw_prediction)))

        payload: dict[str, Any] = {
            'status': 'success',
            'prediction': safe_serialize(prediction),
            'prediction_label': safe_serialize(prediction),
            'cv_score_mean': safe_serialize(bundle.get('cv_score_mean')),
            'cv_score_std': safe_serialize(bundle.get('cv_score_std')),
        }
        if bundle['problem_type'] == 'regression' and bundle.get('rmse') is not None:
            interval = 1.96 * float(bundle['rmse'])
            numeric_prediction = float(raw_prediction)
            payload['confidence_interval'] = {
                'lower': round(numeric_prediction - interval, 6),
                'upper': round(numeric_prediction + interval, 6),
                'margin': round(interval, 6),
            }

        label_encoder: LabelEncoder | None = bundle.get('label_encoder')
        if bundle['problem_type'] == 'classification' and label_encoder is not None:
            label = label_encoder.inverse_transform([int(prediction)])[0]
            payload['prediction_label'] = str(label)

            if hasattr(estimator, 'predict_proba'):
                probabilities = estimator.predict_proba(frame_or_array)[0]
                probability_map: dict[str, float] = {}
                for encoded_class, probability in enumerate(probabilities):
                    label_name = label_encoder.inverse_transform([encoded_class])[0]
                    probability_map[str(label_name)] = round(float(probability), 6)
                payload['probabilities'] = probability_map
                payload['confidence'] = round(float(np.max(probabilities)), 6)
                payload['top_class'] = max(probability_map, key=probability_map.get)

        record_activity(
            request=http_request,
            action='predict',
            status='success',
            model_id=request.model_id,
            detail='Generated a prediction from a trained model.',
            metadata={
                'feature_count': len(request.features),
                'prediction': payload.get('prediction_label', payload.get('prediction')),
                'confidence': payload.get('confidence'),
            },
        )
        return JSONResponse(content=safe_serialize(payload))
    except Exception as error:
        logger.exception('Prediction failed model_id=%s', request.model_id)
        return JSONResponse(
            status_code=200,
            content={
                'status': 'failed',
                'error': f'Prediction failed: {error}',
                'detail': ''.join(traceback.format_exception(type(error), error, error.__traceback__)),
            },
        )


@router.post('/upload-model')
async def upload_model(http_request: Request, file: UploadFile = File(...)) -> JSONResponse:
    filename = file.filename or 'uploaded_model'
    if not filename.lower().endswith(('.joblib', '.pkl', '.pickle')):
        raise HTTPException(status_code=400, detail='Only .joblib, .pkl, and .pickle model files are supported.')

    try:
        content = await file.read()
        raw_bundle = joblib.load(io.BytesIO(content))
        model_bundle = normalize_uploaded_bundle(raw_bundle, filename)
    except HTTPException:
        raise
    except Exception as error:
        raise HTTPException(status_code=400, detail=f'Failed to load model file: {error}') from error

    model_id = str(uuid.uuid4())[:8]
    save_model_bundle(model_id, model_bundle)
    logger.info('Uploaded model loaded successfully model_id=%s model_name=%s', model_id, model_bundle['model_name'])

    response = {
        'model_id': model_id,
        'model_name': model_bundle['model_name'],
        'model_type': model_bundle['model_type'],
        'problem_type': model_bundle['problem_type'],
        'target_column': model_bundle['target_column'],
        'feature_columns': model_bundle['feature_columns'],
        'trained_at': model_bundle['trained_at'],
        'source_filename': filename,
    }
    record_activity(
        request=http_request,
        action='upload_model',
        status='success',
        model_id=model_id,
        file_name=filename,
        detail=f'Uploaded model bundle {model_bundle["model_name"]}.',
        metadata={
            'model_type': model_bundle['model_type'],
            'problem_type': model_bundle['problem_type'],
            'feature_count': len(model_bundle['feature_columns']),
        },
    )
    return JSONResponse(content=safe_serialize(response))


def build_column_info_from_frame(frame: pd.DataFrame) -> list[dict[str, Any]]:
    info: list[dict[str, Any]] = []
    total_rows = len(frame)
    for column in frame.columns:
        series = frame[column]
        non_null = int(series.notna().sum())
        null_count = int(total_rows - non_null)
        if total_rows > LARGE_COL_CUTOFF:
            sample = series.dropna()
            if len(sample) > SAMPLE_SIZE:
                sample = sample.sample(n=SAMPLE_SIZE, random_state=RANDOM_STATE)
            unique_count = int(sample.nunique(dropna=True))
        else:
            unique_count = int(series.nunique(dropna=True))
        role = 'categorical'
        if pd.api.types.is_bool_dtype(series):
            role = 'boolean'
        elif pd.api.types.is_numeric_dtype(series):
            role = 'numeric'
        elif pd.api.types.is_datetime64_any_dtype(series):
            role = 'datetime'
        elif unique_count == total_rows and total_rows > 0:
            role = 'identifier'
        sample_values = series.dropna().head(5).map(str).tolist()
        info.append({
            'name': str(column),
            'dtype': str(series.dtype),
            'nonNull': non_null,
            'nullCount': null_count,
            'uniqueCount': unique_count,
            'role': role,
            'sample': sample_values,
        })
    return info


def build_column_info_from_polars_frame(frame: pl.DataFrame) -> list[dict[str, Any]]:
    info: list[dict[str, Any]] = []
    total_rows = frame.height
    for column in frame.columns:
        series = frame.get_column(column)
        non_null = int(series.len() - series.null_count())
        null_count = int(series.null_count())
        unique_count = int(series.n_unique()) if total_rows > 0 else 0
        dtype = series.dtype
        role = 'categorical'
        if dtype == pl.Boolean:
            role = 'boolean'
        elif dtype.is_numeric():
            role = 'numeric'
        elif dtype in pl.TEMPORAL_DTYPES:
            role = 'datetime'
        elif unique_count == total_rows and total_rows > 0:
            role = 'identifier'
        sample_values = [str(value) for value in series.drop_nulls().head(5).to_list()]
        info.append({
            'name': str(column),
            'dtype': str(dtype),
            'nonNull': non_null,
            'nullCount': null_count,
            'uniqueCount': unique_count,
            'role': role,
            'sample': sample_values,
        })
    return info



def try_parse_datetime_series(series: pd.Series) -> pd.Series | None:
    if pd.api.types.is_datetime64_any_dtype(series):
        return pd.to_datetime(series, errors='coerce')
    if not (pd.api.types.is_object_dtype(series) or pd.api.types.is_string_dtype(series)):
        return None

    sample = series.dropna().astype(str).head(50)
    if sample.empty:
        return None

    parsed = pd.to_datetime(sample, errors='coerce')
    success_ratio = float(parsed.notna().mean()) if len(sample) else 0.0
    if success_ratio < 0.6:
        return None
    return pd.to_datetime(series, errors='coerce')


def clean_cached_dataset(request: ParquetCleaningRequest) -> dict[str, Any]:
    dataset_entry = DATASET_CACHE.get(request.dataset_id)
    if dataset_entry is None:
        raise HTTPException(status_code=400, detail='Cached dataset not found. Please upload the file again.')

    if dataset_entry.get('frame_path') or dataset_entry.get('csv_path') or dataset_entry.get('excel_path'):
        if dataset_entry.get('frame_path'):
            frame = read_cached_frame(dataset_entry)
        elif dataset_entry.get('csv_path'):
            frame = read_cached_csv(dataset_entry)
        else:
            frame = read_cached_excel(dataset_entry)
        frame = normalize_dataframe(pd.DataFrame(frame))
        original_row_count = int(len(frame))
        logs: list[dict[str, Any]] = []

        if request.standardize_names:
            renamed_columns = make_unique_column_names(list(frame.columns))
            if renamed_columns != list(frame.columns):
                frame.columns = renamed_columns
                logs.append({
                    'action': 'Standardized Column Names',
                    'detail': 'Normalized column names for easier analysis and modeling.',
                    'timestamp': datetime.utcnow().isoformat(),
                })

        if request.remove_duplicates:
            before = len(frame)
            frame = frame.drop_duplicates().reset_index(drop=True)
            removed = before - len(frame)
            if removed > 0:
                logs.append({
                    'action': 'Removed Duplicates',
                    'detail': f'Removed {removed} duplicate rows.',
                    'timestamp': datetime.utcnow().isoformat(),
                })

        if request.handle_missing:
            filled_columns: list[str] = []
            for column in frame.columns:
                series = frame[column]
                if not series.isna().any():
                    continue
                if pd.api.types.is_numeric_dtype(series):
                    median = series.median()
                    fill_value = 0 if pd.isna(median) else median
                    frame[column] = series.fillna(fill_value)
                else:
                    mode = series.mode(dropna=True)
                    fill_value = mode.iloc[0] if not mode.empty else 'Unknown'
                    frame[column] = series.fillna(fill_value)
                filled_columns.append(str(column))
            if filled_columns:
                logs.append({
                    'action': 'Handled Missing Values',
                    'detail': f'Filled missing values in {len(filled_columns)} column(s).',
                    'timestamp': datetime.utcnow().isoformat(),
                })

        if request.convert_dates:
            converted_columns: list[str] = []
            for column in frame.columns:
                try:
                    parsed_series = try_parse_datetime_series(frame[column])
                except Exception:
                    logger.warning('Skipping date conversion for column %s during cleaning.', column, exc_info=True)
                    continue
                if parsed_series is None or parsed_series.notna().sum() == 0:
                    continue
                frame[column] = parsed_series.dt.strftime('%Y-%m-%d').where(parsed_series.notna(), None)
                converted_columns.append(str(column))
            if converted_columns:
                logs.append({
                    'action': 'Converted Date Columns',
                    'detail': f'Converted {len(converted_columns)} date-like column(s).',
                    'timestamp': datetime.utcnow().isoformat(),
                })

        dtype_payload: dict[str, Any] | None = None
        if request.infer_dtypes:
            frame, dtype_payload = build_dtype_inference_payload(frame)
            accepted_count = int(sum(1 for item in dtype_payload['audit'] if bool(item.get('accepted'))))
            logs.append({
                'action': 'Inferred Data Types',
                'detail': f'Applied universal dtype inference across {len(frame.columns)} column(s); {accepted_count} column decision(s) accepted.',
                'timestamp': datetime.utcnow().isoformat(),
            })

        updated_dataset_path = persist_inferred_dataset_frame(request.dataset_id, dataset_entry, frame)
        duplicate_rows = int(DATASET_CACHE[request.dataset_id].get('duplicate_count') or 0)
        memory_size = updated_dataset_path.stat().st_size
        preview_frame = frame.head(DATASET_PREVIEW_ROW_LIMIT)
        return {
            'datasetId': request.dataset_id,
            'data': safe_serialize(preview_frame.to_dict(orient='records')),
            'columns': build_column_info_from_frame(frame),
            'rowCount': int(len(frame)),
            'originalRowCount': original_row_count,
            'loadedRowCount': int(len(preview_frame)),
            'previewLoaded': len(frame) > len(preview_frame),
            'duplicates': duplicate_rows,
            'memoryUsage': f'{memory_size / (1024 * 1024):.2f} MB',
            'logs': logs,
            'dtypeInference': dtype_payload,
        }

    if not dataset_entry.get('parquet_path'):
        raise HTTPException(status_code=400, detail='Cached dataset storage is missing. Please upload the file again.')

    frame = read_cached_parquet(dataset_entry, low_memory=True)
    original_row_count = int(frame.height)
    logs: list[dict[str, Any]] = []

    if request.standardize_names:
        renamed_columns = make_unique_column_names(list(frame.columns))
        if renamed_columns != list(frame.columns):
            frame.columns = renamed_columns
            logs.append({
                'action': 'Standardized Column Names',
                'detail': 'Normalized column names for easier analysis and modeling.',
                'timestamp': datetime.utcnow().isoformat(),
            })

    if request.remove_duplicates:
        before = frame.height
        frame = frame.unique(maintain_order=True)
        removed = before - frame.height
        if removed > 0:
            logs.append({
                'action': 'Removed Duplicates',
                'detail': f'Removed {removed} duplicate rows.',
                'timestamp': datetime.utcnow().isoformat(),
            })

    if request.handle_missing:
        fill_expressions: list[pl.Expr] = []
        filled_columns: list[str] = []
        schema = frame.schema
        for column, dtype in schema.items():
            series = frame.get_column(column)
            if series.null_count() == 0:
                continue
            fill_value: Any | None
            if dtype.is_numeric():
                median_value = series.median()
                fill_value = 0 if median_value is None else median_value
            else:
                mode_frame = frame.select(pl.col(column).drop_nulls().mode().first().alias('mode'))
                fill_value = mode_frame.item(0, 0) if mode_frame.height > 0 else None
                if fill_value is None:
                    if dtype in pl.TEMPORAL_DTYPES:
                        non_null_values = series.drop_nulls()
                        fill_value = non_null_values[0] if non_null_values.len() > 0 else None
                    elif dtype.is_(pl.String):
                        fill_value = 'Unknown'
            if fill_value is None:
                continue
            fill_expressions.append(pl.col(column).fill_null(fill_value).alias(column))
            filled_columns.append(str(column))
        if fill_expressions:
            frame = frame.with_columns(fill_expressions)
            logs.append({
                'action': 'Handled Missing Values',
                'detail': f'Filled missing values in {len(filled_columns)} column(s).',
                'timestamp': datetime.utcnow().isoformat(),
            })

    if request.convert_dates:
        converted_columns: list[str] = []
        date_expressions: list[pl.Expr] = []
        for column, dtype in frame.schema.items():
            try:
                parsed_expr = build_polars_datetime_expr(column, dtype)
                sample = frame.select(parsed_expr.alias('__parsed_date')).drop_nulls().head(50).to_series()
                if sample.len() == 0:
                    continue
                success_ratio = float(sample.len() / min(50, max(1, frame.select(pl.col(column).drop_nulls().len()).item()))) if frame.height > 0 else 0.0
                if dtype not in pl.TEMPORAL_DTYPES and success_ratio < 0.6:
                    continue
            except Exception:
                logger.warning('Skipping date conversion for column %s during cleaning.', column, exc_info=True)
                continue
            date_expressions.append(
                pl.when(pl.col(column).is_null())
                .then(None)
                .otherwise(parsed_expr.dt.strftime('%Y-%m-%d'))
                .alias(column)
            )
            converted_columns.append(str(column))
        if date_expressions:
            frame = frame.with_columns(date_expressions)
            logs.append({
                'action': 'Converted Date Columns',
                'detail': f'Converted {len(converted_columns)} date-like column(s).',
                'timestamp': datetime.utcnow().isoformat(),
            })

    if request.infer_dtypes:
        pandas_frame = normalize_dataframe(frame.to_pandas(use_pyarrow_extension_array=False))
        inferred_frame, dtype_payload = build_dtype_inference_payload(pandas_frame)
        updated_dataset_path = persist_inferred_dataset_frame(request.dataset_id, dataset_entry, inferred_frame)
        duplicate_rows = int(DATASET_CACHE[request.dataset_id].get('duplicate_count') or 0)
        memory_size = updated_dataset_path.stat().st_size
        accepted_count = int(sum(1 for item in dtype_payload['audit'] if bool(item.get('accepted'))))
        logs.append({
            'action': 'Inferred Data Types',
            'detail': f'Applied universal dtype inference across {len(inferred_frame.columns)} column(s); {accepted_count} column decision(s) accepted.',
            'timestamp': datetime.utcnow().isoformat(),
        })
        preview_frame = inferred_frame.head(DATASET_PREVIEW_ROW_LIMIT)
        return {
            'datasetId': request.dataset_id,
            'data': safe_serialize(preview_frame.to_dict(orient='records')),
            'columns': build_column_info_from_frame(inferred_frame),
            'rowCount': int(len(inferred_frame)),
            'originalRowCount': original_row_count,
            'loadedRowCount': int(len(preview_frame)),
            'previewLoaded': len(inferred_frame) > len(preview_frame),
            'duplicates': duplicate_rows,
            'memoryUsage': f'{memory_size / (1024 * 1024):.2f} MB',
            'logs': logs,
            'dtypeInference': dtype_payload,
        }

    parquet_buffer = io.BytesIO()
    frame.write_parquet(parquet_buffer)
    updated_dataset_path = write_dataset_file(request.dataset_id, parquet_buffer.getvalue())
    duplicate_rows = int(max(0, frame.height - frame.unique().height))
    DATASET_CACHE[request.dataset_id] = {
        'parquet_path': str(updated_dataset_path),
        'filename': dataset_entry['filename'],
        'row_count': int(frame.height),
        'column_count': int(len(frame.columns)),
        'columns': list(frame.columns),
        'duplicate_count': duplicate_rows,
    }
    memory_size = updated_dataset_path.stat().st_size

    preview_frame = frame.head(DATASET_PREVIEW_ROW_LIMIT)
    return {
        'datasetId': request.dataset_id,
        'data': safe_serialize(preview_frame.to_dicts()),
        'columns': build_column_info_from_polars_frame(frame),
        'rowCount': int(frame.height),
        'originalRowCount': original_row_count,
        'loadedRowCount': int(preview_frame.height),
        'previewLoaded': frame.height > preview_frame.height,
        'duplicates': duplicate_rows,
        'memoryUsage': f'{memory_size / (1024 * 1024):.2f} MB',
        'logs': logs,
    }


def generate_cleaning_justification(request: CleaningJustificationRequest) -> str:
    dataset_label = request.fileName or 'uploaded dataset'
    loaded_rows = request.loadedRowCount or request.totalRows
    scope_line = (
        f"The dataset was uploaded as '{dataset_label}'. A preview of {loaded_rows} rows is currently rendered while cleaning decisions are being applied to the full {request.totalRows}-row dataset."
        if request.previewLoaded and request.totalRows > loaded_rows
        else f"The dataset was uploaded as '{dataset_label}' with {request.totalRows} rows available for direct cleaning review."
    )
    summary_lines = [
        scope_line,
        f"It contains {request.totalColumns} columns, so the cleaning workflow focuses on changes that improve reliability without assuming any specific business domain.",
        'The following cleaning steps were applied to improve data quality:',
    ]
    for log in request.logs:
        summary_lines.append(f"- {log.action}: {log.detail}")
    summary_lines.append('These changes make the uploaded dataset more consistent for EDA, forecasting, machine learning training, and downstream prediction without hard-coding dataset-specific rules.')
    return "\n".join(summary_lines)


def build_report_pdf(payload: ReportPayload) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=34, rightMargin=34, topMargin=30, bottomMargin=30)
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'ReportTitle',
        parent=styles['Title'],
        fontName='Helvetica-Bold',
        fontSize=24,
        leading=28,
        textColor=colors.white,
        spaceAfter=6,
    )
    subtitle_style = ParagraphStyle(
        'ReportSubtitle',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#d1fae5'),
    )
    cover_tag_style = ParagraphStyle(
        'CoverTag',
        parent=styles['BodyText'],
        fontName='Helvetica-Bold',
        fontSize=8,
        leading=10,
        textColor=colors.HexColor('#ccfbf1'),
        spaceAfter=4,
    )
    cover_meta_label_style = ParagraphStyle(
        'CoverMetaLabel',
        parent=styles['BodyText'],
        fontName='Helvetica-Bold',
        fontSize=7,
        leading=9,
        textColor=colors.HexColor('#99f6e4'),
    )
    cover_meta_value_style = ParagraphStyle(
        'CoverMetaValue',
        parent=styles['BodyText'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=13,
        textColor=colors.white,
    )
    section_style = ParagraphStyle(
        'SectionTitle',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=13,
        leading=16,
        textColor=colors.white,
        spaceAfter=0,
    )
    body_style = ParagraphStyle(
        'ReportBody',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=9,
        leading=13,
        textColor=colors.HexColor('#1f2937'),
        spaceAfter=4,
    )
    small_style = ParagraphStyle(
        'ReportSmall',
        parent=body_style,
        fontSize=8,
        leading=11,
        textColor=colors.HexColor('#475569'),
    )
    card_label_style = ParagraphStyle(
        'CardLabel',
        parent=styles['BodyText'],
        fontName='Helvetica-Bold',
        fontSize=8,
        leading=10,
        textColor=colors.HexColor('#0f766e'),
    )
    table_header_style = ParagraphStyle(
        'ReportTableHeader',
        parent=styles['BodyText'],
        fontName='Helvetica-Bold',
        fontSize=8,
        leading=10,
        textColor=colors.white,
    )
    card_value_style = ParagraphStyle(
        'CardValue',
        parent=styles['BodyText'],
        fontName='Helvetica-Bold',
        fontSize=14,
        leading=17,
        textColor=colors.HexColor('#111827'),
    )

    elements: list[Any] = []

    def as_paragraph(text: Any, style: ParagraphStyle = body_style) -> Paragraph:
        return Paragraph(str(text).replace('\n', '<br/>'), style)

    def add_paragraph(text: Any, style: ParagraphStyle = body_style) -> None:
        elements.append(as_paragraph(text, style))

    def add_section(title: str, blurb: str | None = None) -> None:
        header = Table([[as_paragraph(title, section_style)]], colWidths=[540])
        header.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#0f766e')),
            ('LEFTPADDING', (0, 0), (-1, -1), 10),
            ('RIGHTPADDING', (0, 0), (-1, -1), 10),
            ('TOPPADDING', (0, 0), (-1, -1), 7),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 7),
        ]))
        elements.append(header)
        if blurb:
            info = Table([[as_paragraph(blurb, small_style)]], colWidths=[540])
            info.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f0fdfa')),
                ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor('#99f6e4')),
                ('LEFTPADDING', (0, 0), (-1, -1), 10),
                ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            elements.append(info)
        elements.append(Spacer(1, 8))

    def add_table(rows: list[list[Any]], widths: list[int] | None = None, header_bg: str = '#0f766e') -> None:
        if not rows:
            return
        normalized: list[list[Any]] = []
        for row_index, row in enumerate(rows):
            style = table_header_style if row_index == 0 else body_style
            normalized.append([as_paragraph(cell, style) for cell in row])
        table = Table(normalized, colWidths=widths, repeatRows=1)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(header_bg)),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8fafc')]),
            ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#cbd5e1')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ]))
        elements.append(table)

    def add_stat_cards(cards: list[tuple[str, Any]]) -> None:
        row = []
        widths = []
        for label, value in cards:
            card = Table([[as_paragraph(label, card_label_style)], [as_paragraph(value, card_value_style)]], colWidths=[124])
            card.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f8fafc')),
                ('BOX', (0, 0), (-1, -1), 0.75, colors.HexColor('#cbd5e1')),
                ('LEFTPADDING', (0, 0), (-1, -1), 8),
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            row.append(card)
            widths.append(128)
        wrapper = Table([row], colWidths=widths)
        wrapper.setStyle(TableStyle([('VALIGN', (0, 0), (-1, -1), 'TOP')]))
        elements.append(wrapper)

    def decorate_page(canvas: Any, doc_obj: Any) -> None:
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor('#cbd5e1'))
        canvas.setLineWidth(0.5)
        canvas.line(doc.leftMargin, 20, letter[0] - doc.rightMargin, 20)
        canvas.setFont('Helvetica', 8)
        canvas.setFillColor(colors.HexColor('#64748b'))
        canvas.drawString(doc.leftMargin, 8, f'AI-EDA & ML Workflow Report | {payload.fileName}')
        canvas.drawRightString(letter[0] - doc.rightMargin, 8, f'Page {canvas.getPageNumber()}')
        canvas.restoreState()

    workflow_rows = [
        ['Workflow Area', 'Included Details'],
        ['Upload', f'File {payload.fileName}, {payload.totalRows} rows, {len(payload.columns)} columns, memory {payload.memoryUsage}'],
        ['Understanding', 'Dataset quality, preview context, and upload profiling details'],
        ['Cleaning', f'{len(payload.cleaningLogs)} logged operations and cleaned row count {payload.cleanedRowCount}'],
        ['EDA', f'{len(payload.edaStats.numericColumns)} numeric columns, {len(payload.edaStats.categoricalColumns)} categorical columns, schema, statistics, and correlations'],
        ['Sales Forecast', 'Time-series training split, backtest metrics, backtest samples, and future forecast' if payload.salesForecastResult else 'No sales forecast run captured'],
        ['ML', f"Model {payload.selectedModel or 'Not trained'}, problem type {payload.problemType}, metrics and feature importance"],
        ['Prediction', 'Latest prediction, model context, probabilities, and recent prediction history' if payload.predictionResult is not None else 'No prediction captured'],
    ]

    generated_on = datetime.now().strftime('%d %b %Y, %I:%M %p')
    workflow_coverage = f"{7 if payload.salesForecastResult is not None else 6}/7 sections"
    report_status = 'Complete workflow captured' if payload.predictionResult is not None else 'Workflow summary generated'

    cover_meta = Table([
        [
            Table([
                [as_paragraph('DATASET', cover_meta_label_style), as_paragraph('GENERATED ON', cover_meta_label_style), as_paragraph('WORKFLOW COVERAGE', cover_meta_label_style)],
                [as_paragraph(payload.fileName, cover_meta_value_style), as_paragraph(generated_on, cover_meta_value_style), as_paragraph(workflow_coverage, cover_meta_value_style)],
            ], colWidths=[170, 150, 140]),
        ],
        [
            Table([
                [as_paragraph('REPORT STATUS', cover_meta_label_style), as_paragraph('CLEANED ROWS', cover_meta_label_style), as_paragraph('COLUMNS PROFILED', cover_meta_label_style)],
                [as_paragraph(report_status, cover_meta_value_style), as_paragraph(f'{payload.cleanedRowCount:,}', cover_meta_value_style), as_paragraph(str(len(payload.columns)), cover_meta_value_style)],
            ], colWidths=[170, 150, 140]),
        ],
    ], colWidths=[500])
    cover_meta.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#134e4a')),
        ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor('#5eead4')),
        ('INNERGRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#115e59')),
        ('LEFTPADDING', (0, 0), (-1, -1), 10),
        ('RIGHTPADDING', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))

    cover = Table([[
        as_paragraph('AI-Assisted EDA & ML Platform', cover_tag_style),
        as_paragraph('Workflow Report', title_style),
        as_paragraph(
            'A complete view of the dataset journey from upload and cleaning to EDA, forecasting, machine learning, and final prediction outputs.',
            subtitle_style,
        ),
        cover_meta,
    ]], colWidths=[540])
    cover.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#0f766e')),
        ('BOX', (0, 0), (-1, -1), 0.8, colors.HexColor('#14b8a6')),
        ('LEFTPADDING', (0, 0), (-1, -1), 18),
        ('RIGHTPADDING', (0, 0), (-1, -1), 18),
        ('TOPPADDING', (0, 0), (-1, -1), 18),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 18),
    ]))
    elements.append(cover)
    elements.append(Spacer(1, 12))

    add_section('Workflow Coverage', 'This overview mirrors the product tabs so the report reads like the same journey your team followed in the app.')
    add_table(workflow_rows, [125, 415])
    elements.append(Spacer(1, 10))

    add_section('Data Upload', 'Initial dataset intake, scale, and storage footprint at the moment the workflow began.')
    add_stat_cards([
        ('Rows', f'{payload.totalRows:,}'),
        ('Columns', len(payload.columns)),
        ('Duplicates', payload.duplicates),
        ('Memory Usage', payload.memoryUsage),
    ])
    elements.append(Spacer(1, 8))
    add_table([
        ['File Name', 'Cleaned Rows', 'Cleaning Done'],
        [payload.fileName, str(payload.cleanedRowCount), 'Yes' if payload.cleaningDone else 'No'],
    ], [240, 140, 160], header_bg='#115e59')
    elements.append(Spacer(1, 10))

    add_section('Data Understanding', 'This step captures dataset identity, quality checks, preview context, and the initial profiling needed before cleaning and deeper EDA.')
    column_rows = [['Column', 'Type', 'Role', 'Non-null', 'Nulls', 'Unique']]
    for column in payload.columns[:18]:
        column_rows.append([column.name, column.dtype, column.role, str(column.nonNull), str(column.nullCount), str(column.uniqueCount)])
    add_table(column_rows, [165, 70, 80, 60, 50, 50], header_bg='#134e4a')
    if len(payload.columns) > 18:
        add_paragraph(f'Showing the first 18 columns out of {len(payload.columns)} total columns.', small_style)
    elements.append(Spacer(1, 10))

    add_section('Data Cleaning', 'This section records the applied transformations so the report preserves not just the outcome, but also the reasoning trail.')
    add_paragraph(f"Cleaning completed: {'Yes' if payload.cleaningDone else 'No'}. Cleaned row count: {payload.cleanedRowCount}.")
    if payload.cleaningLogs:
        cleaning_rows = [['Action', 'Detail', 'Timestamp']]
        for log in payload.cleaningLogs[:20]:
            cleaning_rows.append([log.action, log.detail, log.timestamp])
        add_table(cleaning_rows, [120, 300, 120], header_bg='#0f766e')
    else:
        add_paragraph('No cleaning steps were recorded for this run.', small_style)
    elements.append(Spacer(1, 10))

    add_section('Exploratory Data Analysis', 'EDA summarizes the dataset schema, descriptive statistics, and strongest numeric relationships for downstream decisions.')
    add_stat_cards([
        ('Numeric Columns', len(payload.edaStats.numericColumns)),
        ('Categorical Columns', len(payload.edaStats.categoricalColumns)),
        ('Correlations', len(payload.edaStats.correlations)),
        ('AI Insight', 'Available' if payload.aiInsights else 'Not captured'),
    ])
    elements.append(Spacer(1, 8))
    if payload.edaStats.numericColumns:
        numeric_rows = [['Numeric Column', 'Mean', 'Std', 'Min', 'Max']]
        for column_name in payload.edaStats.numericColumns[:10]:
            stats = payload.edaStats.stats.get(column_name, {})
            numeric_rows.append([column_name, stats.get('mean', 'N/A'), stats.get('std', 'N/A'), stats.get('min', 'N/A'), stats.get('max', 'N/A')])
        add_table(numeric_rows, [180, 85, 85, 85, 85], header_bg='#115e59')
    if payload.edaStats.correlations:
        elements.append(Spacer(1, 6))
        corr_rows = [['Pair', 'Correlation']]
        for item in payload.edaStats.correlations[:8]:
            corr_rows.append([str(item.get('pair', 'N/A')), str(item.get('correlation', 'N/A'))])
        add_table(corr_rows, [430, 90], header_bg='#115e59')
    if payload.aiInsights:
        elements.append(Spacer(1, 6))
        insight_box = Table([[as_paragraph(payload.aiInsights, body_style)]], colWidths=[540])
        insight_box.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#ecfeff')),
            ('BOX', (0, 0), (-1, -1), 0.6, colors.HexColor('#67e8f9')),
            ('LEFTPADDING', (0, 0), (-1, -1), 10),
            ('RIGHTPADDING', (0, 0), (-1, -1), 10),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ]))
        elements.append(insight_box)
    elements.append(Spacer(1, 10))

    add_section('Sales Forecast', 'Sales forecasting fits best after EDA because it depends on cleaned, time-aware historical patterns rather than the generic supervised ML pipeline.')
    if payload.salesForecastResult is not None:
        forecast = payload.salesForecastResult
        add_stat_cards([
            ('Forecast Model', forecast.training_summary.model_name),
            ('Split', f"{forecast.training_summary.train_percentage}% / {forecast.training_summary.test_percentage}%"),
            ('Forecast Horizon', forecast.training_summary.forecast_periods),
            ('Period Type', forecast.period_label or forecast.frequency or 'Period'),
        ])
        elements.append(Spacer(1, 8))
        add_table([
            ['Date Column', 'Sales Column', 'Train Window', 'Backtest Window'],
            [forecast.date_column, forecast.target_column, f"{forecast.training_summary.train_start} to {forecast.training_summary.train_end}", f"{forecast.training_summary.test_start} to {forecast.training_summary.test_end}"],
        ], [110, 110, 160, 160], header_bg='#115e59')
        elements.append(Spacer(1, 6))
        add_stat_cards([
            ('MAE', forecast.metrics.mae),
            ('RMSE', forecast.metrics.rmse),
            ('MAPE', f"{forecast.metrics.mape}%"),
            ('Observed Periods', forecast.training_summary.total_periods),
        ])
        elements.append(Spacer(1, 8))
        add_paragraph(forecast.analysis)
        if forecast.test_forecast:
            elements.append(Spacer(1, 6))
            backtest_rows = [['Backtest Period', 'Actual', 'Predicted']]
            for item in forecast.test_forecast[:8]:
                backtest_rows.append([item.period, item.actual if item.actual is not None else 'N/A', item.predicted if item.predicted is not None else 'N/A'])
            add_table(backtest_rows, [190, 160, 160], header_bg='#134e4a')
        if forecast.future_forecast:
            elements.append(Spacer(1, 6))
            future_rows = [['Future Period', 'Forecasted Sales']]
            for item in forecast.future_forecast[:8]:
                future_rows.append([item.period, item.predicted if item.predicted is not None else 'N/A'])
            add_table(future_rows, [270, 270], header_bg='#134e4a')
    else:
        add_paragraph('No sales forecasting run was available for this report.', small_style)
    elements.append(Spacer(1, 10))

    add_section('Machine Learning', 'General machine learning follows forecasting in the workflow because it is a broader predictive branch for supervised models and downstream prediction serving.')
    add_stat_cards([
        ('Target', payload.targetColumn or 'Not selected'),
        ('Problem Type', payload.problemType.title()),
        ('Selected Model', payload.selectedModel or 'Not trained'),
        ('Features', len(payload.selectedFeatures)),
    ])
    elements.append(Spacer(1, 8))
    if payload.selectedFeatures:
        add_paragraph('Selected features: ' + ', '.join(payload.selectedFeatures[:20]))
    if payload.modelMetrics:
        metric_rows = [['Metric', 'Value']]
        for key, value in payload.modelMetrics.items():
            metric_rows.append([key, value])
        add_table(metric_rows, [270, 270], header_bg='#134e4a')
    else:
        add_paragraph('No ML metrics were available.', small_style)
    if payload.featureImportance:
        elements.append(Spacer(1, 6))
        importance_rows = [['Rank', 'Feature', 'Importance']]
        for index, item in enumerate(payload.featureImportance[:12], start=1):
            importance_rows.append([index, item.get('name', 'N/A'), item.get('importance', 'N/A')])
        add_table(importance_rows, [50, 360, 130], header_bg='#134e4a')
    elements.append(Spacer(1, 10))

    add_section('Prediction', 'The report closes with the latest scoring output, supporting model context, and recent prediction history when available.')
    if payload.uploadedModel is not None:
        add_table([
            ['Prediction Model', 'Type', 'Target', 'Problem', 'Trained At'],
            [payload.uploadedModel.name, payload.uploadedModel.type, payload.uploadedModel.target, payload.uploadedModel.problem, payload.uploadedModel.trainedAt],
        ], [130, 90, 120, 70, 130], header_bg='#115e59')
        if payload.uploadedModel.features:
            elements.append(Spacer(1, 6))
            add_paragraph('Prediction model features: ' + ', '.join(payload.uploadedModel.features[:20]))
    if payload.predictionResult is not None:
        elements.append(Spacer(1, 6))
        add_stat_cards([
            ('Latest Prediction', payload.predictionResult),
            ('History Entries', len(payload.predictionHistory)),
            ('Probabilities', 'Available' if payload.predictionProbabilities else 'N/A'),
            ('Prediction Analysis', 'Available' if payload.predictionAnalysis else 'N/A'),
        ])
        elements.append(Spacer(1, 8))
        if payload.predictionAnalysis:
            add_paragraph(payload.predictionAnalysis)
        if payload.predictionProbabilities:
            elements.append(Spacer(1, 6))
            prob_rows = [['Class', 'Probability']]
            for label, probability in list(payload.predictionProbabilities.items())[:10]:
                prob_rows.append([label, f'{round(probability * 100, 2)}%'])
            add_table(prob_rows, [270, 270], header_bg='#134e4a')
        if payload.predictionHistory:
            elements.append(Spacer(1, 6))
            history_rows = [['Timestamp', 'Prediction', 'Confidence']]
            for item in payload.predictionHistory[-8:]:
                confidence = 'N/A' if item.confidence is None else f'{round(item.confidence * 100, 2)}%'
                history_rows.append([item.timestamp, item.prediction, confidence])
            add_table(history_rows, [230, 170, 140], header_bg='#134e4a')
    else:
        add_paragraph('No predictions were generated for this report.', small_style)

    doc.build(elements, onFirstPage=decorate_page, onLaterPages=decorate_page)
    return buffer.getvalue()


def build_line_chart_image(
    title: str,
    history: list[dict[str, Any]],
    test_forecast: list[dict[str, Any]],
    future_forecast: list[dict[str, Any]],
    include_interval: bool = False,
) -> Image:
    fig, ax = plt.subplots(figsize=(7.0, 3.2))
    all_periods = [item['period'] for item in history]
    all_periods.extend(item['period'] for item in future_forecast if item['period'] not in all_periods)
    x_lookup = {period: index for index, period in enumerate(all_periods)}
    history_periods = [x_lookup[item['period']] for item in history]
    history_values = [float(item.get('actual', 0) or 0) for item in history]
    ax.plot(history_periods, history_values, label='Actual', color='#0f766e', linewidth=2)

    if test_forecast:
      ax.plot([x_lookup[item['period']] for item in test_forecast], [float(item.get('predicted', 0) or 0) for item in test_forecast], label='Backtest', color='#f59e0b', linestyle='--', linewidth=2)
    if future_forecast:
      periods = [x_lookup[item['period']] for item in future_forecast]
      values = [float(item.get('predicted', 0) or 0) for item in future_forecast]
      ax.plot(periods, values, label='Forecast', color='#2563eb', linewidth=2)
      if include_interval:
          lowers = [float(item.get('lower', item.get('predicted', 0)) or 0) for item in future_forecast]
          uppers = [float(item.get('upper', item.get('predicted', 0)) or 0) for item in future_forecast]
          ax.fill_between(periods, lowers, uppers, color='#93c5fd', alpha=0.3, label='95% interval')

    ax.set_title(title)
    ax.set_xticks(list(x_lookup.values()))
    ax.set_xticklabels(all_periods, rotation=35, fontsize=8)
    ax.tick_params(axis='y', labelsize=8)
    ax.grid(alpha=0.2)
    ax.legend(fontsize=8)
    fig.tight_layout()
    image_buffer = io.BytesIO()
    fig.savefig(image_buffer, format='png', dpi=160, bbox_inches='tight')
    plt.close(fig)
    image_buffer.seek(0)
    return Image(image_buffer, width=480, height=220)


def build_bar_chart_image(title: str, items: list[dict[str, Any]]) -> Image:
    fig, ax = plt.subplots(figsize=(7.0, 3.0))
    trimmed = items[:10]
    names = [str(item.get('name', 'Feature')) for item in trimmed][::-1]
    values = [float(item.get('importance', 0) or 0) for item in trimmed][::-1]
    ax.barh(names, values, color='#0f766e')
    ax.set_title(title)
    ax.grid(axis='x', alpha=0.2)
    ax.tick_params(axis='y', labelsize=8)
    fig.tight_layout()
    image_buffer = io.BytesIO()
    fig.savefig(image_buffer, format='png', dpi=160, bbox_inches='tight')
    plt.close(fig)
    image_buffer.seek(0)
    return Image(image_buffer, width=480, height=220)


def build_correlation_chart_image(correlations: list[dict[str, Any]]) -> Image | None:
    if not correlations:
        return None
    fig, ax = plt.subplots(figsize=(7.0, 3.0))
    trimmed = correlations[:8][::-1]
    names = [str(item.get('pair', 'Pair')) for item in trimmed]
    values = [float(item.get('correlation', 0) or 0) for item in trimmed]
    colors_list = ['#0f766e' if value >= 0 else '#dc2626' for value in values]
    ax.barh(names, values, color=colors_list)
    ax.set_title('EDA Correlation Heatmap Summary')
    ax.grid(axis='x', alpha=0.2)
    ax.tick_params(axis='y', labelsize=8)
    fig.tight_layout()
    image_buffer = io.BytesIO()
    fig.savefig(image_buffer, format='png', dpi=160, bbox_inches='tight')
    plt.close(fig)
    image_buffer.seek(0)
    return Image(image_buffer, width=480, height=220)


def build_image_from_base64(data_uri: str | None, *, max_width: float = 480, max_height: float = 260) -> Image | None:
    if not data_uri:
        return None

    try:
        encoded = data_uri.split(',', 1)[1] if ',' in data_uri else data_uri
        image_bytes = base64.b64decode(encoded)
        image_buffer = io.BytesIO(image_bytes)
        image = Image(image_buffer)
        image.drawWidth = max_width
        image.drawHeight = max_height
        return image
    except Exception:
        logger.exception('Failed to decode base64 image for EDA PDF report.')
        return None


def _chart_to_base64(fig: plt.Figure) -> str:
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=160, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return base64.b64encode(buf.read()).decode('utf-8')


def _build_line_chart_base64(
    title: str,
    history: list[dict[str, Any]],
    test_forecast: list[dict[str, Any]],
    future_forecast: list[dict[str, Any]],
    include_interval: bool = False,
) -> str:
    fig, ax = plt.subplots(figsize=(8.4, 3.0))
    all_periods = [item['period'] for item in history]
    all_periods.extend(item['period'] for item in future_forecast if item['period'] not in all_periods)
    x_lookup = {period: idx for idx, period in enumerate(all_periods)}
    hp = [x_lookup[item['period']] for item in history]
    hv = [float(item.get('actual', 0) or 0) for item in history]
    ax.plot(hp, hv, label='Actual', color='#0f766e', linewidth=2)
    if test_forecast:
        ax.plot([x_lookup[item['period']] for item in test_forecast],
                [float(item.get('predicted', 0) or 0) for item in test_forecast],
                label='Backtest', color='#f59e0b', linestyle='--', linewidth=2)
    if future_forecast:
        p = [x_lookup[item['period']] for item in future_forecast]
        v = [float(item.get('predicted', 0) or 0) for item in future_forecast]
        ax.plot(p, v, label='Forecast', color='#2563eb', linewidth=2)
        if include_interval:
            ax.fill_between(p,
                            [float(item.get('lower', item.get('predicted', 0)) or 0) for item in future_forecast],
                            [float(item.get('upper', item.get('predicted', 0)) or 0) for item in future_forecast],
                            color='#93c5fd', alpha=0.3, label='95% interval')
    ax.set_title(title)
    ax.set_xticks(list(x_lookup.values()))
    ax.set_xticklabels(all_periods, rotation=35, fontsize=8)
    ax.tick_params(axis='y', labelsize=8)
    ax.grid(alpha=0.2)
    ax.legend(fontsize=8)
    fig.tight_layout()
    return _chart_to_base64(fig)


def _build_bar_chart_base64(title: str, items: list[dict[str, Any]]) -> str:
    fig, ax = plt.subplots(figsize=(8.4, 3.0))
    trimmed = items[:10]
    names = [str(item.get('name', 'Feature')) for item in trimmed][::-1]
    values = [float(item.get('importance', 0) or 0) for item in trimmed][::-1]
    ax.barh(names, values, color='#0f766e')
    ax.set_title(title)
    ax.grid(axis='x', alpha=0.2)
    ax.tick_params(axis='y', labelsize=8)
    fig.tight_layout()
    return _chart_to_base64(fig)


def _build_corr_chart_base64(correlations: list[dict[str, Any]]) -> str | None:
    if not correlations:
        return None
    fig, ax = plt.subplots(figsize=(8.4, 3.0))
    trimmed = correlations[:8][::-1]
    names = [str(item.get('pair', 'Pair')) for item in trimmed]
    values = [float(item.get('correlation', 0) or 0) for item in trimmed]
    colors_list = ['#0f766e' if v >= 0 else '#dc2626' for v in values]
    ax.barh(names, values, color=colors_list)
    ax.set_title('Correlation Heatmap Summary')
    ax.grid(axis='x', alpha=0.2)
    ax.tick_params(axis='y', labelsize=8)
    fig.tight_layout()
    return _chart_to_base64(fig)


def _build_loss_chart_base64(loss_rows: list[dict[str, Any]]) -> str:
    fig, ax = plt.subplots(figsize=(8.4, 3.0))
    periods = [str(row.get('period')) for row in loss_rows]
    ax.plot(periods, [float(row.get('total_loss') or 0) for row in loss_rows],
            color='#dc2626', linewidth=2.5, label='Total Loss')
    for key, color, lbl in [
        ('revenue_loss', '#ef4444', 'Revenue'),
        ('operational_loss', '#f97316', 'Operational'),
        ('inventory_loss', '#f59e0b', 'Inventory'),
        ('discount_loss', '#8b5cf6', 'Discount'),
    ]:
        ax.plot(periods, [float(row.get(key) or 0) for row in loss_rows],
                color=color, linewidth=1.5, label=lbl)
    ax.set_title('Loss Trend by Driver')
    ax.tick_params(axis='x', rotation=35, labelsize=7)
    ax.grid(True, alpha=0.25)
    ax.legend(fontsize=7, ncol=5)
    fig.tight_layout()
    return _chart_to_base64(fig)


def _build_profit_chart_base64(
    profit_scenarios: dict[str, list[dict[str, Any]]],
    profit_rows: list[dict[str, Any]],
) -> str:
    fig, ax = plt.subplots(figsize=(8.4, 3.0))
    for scenario_name, color in [('optimistic', '#10b981'), ('baseline', '#2563eb'), ('pessimistic', '#f43f5e')]:
        rows = profit_scenarios.get(scenario_name, [])
        if rows:
            ax.plot([str(row.get('period')) for row in rows],
                    [float(row.get('net_profit') or 0) for row in rows],
                    label=scenario_name.title(), color=color, linewidth=2)
    ax.axhline(0, color='#64748b', linewidth=1)
    ax.set_title('Net Profit Forecast by Scenario')
    ax.tick_params(axis='x', rotation=35, labelsize=7)
    ax.grid(True, alpha=0.25)
    ax.legend(fontsize=8)
    fig.tight_layout()
    return _chart_to_base64(fig)


def _generate_report_metadata(payload: ReportPayload) -> dict[str, str]:
    return {
        'report_id': str(uuid.uuid4())[:12],
        'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'dataset_name': payload.fileName,
        'analysis_type': 'Exploratory Data Analysis with Forecasting and ML',
        'agent_version': 'IDA v2.1.0',
        'report_format': 'Comprehensive Workflow Report',
    }


TEMP_REPORT_STORE: dict[str, dict[str, Any]] = {}


def store_temp_report(report_id: str, pdf_bytes: bytes, html_bytes: bytes, docx_bytes: bytes,
                      payload: ReportPayload) -> str:
    TEMP_REPORT_STORE[report_id] = {
        'pdf': pdf_bytes,
        'html': html_bytes,
        'docx': docx_bytes,
        'payload': payload.model_dump() if hasattr(payload, 'model_dump') else payload,
        'created_at': datetime.now().isoformat(),
    }
    return report_id


def build_eda_pdf(payload: EdaPdfPayload) -> bytes:
    loaded_row_count = payload.loadedRowCount or payload.totalRows
    preview_mode = payload.previewLoaded and payload.totalRows > loaded_row_count
    advanced = payload.advancedAnalysis or {}

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter), leftMargin=30, rightMargin=30, topMargin=26, bottomMargin=24)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('EDA_Title', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=22, leading=26, textColor=colors.HexColor('#0f172a'), spaceAfter=8)
    heading_style = ParagraphStyle('EDA_Heading', parent=styles['Heading2'], fontName='Helvetica-Bold', fontSize=15, leading=18, textColor=colors.HexColor('#0f172a'), spaceAfter=6)
    body_style = ParagraphStyle('EDA_Body', parent=styles['BodyText'], fontName='Helvetica', fontSize=9.2, leading=13, textColor=colors.HexColor('#334155'))
    small_style = ParagraphStyle('EDA_Small', parent=body_style, fontSize=8.2, leading=11, textColor=colors.HexColor('#64748b'))
    label_style = ParagraphStyle('EDA_Label', parent=body_style, fontName='Helvetica-Bold', fontSize=8.2, leading=10, textColor=colors.HexColor('#0f766e'))
    table_header_style = ParagraphStyle('EDA_TableHeader', parent=body_style, fontName='Helvetica-Bold', fontSize=8.2, leading=10, textColor=colors.white)
    value_style = ParagraphStyle('EDA_Value', parent=body_style, fontName='Helvetica-Bold', fontSize=13, leading=16, textColor=colors.HexColor('#0f172a'))
    elements: list[Any] = []
    page_width = landscape(letter)[0]
    content_width = page_width - 60

    def paragraph(text: Any, style: ParagraphStyle = body_style) -> Paragraph:
        return Paragraph(str(text).replace('\n', '<br/>'), style)

    def add_table(rows: list[list[Any]], widths: list[float], header_bg: str = '#0f766e') -> None:
        normalized = []
        for row_index, row in enumerate(rows):
            row_style = table_header_style if row_index == 0 else body_style
            normalized.append([paragraph(cell, row_style) for cell in row])
        table = Table(normalized, colWidths=widths, repeatRows=1)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(header_bg)),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8fbff')]),
            ('GRID', (0, 0), (-1, -1), 0.35, colors.HexColor('#dbe4f0')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ]))
        elements.append(table)

    def add_stat_cards(cards: list[tuple[str, Any]]) -> None:
        row = []
        widths = []
        for label, value in cards:
            card = Table([[paragraph(label, label_style)], [paragraph(value, value_style)]], colWidths=[content_width / max(1, len(cards)) - 8])
            card.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f8fbff')),
                ('BOX', (0, 0), (-1, -1), 0.65, colors.HexColor('#d6e3f1')),
                ('LEFTPADDING', (0, 0), (-1, -1), 10),
                ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ]))
            row.append(card)
            widths.append(content_width / max(1, len(cards)))
        wrapper = Table([row], colWidths=widths)
        wrapper.setStyle(TableStyle([('VALIGN', (0, 0), (-1, -1), 'TOP')]))
        elements.append(wrapper)

    def add_section(title: str, blurb: str) -> None:
        elements.append(paragraph(title, heading_style))
        elements.append(paragraph(blurb, body_style))
        elements.append(Spacer(1, 8))

    def add_chart_section(title: str, description: str, chart_items: list[tuple[str, str | None]], *, subtitle: str | None = None) -> None:
        add_section(title, description)
        if subtitle:
            elements.append(paragraph(subtitle, small_style))
            elements.append(Spacer(1, 6))
        rendered_any = False
        for chart_title, chart_base64 in chart_items[:4]:
            image = build_image_from_base64(chart_base64)
            elements.append(paragraph(chart_title, label_style))
            if image is not None:
                elements.append(image)
                rendered_any = True
            else:
                elements.append(paragraph('No chart available for this item.', small_style))
            elements.append(Spacer(1, 8))
        if not rendered_any and not chart_items:
            elements.append(paragraph('No chart outputs were available for this section.', small_style))

    def decorate_page(canvas: Any, doc_obj: Any) -> None:
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor('#dbe4f0'))
        canvas.line(doc.leftMargin, 18, page_width - doc.rightMargin, 18)
        canvas.setFont('Helvetica', 8)
        canvas.setFillColor(colors.HexColor('#64748b'))
        canvas.drawString(doc.leftMargin, 8, f'EDA PDF | {payload.fileName}')
        canvas.drawRightString(page_width - doc.rightMargin, 8, f'Page {canvas.getPageNumber()}')
        canvas.restoreState()

    elements.append(paragraph('Exploratory Data Analysis PDF', title_style))
    elements.append(paragraph(
        f'This export captures the EDA tab functionality, working flow, descriptive statistics, relationship analysis, and advanced analytical features for {payload.fileName}.',
        body_style,
    ))
    elements.append(Spacer(1, 10))
    add_stat_cards([
        ('Total Rows', f'{payload.totalRows:,}'),
        ('Rows In Workspace', f'{loaded_row_count:,}'),
        ('Columns', len(payload.columns)),
        ('Numeric Fields', len(payload.edaStats.numericColumns)),
        ('Categorical Fields', len(payload.edaStats.categoricalColumns)),
    ])
    elements.append(Spacer(1, 8))
    elements.append(paragraph(
        f'EDA working mode: {"Preview-backed browser analysis with cached backend dataset support." if preview_mode else "Direct workspace analysis across the full loaded dataset."}',
        small_style,
    ))
    elements.append(Spacer(1, 12))

    add_section('EDA Functional Coverage', 'This PDF mirrors the EDA workflow itself: schema review, numeric profiling, correlation discovery, advanced charts, and automated statistical recommendations.')
    add_table([
        ['Feature Area', 'What The EDA Workflow Does'],
        ['Dataset Schema', 'Profiles column type, completeness, uniqueness, and inferred role for each field.'],
        ['Statistical Summary', 'Computes count, mean, spread, quartiles, and extrema for numeric columns.'],
        ['Relationships', 'Highlights the strongest positive and negative numeric correlations.'],
        ['Correlation Heatmap', 'Shows matrix-style correlation strength across the leading numeric fields.'],
        ['Advanced Modules', 'Extends the base EDA with missingness, distributions, categorical analysis, interactions, and automated insights.'],
    ], [content_width * 0.24, content_width * 0.72], header_bg='#115e59')
    elements.append(PageBreak())

    add_section('Dataset Schema', 'The EDA tab begins by establishing the structure and quality envelope of the active dataset.')
    schema_rows = [['Column', 'Type', 'Non-Null', 'Missing', 'Unique', 'Role']]
    for column in payload.columns[:24]:
        schema_rows.append([column.name, column.dtype, column.nonNull, column.nullCount, column.uniqueCount, column.role])
    add_table(schema_rows, [content_width * 0.28, content_width * 0.12, content_width * 0.12, content_width * 0.12, content_width * 0.12, content_width * 0.14], header_bg='#115e59')
    if len(payload.columns) > 24:
        elements.append(Spacer(1, 6))
        elements.append(paragraph(f'Showing the first 24 columns out of {len(payload.columns)} total profiled columns.', small_style))
    elements.append(Spacer(1, 10))

    add_section('Statistical Summary', 'Numeric fields are summarized to expose central tendency, spread, and range before cleaning or modeling.')
    numeric_rows = [['Field', 'Mean', 'Std', 'Min', 'Median', 'Max']]
    for field_name in payload.edaStats.numericColumns[:12]:
        stats = payload.edaStats.stats.get(field_name, {})
        numeric_rows.append([
            field_name,
            stats.get('mean', 'N/A'),
            stats.get('std', 'N/A'),
            stats.get('min', 'N/A'),
            stats.get('median', 'N/A'),
            stats.get('max', 'N/A'),
        ])
    add_table(numeric_rows if len(numeric_rows) > 1 else [['Field', 'Mean', 'Std', 'Min', 'Median', 'Max'], ['N/A', 'N/A', 'N/A', 'N/A', 'N/A', 'N/A']], [content_width * 0.34, content_width * 0.11, content_width * 0.11, content_width * 0.11, content_width * 0.11, content_width * 0.11], header_bg='#115e59')
    elements.append(PageBreak())

    add_section('Relationships and Correlation Working', 'The EDA tab surfaces the strongest numeric relationships so users can quickly assess signal, redundancy, and interaction behavior.')
    corr_image = build_correlation_chart_image(payload.edaStats.correlations)
    if corr_image is not None:
        elements.append(corr_image)
        elements.append(Spacer(1, 8))
    correlation_rows = [['Pair', 'Correlation']]
    for item in payload.edaStats.correlations[:10]:
        correlation_rows.append([item.get('pair', 'N/A'), item.get('correlation', 'N/A')])
    add_table(correlation_rows if len(correlation_rows) > 1 else [['Pair', 'Correlation'], ['N/A', 'N/A']], [content_width * 0.78, content_width * 0.18], header_bg='#115e59')

    insights = ((advanced.get('insights') or {}).get('insights') if isinstance(advanced.get('insights'), dict) else None) or []
    if insights:
        elements.append(Spacer(1, 10))
        add_section('Automated Insights', 'The advanced EDA layer translates statistical anomalies into plain-language recommendations.')
        insight_rows = [['Insight']]
        for item in insights[:10]:
            insight_rows.append([item])
        add_table(insight_rows, [content_width * 0.96], header_bg='#115e59')
    elements.append(PageBreak())

    missingness = advanced.get('missingness') if isinstance(advanced.get('missingness'), dict) else {}
    distributions = advanced.get('distributions') if isinstance(advanced.get('distributions'), dict) else {}
    categorical = advanced.get('categorical') if isinstance(advanced.get('categorical'), dict) else {}
    interactions = advanced.get('interactions') if isinstance(advanced.get('interactions'), dict) else {}

    add_chart_section(
        'Advanced EDA: Data Quality and Missingness',
        'This section documents how the advanced EDA tab checks missing-value concentration and dataset completeness behavior.',
        [('Missingness Intensity Map', missingness.get('chart_base64'))] if missingness else [],
        subtitle=str(missingness.get('message') or '') if missingness else None,
    )
    add_chart_section(
        'Advanced EDA: Distributions and Outliers',
        'These charts show how the EDA tab evaluates numeric spread, skew, and potential outliers.',
        [(str(item.get('column', 'Distribution')), item.get('chart_base64')) for item in (distributions.get('charts') or []) if isinstance(item, dict)],
        subtitle=str(distributions.get('message') or '') if distributions else None,
    )
    add_chart_section(
        'Advanced EDA: Categorical Features',
        'These plots document top-category behavior and warn about high-cardinality features that may affect ML readiness.',
        [(f"{item.get('column', 'Category')} ({item.get('unique_count', 'N/A')} unique)", item.get('chart_base64')) for item in (categorical.get('charts') or []) if isinstance(item, dict)],
        subtitle=str(categorical.get('message') or '') if categorical else None,
    )
    add_chart_section(
        'Advanced EDA: Key Variable Interactions',
        'These interaction views show the strongest numeric pairings explored by the advanced EDA feature set.',
        [(str(item.get('pair', 'Interaction')), item.get('chart_base64')) for item in (interactions.get('plots') or []) if isinstance(item, dict)],
        subtitle=str(interactions.get('message') or '') if interactions else None,
    )

    doc.build(elements, onFirstPage=decorate_page, onLaterPages=decorate_page)
    return buffer.getvalue()


def build_dynamic_report_pdf(payload: ReportPayload) -> bytes:
    session_id = get_session_id(payload.datasetId, payload.sessionId)
    session_state = ensure_session_state(session_id)
    completed_steps = set(payload.forecastingStepsCompleted)
    if session_state['forecast_steps'].get('ts'):
        completed_steps.add(5)
    if session_state['forecast_steps'].get('ml'):
        completed_steps.add(6)
    if session_state['forecast_steps'].get('loss'):
        completed_steps.add(7)
    if session_state['forecast_steps'].get('profit'):
        completed_steps.add(8)

    ts_result_raw = payload.timeSeriesForecastResult or session_state.get('time_series_result')
    ml_result_raw = payload.mlForecastResult or session_state.get('ml_forecast_result')
    ts_result = ts_result_raw.model_dump() if hasattr(ts_result_raw, 'model_dump') else ts_result_raw
    ml_forecast_result = ml_result_raw.model_dump() if hasattr(ml_result_raw, 'model_dump') else ml_result_raw
    loss_forecast_result = payload.lossForecast or session_state.get('loss_forecast_result') or []
    profit_scenarios = payload.scenarios or session_state.get('profit_scenarios') or {}
    selected_profit_result = payload.profitForecast or profit_scenarios.get(payload.reportConfig.scenario, [])
    loss_segments = payload.lossSegments or session_state.get('loss_segments') or []
    breakeven_period = payload.breakevenPeriod or (session_state.get('breakeven') or {}).get('breakeven_period')

    page_size = landscape(letter)
    page_width, page_height = page_size
    content_width = page_width - 64
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=page_size, leftMargin=32, rightMargin=32, topMargin=28, bottomMargin=24)
    styles = getSampleStyleSheet()
    eyebrow_style = ParagraphStyle('IDA_Eyebrow', parent=styles['BodyText'], fontName='Helvetica-Bold', fontSize=9, leading=11, textColor=colors.HexColor('#a5f3fc'), spaceAfter=4)
    hero_title_style = ParagraphStyle('IDA_HeroTitle', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=25, leading=29, textColor=colors.white, spaceAfter=8)
    hero_subtitle_style = ParagraphStyle('IDA_HeroSubtitle', parent=styles['BodyText'], fontName='Helvetica', fontSize=10.5, leading=15, textColor=colors.HexColor('#e0f2fe'))
    title_style = ParagraphStyle('IDA_Title', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=22, leading=25, textColor=colors.HexColor('#0f172a'), spaceAfter=8)
    heading_style = ParagraphStyle('IDA_Heading', parent=styles['Heading2'], fontName='Helvetica-Bold', fontSize=16, leading=19, textColor=colors.HexColor('#0f172a'), spaceAfter=6)
    body_style = ParagraphStyle('IDA_Body', parent=styles['BodyText'], fontName='Helvetica', fontSize=9.2, leading=13.5, textColor=colors.HexColor('#334155'))
    small_style = ParagraphStyle('IDA_Small', parent=body_style, fontSize=8.1, leading=11, textColor=colors.HexColor('#64748b'))
    label_style = ParagraphStyle('IDA_Label', parent=body_style, fontName='Helvetica-Bold', fontSize=8.2, leading=10, textColor=colors.HexColor('#0369a1'))
    table_header_style = ParagraphStyle('IDA_TableHeader', parent=body_style, fontName='Helvetica-Bold', fontSize=8.2, leading=10, textColor=colors.white)
    value_style = ParagraphStyle('IDA_Value', parent=body_style, fontName='Helvetica-Bold', fontSize=14, leading=17, textColor=colors.HexColor('#0f172a'))
    section_label_style = ParagraphStyle('IDA_SectionLabel', parent=styles['BodyText'], fontName='Helvetica-Bold', fontSize=8.5, leading=10, textColor=colors.HexColor('#0284c7'))
    section_blurb_style = ParagraphStyle('IDA_SectionBlurb', parent=body_style, fontSize=9.4, leading=14, textColor=colors.HexColor('#475569'))
    elements: list[Any] = []

    def as_paragraph(text: Any, style: ParagraphStyle = body_style) -> Paragraph:
        return Paragraph(str(text).replace('\n', '<br/>'), style)

    def add_paragraph(text: Any, style: ParagraphStyle = body_style) -> None:
        elements.append(as_paragraph(text, style))

    def add_table(rows: list[list[Any]], widths: list[int], header_bg: str = '#0f766e') -> None:
        normalized: list[list[Any]] = []
        for row_index, row in enumerate(rows):
            style = table_header_style if row_index == 0 else body_style
            normalized.append([as_paragraph(cell, style) for cell in row])
        table = Table(normalized, colWidths=widths, repeatRows=1)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(header_bg)),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8fbff')]),
            ('GRID', (0, 0), (-1, -1), 0.35, colors.HexColor('#dbe4f0')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 7),
            ('RIGHTPADDING', (0, 0), (-1, -1), 7),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        elements.append(table)

    def add_stat_cards(cards: list[tuple[str, Any]]) -> None:
        row = []
        widths = []
        for label, value in cards:
            card = Table([[as_paragraph(label, label_style)], [as_paragraph(value, value_style)]], colWidths=[content_width / max(1, len(cards)) - 8])
            card.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f8fbff')),
                ('BOX', (0, 0), (-1, -1), 0.65, colors.HexColor('#d6e3f1')),
                ('LEFTPADDING', (0, 0), (-1, -1), 10),
                ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ]))
            row.append(card)
            widths.append(content_width / max(1, len(cards)))
        wrapper = Table([row], colWidths=widths)
        wrapper.setStyle(TableStyle([('VALIGN', (0, 0), (-1, -1), 'TOP')]))
        elements.append(wrapper)

    def add_section(title: str, blurb: str) -> None:
        section_card = Table([[
            Paragraph('WORKFLOW SECTION', section_label_style),
            Paragraph(title, heading_style),
            Paragraph(blurb, section_blurb_style),
        ]], colWidths=[content_width])
        section_card.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#ffffff')),
            ('BOX', (0, 0), (-1, -1), 0.65, colors.HexColor('#dbe4f0')),
            ('LEFTPADDING', (0, 0), (-1, -1), 14),
            ('RIGHTPADDING', (0, 0), (-1, -1), 14),
            ('TOPPADDING', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
        ]))
        elements.append(section_card)
        elements.append(Spacer(1, 8))

    def add_callout(title: str, text: str, tone: str = '#eff6ff', border: str = '#93c5fd') -> None:
        callout = Table([[
            Paragraph(f'<b>{title}</b><br/>{text}', body_style)
        ]], colWidths=[content_width])
        callout.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(tone)),
            ('BOX', (0, 0), (-1, -1), 0.7, colors.HexColor(border)),
            ('LEFTPADDING', (0, 0), (-1, -1), 12),
            ('RIGHTPADDING', (0, 0), (-1, -1), 12),
            ('TOPPADDING', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 9),
        ]))
        elements.append(callout)

    def role_count(role_name: str) -> int:
        return sum(1 for column in payload.columns if str(column.role).lower() == role_name)

    def metric_text(value: Any) -> str:
        if value is None:
            return 'N/A'
        if isinstance(value, float):
            return f'{value:,.3f}'
        return str(value)

    loaded_row_count = payload.loadedRowCount or payload.totalRows
    preview_mode = payload.previewLoaded and payload.totalRows > loaded_row_count
    workspace_scope = (
        f'{loaded_row_count:,} preview rows were rendered in the browser while the full {payload.totalRows:,}-row dataset remained cached on the backend.'
        if preview_mode
        else f'The full {payload.totalRows:,}-row dataset was loaded directly into the active workspace.'
    )

    def latest_prediction_feature_summary() -> list[list[Any]] | None:
        if not payload.predictionHistory:
            return None
        latest = payload.predictionHistory[-1]
        features = latest.features or {}
        if not features:
            return None
        rows = [['Feature', 'Latest Scored Value']]
        for key, value in list(features.items())[:12]:
            rows.append([key, value])
        return rows

    def decorate_page(canvas: Any, doc_obj: Any) -> None:
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor('#dbe4f0'))
        canvas.line(doc.leftMargin, 20, page_width - doc.rightMargin, 20)
        canvas.setFont('Helvetica', 8)
        canvas.setFillColor(colors.HexColor('#64748b'))
        canvas.drawString(doc.leftMargin, 10, f'Intelligent Data Assistant | {payload.fileName}')
        canvas.drawRightString(page_width - doc.rightMargin, 10, f'Page {canvas.getPageNumber()}')
        canvas.restoreState()

    workflow_rows = [
        ['Workflow Area', 'Status', 'Coverage'],
        ['Upload', 'Completed' if payload.totalRows > 0 else 'Pending', f'{payload.fileName} with {payload.totalRows:,} rows, {len(payload.columns)} columns, and {"preview-backed caching" if preview_mode else "full workspace loading"}'],
        ['Understanding', 'Completed' if payload.columns else 'Pending', f'Role inference, null counts, unique counts, and schema profiling across {len(payload.columns)} columns'],
        ['EDA', 'Completed' if payload.columns else 'Pending', f'{len(payload.edaStats.numericColumns)} numeric and {len(payload.edaStats.categoricalColumns)} categorical columns summarized with {len(payload.edaStats.correlations)} sampled correlation signals'],
        ['Cleaning', 'Completed' if payload.cleaningDone else 'Pending', f'{len(payload.cleaningLogs)} logged operations and {payload.cleanedRowCount:,} cleaned rows retained'],
        ['Time Series Forecast', 'Completed' if ts_result else 'Skipped', 'Time-driven forecasting, backtest metrics, horizon outputs, and interval-aware charting'],
        ['Machine Learning Forecast', 'Completed' if ml_forecast_result else 'Skipped', 'Feature-engineered forecasting, SHAP importance, generated features, and projected horizon'],
        ['Loss Forecast', 'Completed' if loss_forecast_result else 'Skipped', 'Revenue, operational, inventory, and discount-loss projections with risk scoring'],
        ['Profit Forecast', 'Completed' if selected_profit_result else 'Skipped', 'Scenario-based P&L projection, margins, net profit, and break-even analysis'],
        ['ML Assistant', 'Completed' if payload.modelMetrics else 'Pending', f'Model selection, target setup, {len(payload.selectedFeatures)} features, and training metrics'],
        ['Prediction', 'Completed' if payload.predictionResult is not None else 'Pending', f'{len(payload.predictionHistory)} prediction history entries and latest scoring output'],
    ]
    workflow_status = f"{sum(1 for row in workflow_rows[1:] if row[1] == 'Completed')}/10 workflow areas completed"
    forecast_status = ', '.join(name for name, present in [('TS', bool(ts_result)), ('ML', bool(ml_forecast_result)), ('Loss', bool(loss_forecast_result)), ('Profit', bool(selected_profit_result))] if present) or 'None'

    generated_on = datetime.now().strftime('%d %b %Y, %I:%M %p')
    cover_card = Table([[
        Paragraph('INTELLIGENT DATA ASSISTANT', eyebrow_style),
        Paragraph('Executive Workflow Report', hero_title_style),
        Paragraph(
            'A presentation-ready summary of the end-to-end analytics journey, designed for stakeholder review, project handoff, and decision-making.',
            hero_subtitle_style,
        ),
        Spacer(1, 6),
        Table([
            [Paragraph('Dataset', label_style), Paragraph('Generated', label_style), Paragraph('Workflow Status', label_style), Paragraph('Forecast Paths', label_style)],
            [Paragraph(payload.fileName, value_style), Paragraph(generated_on, value_style), Paragraph(workflow_status, value_style), Paragraph(forecast_status, value_style)],
        ], colWidths=[content_width * 0.28, content_width * 0.22, content_width * 0.25, content_width * 0.17]),
    ]], colWidths=[content_width])
    cover_card.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#0f766e')),
        ('BOX', (0, 0), (-1, -1), 0.9, colors.HexColor('#14b8a6')),
        ('LEFTPADDING', (0, 0), (-1, -1), 18),
        ('RIGHTPADDING', (0, 0), (-1, -1), 18),
        ('TOPPADDING', (0, 0), (-1, -1), 18),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 18),
    ]))
    elements.append(cover_card)
    elements.append(Spacer(1, 12))
    add_stat_cards([
        ('Rows', f'{payload.totalRows:,}'),
        ('Columns', len(payload.columns)),
        ('Cleaned Rows', f'{payload.cleanedRowCount:,}'),
        ('Workflow', workflow_status),
    ])
    elements.append(Spacer(1, 8))
    add_callout(
        'Executive Summary',
        (
            f'The dataset entered the platform as {payload.fileName}. '
            f'{len(payload.edaStats.numericColumns)} numeric fields and {len(payload.edaStats.categoricalColumns)} categorical fields were profiled during EDA, '
            f'{len(payload.cleaningLogs)} cleaning actions were recorded, '
            f'and the latest workflow outcome is {payload.predictionResult if payload.predictionResult is not None else "still pending prediction"}.'
        ),
        tone='#ecfeff',
        border='#67e8f9',
    )
    elements.append(Spacer(1, 8))
    add_table([
        ['Metric', 'Value'],
        ['Dataset File', payload.fileName],
        ['Workspace Scope', workspace_scope],
        ['Duplicates Detected', payload.duplicates],
        ['Estimated Memory', payload.memoryUsage],
        ['Forecast Paths Run', forecast_status],
        ['Prediction Available', 'Yes' if payload.predictionResult is not None else 'No'],
    ], [220, 280], header_bg='#115e59')
    elements.append(PageBreak())

    add_section('Workflow Coverage Map', 'The report follows the same professional workflow sequence used in the application so the exported file reads like a presentation replay of the full analysis.')
    add_table(workflow_rows, [145, 90, content_width - 235], header_bg='#0f766e')
    elements.append(Spacer(1, 10))

    add_section('Data Upload', 'The upload stage establishes dataset identity, scale, and storage footprint before any cleaning or modeling decisions are made.')
    add_stat_cards([
        ('Dataset ID', payload.datasetId or 'Session only'),
        ('Rows Loaded', f'{payload.totalRows:,}'),
        ('Browser Rows', f'{loaded_row_count:,}'),
        ('Columns Found', len(payload.columns)),
        ('Workspace Mode', 'Preview + cached backend' if preview_mode else 'Full in-browser dataset'),
    ])
    elements.append(Spacer(1, 8))
    add_paragraph(
        f'The uploaded dataset entered the application as {payload.fileName}. {workspace_scope} This report is designed to remain dataset-agnostic, so downstream sections rely on detected roles, computed statistics, and executed workflow steps rather than hard-coded assumptions about specific fields.'
    )
    elements.append(Spacer(1, 10))

    add_section('Data Understanding', 'Profiling converts raw columns into usable metadata by estimating types, inferring roles, and quantifying completeness before transformation.')
    add_stat_cards([
        ('Numeric', role_count('numeric')),
        ('Categorical', role_count('categorical')),
        ('Datetime', role_count('datetime')),
        ('Identifiers', role_count('identifier')),
    ])
    elements.append(Spacer(1, 8))
    understanding_rows = [['Column', 'Type', 'Role', 'Non-null', 'Nulls', 'Unique']]
    for column in payload.columns[:20]:
        understanding_rows.append([column.name, column.dtype, column.role, column.nonNull, column.nullCount, column.uniqueCount])
    add_table(understanding_rows, [content_width * 0.28, content_width * 0.12, content_width * 0.12, content_width * 0.12, content_width * 0.11, content_width * 0.11], header_bg='#134e4a')
    if len(payload.columns) > 20:
        add_paragraph(f'Showing the first 20 columns out of {len(payload.columns)} profiled columns.', small_style)
    elements.append(PageBreak())

    add_section('Exploratory Data Analysis', 'EDA summarizes the dataset structure, descriptive behavior, and strongest relationships so later cleaning and modeling choices have context.')
    add_stat_cards([
        ('Numeric Fields', len(payload.edaStats.numericColumns)),
        ('Categorical Fields', len(payload.edaStats.categoricalColumns)),
        ('Correlations', len(payload.edaStats.correlations)),
        ('AI Insight', 'Captured' if payload.aiInsights else 'Not captured'),
    ])
    elements.append(Spacer(1, 8))
    corr_image = build_correlation_chart_image(payload.edaStats.correlations)
    if corr_image is not None:
        elements.append(corr_image)
        elements.append(Spacer(1, 8))
    numeric_rows = [['Field', 'Mean', 'Std', 'Min', 'Median', 'Max']]
    for field_name in payload.edaStats.numericColumns[:10]:
        stats = payload.edaStats.stats.get(field_name, {})
        numeric_rows.append([
            field_name,
            metric_text(stats.get('mean')),
            metric_text(stats.get('std')),
            metric_text(stats.get('min')),
            metric_text(stats.get('median')),
            metric_text(stats.get('max')),
        ])
    add_table(numeric_rows if len(numeric_rows) > 1 else [['Field', 'Mean', 'Std', 'Min', 'Median', 'Max'], ['N/A', 'N/A', 'N/A', 'N/A', 'N/A', 'N/A']], [content_width * 0.3, content_width * 0.11, content_width * 0.11, content_width * 0.11, content_width * 0.11, content_width * 0.11], header_bg='#115e59')
    if payload.edaStats.correlations:
        elements.append(Spacer(1, 8))
        correlation_rows = [['Pair', 'Correlation']]
        for item in payload.edaStats.correlations[:8]:
            correlation_rows.append([item.get('pair', 'N/A'), item.get('correlation', 'N/A')])
        add_table(correlation_rows, [content_width * 0.78, content_width * 0.18], header_bg='#115e59')
    if payload.aiInsights:
        elements.append(Spacer(1, 8))
        add_callout('AI Insight Summary', str(payload.aiInsights), tone='#eff6ff', border='#93c5fd')
    elements.append(PageBreak())

    add_section('Data Cleaning', 'Cleaning follows exploratory analysis and prepares the stable analysis layer used by forecasting, ML training, and final prediction.')
    add_stat_cards([
        ('Cleaning Done', 'Yes' if payload.cleaningDone else 'No'),
        ('Logged Actions', len(payload.cleaningLogs)),
        ('Rows Removed', max(0, payload.totalRows - payload.cleanedRowCount)),
        ('Rows Retained', f'{payload.cleanedRowCount:,}'),
    ])
    elements.append(Spacer(1, 8))
    if payload.cleaningLogs:
        cleaning_rows = [['Action', 'Detail', 'Timestamp']]
        for log in payload.cleaningLogs[:24]:
            cleaning_rows.append([log.action, log.detail, log.timestamp])
        add_table(cleaning_rows, [content_width * 0.18, content_width * 0.57, content_width * 0.18], header_bg='#0f766e')
    else:
        add_paragraph('No cleaning logs were captured for this run. The report still remains valid and summarizes the workflow with the currently available state.', small_style)

    if ts_result:
        elements.append(PageBreak())
        add_section('Time Series Forecast', 'The time-series forecasting tab models chronology directly and is appropriate when the temporal sequence itself carries the predictive signal.')
        # Determine if we have DB-enriched multi-model data or legacy session data
        ts_db_model = ts_result.get('ts_db_model')
        if ts_db_model:
            # Multi-model data from ts_forecast_results DB
            db_mae = ts_result.get('ts_db_mae')
            db_rmse = ts_result.get('ts_db_rmse')
            db_smape = ts_result.get('ts_db_smape')
            db_comparison = ts_result.get('ts_db_model_comparison') or []
            db_future = ts_result.get('ts_db_future_forecast') or []
            db_stationarity = ts_result.get('ts_db_stationarity') or {}
            db_insight = ts_result.get('ts_db_insight') or {}
            add_stat_cards([
                ('Best Model', ts_db_model),
                ('SMAPE', metric_text(db_smape)),
                ('MAE', metric_text(db_mae)),
                ('Horizon', f'{len(db_future)} periods'),
            ])
            elements.append(Spacer(1, 8))
            add_table([
                ['Field', 'Value'],
                ['Selected Model', ts_db_model],
                ['Stationarity', db_stationarity.get('status', 'N/A')],
                ['ADF p-value', metric_text(db_stationarity.get('adf_pvalue'))],
                ['KPSS p-value', metric_text(db_stationarity.get('kpss_pvalue'))],
                ['MAE', metric_text(db_mae)],
                ['RMSE', metric_text(db_rmse)],
                ['SMAPE', metric_text(db_smape)],
            ], [content_width * 0.28, content_width * 0.68], header_bg='#134e4a')
            if db_comparison:
                elements.append(Spacer(1, 8))
                comp_rows = [['Candidate', 'Status', 'MAE', 'RMSE', 'SMAPE']]
                for item in db_comparison[:5]:
                    comp_rows.append([
                        item.get('model', 'N/A'), item.get('status', 'N/A'),
                        metric_text(item.get('mae')), metric_text(item.get('rmse')),
                        metric_text(item.get('smape'))
                    ])
                add_table(comp_rows, [content_width * 0.2, content_width * 0.15, content_width * 0.15, content_width * 0.15, content_width * 0.15], header_bg='#0f766e')
            if db_future:
                elements.append(Spacer(1, 8))
                f_rows = [['Future Period', 'Forecast', 'Lower', 'Upper']]
                for item in db_future[:10]:
                    f_rows.append([item.get('period', 'N/A'), metric_text(item.get('forecast')), metric_text(item.get('lower')), metric_text(item.get('upper'))])
                add_table(f_rows, [content_width * 0.22, content_width * 0.22, content_width * 0.22, content_width * 0.22], header_bg='#115e59')
            elements.append(Spacer(1, 8))
            add_paragraph(db_insight.get('insight_text', 'Time-series forecasting output recorded.'), body_style)
            if db_insight.get('risk_flag'):
                elements.append(Spacer(1, 6))
                add_callout('Risk Flag', str(db_insight['risk_flag']), tone='#fffbeb', border='#f59e0b')
        else:
            # Legacy session state data
            ts_training = ts_result.get('training_summary', {}) or {}
            ts_metrics = ts_result.get('metrics', {}) or {}
            ts_profile = ts_result.get('dataset_profile', {}) or {}
            stationarity = ts_result.get('stationarity_check', {}) or {}
            add_stat_cards([
                ('Model', ts_training.get('model_name', 'N/A')),
                ('Train/Test', f"{ts_training.get('train_percentage', 'N/A')}% / {ts_training.get('test_percentage', 'N/A')}%"),
                ('Horizon', ts_training.get('forecast_periods', 'N/A')),
                ('Frequency', ts_result.get('period_label') or ts_result.get('frequency') or ts_profile.get('detected_frequency', 'Period')),
            ])
            elements.append(Spacer(1, 8))
            add_table([
                ['Field', 'Value'],
                ['Date Column', ts_result.get('date_column', 'N/A')],
                ['Target Column', ts_result.get('target_column', 'N/A')],
                ['Usable Periods', ts_profile.get('usable_periods', 'N/A')],
                ['Data Quality Score', (ts_result.get('data_quality') or {}).get('score', 'N/A')],
                ['Naive Baseline MAE', ((ts_result.get('naive_baseline') or {}).get('metrics') or {}).get('mae', 'N/A')],
                ['MAE Improvement vs Naive', f"{(ts_result.get('naive_baseline') or {}).get('mae_improvement_pct', 'N/A')}%"],
                ['Volatility', metric_text(ts_profile.get('volatility'))],
                ['Stationarity Verdict', stationarity.get('verdict', 'N/A')],
                ['Stationarity Note', stationarity.get('note', 'N/A')],
                ['MAE / RMSE / MAPE', f"{metric_text(ts_metrics.get('mae'))} / {metric_text(ts_metrics.get('rmse'))} / {metric_text(ts_metrics.get('mape'))}"],
            ], [content_width * 0.28, content_width * 0.68], header_bg='#134e4a')
            if ts_result.get('model_comparison'):
                elements.append(Spacer(1, 8))
                comparison_rows = [['Candidate', 'Status', 'MAE', 'RMSE', 'MAPE']]
                for item in ts_result.get('model_comparison', [])[:8]:
                    metrics = item.get('metrics') or {}
                    comparison_rows.append([item.get('model_name', 'N/A'), item.get('status', 'N/A'), metric_text(metrics.get('mae')), metric_text(metrics.get('rmse')), metric_text(metrics.get('mape'))])
                add_table(comparison_rows, [content_width * 0.22, content_width * 0.18, content_width * 0.16, content_width * 0.16, content_width * 0.16], header_bg='#0f766e')
            if ts_result.get('assumptions_audit'):
                elements.append(Spacer(1, 8))
                add_callout('Methodology & Assumptions', '<br/>'.join(str(item) for item in ts_result.get('assumptions_audit', [])[:8]))
            elements.append(Spacer(1, 8))
            elements.append(build_line_chart_image('Time Series Forecast', ts_result.get('history', []), ts_result.get('test_forecast', []), ts_result.get('future_forecast', []), include_interval=True))
            if ts_result.get('future_forecast'):
                elements.append(Spacer(1, 8))
                future_rows = [['Future Period', 'Forecast', 'Lower', 'Upper']]
                for item in ts_result.get('future_forecast', [])[:10]:
                    future_rows.append([item.get('period', 'N/A'), metric_text(item.get('predicted')), metric_text(item.get('lower')), metric_text(item.get('upper'))])
                add_table(future_rows, [content_width * 0.22, content_width * 0.22, content_width * 0.22, content_width * 0.22], header_bg='#115e59')
            elements.append(Spacer(1, 8))
            add_paragraph(ts_result.get('analysis', 'Time-series forecasting output was recorded for this workflow.'), body_style)

    if ml_forecast_result:
        elements.append(PageBreak())
        add_section('Machine Learning Forecast', 'The ML forecasting path transforms time into engineered features, then trains a general-purpose learner to project future periods.')
        ml_training = ml_forecast_result.get('training_summary', {}) or {}
        ml_metrics = ml_forecast_result.get('metrics', {}) or {}
        ml_profile = ml_forecast_result.get('dataset_profile', {}) or {}
        add_stat_cards([
            ('Model', ml_training.get('model_name', 'N/A')),
            ('Generated Features', len(ml_forecast_result.get('generated_features', []))),
            ('Lag Depth', ml_training.get('lag_periods', 'N/A')),
            ('Forecast Horizon', ml_training.get('forecast_periods', 'N/A')),
        ])
        elements.append(Spacer(1, 8))
        add_table([
            ['Field', 'Value'],
            ['Date Column', ml_forecast_result.get('date_column', 'N/A')],
            ['Target Column', ml_forecast_result.get('target_column', 'N/A')],
            ['Detected Frequency', ml_profile.get('detected_frequency', 'N/A')],
            ['Usable Periods', ml_profile.get('usable_periods', 'N/A')],
            ['Data Quality Score', (ml_forecast_result.get('data_quality') or {}).get('score', 'N/A')],
            ['Naive Baseline MAE', ((ml_forecast_result.get('naive_baseline') or {}).get('metrics') or {}).get('mae', 'N/A')],
            ['MAE Improvement vs Naive', f"{(ml_forecast_result.get('naive_baseline') or {}).get('mae_improvement_pct', 'N/A')}%"],
            ['MAE / RMSE / MAPE', f"{metric_text(ml_metrics.get('mae'))} / {metric_text(ml_metrics.get('rmse'))} / {metric_text(ml_metrics.get('mape'))}"],
        ], [content_width * 0.28, content_width * 0.68], header_bg='#134e4a')
        if ml_forecast_result.get('model_comparison'):
            elements.append(Spacer(1, 8))
            comparison_rows = [['Candidate', 'Status', 'MAE', 'RMSE', 'MAPE']]
            for item in ml_forecast_result.get('model_comparison', [])[:8]:
                metrics = item.get('metrics') or {}
                comparison_rows.append([item.get('model_name', 'N/A'), item.get('status', 'N/A'), metric_text(metrics.get('mae')), metric_text(metrics.get('rmse')), metric_text(metrics.get('mape'))])
            add_table(comparison_rows, [content_width * 0.22, content_width * 0.18, content_width * 0.16, content_width * 0.16, content_width * 0.16], header_bg='#0f766e')
        if ml_forecast_result.get('assumptions_audit'):
            elements.append(Spacer(1, 8))
            add_callout('Methodology & Assumptions', '<br/>'.join(str(item) for item in ml_forecast_result.get('assumptions_audit', [])[:8]))
        elements.append(Spacer(1, 8))
        elements.append(build_line_chart_image('ML Forecast', ml_forecast_result.get('history', []), ml_forecast_result.get('test_forecast', []), ml_forecast_result.get('future_forecast', []), include_interval=True))
        shap_items = ml_forecast_result.get('shap_feature_importance', [])
        if shap_items:
            elements.append(Spacer(1, 8))
            elements.append(build_bar_chart_image('SHAP Feature Importance', shap_items))
        feature_rows = [['Generated Feature']]
        for feature in ml_forecast_result.get('generated_features', [])[:16]:
            feature_rows.append([feature])
        elements.append(Spacer(1, 8))
        add_table(feature_rows if len(feature_rows) > 1 else [['Generated Feature'], ['None captured']], [540], header_bg='#115e59')
        preview_rows = ml_forecast_result.get('feature_preview_rows', [])
        if preview_rows:
            preview_columns = list(preview_rows[0].keys())[:6]
            rows = [['Preview Feature Row'] + preview_columns]
            for row_index, item in enumerate(preview_rows[:6], start=1):
                rows.append([f'Row {row_index}'] + [item.get(column_name, 'N/A') for column_name in preview_columns])
            elements.append(Spacer(1, 8))
            add_table(rows, [90] + [max(88, (content_width - 90) / max(1, len(preview_columns)))] * len(preview_columns), header_bg='#134e4a')
        elements.append(Spacer(1, 8))
        add_paragraph(ml_forecast_result.get('analysis', 'ML forecasting output was recorded for this workflow.'), body_style)

    if payload.reportConfig.includeLoss and loss_forecast_result:
        elements.append(PageBreak())
        add_section('Loss Forecast Analysis', 'Loss forecasting quantifies future value erosion across revenue, operational, inventory, and discount drivers.')
        total_loss = sum(float(row.get('total_loss') or 0) for row in loss_forecast_result)
        peak_row = max(loss_forecast_result, key=lambda row: float(row.get('total_loss') or 0))
        avg_risk = sum(float(row.get('loss_risk_score') or 0) for row in loss_forecast_result) / max(1, len(loss_forecast_result))
        driver_totals = {
            'Revenue Loss': sum(float(row.get('revenue_loss') or 0) for row in loss_forecast_result),
            'Operational Loss': sum(float(row.get('operational_loss') or 0) for row in loss_forecast_result),
            'Inventory Loss': sum(float(row.get('inventory_loss') or 0) for row in loss_forecast_result),
            'Discount Loss': sum(float(row.get('discount_loss') or 0) for row in loss_forecast_result),
        }
        top_driver = max(driver_totals.items(), key=lambda item: item[1])
        add_stat_cards([
            ('Total Loss', f'{total_loss:,.0f}'),
            ('Peak Loss Period', peak_row.get('period', 'N/A')),
            ('Avg Risk Score', f'{avg_risk:.1%}'),
            ('Top Loss Driver', f'{top_driver[0]} ({(top_driver[1] / total_loss * 100) if total_loss else 0:.0f}%)'),
        ])
        elements.append(Spacer(1, 8))
        try:
            fig, ax = plt.subplots(figsize=(8.8, 2.8))
            periods = [str(row.get('period')) for row in loss_forecast_result]
            ax.plot(periods, [float(row.get('total_loss') or 0) for row in loss_forecast_result], color='#dc2626', linewidth=2.5, label='Total Loss')
            ax.plot(periods, [float(row.get('revenue_loss') or 0) for row in loss_forecast_result], color='#ef4444', linewidth=1.5, label='Revenue')
            ax.plot(periods, [float(row.get('operational_loss') or 0) for row in loss_forecast_result], color='#f97316', linewidth=1.5, label='Operational')
            ax.plot(periods, [float(row.get('inventory_loss') or 0) for row in loss_forecast_result], color='#f59e0b', linewidth=1.5, label='Inventory')
            ax.plot(periods, [float(row.get('discount_loss') or 0) for row in loss_forecast_result], color='#8b5cf6', linewidth=1.5, label='Discount')
            ax.set_title('Loss Trend by Driver')
            ax.tick_params(axis='x', rotation=35, labelsize=7)
            ax.grid(True, alpha=0.25)
            ax.legend(fontsize=7, ncol=5)
            chart_buffer = io.BytesIO()
            fig.tight_layout()
            fig.savefig(chart_buffer, format='png', dpi=160)
            plt.close(fig)
            chart_buffer.seek(0)
            elements.append(Image(chart_buffer, width=content_width * 0.92, height=190))
            elements.append(Spacer(1, 8))
        except Exception:
            logger.exception('Failed to render loss chart for report.')
        loss_rows = [['Period', 'Revenue', 'Operational', 'Inventory', 'Discount', 'Total', 'Risk']]
        for row in loss_forecast_result[:14]:
            loss_rows.append([
                row.get('period', 'N/A'),
                metric_text(row.get('revenue_loss')),
                metric_text(row.get('operational_loss')),
                metric_text(row.get('inventory_loss')),
                metric_text(row.get('discount_loss')),
                metric_text(row.get('total_loss')),
                f"{float(row.get('loss_risk_score') or 0):.1%} {row.get('risk_label', '')}",
            ])
        add_table(loss_rows, [content_width * 0.15, content_width * 0.13, content_width * 0.14, content_width * 0.13, content_width * 0.13, content_width * 0.13, content_width * 0.13], header_bg='#991b1b')
        if loss_segments:
            elements.append(Spacer(1, 8))
            segment_rows = [['Segment', 'Type', 'Total Loss', 'Risk Score', 'Risk Label']]
            for item in loss_segments[:12]:
                segment_rows.append([item.get('segment', 'N/A'), item.get('segment_type', 'N/A'), metric_text(item.get('total_loss')), f"{float(item.get('risk_score') or 0):.1%}", item.get('risk_label', 'N/A')])
            add_table(segment_rows, [content_width * 0.3, content_width * 0.15, content_width * 0.18, content_width * 0.14, content_width * 0.14], header_bg='#b91c1c')
        driver_sentence = ', '.join(f'{name} ({value / total_loss * 100:.0f}%)' for name, value in sorted(driver_totals.items(), key=lambda item: item[1], reverse=True)[:3]) if total_loss else 'no material drivers'
        add_callout(
            'Loss Forecast Insights',
            f'The top forecasted loss drivers are {driver_sentence}. Peak exposure appears in {peak_row.get("period", "N/A")} with total loss {float(peak_row.get("total_loss") or 0):,.0f}. Recommended action is to focus mitigation on the largest driver before the next planning cycle.',
            tone='#fff1f2',
            border='#fda4af',
        )

    if payload.reportConfig.includeProfit and selected_profit_result:
        elements.append(PageBreak())
        add_section('Profit Forecast & P&L Projection', 'Profit forecasting combines revenue, cost, operating expense, and forecasted losses into scenario-based P&L projections.')
        scenario_names = ['optimistic', 'baseline', 'pessimistic']
        scenario_summary_rows = [['Scenario', 'Total Revenue', 'Total COGS', 'Gross Profit', 'Total Losses', 'Net Profit', 'Net Margin']]
        for scenario_name in scenario_names:
            rows = profit_scenarios.get(scenario_name, [])
            total_revenue = sum(float(row.get('forecasted_revenue') or 0) for row in rows)
            total_cogs = sum(float(row.get('forecasted_cogs') or 0) for row in rows)
            gross_profit = sum(float(row.get('gross_profit') or 0) for row in rows)
            total_losses = sum(float(row.get('total_losses') or 0) for row in rows)
            net_profit = sum(float(row.get('net_profit') or 0) for row in rows)
            net_margin = (net_profit / total_revenue * 100) if total_revenue else 0
            scenario_summary_rows.append([scenario_name.title(), f'{total_revenue:,.0f}', f'{total_cogs:,.0f}', f'{gross_profit:,.0f}', f'{total_losses:,.0f}', f'{net_profit:,.0f}', f'{net_margin:.1f}%'])
        add_table(scenario_summary_rows, [content_width * 0.14, content_width * 0.14, content_width * 0.14, content_width * 0.14, content_width * 0.14, content_width * 0.14, content_width * 0.12], header_bg='#075985')
        elements.append(Spacer(1, 8))
        try:
            fig, ax = plt.subplots(figsize=(8.8, 2.8))
            for scenario_name, color in [('optimistic', '#10b981'), ('baseline', '#2563eb'), ('pessimistic', '#f43f5e')]:
                rows = profit_scenarios.get(scenario_name, [])
                ax.plot([str(row.get('period')) for row in rows], [float(row.get('net_profit') or 0) for row in rows], label=scenario_name.title(), color=color, linewidth=2)
            ax.axhline(0, color='#64748b', linewidth=1)
            ax.set_title('Net Profit Forecast by Scenario')
            ax.tick_params(axis='x', rotation=35, labelsize=7)
            ax.grid(True, alpha=0.25)
            ax.legend(fontsize=8)
            chart_buffer = io.BytesIO()
            fig.tight_layout()
            fig.savefig(chart_buffer, format='png', dpi=160)
            plt.close(fig)
            chart_buffer.seek(0)
            elements.append(Image(chart_buffer, width=content_width * 0.92, height=190))
        except Exception:
            logger.exception('Failed to render profit chart for report.')
        elements.append(Spacer(1, 8))
        selected_rows = selected_profit_result[:14]
        pnl_rows = [['Period', 'Revenue', 'COGS', 'Gross Profit', 'OpEx', 'Losses', 'Net Profit', 'Net Margin']]
        for row in selected_rows:
            pnl_rows.append([
                row.get('period', 'N/A'),
                metric_text(row.get('forecasted_revenue')),
                metric_text(row.get('forecasted_cogs')),
                metric_text(row.get('gross_profit')),
                metric_text(row.get('operating_expenses')),
                metric_text(row.get('total_losses')),
                metric_text(row.get('net_profit')),
                f"{float(row.get('net_margin_pct') or 0):.1f}%",
            ])
        add_table(pnl_rows, [content_width * 0.13, content_width * 0.12, content_width * 0.12, content_width * 0.13, content_width * 0.12, content_width * 0.12, content_width * 0.13, content_width * 0.11], header_bg='#0369a1')
        elements.append(Spacer(1, 8))
        add_callout(
            'Break-even Analysis',
            f'Break-even period for the baseline scenario is {breakeven_period or "not reached in the selected horizon"}. The selected report scenario is {payload.reportConfig.scenario.title()}, and the final period net profit is {float(selected_profit_result[-1].get("net_profit") or 0):,.0f}.',
            tone='#ecfdf5',
            border='#86efac',
        )
        add_callout(
            'Executive Profit Outlook',
            f'The {payload.reportConfig.scenario} scenario projects total revenue of {sum(float(row.get("forecasted_revenue") or 0) for row in selected_profit_result):,.0f}. Key risk periods are those where net profit approaches or falls below zero, especially when losses absorb margin. Recommended actions are to protect margin through cost control, loss-driver mitigation, and scenario monitoring before the forecast horizon closes.',
            tone='#eff6ff',
            border='#93c5fd',
        )

    elements.append(PageBreak())
    add_section('ML Assistant', 'This section summarizes the supervised learning branch, including selected target, modeling objective, chosen algorithm, feature set, and performance evidence.')
    add_stat_cards([
        ('Target', payload.targetColumn or 'N/A'),
        ('Problem Type', str(payload.problemType).title()),
        ('Selected Model', payload.selectedModel or 'Not trained'),
        ('Features Used', len(payload.selectedFeatures)),
    ])
    elements.append(Spacer(1, 8))
    if payload.selectedFeatures:
        add_paragraph('Selected features: ' + ', '.join(payload.selectedFeatures[:24]))
    if payload.modelMetrics:
        metric_rows = [['Metric', 'Value']]
        for key, value in payload.modelMetrics.items():
            metric_rows.append([key, metric_text(value)])
        add_table(metric_rows, [content_width * 0.48, content_width * 0.48], header_bg='#134e4a')
    else:
        add_paragraph('No supervised ML training metrics were available in the current session.', small_style)
    if payload.featureImportance:
        elements.append(Spacer(1, 8))
        importance_rows = [['Rank', 'Feature', 'Importance']]
        for index, item in enumerate(payload.featureImportance[:12], start=1):
            importance_rows.append([index, item.get('name', 'N/A'), metric_text(item.get('importance'))])
        add_table(importance_rows, [60, content_width * 0.62, content_width * 0.22], header_bg='#115e59')

    elements.append(PageBreak())
    add_section('Prediction', 'The final workflow stage captures the application outcome by storing the latest inference result, probability breakdowns when available, and recent prediction history.')
    if payload.uploadedModel:
        add_table([
            ['Model Name', 'Type', 'Target', 'Problem', 'Trained At'],
            [payload.uploadedModel.name, payload.uploadedModel.type, payload.uploadedModel.target, payload.uploadedModel.problem, payload.uploadedModel.trainedAt],
        ], [content_width * 0.23, content_width * 0.14, content_width * 0.2, content_width * 0.12, content_width * 0.22], header_bg='#0f766e')
        if payload.uploadedModel.features:
            elements.append(Spacer(1, 8))
            add_paragraph('Prediction-serving model features: ' + ', '.join(payload.uploadedModel.features[:24]))
    elements.append(Spacer(1, 8))
    add_stat_cards([
        ('Latest Prediction', payload.predictionResult if payload.predictionResult is not None else 'N/A'),
        ('History Entries', len(payload.predictionHistory)),
        ('Probabilities', 'Available' if payload.predictionProbabilities else 'N/A'),
        ('Analysis', 'Available' if payload.predictionAnalysis else 'N/A'),
    ])
    elements.append(Spacer(1, 8))
    if payload.predictionAnalysis:
        add_paragraph(payload.predictionAnalysis)
    if payload.predictionProbabilities:
        probability_rows = [['Outcome', 'Probability']]
        for label, probability in list(payload.predictionProbabilities.items())[:10]:
            probability_rows.append([label, f'{round(probability * 100, 2)}%'])
        add_table(probability_rows, [content_width * 0.48, content_width * 0.48], header_bg='#134e4a')
        elements.append(Spacer(1, 8))
    latest_feature_rows = latest_prediction_feature_summary()
    if latest_feature_rows:
        add_table(latest_feature_rows, [content_width * 0.48, content_width * 0.48], header_bg='#115e59')
        elements.append(Spacer(1, 8))
    if payload.predictionHistory:
        history_rows = [['Timestamp', 'Prediction', 'Confidence']]
        for item in payload.predictionHistory[-10:]:
            history_rows.append([item.timestamp, item.prediction, 'N/A' if item.confidence is None else f'{round(item.confidence * 100, 2)}%'])
        add_table(history_rows, [content_width * 0.46, content_width * 0.28, content_width * 0.2], header_bg='#134e4a')
    else:
        add_paragraph('No prediction history was recorded in the current session.', small_style)

    doc.build(elements, onFirstPage=decorate_page, onLaterPages=decorate_page)
    return buffer.getvalue()


def build_dynamic_report_doc(payload: ReportPayload) -> bytes:
    session_id = get_session_id(payload.datasetId, payload.sessionId)
    session_state = ensure_session_state(session_id)
    ts_result_raw = payload.timeSeriesForecastResult or session_state.get('time_series_result')
    ml_result_raw = payload.mlForecastResult or session_state.get('ml_forecast_result')
    ts_result = ts_result_raw.model_dump() if hasattr(ts_result_raw, 'model_dump') else ts_result_raw
    ml_forecast_result = ml_result_raw.model_dump() if hasattr(ml_result_raw, 'model_dump') else ml_result_raw

    def html_table(headers: list[str], rows: list[list[Any]]) -> str:
        head_html = ''.join(f'<th>{escape(str(header))}</th>' for header in headers)
        body_html = ''.join(
            '<tr>' + ''.join(f'<td>{escape(str(cell))}</td>' for cell in row) + '</tr>'
            for row in rows
        )
        return f'<table><thead><tr>{head_html}</tr></thead><tbody>{body_html}</tbody></table>'

    column_rows = [
        [column.name, column.dtype, column.role, column.nonNull, column.nullCount, column.uniqueCount]
        for column in payload.columns[:20]
    ]
    cleaning_rows = [
        [log.action, log.detail, log.timestamp]
        for log in payload.cleaningLogs[:24]
    ] or [['None', 'No cleaning logs captured', 'N/A']]
    metric_rows = [[key, value] for key, value in (payload.modelMetrics or {}).items()] or [['N/A', 'No ML metrics captured']]
    prediction_rows = [
        [item.timestamp, item.prediction, 'N/A' if item.confidence is None else f'{round(item.confidence * 100, 2)}%']
        for item in payload.predictionHistory[-10:]
    ] or [['N/A', 'No prediction history captured', 'N/A']]
    ts_future_rows = [
        [item.get('period', 'N/A'), item.get('predicted', 'N/A'), item.get('lower', 'N/A'), item.get('upper', 'N/A')]
        for item in (ts_result or {}).get('future_forecast', [])[:10]
    ]
    ml_future_rows = [
        [item.get('period', 'N/A'), item.get('predicted', 'N/A')]
        for item in (ml_forecast_result or {}).get('future_forecast', [])[:10]
    ]
    loaded_row_count = payload.loadedRowCount or payload.totalRows
    preview_mode = payload.previewLoaded and payload.totalRows > loaded_row_count
    workflow_rows = [
        ['Upload', 'Completed' if payload.totalRows > 0 else 'Pending', f'{payload.totalRows:,} total rows; {loaded_row_count:,} browser rows'],
        ['Understanding', 'Completed' if payload.columns else 'Pending', f'{len(payload.columns)} columns profiled'],
        ['EDA', 'Completed' if payload.columns else 'Pending', f'{len(payload.edaStats.numericColumns)} numeric, {len(payload.edaStats.categoricalColumns)} categorical, {len(payload.edaStats.correlations)} correlations'],
        ['Cleaning', 'Completed' if payload.cleaningDone else 'Pending', f'{len(payload.cleaningLogs)} actions, {payload.cleanedRowCount:,} rows retained'],
        ['Time Series Forecast', 'Completed' if ts_result else 'Skipped', 'Chronology-first forecasting branch'],
        ['Machine Learning Forecast', 'Completed' if ml_forecast_result else 'Skipped', 'Feature-engineered forecasting branch'],
        ['ML Assistant', 'Completed' if payload.modelMetrics else 'Pending', f'{len(payload.selectedFeatures)} selected features'],
        ['Prediction', 'Completed' if payload.predictionResult is not None else 'Pending', f'{len(payload.predictionHistory)} prediction records'],
    ]

    workflow_status = 'Complete workflow captured' if payload.predictionResult is not None else 'Workflow summary generated'
    forecast_status = ', '.join(name for name, present in [('Time Series', bool(ts_result)), ('ML Forecast', bool(ml_forecast_result))] if present) or 'No forecasting branch executed'
    prediction_value = escape(str(payload.predictionResult if payload.predictionResult is not None else 'Pending'))
    generated_at = datetime.now().strftime('%Y-%m-%d %H:%M')
    generation_date = datetime.now().strftime('%Y-%m-%d')
    html = f"""<!DOCTYPE html>
<html>
<head>
  <meta charset="utf-8" />
  <title>{escape(payload.fileName)} Workflow Report</title>
  <style>
    @page {{ size: A4 landscape; margin: 0.6in; }}
    body {{ font-family: Arial, Helvetica, sans-serif; color: #1e293b; margin: 0; line-height: 1.5; background: #f8fafc; }}
    .page {{ page-break-after: always; padding: 8px 0 16px; }}
    .page:last-child {{ page-break-after: auto; }}
    .hero {{
      background: #0f172a;
      border-bottom: 8px solid #1e3a5f;
      color: #ffffff;
      padding: 28px;
      border-radius: 14px;
    }}
    .eyebrow {{ font-size: 11px; font-weight: 700; letter-spacing: 0.16em; text-transform: uppercase; color: #e2e8f0; }}
    .hero h1 {{ margin: 10px 0 8px; font-size: 34px; line-height: 1.05; color: #ffffff; }}
    .hero p {{ margin: 0; color: #e2e8f0 !important; font-size: 15px; max-width: 860px; }}
    .hero-grid, .stats {{ width: 100%; border-collapse: separate; border-spacing: 12px; margin-top: 20px; }}
    .hero-card, .stat {{
      background: #ffffff;
      border: 1px solid #e2e8f0;
      border-radius: 12px;
      padding: 14px;
      vertical-align: top;
    }}
    .label {{ font-size: 10px; color: #475569; font-weight: 700; text-transform: uppercase; letter-spacing: 0.12em; }}
    .value {{ font-size: 22px; font-weight: 700; margin-top: 8px; color: #0369a1; }}
    .deck-title {{ font-size: 24px; color: #0f172a; margin: 26px 0 10px; }}
    .summary {{
      background: #ffffff;
      border: 1px solid #e2e8f0;
      border-radius: 12px;
      padding: 18px 20px;
      box-shadow: 0 2px 8px rgba(0,0,0,0.06);
      margin-top: 18px;
    }}
    .summary strong {{ color: #0f172a; }}
    .section {{
      background: #ffffff;
      border: 1px solid #e2e8f0;
      border-radius: 12px;
      padding: 22px;
      box-shadow: 0 2px 8px rgba(0,0,0,0.06);
      margin-top: 18px;
    }}
    .section-label {{ font-size: 10px; font-weight: 700; letter-spacing: 0.16em; text-transform: uppercase; color: #475569; }}
    h2 {{ color: #0f172a; margin: 8px 0 6px; font-size: 24px; }}
    h3 {{ color: #0f172a; margin: 16px 0 8px; font-size: 18px; }}
    p {{ color: #1e293b; }}
    .muted {{ color: #64748b; }}
    .metric-grid {{ width: 100%; border-collapse: separate; border-spacing: 12px; margin: 12px 0 6px; }}
    .metric-card {{
      background: #f8fbff;
      border: 1px solid #e2e8f0;
      border-radius: 12px;
      padding: 14px;
      width: 25%;
      vertical-align: top;
    }}
    table.data {{ width: 100%; border-collapse: collapse; margin: 12px 0 4px; font-size: 13px; }}
    table.data th, table.data td {{ border: 1px solid #e2e8f0; padding: 9px 10px; text-align: left; vertical-align: top; color: #1e293b; }}
    table.data th {{ background: #1e3a5f; color: #ffffff; }}
    table.data tr:nth-child(even) td {{ background: #f8fbff; }}
    .note {{
      background: #eff6ff;
      border-left: 3px solid #2563eb;
      border-top: 1px solid #93c5fd;
      border-right: 1px solid #93c5fd;
      border-bottom: 1px solid #93c5fd;
      border-radius: 8px;
      color: #1e293b;
      padding: 14px 16px;
      margin: 14px 0 0;
    }}
    .footer-note {{
      background: #f8fafc;
      color: #64748b;
      border: 1px solid #e2e8f0;
      border-radius: 8px;
      font-size: 12px;
      margin-top: 18px;
      padding: 12px 14px;
    }}
  </style>
</head>
<body>
  <div class="page">
    <div class="hero">
      <div class="eyebrow">Aroha Technologies</div>
      <h1>Executive Workflow Report</h1>
      <p>Intelligent Data Assistant | Dataset: {escape(payload.fileName)} | Generated: {escape(generated_at)}</p>
      <table class="hero-grid">
        <tr>
          <td class="hero-card"><div class="label">Dataset</div><div class="value">{escape(payload.fileName)}</div></td>
          <td class="hero-card"><div class="label">Workflow Status</div><div class="value">{escape(workflow_status)}</div></td>
          <td class="hero-card"><div class="label">Forecast Paths</div><div class="value">{escape(forecast_status)}</div></td>
          <td class="hero-card"><div class="label">Prediction</div><div class="value">{prediction_value}</div></td>
        </tr>
      </table>
      <table class="stats">
        <tr>
          <td class="stat"><div class="label">Rows</div><div class="value">{payload.totalRows:,}</div></td>
          <td class="stat"><div class="label">Browser Rows</div><div class="value">{loaded_row_count:,}</div></td>
          <td class="stat"><div class="label">Columns</div><div class="value">{len(payload.columns)}</div></td>
          <td class="stat"><div class="label">Workspace Mode</div><div class="value">{escape('Preview + backend cache' if preview_mode else 'Full dataset')}</div></td>
        </tr>
      </table>
    </div>

    <div class="summary">
      <div class="section-label">Executive Summary</div>
      <h2>Workflow Narrative</h2>
      <p>The uploaded dataset entered the application as <strong>{escape(payload.fileName)}</strong>. The current export reflects the executed workflow path across ingestion, profiling, exploratory analysis, cleaning, forecasting, machine learning, and prediction. This version is intended to read more like a presentation deck than a raw technical dump, so the highlights are surfaced first and the operational details follow as structured tables.</p>
      <div class="note">Estimated memory footprint: {escape(payload.memoryUsage)}. Problem type: {escape(payload.problemType)}. Workspace scope: {escape(f'{loaded_row_count:,} preview rows shown in-browser while the backend kept the full dataset cached.' if preview_mode else f'The full {payload.totalRows:,}-row dataset was available directly in the workspace.')}</div>
    </div>
  </div>

  <div class="page">
    <div class="section">
      <div class="section-label">Coverage</div>
      <h2>Workflow Coverage Map</h2>
      {html_table(['Workflow Area', 'Status', 'Coverage'], workflow_rows).replace('<table>', '<table class="data">')}
      <h3>Upload and Understanding</h3>
      <table class="metric-grid">
        <tr>
          <td class="metric-card"><div class="label">Rows Loaded</div><div class="value">{payload.totalRows:,}</div></td>
          <td class="metric-card"><div class="label">Browser Rows</div><div class="value">{loaded_row_count:,}</div></td>
          <td class="metric-card"><div class="label">Columns Profiled</div><div class="value">{len(payload.columns)}</div></td>
          <td class="metric-card"><div class="label">Duplicates</div><div class="value">{payload.duplicates:,}</div></td>
        </tr>
      </table>
      <div class="note">{escape(f'{loaded_row_count:,} preview rows were rendered in the browser while the full {payload.totalRows:,}-row dataset remained cached on the backend.' if preview_mode else f'The full {payload.totalRows:,}-row dataset was loaded directly into the workspace.')}</div>
      {html_table(['Column', 'Type', 'Role', 'Non-null', 'Nulls', 'Unique'], column_rows).replace('<table>', '<table class="data">')}
    </div>
  </div>

  <div class="page">
    <div class="section">
      <div class="section-label">Analysis</div>
      <h2>EDA and Cleaning Overview</h2>
      <table class="metric-grid">
        <tr>
          <td class="metric-card"><div class="label">Numeric Columns</div><div class="value">{len(payload.edaStats.numericColumns)}</div></td>
          <td class="metric-card"><div class="label">Categorical Columns</div><div class="value">{len(payload.edaStats.categoricalColumns)}</div></td>
          <td class="metric-card"><div class="label">Correlation Signals</div><div class="value">{len(payload.edaStats.correlations)}</div></td>
          <td class="metric-card"><div class="label">AI Insight</div><div class="value">{escape('Captured' if payload.aiInsights else 'Not captured')}</div></td>
        </tr>
      </table>
      <div class="note">{escape(payload.aiInsights or 'No AI insight captured for this session.')}</div>
      <h3>Cleaning Trail</h3>
      <p class="muted">Cleaning follows EDA in the current application workflow and prepares the stable dataset used by forecasting, machine learning, and prediction.</p>
      {html_table(['Action', 'Detail', 'Timestamp'], cleaning_rows).replace('<table>', '<table class="data">')}
      <h3>Forecasting Overview</h3>
      <h3>Time Series Forecast</h3>
      <p class="muted">{escape(str((ts_result or {}).get('analysis', 'Time-series forecast was not executed in this session.')))}</p>
      {html_table(['Future Period', 'Forecast', 'Lower', 'Upper'], ts_future_rows or [['N/A', 'N/A', 'N/A', 'N/A']]).replace('<table>', '<table class="data">')}
      <h3>ML Forecast</h3>
      <p class="muted">{escape(str((ml_forecast_result or {}).get('analysis', 'ML forecast was not executed in this session.')))}</p>
      {html_table(['Future Period', 'Forecast'], ml_future_rows or [['N/A', 'N/A']]).replace('<table>', '<table class="data">')}
    </div>
  </div>

  <div class="page">
    <div class="section">
      <div class="section-label">Modeling</div>
      <h2>Machine Learning and Prediction</h2>
      <table class="metric-grid">
        <tr>
          <td class="metric-card"><div class="label">Target Column</div><div class="value">{escape(str(payload.targetColumn or 'N/A'))}</div></td>
          <td class="metric-card"><div class="label">Selected Model</div><div class="value">{escape(str(payload.selectedModel or 'N/A'))}</div></td>
          <td class="metric-card"><div class="label">Feature Count</div><div class="value">{len(payload.selectedFeatures)}</div></td>
          <td class="metric-card"><div class="label">Latest Prediction</div><div class="value">{prediction_value}</div></td>
        </tr>
      </table>
      <h3>ML Metrics</h3>
      {html_table(['Metric', 'Value'], metric_rows).replace('<table>', '<table class="data">')}
      <h3>Prediction Log</h3>
      <p class="muted">{escape(payload.predictionAnalysis or 'No prediction analysis captured.')}</p>
      {html_table(['Timestamp', 'Prediction', 'Confidence'], prediction_rows).replace('<table>', '<table class="data">')}
      <p class="footer-note">Aroha Technologies | hr@aroha.co.in | +91 9886228615 | Generated {escape(generation_date)}</p>
    </div>
  </div>
</body>
</html>"""
    return html.encode('utf-8')


def build_dynamic_report_pdf(payload: ReportPayload) -> bytes:
    session_id = get_session_id(payload.datasetId, payload.sessionId)
    session_state = ensure_session_state(session_id)
    ts_raw = payload.timeSeriesForecastResult or session_state.get('time_series_result')
    ml_raw = payload.mlForecastResult or session_state.get('ml_forecast_result')
    ts_result = ts_raw.model_dump() if hasattr(ts_raw, 'model_dump') else ts_raw
    ml_result = ml_raw.model_dump() if hasattr(ml_raw, 'model_dump') else ml_raw
    # Enrich ts_result from DB if available (multi-model data)
    if not ts_result and ACTIVITY_DB_AVAILABLE and session_id:
        try:
            with get_activity_connection() as conn:
                ts_db = conn.execute(
                    'SELECT best_model, mae, rmse, smape, model_comparison, future_forecast, stationarity_report, insight FROM ts_forecast_results WHERE dataset_id = %s ORDER BY created_at DESC LIMIT 1',
                    [session_id]
                ).fetchone()
        except Exception:
            ts_db = None
        if ts_db:
            ts_result = {
                'ts_db_model': ts_db['best_model'],
                'ts_db_mae': ts_db['mae'],
                'ts_db_rmse': ts_db['rmse'],
                'ts_db_smape': ts_db['smape'],
                'ts_db_model_comparison': ts_db['model_comparison'],
                'ts_db_future_forecast': ts_db['future_forecast'],
                'ts_db_stationarity': ts_db['stationarity_report'],
                'ts_db_insight': ts_db['insight'],
            }
    loss_rows_raw = payload.lossForecast or session_state.get('loss_forecast_result') or []
    profit_scenarios = payload.scenarios or session_state.get('profit_scenarios') or {}
    profit_rows_raw = payload.profitForecast or profit_scenarios.get(payload.reportConfig.scenario, []) or profit_scenarios.get('baseline', [])
    loss_segments = payload.lossSegments or session_state.get('loss_segments') or []
    breakeven_period = payload.breakevenPeriod or (session_state.get('breakeven') or {}).get('breakeven_period')

    page_size = landscape(letter)
    page_width, _ = page_size
    content_width = page_width - 64
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=page_size, leftMargin=32, rightMargin=32, topMargin=28, bottomMargin=24)
    styles = getSampleStyleSheet()
    h_style = ParagraphStyle('PDF_H', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=16, leading=19, textColor=colors.HexColor('#0f172a'))
    tag_style = ParagraphStyle('PDF_Tag', parent=styles['BodyText'], fontName='Helvetica-Bold', fontSize=8, leading=10, textColor=colors.HexColor('#64748b'), alignment=2)
    body_style = ParagraphStyle('PDF_Body', parent=styles['BodyText'], fontName='Helvetica', fontSize=8.5, leading=11.7, textColor=colors.HexColor('#1e293b'))
    muted_style = ParagraphStyle('PDF_Muted', parent=body_style, fontSize=8, leading=10.8, textColor=colors.HexColor('#64748b'))
    label_style = ParagraphStyle('PDF_Label', parent=body_style, fontName='Helvetica-Bold', fontSize=7.6, leading=9, textColor=colors.HexColor('#0369a1'))
    value_style = ParagraphStyle('PDF_Value', parent=body_style, fontName='Helvetica-Bold', fontSize=12, leading=14, textColor=colors.HexColor('#0f172a'))
    table_header_style = ParagraphStyle('PDF_TH', parent=body_style, fontName='Helvetica-Bold', fontSize=7.8, leading=9.5, textColor=colors.white)
    hero_title = ParagraphStyle('PDF_HeroTitle', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=21, leading=25, textColor=colors.white)
    hero_meta = ParagraphStyle('PDF_HeroMeta', parent=body_style, fontSize=9, leading=12, textColor=colors.HexColor('#e2e8f0'))
    elements: list[Any] = []

    def para(value: Any, style: ParagraphStyle = body_style) -> Paragraph:
        return Paragraph(escape(str(value)).replace('\n', '<br/>'), style)

    def fmt(value: Any, digits: int = 3) -> str:
        if value is None:
            return 'N/A'
        if isinstance(value, (int, np.integer)):
            return f'{int(value):,}'
        if isinstance(value, (float, np.floating)):
            return f'{float(value):,.{digits}f}'
        return str(value)

    def money(value: Any) -> str:
        try:
            return f'{float(value):,.0f}'
        except Exception:
            return 'N/A'

    def pct(value: Any) -> str:
        try:
            return f'{float(value):.1f}%'
        except Exception:
            return 'N/A'

    def add_table(rows: list[list[Any]], widths: list[float], header_bg: str = '#1e3a5f') -> None:
        safe_rows = []
        for row_index, row in enumerate(rows):
            row_style = table_header_style if row_index == 0 else body_style
            safe_rows.append([para(cell, row_style) for cell in row])
        table = Table(safe_rows, colWidths=widths, repeatRows=1)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(header_bg)),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8fbff')]),
            ('GRID', (0, 0), (-1, -1), 0.35, colors.HexColor('#e2e8f0')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 5),
            ('RIGHTPADDING', (0, 0), (-1, -1), 5),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ]))
        elements.append(table)

    def start_section(title: str, tag: str) -> None:
        if elements:
            elements.append(PageBreak())
        header = Table([[para(title, h_style), para(tag, tag_style)]], colWidths=[content_width * 0.72, content_width * 0.28])
        header.setStyle(TableStyle([
            ('LINEBELOW', (0, 0), (-1, -1), 0.8, colors.HexColor('#e2e8f0')),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 7),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        elements.append(header)
        elements.append(Spacer(1, 8))

    def add_cards(cards: list[tuple[str, Any]]) -> None:
        cells = []
        widths = []
        for label, value in cards:
            card = Table([[para(label, label_style)], [para(value, value_style)]], colWidths=[content_width / len(cards) - 7])
            card.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.white),
                ('BOX', (0, 0), (-1, -1), 0.65, colors.HexColor('#e2e8f0')),
                ('LEFTPADDING', (0, 0), (-1, -1), 8),
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 7),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 7),
            ]))
            cells.append(card)
            widths.append(content_width / len(cards))
        elements.append(Table([cells], colWidths=widths))

    def add_note(title: str, text: str, tone: str = '#eff6ff', border: str = '#2563eb') -> None:
        note = Table([[Paragraph(f'<b>{escape(title)}</b><br/>{escape(text)}', body_style)]], colWidths=[content_width])
        note.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(tone)),
            ('LINEBEFORE', (0, 0), (-1, -1), 3, colors.HexColor(border)),
            ('BOX', (0, 0), (-1, -1), 0.35, colors.HexColor('#e2e8f0')),
            ('LEFTPADDING', (0, 0), (-1, -1), 10),
            ('RIGHTPADDING', (0, 0), (-1, -1), 10),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ]))
        elements.append(note)

    def skipped(name: str, reason: str, action: str) -> None:
        add_note(name, f'Reason skipped: {reason} Suggested action: {action}', '#fffbeb', '#f59e0b')

    def image_from_fig(fig: Any, width: float = 500, height: float = 170) -> Image | None:
        try:
            img = io.BytesIO()
            fig.tight_layout()
            fig.savefig(img, format='png', dpi=160, bbox_inches='tight')
            plt.close(fig)
            img.seek(0)
            return Image(img, width=width, height=height)
        except Exception:
            logger.exception('Failed to render PDF chart.')
            return None

    def add_image(image: Image | None) -> None:
        if image is not None:
            elements.append(image)
            elements.append(Spacer(1, 7))

    def simple_line(title: str, periods: list[str], series: list[tuple[str, list[float], str]]) -> Image | None:
        if not periods or not series:
            return None
        fig, ax = plt.subplots(figsize=(8.4, 2.6))
        for label, values, color in series:
            if values:
                ax.plot(periods[:len(values)], values, label=label, color=color, linewidth=2)
        ax.set_title(title)
        ax.tick_params(axis='x', rotation=35, labelsize=7)
        ax.tick_params(axis='y', labelsize=7)
        ax.grid(True, alpha=0.25)
        ax.legend(fontsize=7)
        return image_from_fig(fig, content_width * 0.86, 165)

    def role_count(*needles: str) -> int:
        return sum(1 for column in payload.columns if any(needle in str(column.role).lower() or needle in str(column.dtype).lower() for needle in needles))

    def decorate(canvas: Any, doc_obj: Any) -> None:
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor('#e2e8f0'))
        canvas.line(doc.leftMargin, 20, page_width - doc.rightMargin, 20)
        canvas.setFont('Helvetica', 8)
        canvas.setFillColor(colors.HexColor('#64748b'))
        canvas.drawString(doc.leftMargin, 10, f'Aroha Technologies | hr@aroha.co.in | +91 9886228615')
        canvas.drawRightString(page_width - doc.rightMargin, 10, f'Page {canvas.getPageNumber()}')
        canvas.restoreState()

    loaded_rows = payload.loadedRowCount or payload.totalRows
    preview_mode = payload.previewLoaded and payload.totalRows > loaded_rows
    generated_at = datetime.now().strftime('%d %b %Y, %I:%M %p')
    report_id = str(uuid.uuid4())[:12]
    agent_version = 'IDA v2.1.0'
    analysis_type = 'Exploratory Data Analysis with Forecasting and ML'

    start_section('Data Upload', 'Section 1')
    hero = Table([[
        para('Intelligent Data Assistant', hero_meta),
        para('Comprehensive Analysis Report', hero_title),
        para(f'Dataset: {payload.fileName} | Generated: {generated_at}', hero_meta),
        Spacer(1, 4),
        para(f'Report ID: {report_id}  |  Agent: {agent_version}  |  Type: {analysis_type}', ParagraphStyle('HeroMetaSmall', parent=hero_meta, fontSize=8, leading=10, textColor=colors.HexColor('#94a3b8'))),
    ]], colWidths=[content_width])
    hero.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#0f172a')),
        ('LINEBELOW', (0, 0), (-1, -1), 8, colors.HexColor('#1e3a5f')),
        ('LEFTPADDING', (0, 0), (-1, -1), 14),
        ('RIGHTPADDING', (0, 0), (-1, -1), 14),
        ('TOPPADDING', (0, 0), (-1, -1), 12),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
    ]))
    elements.append(hero)
    elements.append(Spacer(1, 8))
    add_cards([('Filename', payload.fileName), ('Dataset ID', payload.datasetId or session_id), ('Rows Loaded', f'{payload.totalRows:,}'), ('Columns', len(payload.columns))])
    elements.append(Spacer(1, 8))
    add_table([
        ['Field', 'Value'],
        ['Estimated Memory Size', payload.memoryUsage or 'N/A'],
        ['Workspace Mode', 'Preview + backend cache' if preview_mode else 'Full workspace load'],
        ['Upload Timestamp', generated_at],
        ['Entry Description', f'{loaded_rows:,} browser rows loaded; {payload.totalRows:,} total rows available.' if preview_mode else f'The file entered the workspace with {payload.totalRows:,} rows available for analysis.'],
    ], [content_width * 0.25, content_width * 0.71])
    elements.append(Spacer(1, 8))
    coverage_rows = [
        ['Workflow Tab', 'Status'],
        ['Data Upload', 'Completed' if payload.totalRows else 'Skipped'],
        ['Data Understanding', 'Completed' if payload.columns else 'Skipped'],
        ['Exploratory Data Analysis', 'Completed' if payload.columns or payload.edaStats.numericColumns else 'Skipped'],
        ['Data Cleaning', 'Completed' if payload.cleaningLogs or payload.cleaningDone else 'Clean - no actions'],
        ['Time Series Forecast', 'Completed' if ts_result else 'Skipped'],
        ['Machine Learning Forecast', 'Completed' if ml_result else 'Skipped'],
        ['Loss Forecast', 'Completed' if loss_rows_raw else 'Skipped'],
        ['Profit Forecast', 'Completed' if profit_rows_raw else 'Skipped'],
        ['ML Assistant', 'Completed' if payload.modelMetrics or payload.selectedModel else 'Skipped'],
        ['Prediction', 'Completed' if payload.predictionResult is not None or payload.predictionHistory else 'Skipped'],
    ]
    coverage_table = Table([[para(cell, table_header_style if row_index == 0 else body_style) for cell in row] for row_index, row in enumerate(coverage_rows)], colWidths=[content_width * 0.55, content_width * 0.25], repeatRows=1)
    coverage_style = [
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e3a5f')),
        ('GRID', (0, 0), (-1, -1), 0.35, colors.HexColor('#e2e8f0')),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 5),
        ('RIGHTPADDING', (0, 0), (-1, -1), 5),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]
    for row_index, row in enumerate(coverage_rows[1:], start=1):
        status = str(row[1]).lower()
        fill = '#ecfdf5' if 'completed' in status else '#f0fdf4' if 'clean' in status else '#fffbeb'
        coverage_style.append(('BACKGROUND', (0, row_index), (-1, row_index), colors.HexColor(fill)))
    coverage_table.setStyle(TableStyle(coverage_style))
    elements.append(coverage_table)

    start_section('Data Understanding', 'Section 2')
    if not payload.columns:
        skipped('Data Understanding', 'Column profiling was not captured.', 'Run Data Understanding after upload.')
    else:
        add_cards([('Numeric', role_count('numeric', 'float', 'int')), ('Categorical', role_count('categorical')), ('Datetime', role_count('datetime', 'date', 'time')), ('Identifier', role_count('identifier'))])
        elements.append(Spacer(1, 8))
        column_rows = [['Column', 'Data Type', 'Role', 'Non-null', 'Null Count', 'Unique Count']]
        column_rows.extend([[c.name, c.dtype, c.role, c.nonNull, c.nullCount, c.uniqueCount] for c in payload.columns])
        add_table(column_rows, [content_width * 0.26, content_width * 0.13, content_width * 0.13, content_width * 0.13, content_width * 0.13, content_width * 0.13])
        elements.append(Spacer(1, 8))
        add_note('AI Schema Summary', f'{len(payload.columns)} columns were profiled across numeric, categorical, datetime, and identifier roles. Null, non-null, and uniqueness metrics define the schema baseline for downstream analysis.')

    start_section('Exploratory Data Analysis', 'Section 3')
    if not payload.columns and not payload.edaStats.numericColumns:
        skipped('Exploratory Data Analysis', 'EDA was not captured.', 'Open the EDA tab after Data Understanding.')
    else:
        add_cards([('Numeric Fields', len(payload.edaStats.numericColumns)), ('Categorical Fields', len(payload.edaStats.categoricalColumns)), ('Correlations', len(payload.edaStats.correlations)), ('AI Insight', 'Captured' if payload.aiInsights else 'Not captured')])
        elements.append(Spacer(1, 8))
        stat_rows = [['Field', 'Mean', 'Std', 'Min', 'Median', 'Max']]
        for field in payload.edaStats.numericColumns:
            stats = payload.edaStats.stats.get(field, {})
            stat_rows.append([field, fmt(stats.get('mean')), fmt(stats.get('std')), fmt(stats.get('min')), fmt(stats.get('median')), fmt(stats.get('max'))])
        add_table(stat_rows if len(stat_rows) > 1 else stat_rows + [['N/A', 'N/A', 'N/A', 'N/A', 'N/A', 'N/A']], [content_width * 0.27, content_width * 0.13, content_width * 0.13, content_width * 0.13, content_width * 0.13, content_width * 0.13])
        elements.append(Spacer(1, 8))
        corr_rows = [['Pair', 'Correlation']]
        corr_rows.extend([[item.get('pair', 'N/A'), fmt(item.get('correlation'))] for item in payload.edaStats.correlations])
        add_table(corr_rows if len(corr_rows) > 1 else corr_rows + [['N/A', 'N/A']], [content_width * 0.72, content_width * 0.22])
        elements.append(Spacer(1, 8))
        add_image(build_correlation_chart_image(payload.edaStats.correlations))
        add_note('AI Insight', payload.aiInsights or 'Not captured', '#eff6ff' if payload.aiInsights else '#f8fafc', '#2563eb' if payload.aiInsights else '#94a3b8')

    start_section('Data Cleaning', 'Section 4')
    rows_removed = max(0, payload.totalRows - (payload.cleanedRowCount or payload.totalRows))
    add_cards([('Cleaning Done', 'Yes' if payload.cleaningDone else 'No'), ('Logged Actions', len(payload.cleaningLogs)), ('Rows Removed', f'{rows_removed:,}'), ('Rows Retained', f'{payload.cleanedRowCount:,}')])
    elements.append(Spacer(1, 8))
    if payload.cleaningLogs:
        clean_rows = [['Action Name', 'Detail Description', 'Timestamp']]
        clean_rows.extend([[log.action, log.detail, log.timestamp] for log in payload.cleaningLogs])
        add_table(clean_rows, [content_width * 0.2, content_width * 0.56, content_width * 0.18])
    else:
        add_note('Data was clean - no actions required', 'No cleaning actions were logged for this session.', '#ecfdf5', '#22c55e')

    start_section('Time Series Forecast', 'Section 5')
    if not ts_result:
        skipped('Time Series Forecast', 'TS Forecast was not run.', 'Run Time Series Forecast with a date column and target.')
    else:
        profile = ts_result.get('dataset_profile') or {}
        training = ts_result.get('training_summary') or {}
        metrics = ts_result.get('metrics') or {}
        add_cards([('Frequency', profile.get('detected_frequency') or ts_result.get('frequency') or 'N/A'), ('Usable Periods', profile.get('usable_periods', 'N/A')), ('Selected Model', 'SARIMA'), ('MAPE', fmt(metrics.get('mape')))])
        elements.append(Spacer(1, 8))
        add_table([
            ['Field', 'Value'],
            ['Volatility', fmt(profile.get('volatility'))],
            ['Training Split', f"{training.get('train_periods', 'N/A')} train / {training.get('test_periods', 'N/A')} test"],
            ['Stationarity Note', (ts_result.get('stationarity_check') or {}).get('note', 'N/A')],
            ['MAE', fmt(metrics.get('mae'))],
            ['RMSE', fmt(metrics.get('rmse'))],
        ], [content_width * 0.25, content_width * 0.71])
        elements.append(Spacer(1, 8))
        future_rows = [['Period', 'Forecast', 'Lower 95%', 'Upper 95%']]
        future_rows.extend([[p.get('period', 'N/A'), fmt(p.get('predicted')), fmt(p.get('lower')), fmt(p.get('upper'))] for p in ts_result.get('future_forecast', [])])
        add_table(future_rows if len(future_rows) > 1 else future_rows + [['N/A', 'N/A', 'N/A', 'N/A']], [content_width * 0.25, content_width * 0.23, content_width * 0.23, content_width * 0.23])
        elements.append(Spacer(1, 8))
        add_image(build_line_chart_image('Historical vs Forecast', ts_result.get('history', []), ts_result.get('test_forecast', []), ts_result.get('future_forecast', []), True))
        add_note('Forecast Insight', ts_result.get('analysis') or 'No forecast insight captured.')

    start_section('Machine Learning Forecast', 'Section 6')
    if not ml_result:
        skipped('Machine Learning Forecast', 'ML Forecast was not run.', 'Run ML Forecast to evaluate machine learning forecast candidates.')
    else:
        profile = ml_result.get('dataset_profile') or {}
        training = ml_result.get('training_summary') or {}
        metrics = ml_result.get('metrics') or {}
        selected_model = (ml_result.get('model_details') or {}).get('model_name') or training.get('model_name') or 'N/A'
        if 'sarima' in str(selected_model).lower():
            selected_model = 'N/A - SARIMA excluded'
        add_cards([('Selected Model', selected_model), ('Generated Features', len(ml_result.get('generated_features', []))), ('Usable Periods', profile.get('usable_periods', 'N/A')), ('Volatility', fmt(profile.get('volatility')))])
        elements.append(Spacer(1, 8))
        if float(metrics.get('mae') or 0) == 0 and float(metrics.get('rmse') or 0) == 0:
            add_note('Metric Warning', 'MAE and RMSE are both 0.000. Review leakage, target construction, and train/test split before relying on this forecast.', '#fffbeb', '#f59e0b')
            elements.append(Spacer(1, 8))
        add_table([
            ['Field', 'Value'],
            ['Date Column', ml_result.get('date_column', 'N/A')],
            ['Target Column', ml_result.get('target_column', 'N/A')],
            ['Detected Frequency', profile.get('detected_frequency') or ml_result.get('frequency') or 'N/A'],
            ['Training Split', f"{training.get('train_periods', 'N/A')} train / {training.get('test_periods', 'N/A')} test"],
            ['Data Quality Score', (ml_result.get('data_quality') or {}).get('score', 'N/A')],
            ['Naive Baseline MAE', ((ml_result.get('naive_baseline') or {}).get('metrics') or {}).get('mae', 'N/A')],
            ['MAE Improvement %', (ml_result.get('naive_baseline') or {}).get('mae_improvement_pct', 'N/A')],
        ], [content_width * 0.25, content_width * 0.71])
        elements.append(Spacer(1, 8))
        comparison = [['Candidate', 'Status', 'MAE', 'RMSE', 'MAPE', 'Availability Note']]
        for item in ml_result.get('model_comparison', []):
            name = str(item.get('model_name') or item.get('model_type') or 'N/A')
            if 'sarima' in name.lower():
                continue
            metrics_row = item.get('metrics') or {}
            comparison.append([name, item.get('status', 'N/A'), fmt(metrics_row.get('mae')), fmt(metrics_row.get('rmse')), fmt(metrics_row.get('mape')), item.get('availability_note') or item.get('skip_reason') or 'Available'])
        for name in ['XGBoost', 'LightGBM', 'Prophet', 'Gradient Boosting']:
            if not any(name.lower().split()[0] in str(row[0]).lower() for row in comparison[1:]):
                comparison.append([name, 'Not evaluated', 'N/A', 'N/A', 'N/A', 'No candidate result returned'])
        add_table(comparison, [content_width * 0.17, content_width * 0.12, content_width * 0.1, content_width * 0.1, content_width * 0.1, content_width * 0.35])
        elements.append(Spacer(1, 8))
        forecast_rows = [['Period', 'Forecast']]
        forecast_rows.extend([[p.get('period', 'N/A'), fmt(p.get('predicted'))] for p in ml_result.get('future_forecast', [])])
        add_table(forecast_rows if len(forecast_rows) > 1 else forecast_rows + [['N/A', 'N/A']], [content_width * 0.48, content_width * 0.48])
        elements.append(Spacer(1, 8))
        add_image(build_line_chart_image('ML Forecast Line Chart', ml_result.get('history', []), ml_result.get('test_forecast', []), ml_result.get('future_forecast', []), False))
        shap = ml_result.get('shap_feature_importance', [])
        add_image(build_bar_chart_image('SHAP Feature Importance', shap) if shap else None)
        if shap:
            add_note('Top Driver', f"{shap[0].get('name', 'N/A')} with importance score {fmt(shap[0].get('importance'))}.")
        feature_rows = [['Generated Feature']]
        feature_rows.extend([[feature] for feature in ml_result.get('generated_features', [])])
        add_table(feature_rows if len(feature_rows) > 1 else feature_rows + [['None captured']], [content_width * 0.96])
        preview = ml_result.get('feature_preview_rows', [])
        if preview:
            columns = list(preview[0].keys())
            rows = [columns] + [[item.get(column, 'N/A') for column in columns] for item in preview]
            elements.append(Spacer(1, 8))
            add_table(rows, [content_width * 0.96 / max(1, len(columns))] * len(columns))
        add_note('Forecast Insight', ml_result.get('analysis') or 'No ML forecast insight captured.')

    start_section('Loss Forecast', 'Section 7')
    if not loss_rows_raw:
        skipped('Loss Forecast', 'Loss Forecast was not run.', 'Run Loss Forecast to quantify loss risk.')
    else:
        total_loss = sum(float(row.get('total_loss') or 0) for row in loss_rows_raw)
        peak = max(loss_rows_raw, key=lambda row: float(row.get('total_loss') or 0))
        avg_risk = sum(float(row.get('loss_risk_score') or 0) for row in loss_rows_raw) / max(1, len(loss_rows_raw))
        drivers = {
            'Revenue Loss': sum(float(row.get('revenue_loss') or 0) for row in loss_rows_raw),
            'Operational Loss': sum(float(row.get('operational_loss') or 0) for row in loss_rows_raw),
            'Inventory Loss': sum(float(row.get('inventory_loss') or 0) for row in loss_rows_raw),
            'Discount Loss': sum(float(row.get('discount_loss') or 0) for row in loss_rows_raw),
        }
        top_driver = max(drivers.items(), key=lambda item: item[1])
        add_cards([('Total Forecasted Loss', money(total_loss)), ('Highest Risk Period', peak.get('period', 'N/A')), ('Average Risk Score', f'{avg_risk:.1%}'), ('Top Driver', f'{top_driver[0]} ({(top_driver[1] / total_loss * 100) if total_loss else 0:.1f}%)')])
        elements.append(Spacer(1, 8))
        periods = [str(row.get('period')) for row in loss_rows_raw]
        add_image(simple_line('Loss Trend Line Chart', periods, [('Total Loss', [float(row.get('total_loss') or 0) for row in loss_rows_raw], '#dc2626')]))
        fig, ax = plt.subplots(figsize=(8.4, 2.6))
        bottoms = np.zeros(len(periods))
        for key, color in [('revenue_loss', '#ef4444'), ('operational_loss', '#f97316'), ('inventory_loss', '#f59e0b'), ('discount_loss', '#8b5cf6')]:
            vals = np.array([float(row.get(key) or 0) for row in loss_rows_raw])
            ax.bar(periods, vals, bottom=bottoms, label=key.replace('_', ' ').title(), color=color)
            bottoms += vals
        ax.set_title('Loss Breakdown Stacked Bar Chart')
        ax.tick_params(axis='x', rotation=35, labelsize=7)
        ax.legend(fontsize=7, ncol=4)
        add_image(image_from_fig(fig, content_width * 0.86, 165))
        fig, ax = plt.subplots(figsize=(4.6, 2.8))
        ax.pie(list(drivers.values()), labels=list(drivers.keys()), autopct='%1.0f%%', textprops={'fontsize': 7})
        ax.set_title('Segment Loss Mix Donut Chart')
        fig.gca().add_artist(plt.Circle((0, 0), 0.55, fc='white'))
        add_image(image_from_fig(fig, 260, 165))
        risk_rows = [['Segment', 'Risk Score', 'Risk Label']]
        risk_rows.extend([[item.get('segment', 'N/A'), f"{float(item.get('risk_score') or 0):.1%}", item.get('risk_label', 'N/A')] for item in loss_segments])
        add_table(risk_rows if len(risk_rows) > 1 else risk_rows + [['N/A', 'N/A', 'N/A']], [content_width * 0.55, content_width * 0.2, content_width * 0.2])
        elements.append(Spacer(1, 8))
        full_loss = [['Period', 'Revenue Loss', 'Operational Loss', 'Inventory Loss', 'Discount Loss', 'Total Loss', 'Risk Score', 'Risk Label']]
        full_loss.extend([[row.get('period', 'N/A'), money(row.get('revenue_loss')), money(row.get('operational_loss')), money(row.get('inventory_loss')), money(row.get('discount_loss')), money(row.get('total_loss')), f"{float(row.get('loss_risk_score') or 0):.1%}", row.get('risk_label', 'N/A')] for row in loss_rows_raw])
        add_table(full_loss, [content_width * 0.12, content_width * 0.12, content_width * 0.13, content_width * 0.12, content_width * 0.12, content_width * 0.12, content_width * 0.12, content_width * 0.11])
        elements.append(Spacer(1, 8))
        driver_rows = [['Driver', 'Total Amount', 'Share']]
        driver_rows.extend([[name, money(value), f'{(value / total_loss * 100) if total_loss else 0:.1f}%'] for name, value in sorted(drivers.items(), key=lambda item: item[1], reverse=True)[:3]])
        add_table(driver_rows, [content_width * 0.36, content_width * 0.3, content_width * 0.28])

    start_section('Profit Forecast', 'Section 8')
    if not profit_rows_raw:
        skipped('Profit Forecast', 'Profit Forecast was not run.', 'Run Profit Forecast to generate scenario P&L projections.')
    else:
        revenue = sum(float(row.get('forecasted_revenue') or 0) for row in profit_rows_raw)
        cogs = sum(float(row.get('forecasted_cogs') or 0) for row in profit_rows_raw)
        gross = sum(float(row.get('gross_profit') or 0) for row in profit_rows_raw)
        losses = sum(float(row.get('total_losses') or 0) for row in profit_rows_raw)
        net = sum(float(row.get('net_profit') or 0) for row in profit_rows_raw)
        add_cards([('Forecasted Revenue', money(revenue)), ('Gross Profit', money(gross)), ('Net Profit', money(net)), ('Net Margin', f'{(net / revenue * 100) if revenue else 0:.1f}%')])
        elements.append(Spacer(1, 8))
        fig, ax = plt.subplots(figsize=(7.2, 2.5))
        ax.bar(['Revenue', 'COGS', 'Loss', 'Net Profit'], [revenue, -cogs, -losses, net], color=['#2563eb', '#f97316', '#dc2626', '#16a34a'])
        ax.axhline(0, color='#64748b', linewidth=1)
        ax.set_title('P&L Waterfall Chart')
        add_image(image_from_fig(fig, content_width * 0.7, 160))
        scenario_periods = [str(row.get('period')) for row in profit_rows_raw]
        scenario_series = []
        for name, color in [('optimistic', '#16a34a'), ('baseline', '#2563eb'), ('pessimistic', '#dc2626')]:
            rows = profit_scenarios.get(name, [])
            if rows:
                scenario_periods = [str(row.get('period')) for row in rows]
            scenario_series.append((name.title(), [float(row.get('net_profit') or 0) for row in rows], color))
        add_image(simple_line('Net Profit Forecast Three-Scenario Line Chart', scenario_periods, scenario_series))
        add_image(simple_line('Gross vs Net Margin Trend Chart', [str(row.get('period')) for row in profit_rows_raw], [('Gross Margin', [float(row.get('gross_margin_pct') or 0) for row in profit_rows_raw], '#0369a1'), ('Net Margin', [float(row.get('net_margin_pct') or 0) for row in profit_rows_raw], '#16a34a')]))
        fig, ax = plt.subplots(figsize=(8.4, 2.6))
        periods = [str(row.get('period')) for row in profit_rows_raw]
        x = np.arange(len(periods))
        width = 0.25
        ax.bar(x - width, [float(row.get('forecasted_revenue') or 0) for row in profit_rows_raw], width, label='Revenue', color='#2563eb')
        ax.bar(x, [float(row.get('forecasted_cogs') or 0) + float(row.get('operating_expenses') or 0) for row in profit_rows_raw], width, label='Cost', color='#f97316')
        ax.bar(x + width, [float(row.get('total_losses') or 0) for row in profit_rows_raw], width, label='Loss', color='#dc2626')
        ax.set_xticks(x)
        ax.set_xticklabels(periods, rotation=35, fontsize=7)
        ax.legend(fontsize=7)
        ax.set_title('Revenue vs Cost vs Loss Bar Chart')
        add_image(image_from_fig(fig, content_width * 0.86, 165))
        add_note('Break-even Analysis', f'Period reached: {breakeven_period or "Not reached"}. Gross margin: {(gross / revenue * 100) if revenue else 0:.1f}%. Net margin: {(net / revenue * 100) if revenue else 0:.1f}%.', '#ecfdf5', '#22c55e')
        pnl = [['Period', 'Revenue', 'COGS', 'Gross Profit', 'OpEx', 'Total Loss', 'Net Profit', 'Gross Margin', 'Net Margin']]
        pnl.extend([[row.get('period', 'N/A'), money(row.get('forecasted_revenue')), money(row.get('forecasted_cogs')), money(row.get('gross_profit')), money(row.get('operating_expenses')), money(row.get('total_losses')), money(row.get('net_profit')), pct(row.get('gross_margin_pct')), pct(row.get('net_margin_pct'))] for row in profit_rows_raw])
        add_table(pnl, [content_width * 0.11, content_width * 0.1, content_width * 0.1, content_width * 0.11, content_width * 0.09, content_width * 0.1, content_width * 0.11, content_width * 0.11, content_width * 0.11])
        elements.append(Spacer(1, 8))
        scenario_rows = [['Scenario', 'Revenue', 'COGS', 'Gross Profit', 'Total Losses', 'Net Profit', 'Net Margin']]
        for name in ['optimistic', 'baseline', 'pessimistic']:
            rows = profit_scenarios.get(name, [])
            s_revenue = sum(float(row.get('forecasted_revenue') or 0) for row in rows)
            s_cogs = sum(float(row.get('forecasted_cogs') or 0) for row in rows)
            s_gross = sum(float(row.get('gross_profit') or 0) for row in rows)
            s_losses = sum(float(row.get('total_losses') or 0) for row in rows)
            s_net = sum(float(row.get('net_profit') or 0) for row in rows)
            scenario_rows.append([name.title(), money(s_revenue), money(s_cogs), money(s_gross), money(s_losses), money(s_net), f'{(s_net / s_revenue * 100) if s_revenue else 0:.1f}%'])
        add_table(scenario_rows, [content_width * 0.14, content_width * 0.14, content_width * 0.13, content_width * 0.14, content_width * 0.14, content_width * 0.14, content_width * 0.11])

    start_section('ML Assistant', 'Section 9')
    if not payload.modelMetrics and not payload.selectedModel:
        skipped('ML Assistant', 'ML Assistant was not run in this session.', 'Select a target, choose features, and train a model.')
    else:
        add_cards([('Target Column', payload.targetColumn or 'N/A'), ('Problem Type', payload.problemType), ('Selected Model', payload.selectedModel or 'N/A'), ('Features Used', len(payload.selectedFeatures))])
        elements.append(Spacer(1, 8))
        metrics = [['Metric', 'Value']]
        metrics.extend([[key, fmt(value)] for key, value in (payload.modelMetrics or {}).items()])
        add_table(metrics if len(metrics) > 1 else metrics + [['N/A', 'No training metrics captured']], [content_width * 0.45, content_width * 0.51])
        if payload.featureImportance:
            elements.append(Spacer(1, 8))
            importance = [['Feature', 'Importance']]
            importance.extend([[item.get('name', 'N/A'), fmt(item.get('importance'))] for item in payload.featureImportance])
            add_table(importance, [content_width * 0.66, content_width * 0.3])
        add_note('Model Summary', f'{payload.selectedModel or "Selected model"} was trained for {payload.problemType} using {len(payload.selectedFeatures)} features targeting {payload.targetColumn or "the selected target"}.')

    start_section('Prediction', 'Section 10')
    if payload.predictionResult is None and not payload.predictionHistory:
        skipped('Prediction', 'No predictions were made in this session.', 'Train or load a model, then run prediction.')
        add_note('Prediction Info', 'No predictions were made in this session', '#f8fafc', '#94a3b8')
    else:
        latest = payload.predictionHistory[-1] if payload.predictionHistory else None
        add_cards([('Latest Prediction', payload.predictionResult if payload.predictionResult is not None else 'N/A'), ('Prediction Timestamp', latest.timestamp if latest else 'N/A'), ('History Entries', len(payload.predictionHistory)), ('Probabilities', 'Available' if payload.predictionProbabilities else 'N/A')])
        elements.append(Spacer(1, 8))
        if latest and latest.features:
            feature_rows = [['Input Feature', 'Value']]
            feature_rows.extend([[key, value] for key, value in latest.features.items()])
            add_table(feature_rows, [content_width * 0.45, content_width * 0.51])
            elements.append(Spacer(1, 8))
        if payload.predictionProbabilities:
            prob_rows = [['Class', 'Probability']]
            prob_rows.extend([[label, f'{round(probability * 100, 2)}%'] for label, probability in payload.predictionProbabilities.items()])
            add_table(prob_rows, [content_width * 0.48, content_width * 0.48])
            elements.append(Spacer(1, 8))
        history = [['Timestamp', 'Prediction', 'Confidence']]
        history.extend([[item.timestamp, item.prediction, 'N/A' if item.confidence is None else f'{round(item.confidence * 100, 2)}%'] for item in payload.predictionHistory])
        add_table(history if len(history) > 1 else history + [['N/A', 'No prediction history captured', 'N/A']], [content_width * 0.45, content_width * 0.3, content_width * 0.21])
        add_note('Analysis Note', payload.predictionAnalysis or 'No prediction analysis note captured.')

    # ---- Section 11: Key Statistical Findings ----
    start_section('Key Statistical Findings', 'Section 11')
    add_para(f'Key statistical findings extracted during the analysis of {payload.fileName}:', muted_style)
    elements.append(Spacer(1, 4))
    stat_items = [
        f'Dataset contains {payload.totalRows:,} rows and {len(payload.columns)} columns',
        f'{len(payload.edaStats.numericColumns)} numeric fields profiled for central tendency and spread',
        f'{len(payload.edaStats.categoricalColumns)} categorical fields with cardinality assessment',
        f'{len(payload.edaStats.correlations)} correlation pairs evaluated',
        f'{len(payload.cleaningLogs)} cleaning operations executed',
    ]
    for item in stat_items:
        elements.append(para(f'- {item}', muted_style))
    elements.append(Spacer(1, 8))
    add_cards([
        ('Total Rows', f'{payload.totalRows:,}'), ('Numeric Fields', len(payload.edaStats.numericColumns)),
        ('Correlation Pairs', len(payload.edaStats.correlations)), ('Cleaning Ops', len(payload.cleaningLogs)),
    ])

    # ---- Section 12: Business Insights ----
    start_section('Business Insights', 'Section 12')
    completeness = (sum(c.nonNull for c in payload.columns) / max(1, sum(c.nonNull + c.nullCount for c in payload.columns)) * 100) if payload.columns else 0
    add_para('Business insights derived from the analysis for stakeholder review:', body_style)
    elements.append(Spacer(1, 4))
    insight_items = [
        f'The dataset provides {len(payload.edaStats.numericColumns)} quantifiable metrics for performance tracking and decision-making.',
        f'Data completeness is at {completeness:.1f}% across all fields.',
    ]
    if payload.edaStats.correlations:
        strong_count = len([c for c in payload.edaStats.correlations if abs(float(c.get('correlation', 0))) > 0.7])
        insight_items.append(f'Correlation analysis revealed {strong_count} strong relationships (|r| > 0.7) that may indicate redundant features or predictive signals.')
    rows_removed = max(0, payload.totalRows - (payload.cleanedRowCount or payload.totalRows))
    if rows_removed > 0:
        insight_items.append(f'Cleaning removed {rows_removed:,} problematic rows, improving overall data quality for downstream analysis.')
    else:
        insight_items.append('Data required minimal cleaning, indicating good upstream data collection practices.')
    for item in insight_items:
        elements.append(para(f'- {item}', body_style))
    elements.append(Spacer(1, 8))
    add_cards([('Completeness', f'{completeness:.1f}%'), ('Strong Correlations', f'{strong_count if payload.edaStats.correlations else 0}'),
               ('Data Quality', 'Good' if rows_removed == 0 else 'Improved'), ('Metrics Available', len(payload.edaStats.numericColumns))])

    # ---- Section 13: Agent Summary ----
    start_section('Agent Summary', 'Section 13')
    agent_discoveries = []
    if payload.edaStats.correlations and any(abs(float(c.get('correlation', 0))) > 0.7 for c in payload.edaStats.correlations):
        agent_discoveries.append(f'{strong_count} strong correlations detected')
    if payload.cleaningLogs:
        agent_discoveries.append(f'{len(payload.cleaningLogs)} data quality issues identified and resolved')
    if ts_result:
        tst_name = (ts_result.get('training_summary') or {}).get('model_name', 'selected algorithm')
        agent_discoveries.append(f'Time-series patterns modeled using {tst_name}')
    if ml_result:
        agent_discoveries.append(f'ML forecast generated with {len(ml_result.get("generated_features", []))} engineered features')
    if payload.modelMetrics:
        agent_discoveries.append(f'{payload.selectedModel or "Trained model"} achieved measurable performance on {payload.problemType or "N/A"} task')
    risks = []
    if payload.duplicates > 0:
        risks.append(f'{payload.duplicates:,} duplicate records detected')
    if any(c.nullCount > 0 for c in payload.columns):
        null_cols = [c.name for c in payload.columns if c.nullCount > 0]
        risks.append(f'{len(null_cols)} columns contain missing values')
    if loss_rows_raw and any(float(r.get('loss_risk_score') or 0) > 0.7 for r in loss_rows_raw):
        risks.append('High-risk periods detected in loss forecast')
    add_cards([('Dataset', payload.fileName), ('Fields Analyzed', len(payload.columns)),
               ('Report ID', report_id), ('Agent Version', agent_version)])
    elements.append(Spacer(1, 6))
    add_para('What Was Analyzed:', body_style)
    add_para(f'The Intelligent Data Assistant performed a comprehensive analysis of {payload.fileName}, covering data profiling, exploratory data analysis, data quality assessment, {", ".join(filter(None, [("time-series forecasting" if ts_result else None), ("ML forecasting" if ml_result else None), ("loss/profit forecasting" if loss_rows_raw or profit_rows_raw else None)]))}, and machine learning model training with prediction capabilities.', small_style)
    elements.append(Spacer(1, 4))
    add_para('Key Discoveries:', body_style)
    for d in agent_discoveries or ['General data profiling completed']:
        elements.append(para(f'- {d}', small_style))
    elements.append(Spacer(1, 4))
    add_para('Risks Detected:', body_style)
    if risks:
        for r in risks:
            elements.append(para(f'- {r}', small_style))
    else:
        add_para('- No significant risks detected during analysis.', small_style)
    elements.append(Spacer(1, 4))
    add_para('Recommended Actions:', body_style)
    actions = ['Review data quality findings and apply recommended cleaning operations for optimal model performance']
    if payload.edaStats.correlations:
        top_pairs = [c.get('pair') for c in payload.edaStats.correlations[:3]]
        actions.append(f'Explore strong correlations ({", ".join(str(p) for p in top_pairs)}) for feature engineering opportunities')
    if ts_result:
        fp = (ts_result.get('training_summary') or {}).get('forecast_periods', 'N/A')
        actions.append(f'Validate time-series forecast and consider scenario planning based on {fp}-period outlook')
    if payload.selectedModel:
        actions.append(f'Deploy trained model ({payload.selectedModel}) for ongoing prediction and monitoring')
    actions.append('Schedule regular re-analysis to track data quality and model performance over time')
    for a in actions:
        elements.append(para(f'- {a}', small_style))

    # ---- Section 14: Recommendations ----
    start_section('Recommendations', 'Section 14')
    add_para(f'Based on the comprehensive analysis of {payload.fileName}, the following recommendations are provided for stakeholder consideration:', body_style)
    elements.append(Spacer(1, 6))
    null_col_count = len([c for c in payload.columns if c.nullCount > 0])
    dq_text = f'Address {null_col_count} columns with missing values and review {payload.duplicates:,} duplicate records.' if null_col_count > 0 else 'Data quality is satisfactory with no critical issues detected.'
    add_note('Data Quality Improvements', dq_text, tone='#ecfdf5', border='#22c55e')
    elements.append(Spacer(1, 4))
    feat_text = f'Leverage {len(payload.edaStats.correlations)} correlation signals for feature selection and consider interaction effects among top correlated pairs.' if payload.edaStats.correlations else 'Continue exploring feature relationships as more data becomes available.'
    add_note('Feature Engineering', feat_text, tone='#eff6ff', border='#2563eb')
    elements.append(Spacer(1, 4))
    model_text = f'The trained {payload.selectedModel or "model"} is ready for deployment with {len(payload.selectedFeatures)} features. Implement monitoring for prediction drift and periodic retraining.' if payload.selectedModel else 'Train a model using the ML Assistant to enable prediction capabilities.'
    add_note('Model Deployment', model_text, tone='#f0fdf4', border='#22c55e')
    elements.append(Spacer(1, 4))
    forecast_text = 'Track actual vs forecasted values regularly and update forecast assumptions as new data becomes available.' if ts_result or ml_result else 'Run forecasting modules to generate forward-looking projections.'
    add_note('Forecast Monitoring', forecast_text, tone='#fffbeb', border='#f59e0b')

    # ---- Section 15: Appendix ----
    start_section('Appendix', 'Section 15')
    add_para(f'Additional technical details and supporting information for the analysis of {payload.fileName}.', body_style)
    elements.append(Spacer(1, 6))
    add_table([
        ['Field', 'Value'],
        ['Report ID', report_id],
        ['Generated At', generated_at],
        ['Dataset', payload.fileName],
        ['Analysis Type', analysis_type],
        ['Agent Version', agent_version],
        ['Rows Analyzed', f'{payload.totalRows:,}'],
        ['Columns Profiled', str(len(payload.columns))],
        ['Numeric Fields', str(len(payload.edaStats.numericColumns))],
        ['Categorical Fields', str(len(payload.edaStats.categoricalColumns))],
        ['Correlations Evaluated', str(len(payload.edaStats.correlations))],
        ['Cleaning Actions', str(len(payload.cleaningLogs))],
        ['Forecasting Paths', ', '.join(filter(None, ['TS' if ts_result else '', 'ML' if ml_result else '', 'Loss' if loss_rows_raw else '', 'Profit' if profit_rows_raw else ''])) or 'None'],
        ['ML Model', payload.selectedModel or 'Not trained'],
        ['Prediction Available', 'Yes' if payload.predictionResult is not None else 'No'],
    ], [content_width * 0.3, content_width * 0.66])
    elements.append(Spacer(1, 8))
    add_para('Report generated by the Intelligent Data Assistant (IDA) agentic layer. Charts rendered using Matplotlib. PDF compiled using ReportLab.', small_style)
    elements.append(Spacer(1, 6))
    add_para(f'Technical support: hr@aroha.co.in | +91 9886228615', small_style)

    doc.build(elements, onFirstPage=decorate, onLaterPages=decorate)
    return buffer.getvalue()


@router.post('/cache-dataset')
def cache_dataset(request: DatasetCacheRequest, http_request: Request) -> JSONResponse:
    if not request.data:
        raise HTTPException(status_code=400, detail='Dataset rows are required.')

    try:
        data_frame = normalize_dataframe(pd.DataFrame(request.data))
    except Exception as error:
        raise HTTPException(status_code=400, detail=f'Failed to cache dataset: {error}') from error

    if data_frame.empty or data_frame.shape[1] == 0:
        raise HTTPException(status_code=400, detail='Dataset must contain at least one row and one column.')

    dataset_id = str(uuid.uuid4())[:8]
    cached_path = write_cached_frame(dataset_id, data_frame)
    DATASET_CACHE[dataset_id] = {
        'frame_path': str(cached_path),
        'filename': request.file_name,
        'row_count': int(len(data_frame)),
        'column_count': int(len(data_frame.columns)),
        'columns': list(data_frame.columns),
    }

    response = {
        'datasetId': dataset_id,
        'rowCount': int(len(data_frame)),
        'loadedRowCount': int(len(data_frame)),
        'columnCount': int(len(data_frame.columns)),
        'columns': list(data_frame.columns),
        'previewLoaded': False,
    }
    record_activity(
        request=http_request,
        action='cache_dataset',
        status='success',
        dataset_id=dataset_id,
        file_name=request.file_name,
        detail=f'Cached {request.file_name} for backend processing.',
        metadata={
            'row_count': int(len(data_frame)),
            'column_count': int(len(data_frame.columns)),
        },
    )
    return JSONResponse(content=response)


@router.get('/dataset-preview')
def get_dataset_preview(
    http_request: Request,
    dataset_id: str = Query(...),
) -> JSONResponse:
    dataset_entry = DATASET_CACHE.get(dataset_id)
    if dataset_entry is None:
        raise HTTPException(status_code=404, detail='Cached dataset not found. Please upload the file again.')

    preview_frame, is_polars_preview = load_cached_preview(dataset_entry, DATASET_PREVIEW_ROW_LIMIT)
    row_count = int(dataset_entry.get('row_count') or (preview_frame.height if is_polars_preview else len(preview_frame)))
    loaded_row_count = int(preview_frame.height if is_polars_preview else len(preview_frame))
    preview_loaded = row_count > loaded_row_count
    duplicate_rows = int(dataset_entry.get('duplicate_count') or 0)

    if is_polars_preview:
        preview_rows = safe_serialize(preview_frame.to_dicts())
        preview_columns = build_column_info_from_polars_frame(preview_frame)
    else:
        preview_rows = safe_serialize(preview_frame.to_dict(orient='records'))
        preview_columns = build_column_info_from_frame(preview_frame)

    response = {
        'datasetId': dataset_id,
        'fileName': dataset_entry.get('filename'),
        'data': preview_rows,
        'columns': preview_columns,
        'rowCount': row_count,
        'loadedRowCount': loaded_row_count,
        'previewLoaded': preview_loaded,
        'duplicates': duplicate_rows,
        'sheetSelection': {
            'availableSheets': dataset_entry.get('workbook_sheets') or [],
            'selectedSheets': dataset_entry.get('selected_sheets') or [],
            'mergeMode': dataset_entry.get('merge_mode') or 'single',
            'requiresSelection': bool(len(dataset_entry.get('workbook_sheets') or []) > 1),
        } if dataset_entry.get('excel_path') else None,
    }
    record_activity(
        request=http_request,
        action='load_dataset_preview',
        status='success',
        dataset_id=dataset_id,
        file_name=str(dataset_entry.get('filename') or ''),
        detail='Loaded a cached dataset preview for workspace restore.',
        metadata={
            'row_count': row_count,
            'loaded_row_count': loaded_row_count,
            'preview_loaded': preview_loaded,
        },
    )
    return JSONResponse(content=response)


@router.post('/infer-dtypes')
def infer_dtypes(request: DtypeInferenceRequest, http_request: Request) -> JSONResponse:
    if request.dataset_id:
        dataset_entry = DATASET_CACHE.get(request.dataset_id)
        if dataset_entry is None:
            raise HTTPException(status_code=404, detail='Cached dataset not found. Please upload the file again.')
        frame = load_full_dataset_frame(request.dataset_id, request.data)
        filename = str(dataset_entry.get('filename') or 'dataset')
    else:
        if not request.data:
            raise HTTPException(status_code=400, detail='Dataset rows are required.')
        dataset_entry = {'filename': 'inline dataset'}
        frame = normalize_dataframe(pd.DataFrame(request.data))
        filename = 'inline dataset'

    try:
        inferred_frame, dtype_payload = build_dtype_inference_payload(frame)
        dataset_id = request.dataset_id
        memory_size = int(inferred_frame.memory_usage(deep=True).sum())
        if request.persist:
            if not dataset_id:
                dataset_id = str(uuid.uuid4())[:8]
            cached_path = persist_inferred_dataset_frame(dataset_id, dataset_entry, inferred_frame)
            memory_size = int(cached_path.stat().st_size)

        preview_frame = inferred_frame.head(DATASET_PREVIEW_ROW_LIMIT)
        response = {
            'datasetId': dataset_id,
            'data': safe_serialize(preview_frame.to_dict(orient='records')),
            'columns': build_column_info_from_frame(inferred_frame),
            'rowCount': int(len(inferred_frame)),
            'loadedRowCount': int(len(preview_frame)),
            'previewLoaded': len(inferred_frame) > len(preview_frame),
            'memoryUsage': f'{memory_size / (1024 * 1024):.2f} MB',
            'dtypeInference': dtype_payload,
        }
        record_activity(
            request=http_request,
            action='infer_dtypes',
            status='success',
            dataset_id=dataset_id,
            file_name=filename,
            detail='Inferred dataset dtypes and produced a column-level audit.',
            metadata={
                'row_count': int(len(inferred_frame)),
                'column_count': int(len(inferred_frame.columns)),
                'memory_saved_bytes': dtype_payload['memorySavedBytes'],
                'persisted': bool(request.persist),
            },
        )
        return JSONResponse(content=safe_serialize(response))
    except HTTPException:
        raise
    except Exception as error:
        logger.exception('Dtype inference failed dataset_id=%s', request.dataset_id)
        raise HTTPException(status_code=400, detail=f'Dtype inference failed: {error}') from error


async def parse_dataset_file(http_request: Request, file: UploadFile = File(...)) -> JSONResponse:
    if not file.filename:
        raise HTTPException(status_code=400, detail='A dataset file is required.')

    file_name = file.filename
    lower_file_name = file_name.lower()
    supported_exts = ('.parquet', '.csv', '.tsv', '.xlsx', '.xls')
    if not lower_file_name.endswith(supported_exts):
        raise HTTPException(status_code=400, detail='Only .csv, .tsv, .xlsx, .xls, and .parquet files are supported.')

    dataset_id = str(uuid.uuid4())[:8]
    file_suffix = Path(file_name).suffix.lower()
    cached_path, _ = await write_uploaded_file(file, dataset_id, suffix=file_suffix)
    dataset_entry: dict[str, Any] = {
        'filename': file_name,
        'row_count': 0,
        'column_count': 0,
        'columns': [],
        'duplicate_count': 0,
    }

    try:
        if lower_file_name.endswith('.parquet'):
            if pq is None:
                raise IngestionFormatError('Parquet support is unavailable because pyarrow is not installed in the backend environment.', issue='pyarrow_unavailable')
            parquet_file = pq.ParquetFile(cached_path)
            total_rows = int(parquet_file.metadata.num_rows)
            column_count = len(parquet_file.schema.names)
            frame = pl.read_parquet(cached_path, n_rows=DATASET_PREVIEW_ROW_LIMIT, low_memory=True)
            dataset_entry.update({'parquet_path': str(cached_path)})
            rows = frame.to_dicts()
            column_info = build_column_info_from_polars_frame(frame)
            preview_duplicate_rows = int(max(0, frame.height - frame.unique().height))
        elif lower_file_name.endswith('.csv') or lower_file_name.endswith('.tsv'):
            delimited_options = sniff_delimited_options(cached_path)
            frame = read_delimited_frame(cached_path, n_rows=DATASET_PREVIEW_ROW_LIMIT, options=delimited_options)
            total_rows = count_delimited_rows_from_path(cached_path, delimited_options)
            column_count = len(frame.columns)
            dataset_entry.update({
                'csv_path': str(cached_path),
                'separator': delimited_options['separator'],
                'encoding': delimited_options['encoding'],
                'header_row': int(delimited_options.get('header_row') or 0),
            })
            rows = frame.where(pd.notna(frame), None).to_dict(orient='records')
            column_info = build_column_info_from_frame(frame)
            preview_duplicate_rows = int(max(0, len(frame) - len(frame.drop_duplicates())))
        else:
            sheet_summaries = build_excel_sheet_summaries(cached_path)
            available_sheets = [str(sheet['name']) for sheet in sheet_summaries]
            selected_sheets = resolve_selected_excel_sheets([], available_sheets)
            selection_payload = build_excel_selection_payload(
                excel_path=cached_path,
                selected_sheets=selected_sheets,
                merge_mode='single',
            )

            frame = selection_payload['frame']
            total_rows = int(selection_payload['total_rows'])
            column_count = len(frame.columns)
            rows = selection_payload['rows']
            column_info = selection_payload['column_info']
            preview_duplicate_rows = int(selection_payload['duplicate_rows'])

            dataset_entry.update({
                'excel_path': str(cached_path),
                'workbook_sheets': sheet_summaries,
                'selected_sheets': selected_sheets,
                'active_sheet': selected_sheets[0],
                'merge_mode': 'single',
            })

        dataset_entry.update({
            'row_count': int(total_rows),
            'column_count': int(column_count),
            'columns': [str(column) for column in frame.columns],
            'duplicate_count': int(preview_duplicate_rows),
        })
        DATASET_CACHE[dataset_id] = dataset_entry

        preview_loaded = int(total_rows) > len(rows)
        response = {
            'datasetId': dataset_id,
            'data': safe_serialize(rows),
            'columns': list(frame.columns),
            'columnInfo': column_info,
            'rowCount': int(total_rows),
            'loadedRowCount': int(len(rows)),
            'columnCount': int(column_count),
            'previewLoaded': preview_loaded,
            'previewLimit': DATASET_PREVIEW_ROW_LIMIT,
            'sheetSelection': {
                'availableSheets': dataset_entry.get('workbook_sheets') or [],
                'selectedSheets': dataset_entry.get('selected_sheets') or [],
                'mergeMode': dataset_entry.get('merge_mode'),
                'requiresSelection': bool(len(dataset_entry.get('workbook_sheets') or []) > 1),
            } if lower_file_name.endswith('.xlsx') or lower_file_name.endswith('.xls') else None,
        }
        record_activity(
            request=http_request,
            action='parse_dataset',
            status='success',
            dataset_id=dataset_id,
            file_name=file_name,
            detail=f'Parsed dataset file {file_name}.',
            metadata={
                'row_count': int(total_rows),
                'loaded_row_count': int(len(rows)),
                'column_count': int(column_count),
                'preview_loaded': preview_loaded,
            },
        )
        return JSONResponse(content=response)
    except HTTPException:
        raise
    except Exception as error:
        logger.exception('Dataset ingestion failed for %s', file_name)
        raise HTTPException(status_code=400, detail=friendly_format_error(error, 'dataset file')) from error




@router.post('/parse-dataset')
async def parse_dataset(http_request: Request, file: UploadFile = File(...)) -> JSONResponse:
    return await parse_dataset_file(http_request, file)


@router.post('/parse-parquet')
async def parse_parquet(http_request: Request, file: UploadFile = File(...)) -> JSONResponse:
    return await parse_dataset_file(http_request, file)


@router.post('/parse-dataset-sheet-selection')
def parse_dataset_sheet_selection(request: DatasetSheetSelectionRequest, http_request: Request) -> JSONResponse:
    try:
        dataset_entry = DATASET_CACHE.get(request.dataset_id)
        if dataset_entry is None:
            raise HTTPException(status_code=404, detail='Cached dataset not found. Please upload the file again.')
        if not dataset_entry.get('excel_path'):
            raise HTTPException(status_code=400, detail='Sheet selection is only available for Excel workbooks.')

        excel_path = Path(str(dataset_entry['excel_path']))
        available_sheet_rows = dataset_entry.get('workbook_sheets') or []
        available_sheets = [str(item.get('name')) for item in available_sheet_rows if item.get('name')]
        if not available_sheets:
            dataset_entry['workbook_sheets'] = build_excel_sheet_summaries(excel_path)
            available_sheets = [str(item.get('name')) for item in dataset_entry['workbook_sheets'] if item.get('name')]

        selected_sheets = resolve_selected_excel_sheets(request.selected_sheets, available_sheets)
        merge_mode: Literal['single', 'stack'] = request.merge_mode
        if merge_mode == 'single':
            selected_sheets = [selected_sheets[0]]

        selection_payload = build_excel_selection_payload(
            excel_path=excel_path,
            selected_sheets=selected_sheets,
            merge_mode=merge_mode,
        )
    except HTTPException:
        raise
    except Exception as error:
        logger.exception('Excel sheet selection failed for dataset_id=%s', request.dataset_id)
        raise HTTPException(status_code=400, detail=friendly_format_error(error, 'Excel workbook')) from error

    frame = selection_payload['frame']
    rows = selection_payload['rows']
    total_rows = int(selection_payload['total_rows'])
    loaded_row_count = int(selection_payload['loaded_row_count'])
    preview_loaded = bool(selection_payload['preview_loaded'])
    duplicate_rows = int(selection_payload['duplicate_rows'])
    column_info = selection_payload['column_info']

    dataset_entry.update({
        'selected_sheets': selected_sheets,
        'active_sheet': selected_sheets[0],
        'merge_mode': merge_mode,
        'columns': [str(column) for column in frame.columns],
        'row_count': total_rows,
        'column_count': int(len(frame.columns)),
        'duplicate_count': duplicate_rows,
    })
    DATASET_CACHE[request.dataset_id] = dataset_entry

    response = {
        'datasetId': request.dataset_id,
        'data': rows,
        'columns': [str(column) for column in frame.columns],
        'columnInfo': column_info,
        'rowCount': total_rows,
        'loadedRowCount': loaded_row_count,
        'columnCount': int(len(frame.columns)),
        'previewLoaded': preview_loaded,
        'previewLimit': DATASET_PREVIEW_ROW_LIMIT,
        'sheetSelection': {
            'availableSheets': dataset_entry.get('workbook_sheets') or [],
            'selectedSheets': selected_sheets,
            'mergeMode': merge_mode,
            'requiresSelection': bool(len(dataset_entry.get('workbook_sheets') or []) > 1),
        },
    }
    record_activity(
        request=http_request,
        action='parse_dataset_sheet_selection',
        status='success',
        dataset_id=request.dataset_id,
        file_name=str(dataset_entry.get('filename') or ''),
        detail='Updated Excel sheet selection for cached dataset.',
        metadata={
            'selected_sheets': selected_sheets,
            'merge_mode': merge_mode,
            'row_count': total_rows,
            'loaded_row_count': loaded_row_count,
            'column_count': int(len(frame.columns)),
        },
    )
    return JSONResponse(content=response)


@router.post('/clean-dataset')
def clean_dataset(request: ParquetCleaningRequest, http_request: Request) -> JSONResponse:
    try:
        result = clean_cached_dataset(request)
        record_activity(
            request=http_request,
            action='clean_dataset',
            status='success',
            dataset_id=request.dataset_id,
            detail='Cleaned cached dataset and persisted the transformed version.',
            metadata={
                'row_count': result.get('rowCount'),
                'original_row_count': result.get('originalRowCount'),
                'logged_actions': len(result.get('logs', [])),
            },
        )
        return JSONResponse(content=result)
    except HTTPException:
        raise
    except Exception as error:
        logger.exception('Dataset cleaning failed dataset_id=%s', request.dataset_id)
        raise HTTPException(status_code=400, detail=f'Dataset cleaning failed: {error}') from error


@router.post('/clean-parquet')
def clean_parquet(request: ParquetCleaningRequest, http_request: Request) -> JSONResponse:
    return clean_dataset(request, http_request)

@router.post('/cleaning-justification')
def cleaning_justification(request: CleaningJustificationRequest, http_request: Request) -> JSONResponse:
    if not request.logs:
        raise HTTPException(status_code=400, detail='No cleaning logs provided.')
    response = {'justification': generate_cleaning_justification(request)}
    record_activity(
        request=http_request,
        action='cleaning_justification',
        status='success',
        detail='Generated AI cleaning justification summary.',
        metadata={
            'log_count': len(request.logs),
            'total_rows': request.totalRows,
            'total_columns': request.totalColumns,
        },
    )
    return JSONResponse(content=response)


@router.post('/eda/advanced')
def advanced_eda(request: AdvancedEdaRequest, http_request: Request) -> JSONResponse:
    try:
        result = safe_serialize(build_advanced_eda_payload(request))
        record_activity(
            request=http_request,
            action='advanced_eda',
            status='success',
            dataset_id=request.dataset_id,
            detail='Generated advanced EDA payload.',
            metadata={'has_cached_dataset': bool(request.dataset_id)},
        )
        return JSONResponse(content=result)
    except HTTPException:
        raise
    except Exception as error:
        logger.exception('Advanced EDA generation failed dataset_id=%s', request.dataset_id)
        raise HTTPException(status_code=400, detail=f'Advanced EDA generation failed: {error}') from error


@router.post('/eda/report')
def generate_eda_report(payload: EdaPdfPayload, http_request: Request) -> Response:
    try:
        report_bytes = build_eda_pdf(payload)
    except Exception as error:
        logger.exception('EDA PDF generation failed file_name=%s', payload.fileName)
        raise HTTPException(status_code=400, detail=f'Failed to generate EDA PDF: {error}') from error

    file_stem = ''.join(ch for ch in payload.fileName.rsplit('.', 1)[0] if ch.isalnum() or ch in ('-', '_', ' ')).strip() or 'dataset'
    record_activity(
        request=http_request,
        action='generate_eda_pdf',
        status='success',
        dataset_id=payload.datasetId,
        file_name=payload.fileName,
        detail='Generated the EDA tab PDF export.',
        metadata={
            'total_rows': payload.totalRows,
            'column_count': len(payload.columns),
            'numeric_columns': len(payload.edaStats.numericColumns),
            'categorical_columns': len(payload.edaStats.categoricalColumns),
            'advanced_analysis_available': bool(payload.advancedAnalysis),
        },
    )
    return Response(
        content=report_bytes,
        media_type='application/pdf',
        headers={'Content-Disposition': f'attachment; filename="{file_stem}_eda_report.pdf"'},
    )


def build_dynamic_report_html(payload: ReportPayload) -> bytes:
    session_id = get_session_id(payload.datasetId, payload.sessionId)
    session_state = ensure_session_state(session_id)
    ts_raw = payload.timeSeriesForecastResult or session_state.get('time_series_result')
    ml_raw = payload.mlForecastResult or session_state.get('ml_forecast_result')
    ts_result = ts_raw.model_dump() if hasattr(ts_raw, 'model_dump') else ts_raw
    ml_result = ml_raw.model_dump() if hasattr(ml_raw, 'model_dump') else ml_raw
    loss_rows = payload.lossForecast or session_state.get('loss_forecast_result') or []
    profit_scenarios = payload.scenarios or session_state.get('profit_scenarios') or {}
    profit_rows = payload.profitForecast or profit_scenarios.get(payload.reportConfig.scenario, []) or profit_scenarios.get('baseline', [])
    loss_segments = payload.lossSegments or session_state.get('loss_segments') or []
    breakeven_period = payload.breakevenPeriod or (session_state.get('breakeven') or {}).get('breakeven_period')
    meta = _generate_report_metadata(payload)
    loaded_rows = payload.loadedRowCount or payload.totalRows
    preview_mode = payload.previewLoaded and payload.totalRows > loaded_rows
    generated_at = datetime.now().strftime('%d %b %Y, %I:%M %p')
    rows_removed = max(0, payload.totalRows - (payload.cleanedRowCount or payload.totalRows))
    agent_version = 'IDA v2.1.0'

    def esc(text: Any) -> str:
        return escape(str(text))

    def fmt(value: Any, digits: int = 3) -> str:
        if value is None: return 'N/A'
        if isinstance(value, (int, np.integer)): return f'{int(value):,}'
        if isinstance(value, (float, np.floating)): return f'{float(value):,.{digits}f}'
        return str(value)

    def mny(value: Any) -> str:
        try: return f'{float(value):,.0f}'
        except: return 'N/A'

    def pct(value: Any) -> str:
        try: return f'{float(value):.1f}%'
        except: return 'N/A'

    def role_count(*needles: str) -> int:
        return sum(1 for c in payload.columns if any(n in str(c.role).lower() or n in str(c.dtype).lower() for n in needles))

    corr_b64 = _build_corr_chart_base64(payload.edaStats.correlations) if payload.edaStats.correlations else None

    ts_chart_b64 = None
    if ts_result:
        ts_chart_b64 = _build_line_chart_base64(
            'Time Series Forecast', ts_result.get('history', []), ts_result.get('test_forecast', []),
            ts_result.get('future_forecast', []), True)

    ml_chart_b64 = ml_bar_b64 = None
    if ml_result:
        ml_chart_b64 = _build_line_chart_base64(
            'ML Forecast', ml_result.get('history', []), ml_result.get('test_forecast', []),
            ml_result.get('future_forecast', []), False)
        shap_items = ml_result.get('shap_feature_importance', [])
        if shap_items:
            ml_bar_b64 = _build_bar_chart_base64('SHAP Feature Importance', shap_items)

    loss_chart_b64 = None
    if loss_rows:
        loss_chart_b64 = _build_loss_chart_base64(loss_rows)

    profit_chart_b64 = None
    if profit_rows:
        profit_chart_b64 = _build_profit_chart_base64(profit_scenarios, profit_rows)

    sections_html = ''
    section_count = 0

    def add_section(title: str, content_html: str) -> None:
        nonlocal section_count, sections_html
        section_count += 1
        sections_html += f'''
        <div class="section page-break">
            <div class="section-header">
                <span class="section-number">Section {section_count}</span>
                <h2>{esc(title)}</h2>
            </div>
            {content_html}
        </div>'''

    def add_card(label: str, value: Any) -> str:
        return f'<div class="stat-card"><div class="stat-label">{esc(label)}</div><div class="stat-value">{esc(value)}</div></div>'

    def add_cards_html(cards: list[tuple[str, Any]]) -> str:
        return f'<div class="card-grid">{"".join(add_card(l, v) for l, v in cards)}</div>'

    def add_table_html(headers: list[str], rows: list[list[Any]]) -> str:
        if not rows: return '<p class="muted">No data available.</p>'
        h = ''.join(f'<th>{esc(h)}</th>' for h in headers)
        r = ''.join(f'<tr>{"".join(f"<td>{esc(c)}</td>" for c in row)}</tr>' for row in rows)
        return f'<table><thead><tr>{h}</tr></thead><tbody>{r}</tbody></table>'

    def add_chart_img(b64: str | None, caption: str = '') -> str:
        if not b64: return ''
        return f'<div class="chart-container"><img src="data:image/png;base64,{b64}" alt="{esc(caption)}"/>{f"<p class=\"chart-caption\">{esc(caption)}</p>" if caption else ""}</div>'

    def add_note_html(title: str, text: str, tone: str = 'blue') -> str:
        return f'<div class="note note-{tone}"><strong>{esc(title)}:</strong> {esc(text)}</div>'

    tsc = ts_result.get('dataset_profile') if ts_result else {}
    tst = ts_result.get('training_summary') if ts_result else {}
    tsm = ts_result.get('metrics') if ts_result else {}
    mlc = ml_result.get('dataset_profile') if ml_result else {}
    mlt = ml_result.get('training_summary') if ml_result else {}
    mlm = ml_result.get('metrics') if ml_result else {}

    section_content = ''
    # ---- Section 1: Executive Summary ----
    completed_count = sum(1 for c in [
        payload.totalRows > 0, bool(payload.columns), bool(payload.cleaningDone),
        bool(ts_result), bool(ml_result), bool(loss_rows), bool(profit_rows),
        bool(payload.modelMetrics), payload.predictionResult is not None
    ] if c)
    exec_summary = f'''
    {add_cards_html([
        ('Dataset', payload.fileName), ('Rows', f'{payload.totalRows:,}'), ('Columns', len(payload.columns)),
        ('Workflow Progress', f'{completed_count}/9 areas completed')
    ])}
    <p>The <strong>{esc(payload.fileName)}</strong> dataset was analyzed through a comprehensive workflow encompassing data profiling, exploratory analysis, cleaning, forecasting (time-series, ML, loss, and profit), machine learning training, and prediction. This report presents all findings in a stakeholder-friendly format.</p>
    <div class="meta-grid">
        <div class="meta-item"><span class="meta-label">Report ID</span><span class="meta-value">{esc(meta["report_id"])}</span></div>
        <div class="meta-item"><span class="meta-label">Generated</span><span class="meta-value">{esc(generated_at)}</span></div>
        <div class="meta-item"><span class="meta-label">Agent Version</span><span class="meta-value">{esc(agent_version)}</span></div>
        <div class="meta-item"><span class="meta-label">Analysis Type</span><span class="meta-value">{esc(meta["analysis_type"])}</span></div>
    </div>
    '''
    section_content += f'''
    <div class="section page-break" id="executive-summary">
        <div class="section-header"><span class="section-number">Section 1</span><h2>Executive Summary</h2></div>
        {exec_summary}
    </div>'''

    # ---- Section 2: Dataset Overview ----
    ds_overview = f'''
    {add_cards_html([
        ('Total Rows', f'{payload.totalRows:,}'), ('Columns', len(payload.columns)),
        ('Duplicates', f'{payload.duplicates:,}'), ('Memory', payload.memoryUsage)
    ])}
    <p>Workspace scope: {"Preview-backed" if preview_mode else "Full dataset"} - {esc(f"{loaded_rows:,} rows available in-browser out of {payload.totalRows:,} total." if preview_mode else f"All {payload.totalRows:,} rows loaded directly.")}</p>
    {add_table_html(
        ['Column', 'Type', 'Role', 'Non-null', 'Nulls', 'Unique'],
        [[c.name, c.dtype, c.role, c.nonNull, c.nullCount, c.uniqueCount] for c in payload.columns[:24]]
    )}
    {f'<p class=\"muted\">Showing first 24 of {len(payload.columns)} columns.</p>' if len(payload.columns) > 24 else ''}
    '''
    section_content += f'''
    <div class="section page-break" id="dataset-overview">
        <div class="section-header"><span class="section-number">Section 2</span><h2>Dataset Overview</h2></div>
        {ds_overview}
    </div>'''

    # ---- Section 3: Data Quality Assessment ----
    dq = f'''
    {add_cards_html([
        ('Numeric Fields', role_count('numeric', 'float', 'int')),
        ('Categorical Fields', role_count('categorical')),
        ('Datetime Fields', role_count('datetime', 'date', 'time')),
        ('Identifier Fields', role_count('identifier'))
    ])}
    <p>Data quality assessment covers schema completeness, null ratios, uniqueness, and role inference. {f"{len(payload.cleaningLogs)} cleaning actions were applied." if payload.cleaningLogs else "No cleaning was required or applied."}</p>
    {add_table_html(
        ['Action', 'Detail', 'Timestamp'],
        [[l.action, l.detail, l.timestamp] for l in payload.cleaningLogs[:24]]
    ) if payload.cleaningLogs else '<p class="muted">No cleaning logs recorded.</p>'}
    {add_cards_html([
        ('Cleaning Status', 'Completed' if payload.cleaningDone else 'Pending'),
        ('Rows Removed', f'{rows_removed:,}'), ('Rows Retained', f'{payload.cleanedRowCount:,}')
    ])}
    '''
    section_content += f'''
    <div class="section page-break" id="data-quality">
        <div class="section-header"><span class="section-number">Section 3</span><h2>Data Quality Assessment</h2></div>
        {dq}
    </div>'''

    # ---- Section 4: Exploratory Data Analysis ----
    stat_rows = [['Field', 'Mean', 'Std', 'Min', 'Median', 'Max']]
    for field in payload.edaStats.numericColumns[:12]:
        s = payload.edaStats.stats.get(field, {})
        stat_rows.append([field, fmt(s.get('mean')), fmt(s.get('std')), fmt(s.get('min')), fmt(s.get('median')), fmt(s.get('max'))])
    corr_rows = [['Pair', 'Correlation']] + [[item.get('pair', 'N/A'), fmt(item.get('correlation'))] for item in payload.edaStats.correlations[:10]]
    eda_content = f'''
    {add_cards_html([
        ('Numeric Fields', len(payload.edaStats.numericColumns)),
        ('Categorical Fields', len(payload.edaStats.categoricalColumns)),
        ('Correlations', len(payload.edaStats.correlations)),
        ('AI Insights', 'Available' if payload.aiInsights else 'Not captured')
    ])}
    <h3>Statistical Summary</h3>
    {add_table_html(stat_rows[0], stat_rows[1:]) if len(stat_rows) > 1 else '<p class="muted">No numeric fields to summarize.</p>'}
    <h3>Correlation Analysis</h3>
    {add_chart_img(corr_b64, 'Correlation Heatmap')}
    {add_table_html(corr_rows[0], corr_rows[1:]) if len(corr_rows) > 1 else ''}
    {add_note_html('AI Insight', payload.aiInsights or 'No AI insight captured.') if payload.aiInsights else ''}
    '''
    section_content += f'''
    <div class="section page-break" id="eda">
        <div class="section-header"><span class="section-number">Section 4</span><h2>Exploratory Data Analysis</h2></div>
        {eda_content}
    </div>'''

    # ---- Section 5: Key Statistical Findings ----
    ksf = f'''
    <p>Key statistical findings extracted during the analysis of <strong>{esc(payload.fileName)}</strong>:</p>
    <ul>
        <li>Dataset contains <strong>{payload.totalRows:,}</strong> rows and <strong>{len(payload.columns)}</strong> columns</li>
        <li><strong>{len(payload.edaStats.numericColumns)}</strong> numeric fields profiled for central tendency and spread</li>
        <li><strong>{len(payload.edaStats.categoricalColumns)}</strong> categorical fields with cardinality assessment</li>
        <li><strong>{len(payload.edaStats.correlations)}</strong> correlation pairs evaluated</li>
        <li><strong>{len(payload.cleaningLogs)}</strong> cleaning operations executed, removing <strong>{rows_removed:,}</strong> rows</li>
    </ul>
    '''
    section_content += f'''
    <div class="section page-break" id="statistical-findings">
        <div class="section-header"><span class="section-number">Section 5</span><h2>Key Statistical Findings</h2></div>
        {ksf}
    </div>'''

    # ---- Section 6: Business Insights ----
    bis = f'''
    <p>Business insights derived from the analysis:</p>
    <ul>
        <li>The dataset provides <strong>{len(payload.edaStats.numericColumns)}</strong> quantifiable metrics for performance tracking and decision-making.</li>
        <li>Data completeness is at <strong>{(sum(c.nonNull for c in payload.columns) / max(1, sum(c.nonNull + c.nullCount for c in payload.columns)) * 100):.1f}%</strong> across all fields.</li>
        <li>{f"Correlation analysis revealed <strong>{len([c for c in payload.edaStats.correlations if abs(float(c.get('correlation', 0))) > 0.7])}</strong> strong relationships (|r| > 0.7) that may indicate redundant features or predictive signals." if payload.edaStats.correlations else "Correlation analysis helps identify feature relationships for model development."}</li>
        <li>{f"Cleaning removed <strong>{rows_removed:,}</strong> problematic rows, improving overall data quality for downstream analysis." if rows_removed > 0 else "Data required minimal cleaning, indicating good upstream data collection practices."}</li>
    </ul>
    '''
    section_content += f'''
    <div class="section page-break" id="business-insights">
        <div class="section-header"><span class="section-number">Section 6</span><h2>Business Insights</h2></div>
        {bis}
    </div>'''

    # ---- Section 7: Forecast Results ----
    forecast_content = ''
    # TS Forecast
    if ts_result:
        selected_model = tst.get('model_name', 'N/A')
        ts_section = f'''
        <h3>Time Series Forecast</h3>
        {add_cards_html([
            ('Model', selected_model), ('Frequency', tsc.get('detected_frequency', ts_result.get('frequency', 'N/A'))),
            ('Horizon', f'{len(ts_result.get("future_forecast", []))} periods'),
            ('MAPE', fmt(tsm.get('mape')))
        ])}
        {add_table_html(
            ['Field', 'Value'],
            [['Volatility', fmt(tsc.get('volatility'))],
             ['Stationarity', (ts_result.get('stationarity_check') or {}).get('note', 'N/A')],
             ['MAE', fmt(tsm.get('mae'))], ['RMSE', fmt(tsm.get('rmse'))],
             ['Training', f"{tst.get('train_periods', 'N/A')} train / {tst.get('test_periods', 'N/A')} test"]]
        )}
        {add_chart_img(ts_chart_b64, 'Time Series Forecast Chart')}
        {add_table_html(
            ['Period', 'Forecast', 'Lower 95%', 'Upper 95%'],
            [[p.get('period', 'N/A'), fmt(p.get('predicted')), fmt(p.get('lower')), fmt(p.get('upper'))]
             for p in ts_result.get('future_forecast', [])[:12]]
        )}
        {add_note_html('Forecast Insight', ts_result.get('analysis') or 'No analysis captured.')}
        '''
        forecast_content += ts_section

    # ML Forecast
    if ml_result:
        ml_selected = (ml_result.get('model_details') or {}).get('model_name') or mlt.get('model_name') or 'N/A'
        ml_section = f'''
        <h3>Machine Learning Forecast</h3>
        {add_cards_html([
            ('Model', ml_selected), ('Features', len(ml_result.get('generated_features', []))),
            ('Horizon', f'{len(ml_result.get("future_forecast", []))} periods'),
            ('MAPE', fmt(mlm.get('mape')))
        ])}
        {add_table_html(
            ['Candidate', 'Status', 'MAE', 'RMSE', 'MAPE'],
            [[item.get('model_name', 'N/A'), item.get('status', 'N/A'),
              fmt((item.get('metrics') or {}).get('mae')), fmt((item.get('metrics') or {}).get('rmse')),
              fmt((item.get('metrics') or {}).get('mape'))]
             for item in ml_result.get('model_comparison', [])[:8]]
        )}
        {add_chart_img(ml_chart_b64, 'ML Forecast Chart')}
        {add_chart_img(ml_bar_b64, 'SHAP Feature Importance')}
        {add_table_html(
            ['Period', 'Forecast'],
            [[p.get('period', 'N/A'), fmt(p.get('predicted'))] for p in ml_result.get('future_forecast', [])[:12]]
        )}
        {add_note_html('Forecast Insight', ml_result.get('analysis') or 'No analysis captured.')}
        '''
        forecast_content += ml_section

    # Loss Forecast
    if loss_rows:
        total_loss = sum(float(r.get('total_loss') or 0) for r in loss_rows)
        peak = max(loss_rows, key=lambda r: float(r.get('total_loss') or 0))
        drivers = {k: sum(float(r.get(k) or 0) for r in loss_rows)
                   for k in ['revenue_loss', 'operational_loss', 'inventory_loss', 'discount_loss']}
        top_driver = max(drivers.items(), key=lambda item: item[1])
        loss_section = f'''
        <h3>Loss Forecast Analysis</h3>
        {add_cards_html([
            ('Total Loss', mny(total_loss)), ('Peak Period', peak.get('period', 'N/A')),
            ('Top Driver', f'{top_driver[0].replace("_"," ").title()}'), ('Risk Score', f'{(sum(float(r.get("loss_risk_score") or 0) for r in loss_rows) / max(1, len(loss_rows))):.1%}')
        ])}
        {add_chart_img(loss_chart_b64, 'Loss Trend by Driver')}
        {add_table_html(
            ['Period', 'Revenue', 'Operational', 'Inventory', 'Discount', 'Total', 'Risk'],
            [[r.get('period', 'N/A'), mny(r.get('revenue_loss')), mny(r.get('operational_loss')),
              mny(r.get('inventory_loss')), mny(r.get('discount_loss')), mny(r.get('total_loss')),
              f"{float(r.get('loss_risk_score') or 0):.1%} {r.get('risk_label', '')}"]
             for r in loss_rows[:14]]
        )}
        '''
        forecast_content += loss_section

    # Profit Forecast
    if profit_rows:
        total_rev = sum(float(r.get('forecasted_revenue') or 0) for r in profit_rows)
        total_net = sum(float(r.get('net_profit') or 0) for r in profit_rows)
        profit_section = f'''
        <h3>Profit Forecast & P&amp;L Projection</h3>
        {add_cards_html([
            ('Scenario', payload.reportConfig.scenario.title() if payload.reportConfig else 'Baseline'),
            ('Revenue', mny(total_rev)), ('Net Profit', mny(total_net)),
            ('Net Margin', f'{(total_net / total_rev * 100) if total_rev else 0:.1f}%')
        ])}
        {add_chart_img(profit_chart_b64, 'Net Profit Forecast by Scenario')}
        {add_table_html(
            ['Period', 'Revenue', 'COGS', 'Gross Profit', 'OpEx', 'Losses', 'Net Profit'],
            [[r.get('period', 'N/A'), mny(r.get('forecasted_revenue')), mny(r.get('forecasted_cogs')),
              mny(r.get('gross_profit')), mny(r.get('operating_expenses')), mny(r.get('total_losses')),
              mny(r.get('net_profit'))] for r in profit_rows[:14]]
        )}
        <h4>Scenario Comparison (Optimistic / Baseline / Pessimistic)</h4>
        {add_table_html(
            ['Scenario', 'Revenue', 'COGS', 'Gross Profit', 'Total Losses', 'Net Profit', 'Net Margin'],
            [[name.title(),
              mny(sum(float(r.get('forecasted_revenue') or 0) for r in profit_scenarios.get(name, []))),
              mny(sum(float(r.get('forecasted_cogs') or 0) for r in profit_scenarios.get(name, []))),
              mny(sum(float(r.get('gross_profit') or 0) for r in profit_scenarios.get(name, []))),
              mny(sum(float(r.get('total_losses') or 0) for r in profit_scenarios.get(name, []))),
              mny(sum(float(r.get('net_profit') or 0) for r in profit_scenarios.get(name, []))),
              f'{(sum(float(r.get("net_profit") or 0) for r in profit_scenarios.get(name, [])) / max(1, sum(float(r.get("forecasted_revenue") or 0) for r in profit_scenarios.get(name, []))) * 100):.1f}%']
             for name in ['optimistic', 'baseline', 'pessimistic']]
        )}
        {add_note_html('Break-even Analysis', f'Period: {breakeven_period or "Not reached"}')}
        <h4>Forecast Assumptions</h4>
        <ul>
            <li><strong>Optimistic:</strong> Revenue growth accelerates with stable costs and minimal losses</li>
            <li><strong>Baseline:</strong> Current trends continue with moderate growth and normalized loss patterns</li>
            <li><strong>Pessimistic:</strong> Revenue contraction with increased costs and elevated losses</li>
        </ul>
        <h4>Confidence Indicators</h4>
        {add_cards_html([
            ('Optimistic Confidence', 'Medium-High'), ('Baseline Confidence', 'High'),
            ('Pessimistic Confidence', 'Medium'), ('Model Reliability', 'Based on historical patterns')
        ])}
        '''
        forecast_content += profit_section

    if forecast_content:
        section_content += f'''
        <div class="section page-break" id="forecast-results">
            <div class="section-header"><span class="section-number">Section 7</span><h2>Forecast Results</h2></div>
            {forecast_content}
        </div>'''

    # ---- Section 8: Machine Learning Results ----
    ml_section_content = ''
    if payload.modelMetrics or payload.selectedModel:
        ml_section_content = f'''
        {add_cards_html([
            ('Target', payload.targetColumn or 'N/A'), ('Problem', payload.problemType.title() if payload.problemType else 'N/A'),
            ('Model', payload.selectedModel or 'N/A'), ('Features', len(payload.selectedFeatures))
        ])}
        <h3>Model Performance Metrics</h3>
        {add_table_html(
            ['Metric', 'Value'],
            [[k, fmt(v)] for k, v in (payload.modelMetrics or {}).items()]
        ) if payload.modelMetrics else '<p class="muted">No metrics captured.</p>'}
        <h3>Feature Importance</h3>
        {add_table_html(
            ['Feature', 'Importance'],
            [[item.get('name', 'N/A'), fmt(item.get('importance'))] for item in (payload.featureImportance or [])[:12]]
        ) if payload.featureImportance else '<p class="muted">No feature importance data available.</p>'}
        {add_note_html('Model Summary', f'{payload.selectedModel or "Selected model"} trained for {payload.problemType or "N/A"} on {len(payload.selectedFeatures)} features targeting {payload.targetColumn or "N/A"}.')}
        '''
    if payload.predictionResult is not None or payload.predictionHistory:
        ml_section_content += f'''
        <h3>Prediction Summary</h3>
        {add_cards_html([
            ('Latest Prediction', fmt(payload.predictionResult)),
            ('History Entries', len(payload.predictionHistory)),
            ('Probabilities', 'Available' if payload.predictionProbabilities else 'N/A'),
        ])}
        {add_table_html(
            ['Timestamp', 'Prediction', 'Confidence'],
            [[i.timestamp, fmt(i.prediction), 'N/A' if i.confidence is None else f'{round(i.confidence * 100, 2)}%']
             for i in payload.predictionHistory[-10:]]
        ) if payload.predictionHistory else ''}
        {add_note_html('Analysis', payload.predictionAnalysis) if payload.predictionAnalysis else ''}
        '''
    if ml_section_content:
        section_content += f'''
        <div class="section page-break" id="ml-results">
            <div class="section-header"><span class="section-number">Section 8</span><h2>Machine Learning Results</h2></div>
            {ml_section_content}
        </div>'''

    # ---- Section 9: Agent Summary ----
    agent_discoveries = []
    if payload.edaStats.correlations:
        strong_corrs = [c for c in payload.edaStats.correlations if abs(float(c.get('correlation', 0))) > 0.7]
        if strong_corrs:
            agent_discoveries.append(f'{len(strong_corrs)} strong correlations detected')
    if payload.cleaningLogs:
        agent_discoveries.append(f'{len(payload.cleaningLogs)} data quality issues identified and resolved')
    if ts_result:
        agent_discoveries.append(f'Time-series patterns modeled using {tst.get("model_name", "selected algorithm")}')
    if ml_result:
        agent_discoveries.append(f'ML forecast generated with {len(ml_result.get("generated_features", []))} engineered features')
    if payload.modelMetrics:
        agent_discoveries.append(f'{payload.selectedModel or "Trained model"} achieved measurable performance on {payload.problemType or "N/A"} task')

    risks = []
    if payload.duplicates > 0:
        risks.append(f'{payload.duplicates:,} duplicate records detected')
    if any(c.nullCount > 0 for c in payload.columns):
        null_cols = [c.name for c in payload.columns if c.nullCount > 0]
        risks.append(f'{len(null_cols)} columns contain missing values')
    if loss_rows:
        high_risk = [r for r in loss_rows if float(r.get('loss_risk_score') or 0) > 0.7]
        if high_risk:
            risks.append(f'{len(high_risk)} high-risk periods in loss forecast')

    agent_summary = f'''
    {add_cards_html([
        ('Dataset', payload.fileName), ('Total Rows', f'{payload.totalRows:,}'),
        ('Fields Analyzed', len(payload.columns)), ('Report ID', meta['report_id'])
    ])}
    <h3>What Was Analyzed</h3>
    <p>The Intelligent Data Assistant performed a comprehensive analysis of <strong>{esc(payload.fileName)}</strong>, covering data profiling, exploratory data analysis, data quality assessment, {f"time-series forecasting, ML forecasting, " if ts_result or ml_result else ""}{f"loss/profit forecasting, " if loss_rows or profit_rows else ""}and machine learning model training with prediction capabilities.</p>
    <h3>Key Discoveries</h3>
    <ul>{"".join(f"<li>{esc(d)}</li>" for d in (agent_discoveries or ['General data profiling completed']))}</ul>
    <h3>Important Trends</h3>
    <ul>
        <li>Data quality score: {f"{(sum(c.nonNull for c in payload.columns) / max(1, sum(c.nonNull + c.nullCount for c in payload.columns)) * 100):.1f}%" if payload.columns else "N/A"} completeness</li>
        <li>{f"{len(payload.edaStats.numericColumns)} metrics available for quantitative analysis" if payload.edaStats.numericColumns else "Analysis focused on qualitative/categorical data"}</li>
        {f"<li>Time-series forecast indicates {tsm.get('mape', 'N/A')}% MAPE accuracy</li>" if ts_result else ""}
        {f"<li>ML forecast generates {len(ml_result.get('future_forecast', []))} period outlook</li>" if ml_result else ""}
    </ul>
    <h3>Risks Detected</h3>
    {f"<ul>{''.join(f'<li>{esc(r)}</li>' for r in risks)}</ul>" if risks else '<p>No significant risks detected during analysis.</p>'}
    <h3>Recommended Actions</h3>
    <ul>
        <li>Review data quality findings and apply recommended cleaning operations for optimal model performance</li>
        {f"<li>Explore strong correlations ({[c.get('pair') for c in (payload.edaStats.correlations or [])[:3]]}) for feature engineering opportunities</li>" if payload.edaStats.correlations else ""}
        {f"<li>Validate time-series forecast and consider scenario planning based on {tst.get('forecast_periods', 'N/A')}-period outlook</li>" if ts_result else ""}
        {f"<li>Review ML forecast feature importance to understand key drivers ({[s.get('name') for s in (ml_result.get('shap_feature_importance', []) or [])[:3]]})</li>" if ml_result and ml_result.get('shap_feature_importance') else ""}
        {f"<li>Monitor loss drivers and implement mitigation strategies for top risk areas</li>" if loss_rows else ""}
        {f"<li>Deploy trained model ({payload.selectedModel}) for ongoing prediction and monitoring</li>" if payload.selectedModel else ""}
        <li>Schedule regular re-analysis to track data quality and model performance over time</li>
    </ul>
    '''
    section_content += f'''
    <div class="section page-break" id="agent-summary">
        <div class="section-header"><span class="section-number">Section 9</span><h2>Agent Summary</h2></div>
        {agent_summary}
    </div>'''

    # ---- Section 10: Recommendations ----
    recs = f'''
    <p>Based on the comprehensive analysis of <strong>{esc(payload.fileName)}</strong>, the following recommendations are provided for stakeholder consideration:</p>
    <div class="recommendation-grid">
        <div class="rec-card">
            <h4>Data Quality Improvements</h4>
            <p>{f"Address {len([c for c in payload.columns if c.nullCount > 0])} columns with missing values and review {payload.duplicates:,} duplicate records." if any(c.nullCount > 0 for c in payload.columns) else "Data quality is satisfactory with no critical issues detected."}</p>
        </div>
        <div class="rec-card">
            <h4>Feature Engineering</h4>
            <p>{f"Leverage {len(payload.edaStats.correlations)} correlation signals for feature selection and consider interaction effects among top correlated pairs." if payload.edaStats.correlations else "Continue exploring feature relationships as more data becomes available."}</p>
        </div>
        <div class="rec-card">
            <h4>Model Deployment</h4>
            <p>{f"The trained {payload.selectedModel or 'model'} is ready for deployment with {len(payload.selectedFeatures)} features. Implement monitoring for prediction drift and periodic retraining." if payload.selectedModel else "Train a model using the ML Assistant to enable prediction capabilities."}</p>
        </div>
        <div class="rec-card">
            <h4>Forecast Monitoring</h4>
            <p>{f"Track actual vs forecasted values regularly and update forecast assumptions as new data becomes available." if ts_result or ml_result else "Run forecasting modules to generate forward-looking projections."}</p>
        </div>
    </div>
    '''
    section_content += f'''
    <div class="section page-break" id="recommendations">
        <div class="section-header"><span class="section-number">Section 10</span><h2>Recommendations</h2></div>
        {recs}
    </div>'''

    # ---- Section 11: Appendix ----
    appendix = f'''
    <p>Additional technical details and supporting information for the analysis of <strong>{esc(payload.fileName)}</strong>.</p>
    <h3>Workflow Coverage</h3>
    {add_table_html(
        ['Workflow Area', 'Status'],
        [['Data Upload', 'Completed' if payload.totalRows else 'Skipped'],
         ['Data Understanding', 'Completed' if payload.columns else 'Skipped'],
         ['EDA', 'Completed' if payload.columns else 'Skipped'],
         ['Data Cleaning', 'Completed' if payload.cleaningDone else 'Pending'],
         ['Time Series Forecast', 'Completed' if ts_result else 'Skipped'],
         ['ML Forecast', 'Completed' if ml_result else 'Skipped'],
         ['Loss Forecast', 'Completed' if loss_rows else 'Skipped'],
         ['Profit Forecast', 'Completed' if profit_rows else 'Skipped'],
         ['ML Assistant', 'Completed' if payload.modelMetrics else 'Skipped'],
         ['Prediction', 'Completed' if payload.predictionResult is not None else 'Skipped']]
    )}
    <h3>Report Metadata</h3>
    {add_table_html(
        ['Field', 'Value'],
        [['Report ID', meta['report_id']], ['Generated At', meta['generated_at']],
         ['Dataset', meta['dataset_name']], ['Analysis Type', meta['analysis_type']],
         ['Agent Version', meta['agent_version']], ['Rows Analyzed', f'{payload.totalRows:,}'],
         ['Columns Profiled', str(len(payload.columns))]]
    )}
    <h3>Technical Environment</h3>
    <p>Report generated by the Intelligent Data Assistant (IDA) agentic layer. Charts rendered using Matplotlib. Report compiled using HTML/CSS with embedded Base64 images for standalone viewing.</p>
    '''
    section_content += f'''
    <div class="section" id="appendix">
        <div class="section-header"><span class="section-number">Section 11</span><h2>Appendix</h2></div>
        {appendix}
    </div>'''

    toc_links = [
        ('executive-summary', 'Executive Summary'),
        ('dataset-overview', 'Dataset Overview'),
        ('data-quality', 'Data Quality Assessment'),
        ('eda', 'Exploratory Data Analysis'),
        ('statistical-findings', 'Key Statistical Findings'),
        ('business-insights', 'Business Insights'),
        ('forecast-results', 'Forecast Results'),
        ('ml-results', 'Machine Learning Results'),
        ('agent-summary', 'Agent Summary'),
        ('recommendations', 'Recommendations'),
        ('appendix', 'Appendix'),
    ]
    toc_html = ''.join(f'<li><a href="#{anchor}">{label}</a></li>' for anchor, label in toc_links)

    html = f'''<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1.0"/>
<title>{esc(payload.fileName)} - Comprehensive Analysis Report</title>
<style>
    *, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
    body {{ font-family: 'Segoe UI', Arial, Helvetica, sans-serif; color: #1e293b; background: #f1f5f9; line-height: 1.6; }}
    .page-break {{ page-break-after: always; }}
    .cover-page {{
        background: linear-gradient(135deg, #0f172a 0%, #1e3a5f 50%, #0f766e 100%);
        color: white; padding: 60px 48px; min-height: 100vh; display: flex; flex-direction: column; justify-content: center;
        page-break-after: always;
    }}
    .cover-page h1 {{ font-size: 42px; font-weight: 700; margin-bottom: 12px; line-height: 1.1; }}
    .cover-page .subtitle {{ font-size: 18px; color: #a5f3fc; margin-bottom: 32px; }}
    .cover-meta {{ display: grid; grid-template-columns: 1fr 1fr; gap: 16px; margin-top: 40px; }}
    .cover-meta-item {{ background: rgba(255,255,255,0.08); border: 1px solid rgba(255,255,255,0.15); border-radius: 10px; padding: 16px 20px; }}
    .cover-meta-item .label {{ font-size: 11px; text-transform: uppercase; letter-spacing: 0.1em; color: #a5f3fc; }}
    .cover-meta-item .value {{ font-size: 16px; font-weight: 600; margin-top: 4px; color: white; }}
    .company-name {{ font-size: 13px; letter-spacing: 0.15em; text-transform: uppercase; color: #5eead4; margin-bottom: 8px; }}
    .toc {{ padding: 40px 48px; background: white; page-break-after: always; }}
    .toc h2 {{ font-size: 28px; color: #0f172a; margin-bottom: 24px; border-bottom: 2px solid #e2e8f0; padding-bottom: 12px; }}
    .toc ul {{ list-style: none; columns: 2; column-gap: 40px; }}
    .toc li {{ margin-bottom: 10px; }}
    .toc a {{ color: #0369a1; text-decoration: none; font-size: 15px; display: block; padding: 6px 10px; border-radius: 6px; transition: background 0.15s; }}
    .toc a:hover {{ background: #f0f9ff; }}
    .toc a::before {{ content: counter(toc-counter) ". "; counter-increment: toc-counter; font-weight: 600; color: #64748b; }}
    .toc ul {{ counter-reset: toc-counter; }}
    .section {{ padding: 32px 48px; background: white; margin: 0 0 2px; }}
    .section-header {{ margin-bottom: 24px; border-bottom: 2px solid #e2e8f0; padding-bottom: 12px; }}
    .section-number {{ font-size: 11px; font-weight: 700; text-transform: uppercase; letter-spacing: 0.15em; color: #0284c7; display: block; margin-bottom: 4px; }}
    .section-header h2 {{ font-size: 26px; color: #0f172a; }}
    h3 {{ font-size: 20px; color: #0f172a; margin: 24px 0 12px; }}
    h4 {{ font-size: 16px; color: #334155; margin: 16px 0 8px; }}
    p {{ margin: 8px 0; color: #475569; }}
    .muted {{ color: #94a3b8; font-size: 13px; }}
    ul, ol {{ margin: 8px 0; padding-left: 24px; }}
    li {{ margin-bottom: 4px; color: #475569; }}
    .card-grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(180px, 1fr)); gap: 12px; margin: 16px 0; }}
    .stat-card {{ background: #f8fafc; border: 1px solid #e2e8f0; border-radius: 10px; padding: 14px 16px; }}
    .stat-label {{ font-size: 11px; font-weight: 600; text-transform: uppercase; letter-spacing: 0.08em; color: #64748b; }}
    .stat-value {{ font-size: 20px; font-weight: 700; color: #0369a1; margin-top: 4px; }}
    .meta-grid {{ display: grid; grid-template-columns: repeat(4, 1fr); gap: 12px; margin: 20px 0; padding: 16px; background: #f0fdf4; border: 1px solid #bbf7d0; border-radius: 10px; }}
    .meta-item {{ }}
    .meta-label {{ font-size: 10px; text-transform: uppercase; letter-spacing: 0.08em; color: #64748b; }}
    .meta-value {{ font-size: 13px; font-weight: 600; color: #0f172a; display: block; margin-top: 2px; }}
    table {{ width: 100%; border-collapse: collapse; margin: 12px 0; font-size: 13px; }}
    th, td {{ border: 1px solid #e2e8f0; padding: 8px 10px; text-align: left; vertical-align: top; }}
    th {{ background: #1e3a5f; color: white; font-weight: 600; font-size: 12px; }}
    tr:nth-child(even) td {{ background: #f8fafc; }}
    .chart-container {{ margin: 16px 0; text-align: center; }}
    .chart-container img {{ max-width: 100%; height: auto; border: 1px solid #e2e8f0; border-radius: 8px; }}
    .chart-caption {{ font-size: 12px; color: #64748b; margin-top: 4px; }}
    .note {{ padding: 12px 16px; border-radius: 8px; margin: 12px 0; font-size: 13px; }}
    .note-blue {{ background: #eff6ff; border-left: 4px solid #2563eb; }}
    .note-green {{ background: #f0fdf4; border-left: 4px solid #22c55e; }}
    .note-amber {{ background: #fffbeb; border-left: 4px solid #f59e0b; }}
    .note-red {{ background: #fef2f2; border-left: 4px solid #dc2626; }}
    .recommendation-grid {{ display: grid; grid-template-columns: repeat(2, 1fr); gap: 16px; margin: 16px 0; }}
    .rec-card {{ background: #f8fafc; border: 1px solid #e2e8f0; border-radius: 10px; padding: 18px; }}
    .rec-card h4 {{ color: #0f172a; margin-bottom: 8px; font-size: 15px; }}
    .rec-card p {{ font-size: 13px; color: #475569; }}
    @media print {{
        body {{ background: white; }}
        .section {{ padding: 24px 32px; margin: 0; }}
        .cover-page {{ min-height: auto; padding: 48px 32px; }}
        .page-break {{ page-break-before: always; }}
        .page-break:first-of-type {{ page-break-before: auto; }}
    }}
    @page {{ size: A4; margin: 0; }}
</style>
</head>
<body>
<div class="cover-page" id="cover">
    <div class="company-name">Intelligent Data Assistant</div>
    <h1>Comprehensive Analysis Report</h1>
    <p class="subtitle">A stakeholder-ready summary of the end-to-end analytics journey for {esc(payload.fileName)}</p>
    <div class="cover-meta">
        <div class="cover-meta-item"><div class="label">Report ID</div><div class="value">{esc(meta['report_id'])}</div></div>
        <div class="cover-meta-item"><div class="label">Generated At</div><div class="value">{esc(generated_at)}</div></div>
        <div class="cover-meta-item"><div class="label">Dataset Name</div><div class="value">{esc(payload.fileName)}</div></div>
        <div class="cover-meta-item"><div class="label">Agent Version</div><div class="value">{esc(agent_version)}</div></div>
        <div class="cover-meta-item"><div class="label">Analysis Type</div><div class="value">{esc(meta['analysis_type'])}</div></div>
        <div class="cover-meta-item"><div class="label">Total Rows</div><div class="value">{payload.totalRows:,}</div></div>
    </div>
</div>

<div class="toc page-break">
    <h2>Table of Contents</h2>
    <ul>
        <li><a href="#executive-summary">Executive Summary</a></li>
        <li><a href="#dataset-overview">Dataset Overview</a></li>
        <li><a href="#data-quality">Data Quality Assessment</a></li>
        <li><a href="#eda">Exploratory Data Analysis</a></li>
        <li><a href="#statistical-findings">Key Statistical Findings</a></li>
        <li><a href="#business-insights">Business Insights</a></li>
        <li><a href="#forecast-results">Forecast Results</a></li>
        <li><a href="#ml-results">Machine Learning Results</a></li>
        <li><a href="#agent-summary">Agent Summary</a></li>
        <li><a href="#recommendations">Recommendations</a></li>
        <li><a href="#appendix">Appendix</a></li>
    </ul>
</div>

{section_content}

<div style="text-align:center;padding:24px;font-size:12px;color:#94a3b8;background:white;border-top:1px solid #e2e8f0;">
    Generated by Intelligent Data Assistant (IDA) v{esc(agent_version)} | {esc(generated_at)} | Report ID: {esc(meta['report_id'])}
</div>
</body>
</html>'''
    return html.encode('utf-8')


def build_dynamic_report_docx(payload: ReportPayload) -> bytes:
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        logger.warning('python-docx not installed, falling back to HTML-based DOC')
        return build_dynamic_report_doc(payload)

    session_id = get_session_id(payload.datasetId, payload.sessionId)
    session_state = ensure_session_state(session_id)
    ts_raw = payload.timeSeriesForecastResult or session_state.get('time_series_result')
    ml_raw = payload.mlForecastResult or session_state.get('ml_forecast_result')
    ts_result = ts_raw.model_dump() if hasattr(ts_raw, 'model_dump') else ts_raw
    ml_result = ml_raw.model_dump() if hasattr(ml_raw, 'model_dump') else ml_raw
    loss_rows = payload.lossForecast or session_state.get('loss_forecast_result') or []
    profit_scenarios = payload.scenarios or session_state.get('profit_scenarios') or {}
    profit_rows = payload.profitForecast or profit_scenarios.get(payload.reportConfig.scenario, []) or profit_scenarios.get('baseline', [])
    breakeven_period = payload.breakevenPeriod or (session_state.get('breakeven') or {}).get('breakeven_period')
    meta = _generate_report_metadata(payload)
    generated_at = datetime.now().strftime('%d %b %Y, %I:%M %p')

    def fmt(value: Any, digits: int = 3) -> str:
        if value is None: return 'N/A'
        if isinstance(value, (int, np.integer)): return f'{int(value):,}'
        if isinstance(value, (float, np.floating)): return f'{float(value):,.{digits}f}'
        return str(value)

    def mny(value: Any) -> str:
        try: return f'{float(value):,.0f}'
        except: return 'N/A'

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    def add_heading(text: str, level: int = 1) -> None:
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(15, 23, 42) if level <= 2 else RGBColor(51, 65, 85)

    def add_para(text: str, bold: bool = False, italic: bool = False, size: int = 10) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = RGBColor(71, 85, 105)

    def add_cards(cards: list[tuple[str, Any]]) -> None:
        p = doc.add_paragraph()
        for label, value in cards:
            run = p.add_run(f'  {label}: {value}  ')
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(3, 105, 161)

    def add_table(headers: list[str], rows: list[list[Any]]) -> None:
        if not rows:
            add_para('No data available.', italic=True, size=9)
            return
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = str(header)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(8)
        for r_idx, row in enumerate(rows):
            for c_idx, cell_value in enumerate(row):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = str(cell_value)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

    def add_chart(b64: str | None, caption: str = '') -> None:
        if not b64: return
        try:
            img_data = base64.b64decode(b64)
            doc.add_picture(io.BytesIO(img_data), width=Inches(5.5))
            if caption:
                add_para(caption, italic=True, size=8)
        except Exception:
            add_para(f'[Chart: {caption}]', italic=True, size=8)

    def add_bullet(text: str) -> None:
        p = doc.add_paragraph(text, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(9)

    # Cover page
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Intelligent Data Assistant')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(94, 234, 212)
    run.bold = True

    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = title2.add_run('Comprehensive Analysis Report')
    run2.font.size = Pt(28)
    run2.bold = True
    run2.font.color.rgb = RGBColor(15, 23, 42)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = subtitle.add_run(f'A stakeholder-ready summary for {payload.fileName}')
    run3.font.size = Pt(12)
    run3.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()
    meta_items = [
        ('Report ID', meta['report_id']), ('Generated', generated_at),
        ('Dataset', payload.fileName), ('Agent Version', meta['agent_version']),
        ('Analysis Type', meta['analysis_type']), ('Total Rows', f'{payload.totalRows:,}'),
    ]
    meta_table = doc.add_table(rows=len(meta_items), cols=2)
    meta_table.style = 'Light Shading Accent 1'
    for i, (k, v) in enumerate(meta_items):
        meta_table.rows[i].cells[0].text = k
        meta_table.rows[i].cells[1].text = v
        for cell in meta_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    doc.add_page_break()

    # TOC
    add_heading('Table of Contents', level=1)
    toc_items = [
        'Executive Summary', 'Dataset Overview', 'Data Quality Assessment',
        'Exploratory Data Analysis', 'Key Statistical Findings', 'Business Insights',
        'Forecast Results', 'Machine Learning Results', 'Agent Summary',
        'Recommendations', 'Appendix',
    ]
    for i, item in enumerate(toc_items, 1):
        add_para(f'{i}. {item}', size=10)
    doc.add_page_break()

    # Section 1: Executive Summary
    add_heading('1. Executive Summary', level=1)
    add_cards([
        ('Dataset', payload.fileName), ('Rows', f'{payload.totalRows:,}'),
        ('Columns', str(len(payload.columns))), ('Report ID', meta['report_id']),
    ])
    add_para(f'The {payload.fileName} dataset was analyzed through a comprehensive workflow encompassing data profiling, exploratory analysis, cleaning, forecasting, machine learning training, and prediction. This report presents all findings for stakeholder review.')
    add_para(f'Report ID: {meta["report_id"]} | Generated: {generated_at} | Agent: {meta["agent_version"]}', italic=True, size=9)
    doc.add_page_break()

    # Section 2: Dataset Overview
    add_heading('2. Dataset Overview', level=1)
    add_cards([('Total Rows', f'{payload.totalRows:,}'), ('Columns', str(len(payload.columns))),
               ('Duplicates', f'{payload.duplicates:,}'), ('Memory', payload.memoryUsage)])
    add_table(['Column', 'Type', 'Role', 'Non-null', 'Nulls', 'Unique'],
              [[c.name, c.dtype, c.role, str(c.nonNull), str(c.nullCount), str(c.uniqueCount)]
               for c in payload.columns[:20]])
    doc.add_page_break()

    # Section 3: Data Quality
    add_heading('3. Data Quality Assessment', level=1)
    add_cards([('Cleaning Status', 'Completed' if payload.cleaningDone else 'Pending'),
               ('Actions', str(len(payload.cleaningLogs))),
               ('Rows Removed', f'{max(0, payload.totalRows - (payload.cleanedRowCount or payload.totalRows)):,}'),
               ('Rows Retained', f'{payload.cleanedRowCount:,}')])
    if payload.cleaningLogs:
        add_table(['Action', 'Detail', 'Timestamp'],
                  [[l.action, l.detail, l.timestamp] for l in payload.cleaningLogs[:20]])
    doc.add_page_break()

    # Section 4: EDA
    add_heading('4. Exploratory Data Analysis', level=1)
    add_cards([('Numeric', str(len(payload.edaStats.numericColumns))),
               ('Categorical', str(len(payload.edaStats.categoricalColumns))),
               ('Correlations', str(len(payload.edaStats.correlations)))])
    stat_rows = [['Field', 'Mean', 'Std', 'Min', 'Median', 'Max']]
    for field in payload.edaStats.numericColumns[:10]:
        s = payload.edaStats.stats.get(field, {})
        stat_rows.append([field, fmt(s.get('mean')), fmt(s.get('std')), fmt(s.get('min')), fmt(s.get('median')), fmt(s.get('max'))])
    if len(stat_rows) > 1:
        add_table(stat_rows[0], stat_rows[1:])
    corr_rows = [[item.get('pair', 'N/A'), fmt(item.get('correlation'))] for item in payload.edaStats.correlations[:10]]
    if corr_rows:
        add_table(['Pair', 'Correlation'], corr_rows)
    doc.add_page_break()

    # Section 5: Key Statistical Findings
    add_heading('5. Key Statistical Findings', level=1)
    add_bullet(f'Dataset contains {payload.totalRows:,} rows and {len(payload.columns)} columns')
    add_bullet(f'{len(payload.edaStats.numericColumns)} numeric fields profiled')
    add_bullet(f'{len(payload.edaStats.categoricalColumns)} categorical fields analyzed')
    add_bullet(f'{len(payload.cleaningLogs)} cleaning operations executed')
    doc.add_page_break()

    # Section 6: Business Insights
    add_heading('6. Business Insights', level=1)
    completeness = (sum(c.nonNull for c in payload.columns) / max(1, sum(c.nonNull + c.nullCount for c in payload.columns)) * 100) if payload.columns else 0
    add_bullet(f'Data completeness: {completeness:.1f}%')
    add_bullet(f'{len(payload.edaStats.numericColumns)} quantifiable metrics available')
    if payload.edaStats.correlations:
        strong = len([c for c in payload.edaStats.correlations if abs(float(c.get('correlation', 0))) > 0.7])
        add_bullet(f'{strong} strong correlations detected')
    doc.add_page_break()

    # Section 7: Forecast Results
    add_heading('7. Forecast Results', level=1)
    if ts_result:
        add_heading('Time Series Forecast', level=2)
        tsm = ts_result.get('metrics') or {}
        tst = ts_result.get('training_summary') or {}
        add_cards([('MAPE', fmt(tsm.get('mape'))), ('MAE', fmt(tsm.get('mae'))),
                   ('Horizon', f'{len(ts_result.get("future_forecast", []))} periods')])
        add_table(['Period', 'Forecast', 'Lower', 'Upper'],
                  [[p.get('period', 'N/A'), fmt(p.get('predicted')), fmt(p.get('lower')), fmt(p.get('upper'))]
                   for p in ts_result.get('future_forecast', [])[:10]])
        ts_chart = _build_line_chart_base64('Time Series Forecast', ts_result.get('history', []),
                                            ts_result.get('test_forecast', []), ts_result.get('future_forecast', []), True)
        add_chart(ts_chart, 'Time Series Forecast Chart')
    if ml_result:
        add_heading('Machine Learning Forecast', level=2)
        mlm = ml_result.get('metrics') or {}
        add_cards([('MAPE', fmt(mlm.get('mape'))), ('MAE', fmt(mlm.get('mae'))),
                   ('Features', str(len(ml_result.get('generated_features', []))))])
        add_table(['Period', 'Forecast'],
                  [[p.get('period', 'N/A'), fmt(p.get('predicted'))] for p in ml_result.get('future_forecast', [])[:10]])
        ml_chart = _build_line_chart_base64('ML Forecast', ml_result.get('history', []),
                                            ml_result.get('test_forecast', []), ml_result.get('future_forecast', []), False)
        add_chart(ml_chart, 'ML Forecast Chart')
        shap = ml_result.get('shap_feature_importance', [])
        if shap:
            shap_chart = _build_bar_chart_base64('SHAP Feature Importance', shap)
            add_chart(shap_chart, 'Feature Importance')
    if loss_rows:
        add_heading('Loss Forecast', level=2)
        total_loss = sum(float(r.get('total_loss') or 0) for r in loss_rows)
        add_cards([('Total Loss', mny(total_loss))])
        loss_chart = _build_loss_chart_base64(loss_rows)
        add_chart(loss_chart, 'Loss Trend by Driver')
    if profit_rows:
        add_heading('Profit Forecast', level=2)
        total_rev = sum(float(r.get('forecasted_revenue') or 0) for r in profit_rows)
        add_cards([('Revenue', mny(total_rev))])
        profit_chart = _build_profit_chart_base64(profit_scenarios, profit_rows)
        add_chart(profit_chart, 'Net Profit Forecast by Scenario')
    doc.add_page_break()

    # Section 8: ML Results
    add_heading('8. Machine Learning Results', level=1)
    if payload.modelMetrics:
        add_cards([('Target', payload.targetColumn or 'N/A'), ('Model', payload.selectedModel or 'N/A'),
                   ('Problem', payload.problemType or 'N/A')])
        add_table(['Metric', 'Value'], [[k, fmt(v)] for k, v in payload.modelMetrics.items()])
    if payload.featureImportance:
        add_table(['Feature', 'Importance'],
                  [[item.get('name', 'N/A'), fmt(item.get('importance'))] for item in payload.featureImportance[:10]])
    if payload.predictionResult is not None:
        add_cards([('Latest Prediction', fmt(payload.predictionResult)),
                   ('History', str(len(payload.predictionHistory)))])
    doc.add_page_break()

    # Section 9: Agent Summary
    add_heading('9. Agent Summary', level=1)
    add_para(f'The Intelligent Data Assistant analyzed {payload.fileName} comprehensively.')
    add_heading('Key Discoveries', level=2)
    if payload.edaStats.correlations:
        add_bullet(f'{len(payload.edaStats.correlations)} correlation pairs evaluated')
    if payload.cleaningLogs:
        add_bullet(f'{len(payload.cleaningLogs)} data quality issues resolved')
    add_heading('Recommended Actions', level=2)
    add_bullet('Review data quality findings and apply recommended cleaning')
    if ts_result:
        add_bullet('Validate time-series forecast and develop scenario plans')
    if payload.selectedModel:
        add_bullet(f'Deploy {payload.selectedModel} for ongoing predictions')
    doc.add_page_break()

    # Section 10: Recommendations
    add_heading('10. Recommendations', level=1)
    null_cols = len([c for c in payload.columns if c.nullCount > 0])
    if null_cols > 0:
        add_bullet(f'Address missing values in {null_cols} columns')
    if payload.duplicates > 0:
        add_bullet(f'Review and remove {payload.duplicates:,} duplicate records')
    if payload.edaStats.correlations:
        add_bullet('Use correlation insights for feature engineering')
    if payload.selectedModel:
        add_bullet(f'Deploy {payload.selectedModel} with monitoring for drift detection')
    else:
        add_bullet('Train a model using the ML Assistant workflow')
    doc.add_page_break()

    # Section 11: Appendix
    add_heading('11. Appendix', level=1)
    add_table(['Field', 'Value'],
              [['Report ID', meta['report_id']], ['Generated', meta['generated_at']],
               ['Dataset', meta['dataset_name']], ['Agent Version', meta['agent_version']],
               ['Rows', f'{payload.totalRows:,}'], ['Columns', str(len(payload.columns))]])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


REPORT_FORMATS = {
    'pdf': ('application/pdf', 'pdf', build_dynamic_report_pdf),
    'html': ('text/html', 'html', build_dynamic_report_html),
    'docx': ('application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'docx', build_dynamic_report_docx),
    'doc': ('application/msword', 'doc', build_dynamic_report_doc),
}


@router.post('/report/generate')
@router.post('/generate-report')
def generate_report(payload: ReportPayload, http_request: Request, format: Literal['pdf', 'html', 'docx', 'doc'] = Query(default='pdf')) -> Response:
    fmt_info = REPORT_FORMATS.get(format)
    if not fmt_info:
        raise HTTPException(status_code=400, detail=f'Unsupported format: {format}')

    media_type, extension, builder = fmt_info
    try:
        report_bytes = builder(payload)
    except Exception as error:
        logger.exception('Report generation failed format=%s file_name=%s', format, payload.fileName)
        raise HTTPException(status_code=400, detail=f'Failed to generate {format.upper()} report: {error}') from error

    file_stem = ''.join(ch for ch in payload.fileName.rsplit('.', 1)[0] if ch.isalnum() or ch in ('-', '_', ' ')).strip() or 'dataset'
    server_session_id = get_session_id(payload.datasetId, payload.sessionId)
    record_activity(
        request=http_request,
        action='generate_report',
        status='success',
        dataset_id=payload.datasetId,
        server_session_id=server_session_id,
        file_name=payload.fileName,
        detail=f'Generated a {format.upper()} workflow report.',
        metadata={
            'format': format,
            'cleaning_done': payload.cleaningDone,
            'problem_type': payload.problemType,
            'selected_model': payload.selectedModel,
            'prediction_available': payload.predictionResult is not None,
        },
    )

    # Store temporarily for re-download
    report_id = str(uuid.uuid4())[:12]
    TEMP_REPORT_STORE[report_id] = {
        'pdf': report_bytes if format == 'pdf' else None,
        'html': report_bytes if format == 'html' else None,
        'docx': report_bytes if format in ('docx', 'doc') else None,
        'payload': payload.model_dump() if hasattr(payload, 'model_dump') else payload,
        'created_at': datetime.now().isoformat(),
    }

    return Response(
        content=report_bytes,
        media_type=media_type,
        headers={
            'Content-Disposition': f'attachment; filename="{file_stem}_analysis_report.{extension}"',
            'X-Report-Id': report_id,
        },
    )


@router.get('/forecast/export/excel/{session_id}')
def export_forecast_excel(session_id: str, http_request: Request) -> Response:
    state = ensure_session_state(session_id)
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
        sheets = {
            'ts_history': (state.get('time_series_result') or {}).get('history', []),
            'ts_future': (state.get('time_series_result') or {}).get('future_forecast', []),
            'ts_model_comparison': (state.get('time_series_result') or {}).get('model_comparison', []),
            'ml_history': (state.get('ml_forecast_result') or {}).get('history', []),
            'ml_future': (state.get('ml_forecast_result') or {}).get('future_forecast', []),
            'ml_model_comparison': (state.get('ml_forecast_result') or {}).get('model_comparison', []),
            'loss_forecast': state.get('loss_forecast_result') or [],
            'profit_baseline': (state.get('profit_scenarios') or {}).get('baseline', []),
            'forecast_versions': state.get('forecast_history_versions') or [],
        }
        for sheet_name, rows in sheets.items():
            serialized_rows = safe_serialize(rows)
            if sheet_name.endswith('model_comparison') and isinstance(serialized_rows, list):
                serialized_rows = [
                    {
                        **{key: value for key, value in row.items() if key not in {'metrics', 'tuning'}},
                        'mae': (row.get('metrics') or {}).get('mae'),
                        'rmse': (row.get('metrics') or {}).get('rmse'),
                        'mape': (row.get('metrics') or {}).get('mape'),
                        'optuna_enabled': (row.get('tuning') or {}).get('enabled'),
                        'tuning_note': (row.get('tuning') or {}).get('note'),
                    }
                    for row in serialized_rows
                    if isinstance(row, dict)
                ]
            frame = pd.DataFrame(serialized_rows)
            if frame.empty:
                frame = pd.DataFrame([{'status': 'not_available'}])
            frame.to_excel(writer, sheet_name=sheet_name[:31], index=False)
        try:
            from openpyxl.chart import BarChart, LineChart, Reference
        except Exception as error:  # pragma: no cover - dependency is declared and validated at runtime
            raise HTTPException(status_code=500, detail='openpyxl chart support is required to build the forecast workbook charts.') from error

        workbook = writer.book
        chart_sheet = workbook.create_sheet('charts')
        chart_sheet['A1'] = 'Forecast Workbook Charts'
        chart_sheet['A2'] = 'Charts are generated from the raw forecast sheets in this workbook.'

        def numeric_chart(
            source_sheet: str,
            title: str,
            value_columns: list[str],
            anchor: str,
            *,
            chart_type: str = 'line',
        ) -> None:
            if source_sheet not in workbook.sheetnames:
                return
            worksheet = workbook[source_sheet]
            if worksheet.max_row < 3:
                return
            headers = [str(cell.value) if cell.value is not None else '' for cell in worksheet[1]]
            if 'period' in headers:
                category_column = headers.index('period') + 1
            elif 'model_name' in headers:
                category_column = headers.index('model_name') + 1
            else:
                category_column = 1
            chart = BarChart() if chart_type == 'bar' else LineChart()
            chart.title = title
            chart.y_axis.title = 'Value'
            chart.x_axis.title = headers[category_column - 1] or 'Category'
            added_series = False
            for column_name in value_columns:
                if column_name not in headers:
                    continue
                column_index = headers.index(column_name) + 1
                values = Reference(worksheet, min_col=column_index, min_row=1, max_row=worksheet.max_row)
                chart.add_data(values, titles_from_data=True)
                added_series = True
            if not added_series:
                return
            categories = Reference(worksheet, min_col=category_column, min_row=2, max_row=worksheet.max_row)
            chart.set_categories(categories)
            chart.height = 8
            chart.width = 18
            chart_sheet.add_chart(chart, anchor)

        numeric_chart('ts_future', 'Time Series Future Forecast', ['predicted', 'lower', 'upper'], 'A4')
        numeric_chart('ml_future', 'Machine Learning Future Forecast', ['predicted', 'lower', 'upper'], 'J4')
        numeric_chart('loss_forecast', 'Loss Forecast by Driver', ['revenue_loss', 'operational_loss', 'inventory_loss', 'discount_loss', 'total_loss'], 'A22')
        numeric_chart('profit_baseline', 'Baseline Profit Forecast', ['forecasted_revenue', 'gross_profit', 'net_profit'], 'J22')
        numeric_chart('ts_model_comparison', 'Time Series Model MAE Comparison', ['mae', 'rmse', 'mape'], 'A40', chart_type='bar')
        numeric_chart('ml_model_comparison', 'ML Forecast Model MAE Comparison', ['mae', 'rmse', 'mape'], 'J40', chart_type='bar')

    buffer.seek(0)
    record_activity(
        request=http_request,
        action='export_forecast_excel',
        status='success',
        dataset_id=session_id,
        server_session_id=session_id,
        detail='Exported forecast workbook with raw numbers, comparisons, and version history.',
    )
    return Response(
        content=buffer.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename="{session_id}_forecast_workbook.xlsx"'},
    )


@router.get('/activities')
def list_activities(
    request: Request,
    dataset_id: str | None = Query(default=None),
    client_session_id: str | None = Query(default=None),
    server_session_id: str | None = Query(default=None),
    limit: int = Query(default=100, ge=1, le=500),
) -> JSONResponse:
    if not ACTIVITY_DB_AVAILABLE:
        return JSONResponse(
            content={
                'activities': [],
                'count': 0,
                'dbAvailable': False,
                'message': 'Activity database is unavailable. Start PostgreSQL to enable persisted activity history.',
            }
        )

    effective_client_session_id = client_session_id or get_client_session_id(request)
    query = '''
        SELECT
            activity_id,
            created_at,
            client_session_id,
            server_session_id,
            dataset_id,
            model_id,
            activity_type,
            action,
            status,
            api_path,
            http_method,
            status_code,
            duration_ms,
            file_name,
            detail,
            metadata_json
        FROM user_activities
        WHERE 1 = 1
    '''
    params: list[Any] = []

    if dataset_id:
        query += ' AND dataset_id = %s'
        params.append(dataset_id)
    if effective_client_session_id:
        query += ' AND client_session_id = %s'
        params.append(effective_client_session_id)
    if server_session_id:
        query += ' AND server_session_id = %s'
        params.append(server_session_id)

    query += ' ORDER BY id DESC LIMIT %s'
    params.append(limit)

    try:
        with get_activity_connection() as connection:
            rows = connection.execute(query, params).fetchall()
    except Exception:
        logger.exception('Failed to query user activities.')
        return JSONResponse(
            content={
                'activities': [],
                'count': 0,
                'dbAvailable': False,
                'message': 'Activity database query failed. Start PostgreSQL to enable persisted activity history.',
            }
        )

    activities = []
    for row in rows:
        metadata_json = row['metadata_json']
        activities.append({
            'activityId': row['activity_id'],
            'createdAt': row['created_at'],
            'clientSessionId': row['client_session_id'],
            'serverSessionId': row['server_session_id'],
            'datasetId': row['dataset_id'],
            'modelId': row['model_id'],
            'activityType': row['activity_type'],
            'action': row['action'],
            'status': row['status'],
            'apiPath': row['api_path'],
            'httpMethod': row['http_method'],
            'statusCode': row['status_code'],
            'durationMs': row['duration_ms'],
            'fileName': row['file_name'],
            'detail': row['detail'],
            'metadata': json.loads(metadata_json) if metadata_json else None,
        })

    return JSONResponse(content={'activities': activities, 'count': len(activities), 'dbAvailable': True})


@router.post('/auth/register')
def register_user(payload: RegisterRequest, request: Request) -> JSONResponse:
    try:
        user = create_app_user(username=payload.username, email=payload.email, password=payload.password)
        _, session_token = create_authenticated_session(user_id=user['userId'], request=request)
    except HTTPException:
        raise
    except Exception as error:
        logger.exception('User registration failed email=%s', payload.email)
        raise HTTPException(status_code=500, detail=f'Failed to register user: {error}') from error

    response = JSONResponse(content={'user': user})
    set_session_cookie(response, session_token)
    record_activity(
        request=request,
        action='user_register',
        status='success',
        activity_type='auth',
        detail=f'User {user["email"]} registered.',
        metadata={
            'user_id': user['userId'],
            'email': user['email'],
            'username': user['username'],
        },
    )
    return response


@router.post('/auth/login')
def login_user(payload: LoginRequest, request: Request) -> JSONResponse:
    try:
        user = authenticate_user(email=payload.email, password=payload.password)
        _, session_token = create_authenticated_session(user_id=user['userId'], request=request)
    except HTTPException:
        raise
    except Exception as error:
        logger.exception('User login failed email=%s', payload.email)
        raise HTTPException(status_code=500, detail=f'Failed to login user: {error}') from error

    response = JSONResponse(content={'user': user})
    set_session_cookie(response, session_token)
    record_activity(
        request=request,
        action='user_login',
        status='success',
        activity_type='auth',
        detail=f'User {user["email"]} signed in.',
        metadata={
            'user_id': user['userId'],
            'email': user['email'],
            'username': user['username'],
        },
    )
    return response


@router.get('/auth/me')
def auth_me(request: Request) -> JSONResponse:
    try:
        user = build_user_payload(get_authenticated_user(request))
    except HTTPException:
        raise
    except Exception as error:
        logger.exception('Auth me failed.')
        raise HTTPException(status_code=500, detail=f'Failed to resolve current session: {error}') from error

    return JSONResponse(content={'user': user})


@router.put('/auth/profile')
async def update_user_profile(
    request: Request,
    username: str = Form(...),
    email: str = Form(...),
    profile_image: UploadFile | None = File(default=None),
) -> JSONResponse:
    try:
        user = await update_authenticated_user_profile(
            request=request,
            username=username,
            email=email,
            profile_image=profile_image,
        )
    except HTTPException:
        raise
    except Exception as error:
        logger.exception('Profile update failed.')
        raise HTTPException(status_code=500, detail=f'Failed to update profile: {error}') from error

    record_activity(
        request=request,
        action='user_profile_update',
        status='success',
        activity_type='auth',
        detail=f'Updated profile for {user["email"]}.',
        metadata={
            'user_id': user['userId'],
            'email': user['email'],
            'username': user['username'],
            'profile_image_updated': profile_image is not None,
        },
    )
    return JSONResponse(content={'user': user})


@router.post('/auth/logout')
def logout_user(request: Request) -> JSONResponse:
    session_token = request.cookies.get(SESSION_COOKIE_NAME)
    if session_token:
        try:
            revoke_session(session_token)
        except Exception:
            logger.exception('Failed to revoke user session during logout.')

    response = JSONResponse(content={'success': True})
    clear_session_cookie(response)
    record_activity(
        request=request,
        action='user_logout',
        status='success',
        activity_type='auth',
        detail='User logged out.',
    )
    return response


@router.get('/health')
def health() -> dict[str, str]:
    return {
        'status': 'healthy',
        'activityDb': 'available' if ACTIVITY_DB_AVAILABLE else 'unavailable',
    }


@app.get('/')
def root() -> dict[str, Any]:
    return {'service': 'AI-Assisted EDA & ML Backend', 'docs': '/docs', 'api': '/api'}


# AGENTIC LAYER START
from agentic.agentic_adapter import agentic_router

app.include_router(agentic_router)
# AGENTIC LAYER END
app.include_router(router)


if __name__ == '__main__':
    import uvicorn

    port = int(os.environ.get('ML_PORT', '3004'))
    uvicorn.run(app, host='0.0.0.0', port=port)




